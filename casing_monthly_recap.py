#!/usr/bin/env python3
"""
CASING DIVISION — MONTHLY HSE RECAP
=====================================
Generates a yard-focused monthly HSE recap for Casing Division management.

Sections:
1. Executive Summary (headline KPIs with MoM arrows)
2. Fleet Mileage by Yard (Motive IFTA trips)
3. Pre/Post Trip Inspections (flagged items)
4. Speeding Summary (Motive speeding events)
5. Camera Events Summary (Motive AI dashcams)
6. Field Assessments (form 381707 detail)
7. Observations (shared HSE form, Casing-filtered)
8. Incident Analysis (classified by OSHA type + RCA accountability)
9. Training Compliance (employee table with incomplete programs)
10. Yard Comparison Scorecard
11. Action Items & Takeaways

Data sources:
- Motive API v1: vehicles, IFTA trips, speeding events
- Motive API v2: camera events (driver performance events)
- KPA API: 21 active CSG forms + shared HSE forms filtered to Casing

Usage:
    python casing_monthly_recap.py --month 2026-02

Output: Casing_Monthly_HSE_Recap_YYYY-MM.docx
"""

import argparse
import calendar
import csv
import json
import os
import sys
import time
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from io import StringIO

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

try:
    from zoneinfo import ZoneInfo
    CENTRAL_TZ = ZoneInfo("America/Chicago")
except Exception:
    CENTRAL_TZ = timezone(timedelta(hours=-6))

# ==============================================================================
# CONFIGURATION
# ==============================================================================

MOTIVE_API_KEY = os.environ.get("MOTIVE_API_KEY")
KPA_API_TOKEN = os.environ.get("KPA_API_TOKEN")
MOTIVE_BASE_V1 = "https://api.gomotive.com/v1"
MOTIVE_BASE_V2 = "https://api.gomotive.com/v2"
KPA_BASE_URL = "https://api.kpaehs.com/v1"
KMH_TO_MPH = 0.621371

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGOS_DIR = os.path.join(SCRIPT_DIR, "logos")
DATA_DIR = os.path.join(SCRIPT_DIR, "casing_monthly_data")

# Branding — Butch's red/black
DARK_RED = RGBColor(0x8B, 0, 0)         # #8B0000
GREEN = RGBColor(0, 0x80, 0)            # #008000
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x66, 0x66, 0x66)       # #666666 — readable on print
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = "F2F2F2"
CALIBRI = "Calibri"

# Alignment shortcuts for table columns
_L = WD_ALIGN_PARAGRAPH.LEFT
_C = WD_ALIGN_PARAGRAPH.CENTER
_R = WD_ALIGN_PARAGRAPH.RIGHT

# ==============================================================================
# CASING YARDS (from Motive groups)
# ==============================================================================

CASING_GROUP_IDS = {
    167175: "Midland",
    169090: "Bryan",
    169092: "Kilgore",
    186740: "Hobbs",
    169091: "Jourdanton",
    186739: "Laredo",
    186746: "",  # Parent "Casing" group
}

ALL_CASING_GROUP_IDS = set(CASING_GROUP_IDS.keys())

YARD_ORDER = ["Midland", "Bryan", "Kilgore", "Hobbs", "Jourdanton", "Laredo"]

YARD_INFO = {
    "Midland": {"safety_reps": "Michael Hancock & Michael Salazar", "manager": "Richie Bentley"},
    "Bryan": {"safety_reps": "Justin Conrad", "manager": "Danny Lohse"},
    "Kilgore": {"safety_reps": "James Barnett (J.P.)", "manager": "Frankie Balderas"},
    "Hobbs": {"safety_reps": "Allen Batts", "manager": "Clifton Eaves"},
    "Jourdanton": {"safety_reps": "Joey Speyrer", "manager": "Enrique Flores"},
    "Laredo": {"safety_reps": "Joey Speyrer", "manager": "Chris Jacobo"},
}

# KPA Line-of-Business ID for Casing
CASING_LOB_ID = "6009f696823a6201bbc9b056"

# Man hours baseline (from 2025 Q1 payroll export -- Jan+Feb avg)
MONTHLY_MAN_HOURS_BY_YARD = {
    "Midland": 77_921,
    "Kilgore": 22_002,
    "Hobbs": 21_488,
    "Jourdanton": 18_745,
    "Bryan": 16_107,
    "Laredo": 13_896,
}
MONTHLY_MAN_HOURS = 183_990  # total across all yards

# KPA headcount by yard (active employees)
HEADCOUNT_BY_YARD = {
    "Midland": 143,
    "Kilgore": 61,
    "Hobbs": 54,
    "Jourdanton": 46,
    "Bryan": 46,
    "Laredo": 36,
}
EMPLOYEE_COUNT = 400  # total active

# Casing-specific Non-DOT Pre/Post Trip form
VEHICLE_INSPECTION_FORM = 229645

# Observation detail hash IDs (from shared HSE observation form 151085)
OBS_TYPE_HASH = "bff8m4x6xbc033kg"
OBS_DESC_HASH = "uncbcge9x8vow9pn"
OBS_LOCATION_HASH = "lg5pnj4chjadnv46"
OBS_NAME_HASH = "fcxf2kaly9s04csq"       # Employee Name (not observer/submitter)


def _obs_employee(row):
    """Get the actual employee name from an observation row.
    Prefers the Name field (employee observed) over observer (submitter)."""
    name = row.get(OBS_NAME_HASH, "").strip()
    if name and name.lower() != "name":
        return name
    return row.get("observer", row.get("_observer", "Unknown"))

# Incident form 151622 hash IDs
INC_TYPE_HASH = "nojcquy0tfl9hqih"        # Incident Type (multi-select)
INC_DESC_HASH = "313e9txgrof0uute"        # Incident Description
INC_DATE_HASH = "eo0bl6w5ouq06s86"        # Date/Time Occurred
INC_EMPLOYEE_HASH = "g5w4b0uh15wxqykt"    # Employee
INC_YARD_HASH = "pk6qj0kiu9vek20v"        # District
INC_LOCATION_HASH = "9ohdd2lwvl7p0oc6"    # Rig/Well/Location

# RCA form 180243 hash IDs
RCA_FORM = 180243
RCA_SERVICE_LINE_HASH = "tqqdm1br3v0kgfq0"
RCA_INC_DATE_HASH = "14ptd2snr25dtm52"
RCA_COMPLETE_DATE_HASH = "xx25cruwjxepechp"
RCA_DESC_HASH = "vfrkwqi09yqjtjd2"
RCA_CAUSES_HASH = "ljnvmnhu1k6k2qu1"
RCA_ACTIONS_HASH = "3jpailc2332vwy1k"
RCA_YARD_HASH = "0pfengduarubztqh"

# ==============================================================================
# KPA FORM IDS — Active in both Jan + Feb 2026
# ==============================================================================

# CSG-specific forms (Casing folder in KPA)
CSG_ACTIVE_FORMS = {
    169922: "PM -- IOI -- Bails",
    179551: "CSG - SSE Hands-on Simulator Evaluation Form",
    187860: "CSG - Workplace Inspection Checklist",
    194580: "CSG - Truck Inspection for Crew Haulers",
    199992: "MIDLAND CSG - Repair Form",
    202292: "CSG - Shop Load Out",
    225522: "CSG - Management Field Audit Form",
    239074: "PM - Daily Forklift Inspection Checklist",
    240113: "PM -- IOI -- Bowls",
    240496: "PM -- IOI -- Spiders",
    240521: "PM -- IOI -- Slips",
    249382: "PM -- Manual Backup Tongs",
    265439: "CSG - TRANSFER FORM",
    329638: "CSG - Supervisor Checklist",
    350996: "PM - Vehicle Service Record",
    381707: "CSG - Safety Casing Field Assessment",
    385265: "CSG - CRT Job Loadout Sheet (TEMP)",
    388792: "CSG - Vendor Pick List",
    472040: "CSG - Backing - Practical Evaluation Checklist",
    476428: "CSG - Monthly Bump Test Form",
}

# Shared HSE forms — pull all, filter to Casing
SHARED_HSE_FORMS = {
    151085: "HSE - Observation Card",
    151622: "HSE - Incident Reporting",
}

FIELD_ASSESSMENT_FORM = 381707
OBSERVATION_FORM = 151085
INCIDENT_FORM = 151622

# ==============================================================================
# CAMERA EVENT CLASSIFICATION (from daily_casing_camera_report.py)
# ==============================================================================

EVENT_TYPE_NORMALIZE = {
    "distraction": "distraction", "distracted_driving": "distraction",
    "driver_distraction": "distraction",
    "cell_phone": "cell_phone", "cell_phone_usage": "cell_phone",
    "phone_use": "cell_phone", "cellphone": "cell_phone",
    "drowsiness": "drowsiness", "drowsy": "drowsiness",
    "drowsy_driving": "drowsiness", "fatigue": "drowsiness",
    "driver_drowsiness": "drowsiness",
    "close_following": "close_following", "following_distance": "close_following",
    "tailgating": "close_following",
    "forward_collision_warning": "forward_collision_warning",
    "forward_collision": "forward_collision_warning",
    "collision": "collision", "crash": "collision",
    "near_collision": "near_collision", "near_crash": "near_collision",
    "stop_sign_violation": "stop_sign_violation", "stop_sign": "stop_sign_violation",
    "unsafe_lane_change": "unsafe_lane_change", "lane_change": "unsafe_lane_change",
    "aggregated_lane_swerving": "lane_swerving", "lane_swerving": "lane_swerving",
    "hard_brake": "hard_brake", "hard_braking": "hard_brake",
    "harsh_brake": "hard_brake",
    "seat_belt_violation": "seat_belt_violation", "seatbelt": "seat_belt_violation",
    "seatbelt_violation": "seat_belt_violation", "no_seatbelt": "seat_belt_violation",
    "seat_belt": "seat_belt_violation",
    "camera_obstruction": "camera_obstruction", "obstruction": "camera_obstruction",
    "driver_facing_cam_obstruction": "camera_obstruction",
    "road_facing_cam_obstruction": "camera_obstruction",
    "smoking": "smoking", "vaping": "smoking",
    "hard_accel": "hard_accel", "hard_acceleration": "hard_accel",
    "harsh_acceleration": "hard_accel",
    "hard_corner": "hard_corner", "hard_cornering": "hard_corner",
    "hard_turn": "hard_corner",
    "speed_violation": "speed_violation", "speeding": "speed_violation",
}

RED_TYPES = {
    "distraction", "cell_phone", "drowsiness", "close_following",
    "forward_collision_warning", "collision", "near_collision",
    "stop_sign_violation", "unsafe_lane_change", "lane_swerving",
}
ORANGE_TYPES = {
    "hard_brake", "seat_belt_violation", "camera_obstruction", "smoking",
}
YELLOW_TYPES = {
    "hard_accel", "hard_corner", "speed_violation",
}

EVENT_DISPLAY_NAMES = {
    "distraction": "Distraction", "cell_phone": "Cell Phone",
    "drowsiness": "Drowsiness", "close_following": "Close Following",
    "forward_collision_warning": "Forward Collision Warning",
    "collision": "Collision", "near_collision": "Near Collision",
    "stop_sign_violation": "Stop Sign Violation",
    "unsafe_lane_change": "Unsafe Lane Change", "lane_swerving": "Lane Swerving",
    "hard_brake": "Hard Brake", "seat_belt_violation": "Seatbelt Violation",
    "camera_obstruction": "Camera Obstruction", "smoking": "Smoking",
    "hard_accel": "Hard Acceleration", "hard_corner": "Hard Corner",
    "speed_violation": "Speed Violation",
}


def _normalize_event_type(raw_type):
    if not raw_type:
        return "unknown"
    key = raw_type.lower().strip().replace(" ", "_").replace("-", "_")
    return EVENT_TYPE_NORMALIZE.get(key, key)


def _classify_tier(event_type):
    if event_type in RED_TYPES:
        return "RED"
    elif event_type in ORANGE_TYPES:
        return "ORANGE"
    elif event_type in YELLOW_TYPES:
        return "YELLOW"
    return "ORANGE"


# ==============================================================================
# DATE HELPERS
# ==============================================================================

def parse_month(month_str):
    dt = datetime.strptime(month_str, "%Y-%m")
    year, month = dt.year, dt.month
    last_day = calendar.monthrange(year, month)[1]
    start_ct = datetime(year, month, 1, 0, 0, 0, tzinfo=CENTRAL_TZ)
    end_ct = datetime(year, month, last_day, 23, 59, 59, tzinfo=CENTRAL_TZ)
    return start_ct, end_ct


def month_label(start_ct, end_ct):
    return f"{start_ct.strftime('%B')} {start_ct.day} \u2013 {end_ct.strftime('%B')} {end_ct.day}, {end_ct.year}"


def prev_month_str(month_str):
    dt = datetime.strptime(month_str, "%Y-%m")
    if dt.month == 1:
        return f"{dt.year - 1}-12"
    return f"{dt.year}-{dt.month - 1:02d}"


def yoy_month_str(month_str):
    """Return same month from the prior year (e.g. '2026-03' -> '2025-03')."""
    dt = datetime.strptime(month_str, "%Y-%m")
    return f"{dt.year - 1}-{dt.month:02d}"


def _format_display_date(raw):
    """Format a raw date string to human-readable (e.g. 'Mar 23, 2026')."""
    if not raw or not isinstance(raw, str):
        return raw or ""
    raw = raw.strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%m/%d/%Y", "%Y-%m-%d",
                "%m/%d/%Y %I:%M %p", "%m/%d/%Y %H:%M %p"):
        try:
            dt = datetime.strptime(raw.split(".")[0], fmt)
            return dt.strftime("%b %d, %Y")
        except ValueError:
            continue
    return raw


def mom_delta(current, previous):
    if previous is None:
        return (current, "N/A", "")
    if previous == 0 and current == 0:
        return (0, "--", "")
    if previous == 0:
        return (current, "\u2191 New", "\u2191")
    change = current - previous
    pct = (change / abs(previous)) * 100
    if change == 0:
        return (0, "--", "")
    arrow = "\u2191" if change > 0 else "\u2193"
    pct_str = f"{arrow} {abs(pct):.1f}%"
    return (change, pct_str, arrow)


# ==============================================================================
# MOTIVE API — VEHICLES
# ==============================================================================

def get_casing_vehicles():
    """Fetch all Casing vehicles from Motive. Returns dict {vehicle_number: vehicle_data}."""
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    vehicles = {}
    page = 1

    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_V1}/vehicles",
                headers=headers,
                params={"per_page": 100, "page_no": page},
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            vlist = data.get("vehicles", [])
            if not vlist:
                break

            for wrapper in vlist:
                v = wrapper.get("vehicle", wrapper)
                num = v.get("number", "")
                if not num:
                    continue
                group_ids = set(v.get("group_ids", []))
                if not group_ids.intersection(ALL_CASING_GROUP_IDS):
                    continue

                # Determine yard
                yard = "Unassigned"
                for gid in group_ids:
                    y = CASING_GROUP_IDS.get(gid, "")
                    if y:
                        yard = y
                        break

                # Driver name
                driver_name = None
                for field in ("current_driver", "permanent_driver"):
                    d = v.get(field)
                    if d and isinstance(d, dict):
                        name = f"{d.get('first_name', '')} {d.get('last_name', '')}".strip()
                        if name:
                            driver_name = name
                            break

                vehicles[num] = {
                    "id": v.get("id"),
                    "number": num,
                    "driver": driver_name or "Unassigned",
                    "yard": yard,
                    "group_ids": list(group_ids),
                    "make": v.get("make", "") or "",
                    "model": v.get("model", "") or "",
                }

            pag = data.get("pagination", {})
            if page * 100 >= pag.get("total", 0):
                break
            page += 1
        except Exception as e:
            print(f"  Warning: vehicle page {page} failed: {e}")
            break

    print(f"  Found {len(vehicles)} Casing vehicles")
    for yard in YARD_ORDER:
        count = sum(1 for v in vehicles.values() if v["yard"] == yard)
        if count:
            print(f"    {yard}: {count}")
    return vehicles


# ==============================================================================
# MOTIVE API — MILEAGE (IFTA)
# ==============================================================================

def get_casing_mileage(vehicles, start_ct, end_ct):
    """Pull IFTA trips for Casing vehicles, group by yard."""
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    vehicle_miles = defaultdict(float)
    page = 1

    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_V1}/ifta/trips",
                headers=headers,
                params={
                    "per_page": 100,
                    "page_no": page,
                    "start_date": start_ct.strftime("%Y-%m-%d"),
                    "end_date": end_ct.strftime("%Y-%m-%d"),
                },
                timeout=60,
            )
            resp.raise_for_status()
            data = resp.json()
            trips = data.get("ifta_trips", [])
            if not trips:
                break

            for wrapper in trips:
                trip = wrapper.get("ifta_trip", wrapper)
                vehicle = trip.get("vehicle", {})
                vnum = vehicle.get("number", "") if isinstance(vehicle, dict) else str(vehicle)
                if vnum in vehicles:
                    vehicle_miles[vnum] += trip.get("distance", 0) or 0

            pag = data.get("pagination", {})
            if page * 100 >= pag.get("total", 0):
                break
            page += 1
        except Exception as e:
            print(f"  Warning: IFTA trips page {page} failed: {e}")
            # Retry once after a short pause
            time.sleep(5)
            try:
                resp = requests.get(
                    f"{MOTIVE_BASE_V1}/ifta/trips",
                    headers=headers,
                    params={
                        "per_page": 100,
                        "page_no": page,
                        "start_date": start_ct.strftime("%Y-%m-%d"),
                        "end_date": end_ct.strftime("%Y-%m-%d"),
                    },
                    timeout=60,
                )
                resp.raise_for_status()
                data = resp.json()
                trips = data.get("ifta_trips", [])
                if not trips:
                    break
                for wrapper in trips:
                    trip = wrapper.get("ifta_trip", wrapper)
                    vehicle = trip.get("vehicle", {})
                    vnum = vehicle.get("number", "") if isinstance(vehicle, dict) else str(vehicle)
                    if vnum in vehicles:
                        vehicle_miles[vnum] += trip.get("distance", 0) or 0
                pag = data.get("pagination", {})
                if page * 100 >= pag.get("total", 0):
                    break
                page += 1
                print(f"  Retry succeeded for page {page - 1}")
            except Exception as e2:
                print(f"  Retry also failed: {e2}")
                break

    # Group by yard
    by_yard = defaultdict(lambda: {"miles": 0, "trucks": 0})
    total_miles = 0
    for vnum, miles in vehicle_miles.items():
        if miles > 0:
            yard = vehicles.get(vnum, {}).get("yard", "Unassigned")
            by_yard[yard]["miles"] += miles
            by_yard[yard]["trucks"] += 1
            total_miles += miles

    for yard in by_yard:
        yd = by_yard[yard]
        yd["avg"] = round(yd["miles"] / yd["trucks"], 1) if yd["trucks"] else 0
        yd["miles"] = round(yd["miles"], 1)

    active_trucks = sum(d["trucks"] for d in by_yard.values())
    result = {
        "total_miles": round(total_miles, 1),
        "active_trucks": active_trucks,
        "avg_miles": round(total_miles / active_trucks, 1) if active_trucks else 0,
        "by_yard": dict(by_yard),
    }
    print(f"  Fleet mileage: {result['total_miles']:,.0f} miles across {active_trucks} trucks")
    return result


# ==============================================================================
# MOTIVE API — SPEEDING
# ==============================================================================

def get_casing_speeding(vehicles, start_ct, end_ct):
    """Pull speeding events for Casing vehicles, group by yard."""
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    start_utc = start_ct.astimezone(timezone.utc)
    end_utc = end_ct.astimezone(timezone.utc)

    raw_events = []
    page = 1
    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_V1}/speeding_events",
                headers=headers,
                params={
                    "per_page": 100,
                    "page_no": page,
                    "start_date": start_utc.strftime("%Y-%m-%d"),
                    "end_date": end_utc.strftime("%Y-%m-%d"),
                },
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            events = data.get("speeding_events", [])
            if not events:
                break
            raw_events.extend(events)
            if page * 100 >= data.get("total", 0):
                break
            page += 1
        except Exception as e:
            print(f"  Warning: speeding page {page} failed: {e}")
            break

    filtered = []
    for wrapper in raw_events:
        evt = wrapper.get("speeding_event", wrapper)
        vehicle_obj = evt.get("vehicle", {})
        vnum = vehicle_obj.get("number", "") if isinstance(vehicle_obj, dict) else str(vehicle_obj)
        if vnum not in vehicles:
            continue

        # Time window check
        evt_time_str = evt.get("start_time", "")
        try:
            evt_utc = datetime.fromisoformat(evt_time_str.replace("Z", "+00:00"))
            evt_central = evt_utc.astimezone(CENTRAL_TZ)
            if not (start_ct <= evt_central <= end_ct):
                continue
            date_str = evt_central.strftime("%m/%d/%Y %I:%M %p")
        except Exception:
            date_str = evt_time_str

        max_over_kmh = evt.get("max_over_speed_in_kph") or evt.get("avg_over_speed_in_kph") or 0
        max_exceeded = round(max_over_kmh * KMH_TO_MPH, 1)
        max_speed_kmh = evt.get("max_vehicle_speed") or 0
        max_speed = round(max_speed_kmh * KMH_TO_MPH, 1)

        if max_exceeded >= 20 or max_speed >= 90:
            severity = "Critical"
        elif max_exceeded >= 15:
            severity = "High"
        else:
            severity = "Medium"

        driver_name = vehicles.get(vnum, {}).get("driver", "Unknown")
        drv = evt.get("driver")
        if (not driver_name or driver_name == "Unassigned") and drv and isinstance(drv, dict):
            name = f"{drv.get('first_name', '')} {drv.get('last_name', '')}".strip()
            if name:
                driver_name = name

        filtered.append({
            "driver": driver_name,
            "vehicle": vnum,
            "yard": vehicles.get(vnum, {}).get("yard", "Unassigned"),
            "date": date_str,
            "severity": severity,
            "max_speed": max_speed,
            "max_exceeded": max_exceeded,
        })

    print(f"  Casing speeding events: {len(filtered)} (from {len(raw_events)} total)")
    return filtered


def process_speeding(events):
    """Summarize speeding events."""
    if not events:
        return {"total": 0, "critical": 0, "high": 0, "medium": 0,
                "by_yard": {}, "top_drivers": []}

    total = len(events)
    critical = sum(1 for e in events if e["severity"] == "Critical")
    high = sum(1 for e in events if e["severity"] == "High")
    medium = total - critical - high

    by_yard = defaultdict(int)
    by_yard_max_over = defaultdict(float)
    driver_counts = Counter()
    for e in events:
        by_yard[e["yard"]] += 1
        if e["max_exceeded"] > by_yard_max_over[e["yard"]]:
            by_yard_max_over[e["yard"]] = e["max_exceeded"]
        driver_counts[e["driver"]] += 1

    # Separate named drivers from "Unassigned" -- show named first, then unassigned
    unassigned_count = driver_counts.pop("Unassigned", 0)
    top_drivers = []
    for driver, count in driver_counts.most_common(5):
        yard = next((e["yard"] for e in events if e["driver"] == driver), "?")
        max_over = max(e["max_exceeded"] for e in events if e["driver"] == driver)
        top_drivers.append({"name": driver, "yard": yard, "events": count, "max_over": max_over})

    # Track unassigned separately for footnote
    unassigned_info = None
    if unassigned_count > 0:
        max_over_unassigned = max((e["max_exceeded"] for e in events if e["driver"] == "Unassigned"), default=0)
        unassigned_info = {"events": unassigned_count, "max_over": max_over_unassigned}

    return {
        "total": total, "critical": critical, "high": high, "medium": medium,
        "by_yard": dict(by_yard), "by_yard_max_over": dict(by_yard_max_over),
        "top_drivers": top_drivers,
        "unassigned": unassigned_info,
    }


# ==============================================================================
# MOTIVE API — CAMERA EVENTS
# ==============================================================================

def get_casing_camera_events(vehicles, start_ct, end_ct):
    """Pull camera events for Casing vehicles from Motive v2 API."""
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    start_utc = start_ct.astimezone(timezone.utc)
    end_utc = end_ct.astimezone(timezone.utc)

    raw_events = []
    cursor = None
    while True:
        try:
            params = {
                "per_page": 100,
                "start_date": start_utc.strftime("%Y-%m-%dT%H:%M:%SZ"),
                "end_date": end_utc.strftime("%Y-%m-%dT%H:%M:%SZ"),
            }
            if cursor:
                params["page_cursor"] = cursor

            resp = requests.get(
                f"{MOTIVE_BASE_V2}/driver_performance_events",
                headers=headers,
                params=params,
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            events = data.get("driver_performance_events", [])
            if not events:
                break

            raw_events.extend(events)

            pag = data.get("pagination", {})
            cursor = pag.get("next_cursor")
            if not cursor:
                break
        except Exception as e:
            print(f"  Warning: camera events failed: {e}")
            break

    filtered = []
    for wrapper in raw_events:
        evt = wrapper.get("driver_performance_event", wrapper)
        vehicle = evt.get("vehicle", {})
        vnum = vehicle.get("number", "") if isinstance(vehicle, dict) else ""
        if vnum not in vehicles:
            continue

        raw_type = evt.get("type", "") or evt.get("event_type", "") or ""
        event_type = _normalize_event_type(raw_type)
        tier = _classify_tier(event_type)

        driver = evt.get("driver", {})
        if isinstance(driver, dict):
            driver_name = f"{driver.get('first_name', '')} {driver.get('last_name', '')}".strip()
        else:
            driver_name = vehicles.get(vnum, {}).get("driver", "Unknown")

        speed_kmh = evt.get("speed") or evt.get("vehicle_speed") or 0
        speed_mph = round(speed_kmh * KMH_TO_MPH, 1) if speed_kmh else 0

        filtered.append({
            "driver": driver_name or "Unknown",
            "vehicle": vnum,
            "yard": vehicles.get(vnum, {}).get("yard", "Unassigned"),
            "event_type": event_type,
            "tier": tier,
            "display_name": EVENT_DISPLAY_NAMES.get(event_type, raw_type.replace("_", " ").title()),
            "speed": speed_mph,
        })

    print(f"  Casing camera events: {len(filtered)} (from {len(raw_events)} total)")
    return filtered


def process_camera_events(events):
    """Summarize camera events."""
    if not events:
        return {"total": 0, "red": 0, "orange": 0, "yellow": 0,
                "by_yard": {}, "by_type": {}, "repeat_offenders": []}

    total = len(events)
    red = sum(1 for e in events if e["tier"] == "RED")
    orange = sum(1 for e in events if e["tier"] == "ORANGE")
    yellow = total - red - orange

    by_yard = defaultdict(lambda: {"red": 0, "orange": 0, "yellow": 0, "total": 0})
    by_type = Counter()
    driver_counts = Counter()
    for e in events:
        by_yard[e["yard"]][e["tier"].lower()] += 1
        by_yard[e["yard"]]["total"] += 1
        by_type[e["display_name"]] += 1
        driver_counts[e["driver"]] += 1

    # Track unassigned camera events separately
    cam_unassigned = driver_counts.pop("Unassigned", 0)
    cam_unassigned_info = {"events": cam_unassigned} if cam_unassigned > 0 else None

    repeat_offenders = [
        {"name": d, "events": c, "yard": next((e["yard"] for e in events if e["driver"] == d), "?")}
        for d, c in driver_counts.most_common(5) if c >= 2
    ]

    # Extract drowsiness events for life-safety callout
    drowsiness_events = [
        {"driver": e["driver"], "vehicle": e["vehicle"], "yard": e["yard"]}
        for e in events if e.get("event_type") == "drowsiness"
    ]

    return {
        "total": total, "red": red, "orange": orange, "yellow": yellow,
        "by_yard": {k: dict(v) for k, v in by_yard.items()},
        "by_type": dict(by_type.most_common(10)),
        "repeat_offenders": repeat_offenders,
        "unassigned": cam_unassigned_info,
        "drowsiness_events": drowsiness_events,
    }


# ==============================================================================
# KPA API
# ==============================================================================

KPA_CALL_DELAY = 1.5  # seconds between calls to avoid rate limiting


def call_kpa(endpoint, params):
    """Make request to KPA API with rate limit retry."""
    url = f"{KPA_BASE_URL}/{endpoint}"
    payload = {"token": KPA_API_TOKEN, "limit": 1000}
    payload.update(params)
    for attempt in range(3):
        try:
            r = requests.post(url, json=payload, timeout=60)
            text = r.text.strip()
            if "rate_limit" in text:
                wait = 30 * (attempt + 1)
                print(f"    Rate limited — waiting {wait}s...")
                time.sleep(wait)
                continue
            return text if text else None
        except Exception as e:
            print(f"    KPA API error: {e}")
            return None
    return None


def call_kpa_paginated(endpoint, params):
    """Make paginated KPA requests, returning all CSV rows."""
    all_rows = []
    headers = None
    page = 1
    while True:
        p = dict(params)
        p["page"] = page
        text = call_kpa(endpoint, p)
        time.sleep(KPA_CALL_DELAY)
        if not text:
            break
        try:
            reader = csv.DictReader(StringIO(text))
            rows = list(reader)
            data = [r for r in rows if r.get("date", "") != "Date"]
            if not data:
                break
            if headers is None:
                headers = reader.fieldnames
            all_rows.extend(data)
            if len(rows) < 1000:
                break
            page += 1
        except Exception:
            break
    return headers or [], all_rows


def _is_casing_kpa(row):
    """Check if a KPA row belongs to Casing division."""
    for field_val in row.values():
        if isinstance(field_val, str):
            val_lower = field_val.lower()
            if "casing" in val_lower or "csg" in val_lower:
                return True
    return False


def _get_kpa_yard(row):
    """Try to extract yard from KPA row."""
    # Check known field hash + common column names
    yard_field = row.get('7vj2l992y7fwqhwz', '') or row.get('yard', '') or row.get('location', '')
    for yard in YARD_ORDER:
        if yard.lower() in yard_field.lower():
            return yard
    # Scan all values
    for val in row.values():
        if isinstance(val, str):
            for yard in YARD_ORDER:
                if yard.lower() in val.lower():
                    return yard
    return "Unassigned"


def _parse_kpa_date(row):
    """Extract and parse date from KPA row. Returns datetime or None."""
    for key in ("date", "Date", "created_at", "updated_at"):
        val = row.get(key, "").strip()
        if val:
            for fmt in ("%Y-%m-%d %H:%M:%S", "%m/%d/%Y", "%Y-%m-%d", "%m/%d/%Y %I:%M %p", "%Y-%m-%dT%H:%M:%S"):
                try:
                    return datetime.strptime(val.split(".")[0], fmt)
                except ValueError:
                    continue
    return None


def get_kpa_form_activity(start_ct, end_ct):
    """Pull submission counts for all active CSG forms."""
    start_ms = int(start_ct.timestamp() * 1000)
    results = {}

    for form_id, form_name in sorted(CSG_ACTIVE_FORMS.items()):
        _, rows = call_kpa_paginated("responses.flat", {
            "form_id": form_id,
            "format": "csv",
            "updated_after": start_ms,
        })

        # Filter to date range (strict -- only count rows with parseable in-range dates)
        in_range = []
        for row in rows:
            dt = _parse_kpa_date(row)
            if dt and start_ct.replace(tzinfo=None) <= dt <= end_ct.replace(tzinfo=None):
                in_range.append(row)

        by_yard = Counter()
        for row in in_range:
            yard = _get_kpa_yard(row)
            by_yard[yard] += 1

        results[form_id] = {
            "name": form_name,
            "total": len(in_range),
            "by_yard": dict(by_yard),
        }
        print(f"    {form_name}: {len(in_range)} submissions")

    return results


def get_kpa_observations(start_ct, end_ct):
    """Pull Observation Cards (form 151085), filter to Casing."""
    start_ms = int(start_ct.timestamp() * 1000)
    _, rows = call_kpa_paginated("responses.flat", {
        "form_id": OBSERVATION_FORM,
        "format": "csv",
        "updated_after": start_ms,
    })

    casing_obs = []
    for row in rows:
        dt = _parse_kpa_date(row)
        if dt and not (start_ct.replace(tzinfo=None) <= dt <= end_ct.replace(tzinfo=None)):
            continue
        if _is_casing_kpa(row):
            row["_yard"] = _get_kpa_yard(row)
            casing_obs.append(row)

    print(f"    Casing observations: {len(casing_obs)} (from {len(rows)} total)")
    return casing_obs


def get_kpa_incidents(start_ct, end_ct):
    """Pull Incident Reports (form 151622), filter to Casing."""
    start_ms = int(start_ct.timestamp() * 1000)
    _, rows = call_kpa_paginated("responses.flat", {
        "form_id": INCIDENT_FORM,
        "format": "csv",
        "updated_after": start_ms,
    })

    date_fmts = ["%Y-%m-%d %H:%M:%S", "%m/%d/%Y", "%Y-%m-%d",
                 "%m/%d/%Y %I:%M %p", "%Y-%m-%dT%H:%M:%S"]
    casing_incidents = []
    for row in rows:
        if not _is_casing_kpa(row):
            continue

        # Prefer incident occurrence date over KPA submission date
        inc_date_raw = row.get(INC_DATE_HASH, "").strip()
        inc_dt = None
        if inc_date_raw:
            for fmt in date_fmts:
                try:
                    inc_dt = datetime.strptime(inc_date_raw.split(".")[0], fmt)
                    break
                except ValueError:
                    continue

        # Fall back to KPA row date if no occurrence date
        if not inc_dt:
            inc_dt = _parse_kpa_date(row)

        if inc_dt and not (start_ct.replace(tzinfo=None) <= inc_dt <= end_ct.replace(tzinfo=None)):
            continue

        row["_yard"] = _get_kpa_yard(row)
        casing_incidents.append(row)

    print(f"    Casing incidents: {len(casing_incidents)} (from {len(rows)} total)")
    return casing_incidents


def get_casing_rcas(start_ct, end_ct):
    """Pull Root Cause Analysis reports (form 180243), filter to Casing.

    Deduplicates by report number (same sub-row pattern as assessments).
    Filters by incident date field within reporting period.
    """
    start_ms = int(start_ct.timestamp() * 1000)
    _, rows = call_kpa_paginated("responses.flat", {
        "form_id": RCA_FORM,
        "format": "csv",
        "updated_after": start_ms,
    })

    date_fmts = ["%Y-%m-%d %H:%M:%S", "%m/%d/%Y", "%Y-%m-%d",
                 "%m/%d/%Y %I:%M %p", "%Y-%m-%dT%H:%M:%S"]

    # Dedup by report number, keep parent rows only
    seen_reports = set()
    casing_rcas = []
    for row in rows:
        rpt = row.get("report number", "").strip()
        if not rpt or rpt in seen_reports:
            continue
        date_val = row.get("date", "").strip()
        if not date_val:
            continue
        seen_reports.add(rpt)

        # Filter to Casing service line
        if row.get(RCA_SERVICE_LINE_HASH, "").strip() != "Casing":
            continue

        # Filter by incident date within reporting period
        inc_date_raw = row.get(RCA_INC_DATE_HASH, "").strip()
        inc_dt = None
        if inc_date_raw:
            for fmt in date_fmts:
                try:
                    inc_dt = datetime.strptime(inc_date_raw.split(".")[0], fmt)
                    break
                except ValueError:
                    continue

        if inc_dt and not (start_ct.replace(tzinfo=None) <= inc_dt <= end_ct.replace(tzinfo=None)):
            continue

        # Parse RCA completion date
        rca_date_raw = row.get(RCA_COMPLETE_DATE_HASH, "").strip()
        rca_dt = None
        if rca_date_raw:
            for fmt in date_fmts:
                try:
                    rca_dt = datetime.strptime(rca_date_raw.split(".")[0], fmt)
                    break
                except ValueError:
                    continue

        # Calculate turnaround
        turnaround = None
        if inc_dt and rca_dt:
            turnaround = (rca_dt - inc_dt).days

        row["_inc_date"] = inc_dt
        row["_rca_date"] = rca_dt
        row["_turnaround_days"] = turnaround
        row["_rca_status"] = "Complete" if rca_dt else "Pending"
        row["_yard"] = row.get(RCA_YARD_HASH, "").strip()
        casing_rcas.append(row)

    print(f"    Casing RCAs: {len(casing_rcas)} (from {len(rows)} rows, "
          f"{len(seen_reports)} unique reports)")
    return casing_rcas


def classify_incidents(incidents):
    """Classify incidents by OSHA type and return structured breakdown.

    KPA incident type field (nojcquy0tfl9hqih) is multi-select -- values can
    be combos like 'Equipment Damage/Property Damage, Recordable'.
    An incident is OSHA recordable if 'Recordable' appears anywhere in the type.
    """
    result = {
        "total": 0,
        "recordable": 0,
        "first_aid": 0,
        "near_miss": 0,
        "vehicle_at_fault": 0,
        "vehicle_not_at_fault": 0,
        "equipment_damage": 0,
        "report_only": 0,
        "personal_illness": 0,
        "by_yard": Counter(),
        "detail": [],
    }
    if not incidents:
        return result

    for inc in incidents:
        inc_type = inc.get(INC_TYPE_HASH, "").strip()
        yard = inc.get("_yard", "Unassigned")
        employee = inc.get(INC_EMPLOYEE_HASH, "").strip()
        desc = inc.get(INC_DESC_HASH, "").strip()
        if len(desc) > 80:
            desc = desc[:77] + "..."
        location = inc.get(INC_LOCATION_HASH, "").strip()

        # Parse incident occurrence date (not KPA submission date)
        inc_date_raw = inc.get(INC_DATE_HASH, "").strip()
        inc_date_display = ""
        if inc_date_raw:
            inc_date_display = _format_display_date(inc_date_raw)
        if not inc_date_display:
            inc_date_display = _format_display_date(inc.get("date", inc.get("Date", "")))

        is_recordable = "Recordable" in inc_type
        type_lower = inc_type.lower()

        result["total"] += 1
        result["by_yard"][yard] += 1

        if is_recordable:
            result["recordable"] += 1
        if "first aid" in type_lower:
            result["first_aid"] += 1
        if "near miss" in type_lower:
            result["near_miss"] += 1
        if "at-fault vehicle" in type_lower and "not at-fault" not in type_lower:
            result["vehicle_at_fault"] += 1
        if "not at-fault" in type_lower:
            result["vehicle_not_at_fault"] += 1
        if "equipment damage" in type_lower or "property damage" in type_lower:
            result["equipment_damage"] += 1
        if "report only" in type_lower:
            result["report_only"] += 1
        if "personal illness" in type_lower:
            result["personal_illness"] += 1

        result["detail"].append({
            "date": inc_date_display,
            "yard": yard,
            "type": inc_type or "Unclassified",
            "employee": employee or "Unknown",
            "description": desc or "See KPA for details",
            "location": location,
            "is_recordable": is_recordable,
            "rca_status": "No RCA",
            "rca_turnaround": None,
        })

    result["by_yard"] = dict(result["by_yard"])
    return result


def cross_reference_rcas(incident_classification, rcas):
    """Match incidents to RCAs and update rca_status on each incident detail.

    Matching logic:
    1. Match incident occurrence date to RCA incident date
    2. Fallback: match first 30 chars of description
    """
    if not rcas or not incident_classification.get("detail"):
        return incident_classification

    for detail in incident_classification["detail"]:
        for rca in rcas:
            # Try date match first
            rca_inc_date = rca.get(RCA_INC_DATE_HASH, "").strip()
            rca_inc_display = _format_display_date(rca_inc_date) if rca_inc_date else ""

            matched = False
            if detail["date"] and rca_inc_display and detail["date"] == rca_inc_display:
                matched = True

            # Fallback: description match
            if not matched:
                rca_desc = rca.get(RCA_DESC_HASH, "").strip()
                inc_desc = detail["description"]
                if inc_desc and rca_desc and len(inc_desc) > 15:
                    if inc_desc[:30] in rca_desc or rca_desc[:30] in inc_desc:
                        matched = True

            if matched:
                detail["rca_status"] = rca["_rca_status"]
                if rca["_turnaround_days"] is not None:
                    detail["rca_turnaround"] = rca["_turnaround_days"]
                    detail["rca_status"] = f"Complete ({rca['_turnaround_days']}d)"
                detail["_rca_causes"] = rca.get(RCA_CAUSES_HASH, "").strip()
                detail["_rca_actions"] = rca.get(RCA_ACTIONS_HASH, "").strip()
                break

    return incident_classification


def get_kpa_assessments(start_ct, end_ct):
    """Pull Field Assessments (form 381707) for Casing.

    KPA flat exports include crew sub-rows (repeating sections) that share the
    same report number but have empty date/observer fields.  We deduplicate by
    report number and only keep the primary (first) row for each assessment.
    """
    start_ms = int(start_ct.timestamp() * 1000)
    _, rows = call_kpa_paginated("responses.flat", {
        "form_id": FIELD_ASSESSMENT_FORM,
        "format": "csv",
        "updated_after": start_ms,
    })

    # Group rows by report number, keep only the first (parent) row per report
    seen_reports = set()
    in_range = []
    for row in rows:
        rpt = row.get("report number", "").strip()
        if not rpt or rpt in seen_reports:
            continue  # skip sub-rows (same report number) and empty rows

        # Only mark as seen if this row has a date (true parent row)
        date_val = row.get("date", "").strip()
        if not date_val:
            continue  # sub-row with no date -- skip

        seen_reports.add(rpt)

        dt = _parse_kpa_date(row)
        if dt and not (start_ct.replace(tzinfo=None) <= dt <= end_ct.replace(tzinfo=None)):
            continue

        row["_yard"] = _get_kpa_yard(row)
        row["_observer"] = (
            row.get("observer", "").strip()
            or row.get("Observer", "").strip()
            or row.get("Name", "").strip()
            or row.get("name", "").strip()
            or row.get("assessor", "").strip()
            or row.get("Assessor", "").strip()
            or row.get("submitted_by", "").strip()
            or row.get("Submitted By", "").strip()
            or "Unknown"
        )
        in_range.append(row)

    print(f"    Field assessments: {len(in_range)} (from {len(rows)} rows, {len(seen_reports)} unique reports)")
    return in_range


def get_casing_vehicle_inspections(start_ct, end_ct):
    """Pull Non-DOT Pre/Post Trip Inspections (form 229645, Casing-specific)."""
    start_ms = int(start_ct.timestamp() * 1000)
    _, rows = call_kpa_paginated("responses.flat", {
        "form_id": VEHICLE_INSPECTION_FORM,
        "format": "csv",
        "updated_after": start_ms,
    })

    casing_inspections = []
    for row in rows:
        dt = _parse_kpa_date(row)
        if dt and not (start_ct.replace(tzinfo=None) <= dt <= end_ct.replace(tzinfo=None)):
            continue
        # Form 229645 is Casing-specific, no service line filter needed
        row["_yard"] = _get_kpa_yard(row)
        casing_inspections.append(row)

    print(f"    Casing pre/post trip inspections: {len(casing_inspections)} (from {len(rows)} total)")
    return casing_inspections


# ==============================================================================
# KPA JSON API -- TRAINING COMPLIANCE
# ==============================================================================

def call_kpa_json_single(endpoint, data_key):
    """Call a KPA JSON endpoint that returns all data in one response."""
    url = f"{KPA_BASE_URL}/{endpoint}"
    payload = {"token": KPA_API_TOKEN}
    try:
        r = requests.post(url, json=payload, timeout=120)
        data = json.loads(r.text.strip())
        return data.get(data_key, [])
    except Exception as e:
        print(f"    KPA JSON API error ({endpoint}): {e}")
        return []


def call_kpa_json_paginated(endpoint, data_key="employees", max_pages=50):
    """Call a paginated KPA JSON endpoint."""
    all_rows = []
    page = 1
    while True:
        payload = {"token": KPA_API_TOKEN, "limit": 500, "page": page}
        url = f"{KPA_BASE_URL}/{endpoint}"
        try:
            r = requests.post(url, json=payload, timeout=120)
            text = r.text.strip()
            if "rate_limit" in text:
                time.sleep(30)
                continue
            data = json.loads(text)
        except Exception as e:
            print(f"    KPA JSON API error ({endpoint}): {e}")
            break
        items = data.get(data_key, [])
        if not items:
            break
        all_rows.extend(items)
        last_page = data.get("paging", {}).get("last_page", 1)
        if page >= last_page or page >= max_pages:
            break
        page += 1
        time.sleep(1.5)
    return all_rows


def get_casing_training_compliance(end_date_str):
    """Fetch KPA training compliance data filtered to Casing employees.

    Returns dict with employees list (each with 'yard' field),
    overall_pct, total_employees, compliant_count, overdue_count.
    """
    print("    Fetching KPA users...")
    all_users = call_kpa_json_single("users.list", "users")
    casing_user_ids = set()
    user_names = {}
    user_yards = {}

    # KPA field office ID -> yard mapping (from fieldoffices.list)
    FO_YARD_MAP = {
        "671017668ee2a10019b2f7f0": "Midland",   # Midland Yukon
        "5d166f31efd5700017316be4": "Midland",   # Midland (generic)
        "5d2cf0da6f00c900179d969b": "Kilgore",
        "6009f55901f3bb0142271514": "Hobbs",
        "5d166ec1d57b5c00178cfab0": "Jourdanton",
        "5cddbce7cc6e850017e270a1": "Bryan",
        "6009f55901f3bb0142271515": "Laredo",
        "6009f55901f3bb014227151c": "Midland",   # San Angelo -> merged to Midland
        "671017898ee2a10019b2fc9a": "Midland",   # Overhead BRHAS
    }

    for u in all_users:
        uid = u.get("id", "")
        if not uid or u.get("terminationDate"):
            continue
        first = u.get("firstname", "")
        last = u.get("lastname", "")
        user_names[uid] = f"{first} {last}".strip()
        lobs = u.get("lineOfBusiness_id", [])
        if not isinstance(lobs, list):
            lobs = [lobs]
        if CASING_LOB_ID in lobs:
            casing_user_ids.add(uid)
            # Determine yard from field office ID (array of strings)
            fo_ids = u.get("fieldOffice_id", [])
            if isinstance(fo_ids, str):
                fo_ids = [fo_ids]
            elif not isinstance(fo_ids, list):
                fo_ids = []

            # Try ID lookup against known mapping
            yard = ""
            for fid in fo_ids:
                yard = FO_YARD_MAP.get(fid, "")
                if yard:
                    break
            if not yard:
                yard = "Unassigned"
            user_yards[uid] = yard

    print(f"    Casing employees: {len(casing_user_ids)} (from {len(all_users)} total users)")
    # Show yard distribution for debugging
    yard_dist = Counter(user_yards.values())
    for y in YARD_ORDER:
        if yard_dist.get(y, 0) > 0:
            print(f"      {y}: {yard_dist[y]}")
    if yard_dist.get("Unassigned", 0) > 0:
        print(f"      Unassigned: {yard_dist['Unassigned']}")

    print("    Fetching training programs...")
    programs = call_kpa_json_single("trainings.v2.list", "trainings")
    training_lookup = {}
    training_created = {}
    for p in programs:
        tid = p.get("id")
        name = p.get("title", p.get("name", ""))
        if tid is not None and name:
            training_lookup[tid] = name
            created_ms = p.get("created_on")
            if created_ms:
                training_created[tid] = created_ms
    print(f"    Training programs: {len(training_lookup)}")

    print("    Fetching training employee status...")
    all_status = call_kpa_json_paginated(
        "training-employee-status.list", data_key="employees"
    )
    casing_status = [r for r in all_status if r.get("m_user_id") in casing_user_ids]
    print(f"    Casing training records: {len(casing_status)} (from {len(all_status)} total)")

    cutoff_ms = None
    if end_date_str:
        try:
            dt_cut = datetime.strptime(end_date_str, "%Y-%m-%d").replace(
                hour=23, minute=59, second=59, tzinfo=timezone.utc)
            cutoff_ms = int(dt_cut.timestamp() * 1000)
        except ValueError:
            pass

    employees = []
    for row in casing_status:
        uid = row.get("m_user_id", "")
        name = user_names.get(uid, "Unknown")
        yard = user_yards.get(uid, "Unassigned")

        incomplete_ids = row.get("incomplete_training_ids", []) or []
        complete_ids = row.get("complete_training_ids", []) or []

        if cutoff_ms:
            incomplete_ids = [tid for tid in incomplete_ids
                              if training_created.get(tid, 0) <= cutoff_ms]
            complete_ids = [tid for tid in complete_ids
                            if training_created.get(tid, 0) <= cutoff_ms]

        incomplete_names = [training_lookup.get(tid, f"Program #{tid}")
                            for tid in incomplete_ids]
        complete_names = [training_lookup.get(tid, f"Program #{tid}")
                          for tid in complete_ids]

        total = len(incomplete_ids) + len(complete_ids)
        pct = round(len(complete_ids) / total * 100) if total > 0 else 100

        status = "Complete"
        if pct < 100:
            status = "Overdue" if row.get("status") == "overdue" else "In Progress"

        # Calculate max days since assignment for incomplete training
        max_days_since = 0
        now_ms = int(datetime.now(tz=timezone.utc).timestamp() * 1000)
        for tid in incomplete_ids:
            created_ms = training_created.get(tid, 0)
            if created_ms > 0:
                days = (now_ms - created_ms) // (1000 * 86400)
                if days > max_days_since:
                    max_days_since = days

        employees.append({
            "employee_name": name,
            "yard": yard,
            "percent_complete": pct,
            "incomplete_training_names": incomplete_names,
            "complete_training_names": complete_names,
            "status": status,
            "days_since_assignment": max_days_since,
        })

    total_emp = len(employees)
    compliant = sum(1 for e in employees if e["percent_complete"] >= 100)
    overdue = sum(1 for e in employees if e["status"] == "Overdue")
    overall_pct = round(compliant / total_emp * 100, 1) if total_emp > 0 else 0

    # Live headcount from KPA active employees
    headcount_by_yard = dict(yard_dist)
    headcount_by_yard.pop("Unassigned", None)
    headcount_total = len(casing_user_ids)

    return {
        "employees": employees,
        "overall_pct": overall_pct,
        "total_employees": total_emp,
        "compliant_count": compliant,
        "overdue_count": overdue,
        "headcount_by_yard": headcount_by_yard,
        "headcount_total": headcount_total,
    }


def analyze_assessments(rows):
    """Analyze field assessments -- by yard, by rep, findings."""
    META_FIELDS = {
        'report number', 'date', 'observer', 'status', 'link', 'kpa_link',
        'name', 'Name', 'form', 'form_id', 'updated_at', 'created_at',
        'report', 'id', 'response_id', '_yard', '_observer',
        '7vj2l992y7fwqhwz', 'yard', 'location',
        'updated', 'updated_time', 'version', 'observer-emp-num',
        'duration', 'latitude', 'longitude', 'temperature', 'wind-speed',
        'weather', 'parentrepnum', 'parentlink', 'surrogate',
        'select-yes', 'select-no', 'select-n/a', 'select-na',
    }

    by_yard = Counter()
    by_rep = Counter()
    with_findings = []
    clean = []

    for row in rows:
        yard = row.get("_yard", "Unassigned")
        rep = row.get("_observer", "Unknown")
        by_yard[yard] += 1
        by_rep[rep] += 1

        has_finding = False

        # Method 1: KPA summary 'select-no' count
        select_no = row.get("select-no", "0").strip()
        try:
            if int(select_no) > 0:
                has_finding = True
        except (ValueError, TypeError):
            pass

        # Method 2: Scan fields for "No" or negative keywords
        if not has_finding:
            for key, val in row.items():
                if key.lower() in META_FIELDS or not val or not isinstance(val, str):
                    continue
                vl = val.strip().lower()
                if vl in ("", "n/a", "na", "none", "yes", "good", "ok",
                          "pass", "satisfactory"):
                    continue
                if vl == "no":
                    has_finding = True
                    break
                if any(kw in vl for kw in ("unsatisfactory", "fail", "deficien",
                                            "corrective", "needs", "issue",
                                            "concern", "damage", "broken")):
                    has_finding = True
                    break

        if has_finding:
            with_findings.append(row)
        else:
            clean.append(row)

    return {
        "total": len(rows),
        "with_findings": len(with_findings),
        "clean": len(clean),
        "by_yard": dict(by_yard),
        "by_rep": dict(by_rep),
    }


def get_form_field_labels(form_id):
    """Fetch field hash -> human-readable label mapping from KPA JSON API."""
    if not KPA_API_TOKEN:
        return {}
    try:
        r = requests.post(
            f"{KPA_BASE_URL}/responses.flat",
            json={"token": KPA_API_TOKEN, "form_id": form_id,
                  "format": "json", "limit": 1},
            timeout=30,
        )
        data = json.loads(r.text.strip())
        responses = data.get("responses", [])
        if responses:
            return {k.lower(): str(v) for k, v in responses[0].items() if v}
    except Exception as exc:
        print(f"  Warning: could not fetch field labels for form {form_id}: {exc}")
    return {}


# ==============================================================================
# DOCX FORMATTING HELPERS
# ==============================================================================

def _apply_table_polish(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl_pr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'))
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(borders_xml))
    cell_mar_xml = (
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="45" w:type="dxa"/><w:bottom w:w="45" w:type="dxa"/>'
        f'<w:left w:w="80" w:type="dxa"/><w:right w:w="80" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tbl_pr.append(parse_xml(cell_mar_xml))
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0


def _add_page_number_field(paragraph):
    """Add a PAGE / NUMPAGES field to a paragraph (Page X of Y)."""
    run1 = paragraph.add_run("Page ")
    run1.font.size = Pt(8)
    run1.font.name = CALIBRI
    run1.font.color.rgb = GRAY

    # PAGE field
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    run_page = paragraph.add_run()
    run_page._r.append(fld_char1)
    instr1 = OxmlElement('w:instrText')
    instr1.set(qn('xml:space'), 'preserve')
    instr1.text = ' PAGE '
    run_instr1 = paragraph.add_run()
    run_instr1._r.append(instr1)
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run_end1 = paragraph.add_run()
    run_end1._r.append(fld_char2)

    run_of = paragraph.add_run(" of ")
    run_of.font.size = Pt(8)
    run_of.font.name = CALIBRI
    run_of.font.color.rgb = GRAY

    # NUMPAGES field
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'begin')
    run_np = paragraph.add_run()
    run_np._r.append(fld_char3)
    instr2 = OxmlElement('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = ' NUMPAGES '
    run_instr2 = paragraph.add_run()
    run_instr2._r.append(instr2)
    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')
    run_end2 = paragraph.add_run()
    run_end2._r.append(fld_char4)


def _setup_header_footer(doc, report_title, footer_label):
    """Add running header and footer with page numbers to all pages."""
    for section in doc.sections:
        # --- Running Header ---
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(report_title)
        hr.font.size = Pt(8)
        hr.font.name = CALIBRI
        hr.font.color.rgb = GRAY
        hr.font.italic = True
        hp.paragraph_format.space_after = Pt(0)

        # --- Footer: label left, page numbers right ---
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)

        fr = fp.add_run(f"{footer_label}  |  Butch's Resources  |  ")
        fr.font.size = Pt(8)
        fr.font.name = CALIBRI
        fr.font.color.rgb = GRAY

        _add_page_number_field(fp)


def set_cell_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def style_header_row(row, bg_hex="8B0000"):
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = CALIBRI


def style_data_row(row, row_idx, font_size=9, col_align=None):
    if row_idx % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, LIGHT_GRAY)
    for ci, cell in enumerate(row.cells):
        align = WD_ALIGN_PARAGRAPH.CENTER
        if col_align and ci < len(col_align):
            align = col_align[ci]
        for p in cell.paragraphs:
            p.alignment = align
            for run in p.runs:
                run.font.size = Pt(font_size)
                run.font.name = CALIBRI


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = DARK_RED
    run.font.name = CALIBRI
    pf = heading.paragraph_format
    pf.keep_with_next = True
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(6)
    else:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
    return heading


def add_data_table(doc, headers, rows_data, font_size=9, col_align=None):
    if not rows_data:
        return None
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        hdr.cells[j].text = h
    style_header_row(hdr)
    for i, row_data in enumerate(rows_data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
        style_data_row(row, i, font_size, col_align=col_align)
    _apply_table_polish(table)
    return table


def add_metric_table(doc, metrics_list):
    num_cols = len(metrics_list[0])
    table = doc.add_table(rows=len(metrics_list), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    for ci in range(num_cols):
        hdr.cells[ci].text = str(metrics_list[0][ci])
    style_header_row(hdr)
    for i, row_data in enumerate(metrics_list):
        if i == 0:
            continue
        row = table.rows[i]
        for ci in range(num_cols):
            row.cells[ci].text = str(row_data[ci]) if ci < len(row_data) else ""
        for p in row.cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = CALIBRI
                run.font.bold = True
        for ci in range(1, num_cols):
            for p in row.cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = CALIBRI
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
    _apply_table_polish(table)
    return table


def _add_text(doc, text, size=11, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = CALIBRI
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return p


# ==============================================================================
# NARRATIVE & TAKEAWAYS
# ==============================================================================

def _build_narrative(mileage, speeding, assessments, observations,
                     incidents, inspections, training_compliance,
                     camera, prev_data=None, yard=None, yoy_data=None):
    """Build an executive narrative paragraph for senior leadership.
    Leads with red flags and compliance gaps; closes with zero-recordable confirmation."""
    red_flag_parts = []
    body_parts = []
    close_parts = []
    pm = prev_data or {}
    yoy = yoy_data or {}
    pm_miles = pm.get("mileage", {}).get("total_miles")
    tc = training_compliance or {}

    scope = f"Casing {yard}" if yard else "Casing Division"
    man_hours = MONTHLY_MAN_HOURS_BY_YARD.get(yard, MONTHLY_MAN_HOURS) if yard else MONTHLY_MAN_HOURS
    headcount = HEADCOUNT_BY_YARD.get(yard, EMPLOYEE_COUNT) if yard else EMPLOYEE_COUNT
    manager = YARD_INFO.get(yard, {}).get("manager", "Division Management") if yard else "Division Management"

    # --- RED FLAGS / COMPLIANCE GAPS FIRST ---
    # Training compliance gap
    if tc and tc.get("overall_pct", 100) < 95:
        non_compliant = tc.get("total_employees", 0) - tc.get("compliant_count", 0)
        red_flag_parts.append(f"Training compliance at {tc['overall_pct']:.1f}% "
                    f"({non_compliant} employees non-compliant, target 95%). "
                    f"{manager} accountable for resolution.")

    if not inspections:
        red_flag_parts.append("Zero pre/post trip inspections documented this period.")

    # Camera RED events
    if camera and camera.get("red", 0) > 0:
        red_flag_parts.append(f"{camera['red']} RED-tier camera events require coaching follow-up.")

    # Assessments with findings
    if assessments["total"] > 0 and assessments["with_findings"] > 0:
        pct = round(assessments["with_findings"] / assessments["total"] * 100)
        red_flag_parts.append(f"{assessments['with_findings']} of {assessments['total']} field assessments "
                     f"identified deficiencies ({pct}% finding rate).")

    # --- BODY: operational context ---
    # Fleet activity
    if mileage["total_miles"] > 0:
        speed_rate = round(speeding["total"] / (mileage["total_miles"] / 10000), 1) if mileage["total_miles"] > 0 else 0
        miles_ctx = ""
        if pm_miles and pm_miles > 0:
            _, pct_str, _ = mom_delta(mileage["total_miles"], pm_miles)
            miles_ctx = f" ({pct_str} vs. prior month)"
        body_parts.append(f"The fleet operated {mileage['active_trucks']} trucks covering "
                     f"{mileage['total_miles']:,.0f} miles{miles_ctx}, "
                     f"with a speeding rate of {speed_rate} events per 10,000 miles.")

        # Fleet growth context for extreme YoY
        yoy_miles = yoy.get("mileage", {}).get("total_miles")
        if yoy_miles and yoy_miles > 0:
            yoy_pct = ((mileage["total_miles"] - yoy_miles) / yoy_miles) * 100
            if yoy_pct > 200:
                yoy_trucks = yoy.get("mileage", {}).get("active_trucks", 0)
                body_parts.append(f"(YoY fleet growth reflects division expansion from "
                             f"{yoy_trucks} to {mileage['active_trucks']} active trucks.)")

    # Observations
    obs_count = len(observations) if isinstance(observations, list) else observations
    if obs_count > 0 and headcount > 0:
        obs_rate = round(obs_count / headcount, 2)
        body_parts.append(f"{obs_count} observation cards submitted ({obs_rate} per employee).")

    # Incidents body (types breakdown)
    ic = incidents if isinstance(incidents, dict) else None
    if ic:
        rec = ic.get("recordable", 0)
        total = ic.get("total", 0)
        fa = ic.get("first_aid", 0)
        nm = ic.get("near_miss", 0)
        if rec > 0:
            red_flag_parts.insert(0, f"This period included {rec} OSHA recordable incident(s), "
                         "requiring immediate executive review of root cause analysis and corrective actions.")
        if total > 0:
            type_parts = []
            if fa > 0:
                type_parts.append(f"{fa} first aid")
            if nm > 0:
                type_parts.append(f"{nm} near miss")
            if ic.get("vehicle_at_fault", 0) > 0:
                type_parts.append(f"{ic['vehicle_at_fault']} at-fault vehicle")
            remaining = total - rec - fa - nm - ic.get("vehicle_at_fault", 0)
            if remaining > 0:
                type_parts.append(f"{remaining} other")
            if type_parts:
                body_parts.append(f"{total} total incident reports filed ({', '.join(type_parts)}).")
            if nm > 0:
                body_parts.append(f"{nm} near miss report(s) indicate active hazard recognition.")

        # --- CLOSE: zero recordables last ---
        if rec == 0:
            close_parts.append(f"{scope} completed the reporting period with zero OSHA recordable incidents.")
    elif incidents:
        body_parts.append(f"This period included {len(incidents)} incident(s) requiring review.")
    else:
        close_parts.append(f"{scope} completed the reporting period with zero incidents.")

    return " ".join(red_flag_parts + body_parts + close_parts)


def _generate_takeaways(mileage, speeding, camera, assessments, observations,
                         incidents, inspections, training_compliance, prev_data,
                         yard=None, man_hours=None):
    """Generate actionable takeaway bullets with owner names."""
    items = []
    pm = prev_data or {}
    tc = training_compliance or {}
    if man_hours is None:
        man_hours = MONTHLY_MAN_HOURS_BY_YARD.get(yard, MONTHLY_MAN_HOURS) if yard else MONTHLY_MAN_HOURS
    headcount = HEADCOUNT_BY_YARD.get(yard, EMPLOYEE_COUNT) if yard else EMPLOYEE_COUNT
    manager = YARD_INFO.get(yard, {}).get("manager", "Division Management") if yard else "Division Management"
    safety_rep = YARD_INFO.get(yard, {}).get("safety_reps", "Safety Team") if yard else "Safety Team"

    # 1. Incidents -- use classification if available
    ic = incidents if isinstance(incidents, dict) else None
    if ic:
        rec = ic.get("recordable", 0)
        total = ic.get("total", 0)
        fa = ic.get("first_aid", 0)
        nm = ic.get("near_miss", 0)
        vaf = ic.get("vehicle_at_fault", 0)

        if rec > 0:
            items.append(f"{rec} OSHA recordable incident(s). {manager} to review root cause "
                         "analysis and confirm all corrective actions are closed.")
        elif total == 0:
            items.append("Zero OSHA recordable incidents documented this period. "
                         "Zero total incident reports filed.")
        else:
            items.append(f"Zero OSHA recordable incidents documented this period "
                         f"({total} total reports filed).")

        if fa > 0:
            fafr = round(fa * 200000 / man_hours, 2) if man_hours > 0 else 0
            items.append(f"{fa} first aid incident(s) (FAFR: {fafr}). {safety_rep} to review "
                         "for trend patterns and escalation risk.")

        if nm > 0:
            items.append(f"{nm} near miss report(s) filed -- indicates active hazard "
                         "recognition culture. Continue encouraging near miss reporting.")

        if vaf > 0:
            items.append(f"{vaf} at-fault vehicle incident(s). {manager} to review with "
                         "drivers and reinforce defensive driving expectations.")

        # RCA accountability
        rca_with = sum(1 for d in ic.get("detail", []) if "No RCA" not in d.get("rca_status", "No RCA"))
        if total > 0:
            items.append(f"RCA completion: {rca_with} of {total} incidents have RCA on file. "
                         f"{manager} accountable for timely RCA completion on all incidents.")
    elif incidents:
        inc_count = len(incidents) if isinstance(incidents, list) else incidents
        items.append(f"{inc_count} incident(s) reported. {manager} to review root cause "
                     "analysis and confirm all corrective actions are closed.")
    else:
        items.append("Zero OSHA recordable incidents documented this period. "
                     "Zero total incident reports filed.")

    # 2. Speeding
    if speeding["total"] > 0:
        msg = f"{speeding['total']} speeding events ({speeding.get('critical', 0)} critical)."
        if speeding.get("top_drivers"):
            worst = speeding["top_drivers"][0]
            msg += f" {manager} to address driver {worst['name']} ({worst['events']} events)."
        if speeding.get("unassigned"):
            ua = speeding["unassigned"]
            msg += (f" {ua['events']} event(s) from unassigned vehicles -- "
                    f"{manager} to verify driver assignments in Motive.")
        items.append(msg)

    # 3. Camera RED events
    if camera and camera.get("red", 0) > 0:
        msg = f"{camera['red']} RED-tier camera events require immediate coaching."
        if camera.get("repeat_offenders"):
            names = ", ".join(o["name"] for o in camera["repeat_offenders"][:2])
            msg += f" Repeat offenders: {names}."
        items.append(msg)

    # 4. Inspections compliance gap
    insp_count = len(inspections) if isinstance(inspections, list) else inspections
    if insp_count == 0:
        items.append(f"REQUIRED: Zero pre/post trip inspections on file. {manager} to ensure "
                     "inspections are submitted through KPA for all active trucks.")

    # 5. Training compliance -- severity-scaled
    if tc and tc.get("overall_pct", 100) < 70:
        non_compliant = tc.get("total_employees", 0) - tc.get("compliant_count", 0)
        items.append(f"CRITICAL: Training compliance at {tc['overall_pct']:.1f}% "
                     f"({non_compliant} employees non-compliant). {manager} and {safety_rep} to push "
                     "outstanding training completions immediately.")
    elif tc and tc.get("overall_pct", 100) < 95:
        non_compliant = tc.get("total_employees", 0) - tc.get("compliant_count", 0)
        items.append(f"Training compliance at {tc['overall_pct']:.1f}% "
                     f"({non_compliant} employees non-compliant). {manager} and {safety_rep} to push "
                     "outstanding training completions to meet 95% target.")
    elif tc and tc.get("overall_pct", 0) >= 95:
        items.append(f"Training compliance at {tc['overall_pct']:.1f}%. Meets 95% target.")

    # 6. Assessments
    if assessments["total"] > 0 and assessments["with_findings"] > 0:
        pct = round(assessments["with_findings"] / assessments["total"] * 100)
        items.append(f"{assessments['with_findings']} of {assessments['total']} field assessments "
                     f"had findings ({pct}%). {safety_rep} to follow up on corrective actions.")
    elif assessments["total"] == 0 and yard:
        items.append(f"REQUIRED: Zero field assessments conducted. {safety_rep} to schedule "
                     "minimum 2 field assessments for the upcoming period.")

    # 7. Low observation rate (yard reports only)
    obs_list = observations if isinstance(observations, list) else []
    obs_count = len(obs_list)
    if yard and headcount > 0:
        obs_rate = obs_count / headcount
        if obs_rate < 0.5:
            items.append(f"OBSERVATION GAP: {obs_count} observation cards ({obs_rate:.2f} per employee) "
                         f"is well below target. {manager} and {safety_rep} to drive crew participation "
                         "in observation program.")

    # 8. Single-employee concentration (yard reports only)
    if yard and obs_list and obs_count >= 5:
        obs_by_emp = Counter(_obs_employee(r) for r in obs_list)
        if obs_by_emp:
            top_emp, top_count = obs_by_emp.most_common(1)[0]
            concentration_pct = round(top_count / obs_count * 100)
            if concentration_pct > 50:
                items.append(f"OBSERVATION CONCENTRATION: {top_emp} appears in {top_count} of "
                             f"{obs_count} cards ({concentration_pct}%). {manager} to broaden "
                             "crew participation -- observation programs require diverse input.")

    return items[:8]


# ==============================================================================
# REPORT GENERATION
# ==============================================================================

def generate_report(month_str, mileage, speeding, camera, form_activity,
                    assessments, assessment_analysis, observations, incidents,
                    inspections, training_compliance,
                    prev_data, output_path, field_labels=None, yard=None,
                    incident_classification=None, rcas=None, yoy_data=None):
    """Generate the full DOCX report. If yard is set, generates a yard-specific report."""
    start_ct, end_ct = parse_month(month_str)
    mlabel = month_label(start_ct, end_ct)
    fl = field_labels or {}

    # Scope-specific constants
    scope = f"Casing - {yard}" if yard else "Casing Division"
    man_hours = MONTHLY_MAN_HOURS_BY_YARD.get(yard, MONTHLY_MAN_HOURS) if yard else MONTHLY_MAN_HOURS
    headcount = HEADCOUNT_BY_YARD.get(yard, EMPLOYEE_COUNT) if yard else EMPLOYEE_COUNT
    tc = training_compliance or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)

    # Running header + footer with page numbers
    month_display = datetime.strptime(month_str, "%Y-%m").strftime("%B %Y")
    header_title = f"Casing {yard} HSE Recap -- {month_display}" if yard else f"Casing Division HSE Recap -- {month_display}"
    footer_label = f"Casing {yard} HSE" if yard else "Casing Division HSE"
    _setup_header_footer(doc, header_title, footer_label)

    # ===== COVER PAGE =====
    logo_path = os.path.join(LOGOS_DIR, "Butchs.jpg")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("CASING DIVISION")
    run.font.size = Pt(28)
    run.font.name = CALIBRI
    run.font.bold = True
    run.font.color.rgb = DARK_RED

    if yard:
        yard_title = doc.add_paragraph()
        yard_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        yard_title.paragraph_format.space_before = Pt(2)
        yard_title.paragraph_format.space_after = Pt(4)
        run = yard_title.add_run(f"{yard} Yard")
        run.font.size = Pt(22)
        run.font.name = CALIBRI
        run.font.bold = True
        run.font.color.rgb = DARK_RED

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("MONTHLY HSE RECAP")
    run.font.size = Pt(20)
    run.font.name = CALIBRI
    run.font.bold = True
    run.font.color.rgb = GRAY

    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_before = Pt(4)
    period.paragraph_format.space_after = Pt(18)
    run = period.add_run(mlabel)
    run.font.size = Pt(14)
    run.font.name = CALIBRI
    run.font.color.rgb = BLACK

    # Manager and Safety Rep on cover for yard reports
    if yard:
        yi = YARD_INFO.get(yard, {})
        mgr_line = doc.add_paragraph()
        mgr_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mgr_line.paragraph_format.space_before = Pt(4)
        mgr_line.paragraph_format.space_after = Pt(4)
        run = mgr_line.add_run(f"Manager: {yi.get('manager', 'TBD')}  |  "
                               f"Safety Rep: {yi.get('safety_reps', 'TBD')}")
        run.font.size = Pt(12)
        run.font.name = CALIBRI
        run.font.color.rgb = BLACK

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen.paragraph_format.space_before = Pt(8)
    run = gen.add_run(f"Prepared by Kelly Rhodes  |  Generated {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.name = CALIBRI
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # Previous month data for MoM
    pm = prev_data or {}
    pm_mileage = pm.get("mileage", {})
    pm_speeding = pm.get("speeding", {})
    pm_camera = pm.get("camera", {})

    # Year-over-year data
    yoy = yoy_data or {}

    # Incident classification (used by dashboard and exec summary)
    ic = incident_classification or {}

    # ===== RED FLAG SUMMARY (company-wide only, one page after cover) =====
    if not yard:
        month_display_dash = datetime.strptime(month_str, "%Y-%m").strftime("%B %Y")
        report_date = start_ct.replace(tzinfo=None)
        due_date_30 = (report_date + timedelta(days=30)).strftime("%B %d, %Y")
        tc_dash = training_compliance or {}

        add_section_heading(doc, f"Red Flag Summary -- {month_display_dash}")

        red_flag_rows = []

        # Training compliance by yard
        tc_by_yard_dash = {}
        for emp in tc_dash.get("employees", []):
            y = emp.get("yard", "Unassigned")
            if y not in tc_by_yard_dash:
                tc_by_yard_dash[y] = {"total": 0, "compliant": 0}
            tc_by_yard_dash[y]["total"] += 1
            if emp["percent_complete"] >= 100:
                tc_by_yard_dash[y]["compliant"] += 1

        # 1. Training <70%
        for y in YARD_ORDER:
            tc_yd = tc_by_yard_dash.get(y, {"total": 0, "compliant": 0})
            tc_pct = round(tc_yd["compliant"] / tc_yd["total"] * 100, 1) if tc_yd["total"] > 0 else 0
            if tc_pct < 70:
                mgr = YARD_INFO.get(y, {}).get("manager", "Yard Manager")
                red_flag_rows.append([y, f"Training compliance at {tc_pct}%", mgr])

        # 2. Zero assessments
        aa_dash = assessment_analysis if isinstance(assessment_analysis, dict) else {}
        for y in YARD_ORDER:
            cur_assess = aa_dash.get("by_yard", {}).get(y, 0)
            if cur_assess == 0:
                srep = YARD_INFO.get(y, {}).get("safety_reps", "Safety Rep")
                red_flag_rows.append([y, "Zero field assessments conducted", srep])

        # 3. Zero or near-zero observations
        obs_by_yard_rf = Counter(r.get("_yard", "Unassigned") for r in (observations if isinstance(observations, list) else []))
        for y in YARD_ORDER:
            y_obs = obs_by_yard_rf.get(y, 0)
            y_hc = HEADCOUNT_BY_YARD.get(y, 1)
            obs_rate = y_obs / y_hc if y_hc > 0 else 0
            if obs_rate < 0.1:
                mgr = YARD_INFO.get(y, {}).get("manager", "Yard Manager")
                red_flag_rows.append([y, f"Observations: {y_obs} ({obs_rate:.2f}/employee)", mgr])

        # 4. Unassigned speeding
        if speeding.get("unassigned"):
            ua = speeding["unassigned"]
            red_flag_rows.append(["Division", f"{ua['events']} speeding events from unassigned vehicles", "Fleet Management"])

        if red_flag_rows:
            add_data_table(doc, ["Yard", "Issue", "Owner"], red_flag_rows,
                           font_size=10, col_align=[_L, _L, _L])
        else:
            _add_text(doc, "No red flags identified this period.",
                      size=11, italic=True, color=GREEN)

        doc.add_page_break()

    # ===== SECTION 1: EXECUTIVE SUMMARY =====
    add_section_heading(doc, "1. Executive Summary")

    # Narrative paragraph -- pass classification dict as incidents
    narrative = _build_narrative(
        mileage, speeding, assessment_analysis, observations,
        ic or incidents, inspections, tc, camera, prev_data, yard,
        yoy_data=yoy_data,
    )
    np = doc.add_paragraph()
    np.paragraph_format.space_after = Pt(8)
    np.paragraph_format.line_spacing = 1.15
    nr = np.add_run(narrative)
    nr.font.size = Pt(11)
    nr.font.name = CALIBRI

    # OSHA rates -- use classification for accurate counts
    recordable = ic.get("recordable", 0) if ic else 0
    first_aid = ic.get("first_aid", 0) if ic else 0
    near_miss = ic.get("near_miss", 0) if ic else 0
    total_incidents = ic.get("total", 0) if ic else (len(incidents) if isinstance(incidents, list) else 0)
    trir = round(recordable * 200000 / man_hours, 2) if man_hours > 0 else 0
    dart = trir  # assume all recordables are DART until lost-time data available
    fafr = round(first_aid * 200000 / man_hours, 2) if man_hours > 0 and first_aid > 0 else 0
    prev_trir = pm.get("trir")
    _, trir_mom, _ = mom_delta(trir, prev_trir)
    yoy_trir = yoy.get("trir")
    _, trir_yoy, _ = mom_delta(trir, yoy_trir)

    # Lagging Indicators table
    add_section_heading(doc, "Lagging Indicators", level=2)
    lagging = [
        ("Metric", "Value", "MoM", "YoY"),
        ("Recordable Incidents", str(recordable),
         mom_delta(recordable, pm.get("incidents_recordable"))[1] if pm.get("incidents_recordable") is not None else "N/A",
         mom_delta(recordable, yoy.get("incidents_recordable"))[1] if yoy.get("incidents_recordable") is not None else "N/A"),
        ("TRIR (per 200k hrs)", f"{trir:.2f}", trir_mom, trir_yoy),
        ("DART Rate (per 200k hrs)", f"{dart:.2f}",
         mom_delta(dart, pm.get("dart"))[1] if pm.get("dart") is not None else "--",
         mom_delta(dart, yoy.get("dart"))[1] if yoy.get("dart") is not None else "--"),
        ("First Aid Incidents", str(first_aid),
         mom_delta(first_aid, pm.get("incidents_first_aid"))[1] if pm.get("incidents_first_aid") is not None else "N/A",
         mom_delta(first_aid, yoy.get("incidents_first_aid"))[1] if yoy.get("incidents_first_aid") is not None else "N/A"),
        ("FAFR (per 200k hrs)", str(fafr) if fafr > 0 else "0.00",
         mom_delta(fafr, pm.get("fafr"))[1] if pm.get("fafr") is not None else "N/A",
         mom_delta(fafr, yoy.get("fafr"))[1] if yoy.get("fafr") is not None else "N/A"),
        ("Near Misses", str(near_miss),
         mom_delta(near_miss, pm.get("incidents_near_miss"))[1] if pm.get("incidents_near_miss") is not None else "N/A",
         mom_delta(near_miss, yoy.get("incidents_near_miss"))[1] if yoy.get("incidents_near_miss") is not None else "N/A"),
        ("Total Incident Reports", str(total_incidents),
         mom_delta(total_incidents, pm.get("incidents_total"))[1] if pm.get("incidents_total") is not None else "N/A",
         mom_delta(total_incidents, yoy.get("incidents_total"))[1] if yoy.get("incidents_total") is not None else "N/A"),
    ]
    add_metric_table(doc, lagging)
    _add_text(doc, f"Based on {headcount:,} employees  \u2014  {man_hours:,} estimated monthly man hours",
              size=9, italic=True, color=GRAY)

    # Leading Indicators table
    add_section_heading(doc, "Leading Indicators", level=2)
    obs_count = len(observations) if isinstance(observations, list) else 0
    insp_count = len(inspections) if isinstance(inspections, list) else 0
    leading = [
        ("Metric", "Value", "MoM", "YoY"),
        ("Observations", str(obs_count),
         mom_delta(obs_count, pm.get("observations_total"))[1] if pm.get("observations_total") is not None else "N/A",
         mom_delta(obs_count, yoy.get("observations_total"))[1] if yoy.get("observations_total") is not None else "N/A"),
        ("Field Assessments", str(assessment_analysis["total"]),
         mom_delta(assessment_analysis["total"], pm.get("assessments_total"))[1] if pm.get("assessments_total") is not None else "N/A",
         mom_delta(assessment_analysis["total"], yoy.get("assessments_total"))[1] if yoy.get("assessments_total") is not None else "N/A"),
        ("Pre/Post Trip Inspections", str(insp_count),
         mom_delta(insp_count, pm.get("inspections_total"))[1] if pm.get("inspections_total") is not None else "N/A",
         mom_delta(insp_count, yoy.get("inspections_total"))[1] if yoy.get("inspections_total") is not None else "N/A"),
        ("Training Compliance", f"{tc.get('overall_pct', 0):.1f}%",
         mom_delta(tc.get('overall_pct', 0), pm.get("training_compliance_pct"))[1] if pm.get("training_compliance_pct") is not None else "--",
         mom_delta(tc.get('overall_pct', 0), yoy.get("training_compliance_pct"))[1] if yoy.get("training_compliance_pct") is not None else "--"),
    ]
    add_metric_table(doc, leading)

    # Fleet Summary table
    add_section_heading(doc, "Fleet Summary", level=2)
    yoy_mileage = yoy.get("mileage", {})
    yoy_speeding = yoy.get("speeding", {})
    yoy_camera = yoy.get("camera", {})
    _, miles_mom, _ = mom_delta(mileage["total_miles"], pm_mileage.get("total_miles"))
    _, speed_mom, _ = mom_delta(speeding["total"], pm_speeding.get("total"))
    _, cam_mom, _ = mom_delta(camera["total"], pm_camera.get("total"))
    _, miles_yoy, _ = mom_delta(mileage["total_miles"], yoy_mileage.get("total_miles"))
    _, speed_yoy, _ = mom_delta(speeding["total"], yoy_speeding.get("total"))
    _, cam_yoy, _ = mom_delta(camera["total"], yoy_camera.get("total"))
    _, crit_mom, _ = mom_delta(speeding.get("critical", 0), pm_speeding.get("critical"))
    _, crit_yoy, _ = mom_delta(speeding.get("critical", 0), yoy_speeding.get("critical"))
    _, camred_mom, _ = mom_delta(camera.get("red", 0), pm_camera.get("red"))
    _, camred_yoy, _ = mom_delta(camera.get("red", 0), yoy_camera.get("red"))
    fleet = [
        ("Metric", "Value", "MoM", "YoY"),
        ("Total Fleet Miles", f"{mileage['total_miles']:,.0f}", miles_mom, miles_yoy),
        ("Speeding Events", str(speeding['total']), speed_mom, speed_yoy),
        ("Critical Speeding", str(speeding.get("critical", 0)), crit_mom, crit_yoy),
        ("Camera Events", str(camera['total']), cam_mom, cam_yoy),
        ("Camera RED Events", str(camera.get("red", 0)), camred_mom, camred_yoy),
    ]
    add_metric_table(doc, fleet)
    _add_text(doc, f"Active Fleet: {mileage['active_trucks']} trucks  \u2014  "
              f"Avg: {mileage['avg_miles']:,.0f} miles/truck",
              size=9, italic=True, color=GRAY)

    # ===== SECTION 2: FLEET MILEAGE =====
    doc.add_page_break()
    add_section_heading(doc, "2. Fleet Mileage")

    if yard:
        # Single yard -- show top vehicles
        yd = mileage["by_yard"].get(yard, {"miles": 0, "trucks": 0, "avg": 0})
        _add_text(doc, f"Total: {yd.get('miles', 0):,.0f} miles  \u2014  "
                  f"Trucks: {yd.get('trucks', 0)}  \u2014  "
                  f"Avg/Truck: {yd.get('avg', 0):,.0f} miles", size=11, bold=True)
    else:
        mile_rows = []
        for y in YARD_ORDER:
            yd = mileage["by_yard"].get(y, {"miles": 0, "trucks": 0, "avg": 0})
            prev_yd = pm_mileage.get("by_yard", {}).get(y, {})
            _, mom_str, _ = mom_delta(yd.get("miles", 0), prev_yd.get("miles"))
            mile_rows.append([
                y,
                f"{yd.get('miles', 0):,.0f}",
                str(yd.get("trucks", 0)),
                f"{yd.get('avg', 0):,.0f}",
                mom_str,
            ])
        add_data_table(doc, ["Yard", "Miles", "Trucks", "Avg/Truck", "MoM"], mile_rows,
                       col_align=[_L, _R, _R, _R, _C])

    # ===== SECTION 3: PRE/POST TRIP INSPECTIONS =====
    add_section_heading(doc, "3. Pre/Post Trip Inspections")

    insp_list = inspections if isinstance(inspections, list) else []
    _add_text(doc, f"Total Inspections: {len(insp_list)}", size=11, bold=True)

    if insp_list:
        # Look for flagged items
        flagged = []
        for insp in insp_list:
            for key, val in insp.items():
                if not val or not isinstance(val, str):
                    continue
                vl = val.strip().lower()
                if any(kw in vl for kw in ("fail", "unsatisfactory", "deficient",
                                            "needs repair", "out of service",
                                            "damage", "broken")):
                    date_str = ""
                    for dk in ("date", "Date"):
                        if insp.get(dk):
                            date_str = _format_display_date(insp[dk])
                            break
                    label = fl.get(key.lower(), key)
                    flagged.append({
                        "date": date_str,
                        "yard": insp.get("_yard", "Unassigned"),
                        "item": label,
                        "finding": val[:100],
                    })
                    break

        if flagged:
            add_section_heading(doc, "Flagged Items", level=2)
            flag_rows = [[f["date"], f["yard"], f["item"], f["finding"]]
                         for f in flagged]
            add_data_table(doc, ["Date", "Yard", "Item", "Finding"], flag_rows, font_size=9,
                           col_align=[_C, _C, _L, _L])
        else:
            _add_text(doc, "No flagged items identified in inspections.", italic=True)
    else:
        _add_text(doc, "COMPLIANCE GAP: No pre/post trip inspections submitted this period. "
                  "All active trucks must have documented inspections.",
                  size=10, bold=True, color=DARK_RED)

    # ===== SECTION 4: SPEEDING =====
    add_section_heading(doc, "4. Speeding Summary")

    if yard:
        _add_text(doc, f"Total: {speeding['total']}  \u2014  Critical: {speeding.get('critical', 0)}  \u2014  "
                  f"High: {speeding.get('high', 0)}  \u2014  Medium: {speeding.get('medium', 0)}",
                  size=11, bold=True)
    else:
        speed_yard_rows = []
        for y in YARD_ORDER:
            count = speeding["by_yard"].get(y, 0)
            prev_count = pm_speeding.get("by_yard", {}).get(y, 0)
            _, mom_str, _ = mom_delta(count, prev_count)
            speed_yard_rows.append([y, str(count), mom_str])
        # Unassigned vehicles as dedicated row in main table
        if speeding.get("unassigned"):
            ua = speeding["unassigned"]
            speed_yard_rows.append(["Unassigned Vehicles", str(ua["events"]), "--"])
        add_data_table(doc, ["Yard", "Events", "MoM"], speed_yard_rows,
                       col_align=[_L, _R, _C])

    # Flag worst yard max overspeed in bold (e.g., Laredo 40.8 mph)
    if not yard and speeding.get("by_yard_max_over"):
        worst_yard = max(speeding["by_yard_max_over"].items(), key=lambda x: x[1])
        if worst_yard[1] > 0:
            _add_text(doc, f"{worst_yard[0]}: max overspeed {worst_yard[1]} mph",
                      size=10, bold=True, color=DARK_RED)

    if speeding.get("top_drivers"):
        add_section_heading(doc, "Top 5 Speeding Drivers", level=2)
        driver_rows = [[d["name"], d.get("yard", ""), str(d["events"]), f"{d['max_over']} mph"]
                       for d in speeding["top_drivers"]]
        add_data_table(doc, ["Driver", "Yard", "Events", "Max Over"], driver_rows,
                       col_align=[_L, _C, _R, _R])

    # ===== SECTION 5: CAMERA EVENTS =====
    add_section_heading(doc, "5. Camera Events Summary")

    _add_text(doc, f"Total: {camera['total']}  \u2014  RED: {camera.get('red', 0)}  \u2014  "
              f"ORANGE: {camera.get('orange', 0)}  \u2014  YELLOW: {camera.get('yellow', 0)}",
              size=11, bold=True)

    # Drowsiness Events Callout (life-safety)
    drowsiness_events = camera.get("drowsiness_events", [])
    if drowsiness_events:
        _add_text(doc, "LIFE-SAFETY ALERT: Drowsiness Events Detected",
                  size=12, bold=True, color=DARK_RED)
        drows_rows = [[d["driver"], d["vehicle"], d.get("yard", "--")]
                      for d in drowsiness_events]
        add_data_table(doc, ["Driver", "Vehicle", "Yard"], drows_rows,
                       font_size=9, col_align=[_L, _C, _C])
        manager_name = (YARD_INFO.get(yard, {}).get("manager", "Division Management")
                        if yard else "Division Management")
        _add_text(doc, f"Action: {manager_name} to conduct fitness-for-duty review "
                  "with each driver listed above.",
                  size=10, bold=True, color=DARK_RED)
    else:
        _add_text(doc, "No drowsiness events detected this period.",
                  size=10, italic=True, color=GREEN)

    if not yard:
        cam_yard_rows = []
        for y in YARD_ORDER:
            yd = camera["by_yard"].get(y, {"red": 0, "orange": 0, "yellow": 0, "total": 0})
            cam_yard_rows.append([
                y, str(yd.get("red", 0)), str(yd.get("orange", 0)),
                str(yd.get("yellow", 0)), str(yd.get("total", 0)),
            ])
        add_data_table(doc, ["Yard", "RED", "ORANGE", "YELLOW", "Total"], cam_yard_rows,
                       col_align=[_L, _R, _R, _R, _R])

    if camera.get("by_type"):
        add_section_heading(doc, "Top Event Types", level=2)
        type_rows = [[etype, str(count)] for etype, count in camera["by_type"].items()]
        add_data_table(doc, ["Event Type", "Count"], type_rows,
                       col_align=[_L, _R])

    if camera.get("repeat_offenders"):
        add_section_heading(doc, "Repeat Offenders (2+ events)", level=2)
        offender_rows = [[o["name"], o.get("yard", ""), str(o["events"])]
                         for o in camera["repeat_offenders"]]
        add_data_table(doc, ["Driver", "Yard", "Events"], offender_rows,
                       col_align=[_L, _C, _R])

    # Unassigned camera events footnote
    if camera.get("unassigned"):
        ua = camera["unassigned"]
        _add_text(doc, f"Note: {ua['events']} camera event(s) from vehicles without "
                  "assigned drivers excluded from repeat offender tracking.",
                  size=9, italic=True, color=GRAY)

    # ===== SECTION 6: FIELD ASSESSMENTS =====
    doc.add_page_break()
    add_section_heading(doc, "6. Field Assessments")

    aa = assessment_analysis
    _add_text(doc, f"Total: {aa['total']}  \u2014  With Findings: {aa['with_findings']}  \u2014  "
              f"Clean: {aa['clean']}", size=11, bold=True)

    if not yard and aa["by_yard"]:
        add_section_heading(doc, "By Yard", level=2)
        assess_yard_rows = []
        for y in YARD_ORDER:
            count = aa["by_yard"].get(y, 0)
            if count > 0:
                assess_yard_rows.append([y, str(count)])
        if assess_yard_rows:
            add_data_table(doc, ["Yard", "Assessments"], assess_yard_rows,
                           col_align=[_L, _R])

    if aa["by_rep"]:
        add_section_heading(doc, "By Safety Rep", level=2)
        rep_rows = [[rep, str(count)] for rep, count in
                    sorted(aa["by_rep"].items(), key=lambda x: x[1], reverse=True)]
        add_data_table(doc, ["Safety Rep", "Assessments"], rep_rows,
                       col_align=[_L, _R])

    # Assessment findings summary (narrative, not line-item detail)
    assess_list = assessments if isinstance(assessments, list) else []
    META_FIELDS_SUMMARY = {
        'report number', 'date', 'observer', 'status', 'link', 'kpa_link',
        'name', 'Name', 'form', 'form_id', 'updated_at', 'created_at',
        'report', 'id', 'response_id', '_yard', '_observer',
        '7vj2l992y7fwqhwz', 'yard', 'location',
        'updated', 'updated_time', 'version', 'observer-emp-num',
        'duration', 'latitude', 'longitude', 'temperature', 'wind-speed',
        'weather', 'parentrepnum', 'parentlink', 'surrogate',
        'select-yes', 'select-no', 'select-n/a', 'select-na',
    }
    finding_categories = Counter()
    for row in assess_list:
        for key, val in row.items():
            if not val or not isinstance(val, str):
                continue
            if key.lower() in META_FIELDS_SUMMARY:
                continue
            vl = val.strip().lower()
            if vl == "no" or any(kw in vl for kw in ("unsatisfactory", "fail", "deficien",
                                                       "corrective", "needs", "damage", "broken")):
                label = fl.get(key.lower(), key)
                if label in ("select-no", "select-yes", "select-n/a"):
                    continue
                finding_categories[label] += 1

    if finding_categories:
        add_section_heading(doc, "Common Findings", level=2)
        top_findings = finding_categories.most_common(5)
        summary_parts = [f"{label} ({count})" for label, count in top_findings]
        _add_text(doc, f"Top finding areas: {', '.join(summary_parts)}. "
                  f"Full assessment detail available upon request.",
                  size=10)

        # V Door systemic finding callout
        for label, count in finding_categories.items():
            if "v door" in label.lower() or "v-door" in label.lower():
                if aa["total"] > 0 and count == aa["total"]:
                    _add_text(doc, f"SYSTEMIC ISSUE: \"{label}\" flagged on {count} of "
                              f"{aa['total']} assessments (100%). This is a division-level "
                              "corrective action item requiring standardized V Door "
                              "operator protocols across all yards.",
                              size=10, bold=True, color=DARK_RED)
                elif count >= aa["total"] * 0.5 and aa["total"] > 0:
                    _add_text(doc, f"RECURRING: \"{label}\" flagged on {count} of "
                              f"{aa['total']} assessments ({round(count/aa['total']*100)}%). "
                              "Requires targeted corrective action.",
                              size=10, bold=True, color=DARK_RED)

    # Days since last assessment for zero-assessment yards
    zero_assess_yards = []
    if not yard:
        for y in YARD_ORDER:
            if aa["by_yard"].get(y, 0) == 0:
                # Check prior month JSON for last assessment
                pm_assess = pm.get("by_yard", {}).get(y, {}).get("assessments_total", 0)
                if pm_assess and pm_assess > 0:
                    pm_month = pm.get("month", "")
                    if pm_month:
                        try:
                            pm_dt = datetime.strptime(pm_month, "%Y-%m")
                            last_str = pm_dt.strftime("%B %Y")
                            days_ago = (start_ct.replace(tzinfo=None) - pm_dt).days
                            zero_assess_yards.append(f"{y}: last assessment {last_str} ({days_ago} days ago)")
                        except ValueError:
                            zero_assess_yards.append(f"{y}: no field assessments conducted")
                    else:
                        zero_assess_yards.append(f"{y}: no field assessments conducted")
                else:
                    zero_assess_yards.append(f"{y}: no field assessments on record")
    elif yard and aa["total"] == 0:
        pm_assess = pm.get("assessments_total", 0)
        if pm_assess and pm_assess > 0:
            pm_month = pm.get("month", "")
            if pm_month:
                try:
                    pm_dt = datetime.strptime(pm_month, "%Y-%m")
                    last_str = pm_dt.strftime("%B %Y")
                    days_ago = (start_ct.replace(tzinfo=None) - pm_dt).days
                    zero_assess_yards.append(f"Last assessment: {last_str} ({days_ago} days ago)")
                except ValueError:
                    zero_assess_yards.append("No field assessments on record")
            else:
                zero_assess_yards.append("No field assessments on record")
        else:
            zero_assess_yards.append("No field assessments on record")

    if zero_assess_yards:
        for note in zero_assess_yards:
            _add_text(doc, note, size=10, italic=True, color=DARK_RED)

    # ===== SECTION 7: OBSERVATIONS =====
    doc.add_page_break()

    obs_list = observations if isinstance(observations, list) else []
    obs_by_type = Counter(r.get(OBS_TYPE_HASH, "").strip() for r in obs_list)
    obs_by_type.pop("", None)  # remove blanks

    # Distinct employee count (by Name field, not KPA submitter)
    obs_by_employee = Counter(_obs_employee(r) for r in obs_list)
    distinct_employees = len(obs_by_employee)

    add_section_heading(doc, f"7. Observations ({len(obs_list)} total -- "
                             f"{distinct_employees} distinct employees)")

    # Observation rate per employee
    obs_rate = round(len(obs_list) / headcount, 2) if headcount > 0 else 0
    _add_text(doc, f"Total Observations: {len(obs_list)}  \u2014  "
              f"Rate: {obs_rate} per employee  \u2014  "
              f"Headcount: {headcount:,}", size=11, bold=True)

    # Employee concentration flag
    if obs_list and obs_by_employee:
        top_employee, top_count = obs_by_employee.most_common(1)[0]
        concentration_pct = round(top_count / len(obs_list) * 100)
        if concentration_pct > 50:
            _add_text(doc, f"NOTE: {concentration_pct}% of observations involve "
                      f"{top_employee} -- broaden crew participation.",
                      size=10, bold=True, color=DARK_RED)

    # Division numbers excluding Midland (company-wide only)
    if not yard and obs_list:
        obs_by_yard_check = Counter(r.get("_yard", "Unassigned") for r in obs_list)
        midland_obs = obs_by_yard_check.get("Midland", 0)
        non_midland_obs = len(obs_list) - midland_obs
        non_midland_hc = EMPLOYEE_COUNT - HEADCOUNT_BY_YARD.get("Midland", 0)
        non_midland_rate = round(non_midland_obs / non_midland_hc, 2) if non_midland_hc > 0 else 0
        _add_text(doc, f"NOTE: Excluding Midland ({midland_obs} obs), remaining yards "
                  f"submitted {non_midland_obs} observations ({non_midland_rate} per employee). "
                  f"Three yards are carrying the division average.",
                  size=10, italic=True, color=GRAY)

    if obs_list:
        # --- Observation Type Breakdown with MoM + YoY ---
        add_section_heading(doc, "Observation Type Breakdown", level=2)

        # Previous month + YoY type data
        pm_obs_types = pm.get("observations_by_type", {})
        yoy_obs_types = yoy.get("observations_by_type", {})

        OBS_TYPE_ORDER = ["Recognition", "At-Risk Condition", "At-Risk Procedure",
                          "At-Risk Behavior", "Near Miss", "Suggestion"]
        type_rows = []
        for otype in OBS_TYPE_ORDER:
            count = obs_by_type.get(otype, 0)
            pct = round(count / len(obs_list) * 100, 1) if len(obs_list) > 0 else 0
            prev_count = pm_obs_types.get(otype, None)
            mom_str = mom_delta(count, prev_count)[1] if prev_count is not None else "N/A"
            yoy_count = yoy_obs_types.get(otype, None)
            yoy_str = mom_delta(count, yoy_count)[1] if yoy_count is not None else "N/A"
            type_rows.append([otype, str(count), f"{pct}%", mom_str, yoy_str])

        # Add any types not in standard order
        for otype, count in obs_by_type.most_common():
            if otype not in OBS_TYPE_ORDER:
                pct = round(count / len(obs_list) * 100, 1)
                type_rows.append([otype, str(count), f"{pct}%", "N/A", "N/A"])

        mom_total = (mom_delta(len(obs_list), pm.get("observations_total"))[1]
                     if pm.get("observations_total") is not None else "N/A")
        yoy_total = (mom_delta(len(obs_list), yoy.get("observations_total"))[1]
                     if yoy.get("observations_total") is not None else "N/A")
        type_rows.append(["Total", str(len(obs_list)), "100%", mom_total, yoy_total])
        add_data_table(doc, ["Type", "Count", "% of Total", "MoM", "YoY"], type_rows,
                       col_align=[_L, _R, _R, _C, _C])

        # Recognition ratio -- key BBS health indicator
        recognition = obs_by_type.get("Recognition", 0)
        at_risk_total = (obs_by_type.get("At-Risk Condition", 0) +
                         obs_by_type.get("At-Risk Procedure", 0) +
                         obs_by_type.get("At-Risk Behavior", 0))
        if at_risk_total > 0:
            ratio = round(recognition / at_risk_total, 2)
            _add_text(doc, f"Recognition-to-At-Risk Ratio: {ratio}:1  "
                      f"(Recognition: {recognition}, At-Risk: {at_risk_total})",
                      size=10, italic=True)

        # --- By Yard (company-wide) ---
        if not yard:
            add_section_heading(doc, "Observations by Yard", level=2)
            obs_by_yard = Counter(r.get("_yard", "Unassigned") for r in obs_list)

            # Yard x Type matrix
            yard_type_matrix = defaultdict(Counter)
            for r in obs_list:
                y = r.get("_yard", "Unassigned")
                t = r.get(OBS_TYPE_HASH, "").strip() or "Other"
                yard_type_matrix[y][t] += 1

            # Build matrix table: Yard | Total | Recognition | At-Risk Cond | At-Risk Proc | At-Risk Behav | Near Miss | Suggestion
            matrix_headers = ["Yard", "Total", "Recog.", "At-Risk C", "At-Risk P", "At-Risk B", "Near Miss", "Suggest."]
            matrix_rows = []
            for y in YARD_ORDER:
                if obs_by_yard.get(y, 0) == 0:
                    continue
                ytc = yard_type_matrix[y]
                matrix_rows.append([
                    y, str(obs_by_yard[y]),
                    str(ytc.get("Recognition", 0)),
                    str(ytc.get("At-Risk Condition", 0)),
                    str(ytc.get("At-Risk Procedure", 0)),
                    str(ytc.get("At-Risk Behavior", 0)),
                    str(ytc.get("Near Miss", 0)),
                    str(ytc.get("Suggestion", 0)),
                ])
            if matrix_rows:
                add_data_table(doc, matrix_headers, matrix_rows, font_size=9,
                               col_align=[_L, _R, _R, _R, _R, _R, _R, _R])

        # --- Top Repeat Observations ---
        add_section_heading(doc, "Top Repeat Observations", level=2)
        # Count observation descriptions to find recurring issues
        obs_desc_counter = Counter()
        obs_desc_type = {}  # track type for each description
        for obs in obs_list:
            desc = obs.get(OBS_DESC_HASH, obs.get("description", "")).strip()
            if not desc or desc == "--":
                continue
            # Normalize: lowercase, strip trailing punctuation
            desc_key = desc.lower().rstrip(".!,;: ")
            if len(desc_key) < 5:
                continue
            obs_desc_counter[desc_key] += 1
            if desc_key not in obs_desc_type:
                obs_desc_type[desc_key] = obs.get(OBS_TYPE_HASH, "").strip() or "--"
        # Show top 10 repeat issues (count >= 2, or top 10 regardless)
        top_issues = obs_desc_counter.most_common(10)
        if top_issues:
            issue_rows = []
            for desc_key, count in top_issues:
                # Capitalize first letter for display
                display_desc = desc_key[:1].upper() + desc_key[1:]
                if len(display_desc) > 100:
                    display_desc = display_desc[:97] + "..."
                otype = obs_desc_type.get(desc_key, "--")
                issue_rows.append([display_desc, otype, str(count)])
            add_data_table(doc, ["Observation", "Type", "Count"], issue_rows,
                           font_size=9, col_align=[_L, _L, _R])
    else:
        _add_text(doc, "No observations reported this period.", italic=True)

    # ===== SECTION 8: INCIDENT ANALYSIS =====
    doc.add_page_break()
    add_section_heading(doc, "8. Incident Analysis")

    if ic and ic.get("total", 0) > 0:
        _add_text(doc, f"Total Incident Reports: {ic['total']}  \u2014  "
                  f"Recordable: {ic.get('recordable', 0)}  \u2014  "
                  f"First Aid: {ic.get('first_aid', 0)}  \u2014  "
                  f"Near Miss: {ic.get('near_miss', 0)}",
                  size=11, bold=True,
                  color=DARK_RED if ic.get("recordable", 0) > 0 else BLACK)

        # Incident Breakdown table
        add_section_heading(doc, "Incident Breakdown by Type", level=2)
        breakdown_rows = []
        type_map = [
            ("Recordable", ic.get("recordable", 0)),
            ("First Aid Only", ic.get("first_aid", 0)),
            ("Near Miss", ic.get("near_miss", 0)),
            ("Vehicle - At-Fault", ic.get("vehicle_at_fault", 0)),
            ("Vehicle - Not At-Fault", ic.get("vehicle_not_at_fault", 0)),
            ("Equipment/Property Damage", ic.get("equipment_damage", 0)),
            ("Report Only", ic.get("report_only", 0)),
            ("Personal Illness", ic.get("personal_illness", 0)),
        ]
        for label, count in type_map:
            if count > 0:
                breakdown_rows.append([label, str(count)])
        breakdown_rows.append(["Total", str(ic["total"])])
        add_data_table(doc, ["Incident Type", "Count"], breakdown_rows,
                       col_align=[_L, _R])

        # Incident Detail table
        add_section_heading(doc, "Incident Detail", level=2)
        inc_detail_rows = []
        for d in ic.get("detail", []):
            inc_detail_rows.append([
                d["date"],
                d["yard"],
                d["type"][:30] if len(d["type"]) <= 30 else d["type"][:27] + "...",
                d["employee"],
                d["description"],
                d.get("rca_status", "No RCA"),
            ])
        if inc_detail_rows:
            tbl = add_data_table(doc, ["Date", "Yard", "Type", "Employee", "Description", "RCA"],
                                 inc_detail_rows, font_size=9,
                                 col_align=[_C, _C, _L, _L, _L, _C])
            # Bold + red for recordable rows
            if tbl:
                for i, d in enumerate(ic.get("detail", [])):
                    if d.get("is_recordable") and i + 1 < len(tbl.rows):
                        for cell in tbl.rows[i + 1].cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = DARK_RED

        # RCA Accountability
        rca_list = rcas or []
        if rca_list:
            add_section_heading(doc, "RCA Accountability", level=2)
            rca_with = sum(1 for d in ic.get("detail", [])
                           if "No RCA" not in d.get("rca_status", "No RCA"))
            turnarounds = [d.get("rca_turnaround") for d in ic.get("detail", [])
                           if d.get("rca_turnaround") is not None]
            avg_turn = round(sum(turnarounds) / len(turnarounds)) if turnarounds else 0

            _add_text(doc, f"RCA Completion: {rca_with} of {ic['total']} incidents  \u2014  "
                      f"Avg Turnaround: {avg_turn} days" if turnarounds else
                      f"RCA Completion: {rca_with} of {ic['total']} incidents",
                      size=11, bold=True)

            # RCA detail table
            rca_detail_rows = []
            for d in ic.get("detail", []):
                if "No RCA" in d.get("rca_status", "No RCA"):
                    continue
                causes = d.get("_rca_causes", "")
                if len(causes) > 80:
                    causes = causes[:77] + "..."
                actions = d.get("_rca_actions", "")
                if len(actions) > 80:
                    actions = actions[:77] + "..."
                rca_detail_rows.append([
                    d["date"],
                    d["yard"],
                    causes or "Available upon request",
                    actions or "Available upon request",
                    d.get("rca_status", ""),
                ])
            if rca_detail_rows:
                add_data_table(doc, ["Inc. Date", "Yard", "Root Cause", "Corrective Actions", "Status"],
                               rca_detail_rows, font_size=9,
                               col_align=[_C, _C, _L, _L, _C])
    else:
        _add_text(doc, "No Casing incidents reported this period.", italic=True)

    # ===== SECTION 9: TRAINING COMPLIANCE =====
    doc.add_page_break()
    add_section_heading(doc, "9. Training Compliance")

    if tc and tc.get("total_employees", 0) > 0:
        overall_pct = tc.get("overall_pct", 0)
        pct_color = BLACK if overall_pct >= 95 else DARK_RED
        _add_text(doc, f"Overall Compliance: {overall_pct:.1f}%  \u2014  "
                  f"Employees: {tc['total_employees']:,}  \u2014  "
                  f"Compliant: {tc['compliant_count']:,}  \u2014  "
                  f"Overdue: {tc['overdue_count']}", size=11, bold=True, color=pct_color)

        # Table of non-compliant employees
        non_compliant = [e for e in tc.get("employees", []) if e["percent_complete"] < 100]
        non_compliant.sort(key=lambda e: e.get("days_since_assignment", 0), reverse=True)
        if non_compliant:
            add_section_heading(doc, "Non-Compliant Employees", level=2)
            nc_rows = []
            for emp in non_compliant[:40]:  # cap at 40
                incomplete = ", ".join(emp.get("incomplete_training_names", [])[:3])
                if len(emp.get("incomplete_training_names", [])) > 3:
                    incomplete += f" (+{len(emp['incomplete_training_names']) - 3} more)"
                days = emp.get("days_since_assignment", 0)
                nc_rows.append([
                    emp["employee_name"],
                    emp.get("yard", ""),
                    f"{emp['percent_complete']}%",
                    str(days) if days > 0 else "--",
                    incomplete,
                ])
            tbl = add_data_table(doc,
                                 ["Employee", "Yard", "Complete", "Days Since Assign.", "Incomplete Programs"],
                                 nc_rows, font_size=9,
                                 col_align=[_L, _C, _R, _R, _L])
            # Color code days: >90 = bold red, 60-90 = dark orange
            if tbl:
                for i, emp in enumerate(non_compliant[:40]):
                    days = emp.get("days_since_assignment", 0)
                    if days > 60 and i + 1 < len(tbl.rows):
                        day_cell = tbl.rows[i + 1].cells[3]
                        for p in day_cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True
                                if days > 90:
                                    run.font.color.rgb = DARK_RED
                                else:
                                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
        else:
            _add_text(doc, "All employees are fully compliant with assigned training.", italic=True)
    else:
        _add_text(doc, "Training compliance data not available.", italic=True)

    # ===== SECTION 10: YARD COMPARISON SCORECARD (company-wide only) =====
    if not yard:
        doc.add_page_break()
        add_section_heading(doc, "10. Yard Comparison Scorecard")

        sc_headers = ["Yard", "Miles", "Speed Evt", "Cam RED", "Assessments",
                      "Observations", "Incidents", "Training %", "Tier"]
        sc_rows = []
        sc_tiers = []  # track tier per row for coloring
        inc_by_yard = ic.get("by_yard", {}) if ic else Counter(r.get("_yard", "Unassigned") for r in (incidents if isinstance(incidents, list) else []))
        obs_by_yard_sc = Counter(r.get("_yard", "Unassigned") for r in obs_list)

        # Training compliance by yard
        tc_by_yard = {}
        for emp in tc.get("employees", []):
            y = emp.get("yard", "Unassigned")
            if y not in tc_by_yard:
                tc_by_yard[y] = {"total": 0, "compliant": 0}
            tc_by_yard[y]["total"] += 1
            if emp["percent_complete"] >= 100:
                tc_by_yard[y]["compliant"] += 1

        for y in YARD_ORDER:
            ym = mileage["by_yard"].get(y, {})
            tc_yd = tc_by_yard.get(y, {"total": 0, "compliant": 0})
            tc_pct = round(tc_yd["compliant"] / tc_yd["total"] * 100, 1) if tc_yd["total"] > 0 else 0
            y_trucks = ym.get("trucks", 0)
            y_speed = speeding["by_yard"].get(y, 0)
            y_cam_red = camera["by_yard"].get(y, {}).get("red", 0)
            y_assess = aa["by_yard"].get(y, 0)
            y_obs = obs_by_yard_sc.get(y, 0)
            y_inc = inc_by_yard.get(y, 0)
            y_hc = HEADCOUNT_BY_YARD.get(y, 1)

            # Performance tier: worst single metric = overall tier
            metric_tiers = []
            speed_per_truck = y_speed / y_trucks if y_trucks > 0 else 0
            obs_per_emp = y_obs / y_hc if y_hc > 0 else 0

            # Training: >=95%=GREEN, <70%=RED, else YELLOW
            if tc_pct >= 95:
                metric_tiers.append("GREEN")
            elif tc_pct < 70:
                metric_tiers.append("RED")
            else:
                metric_tiers.append("YELLOW")
            # Observations per employee: >=1.0=GREEN, <0.1=RED, else YELLOW
            if obs_per_emp >= 1.0:
                metric_tiers.append("GREEN")
            elif obs_per_emp < 0.1:
                metric_tiers.append("RED")
            else:
                metric_tiers.append("YELLOW")
            # Assessments: >=1=GREEN (completed), 0=RED
            if y_assess >= 1:
                metric_tiers.append("GREEN")
            else:
                metric_tiers.append("RED")
            # Speeding per truck: <=3.0=GREEN, >7.0=RED, else YELLOW
            if speed_per_truck <= 3.0:
                metric_tiers.append("GREEN")
            elif speed_per_truck > 7.0:
                metric_tiers.append("RED")
            else:
                metric_tiers.append("YELLOW")

            # Overall = worst tier
            if "RED" in metric_tiers:
                overall_tier = "RED"
            elif "YELLOW" in metric_tiers:
                overall_tier = "YELLOW"
            else:
                overall_tier = "GREEN"

            sc_tiers.append(overall_tier)
            sc_rows.append([
                y,
                f"{ym.get('miles', 0):,.0f}",
                str(y_speed),
                str(y_cam_red),
                str(y_assess),
                str(y_obs),
                str(y_inc),
                f"{tc_pct}%",
                overall_tier,
            ])
        tbl = add_data_table(doc, sc_headers, sc_rows,
                             col_align=[_L, _R, _R, _R, _R, _R, _R, _R, _C])
        # Color the Tier column text
        if tbl:
            tier_colors = {
                "RED": DARK_RED,
                "YELLOW": RGBColor(0xCC, 0x66, 0x00),
                "GREEN": GREEN,
            }
            for i, tier in enumerate(sc_tiers):
                if i + 1 < len(tbl.rows):
                    tier_cell = tbl.rows[i + 1].cells[8]  # Tier column
                    for p in tier_cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = tier_colors.get(tier, BLACK)

        # Tier legend
        _add_text(doc, "Tier: GREEN = Training >=95%, Obs >=1.0/emp, Assessments completed, "
                  "Speeding <=3.0/truck | RED = Training <70%, Obs <0.1/emp, "
                  "Zero assessments, Speeding >7.0/truck | YELLOW = everything else",
                  size=9, italic=True, color=GRAY)

    # ===== SECTION 11 (or 10 for yard): ACTION ITEMS =====
    doc.add_page_break()
    section_num = "10" if yard else "11"
    add_section_heading(doc, f"{section_num}. Action Items & Takeaways")

    # Prior Month Action Item Status
    pm_action_items = pm.get("action_items", [])
    if pm_action_items:
        add_section_heading(doc, "Prior Month Action Item Status", level=2)
        pm_ai_rows = []
        for ai in pm_action_items:
            pm_ai_rows.append([
                ai.get("item", ""),
                ai.get("owner", ""),
                ai.get("due_date", ""),
                ai.get("status", "Carried Over"),
            ])
        add_data_table(doc, ["Item", "Owner", "Due Date", "Status"], pm_ai_rows,
                       font_size=9, col_align=[_L, _L, _C, _C])
    else:
        _add_text(doc, "No prior month action items on file.",
                  size=10, italic=True, color=GRAY)

    add_section_heading(doc, "Current Period Action Items", level=2)

    # Calculate due date for action items (30 days from report generation)
    report_gen_date = datetime.now()
    action_due_date = (report_gen_date + timedelta(days=30)).strftime("%B %d, %Y")

    takeaways = _generate_takeaways(
        mileage, speeding, camera, assessment_analysis,
        obs_list, ic or incidents, inspections, tc, pm, yard,
        man_hours=man_hours,
    )
    # Change 9: Add due dates to every action item
    for item in takeaways:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{item} [Due: {action_due_date}]")
        run.font.size = Pt(11)
        run.font.name = CALIBRI

    doc.save(output_path)
    print(f"  Report saved: {output_path}")


# ==============================================================================
# YARD FILTERING HELPERS
# ==============================================================================

def _filter_mileage_to_yard(mileage, yard):
    """Extract single yard's mileage data."""
    yd = mileage["by_yard"].get(yard, {"miles": 0, "trucks": 0, "avg": 0})
    return {
        "total_miles": yd.get("miles", 0),
        "active_trucks": yd.get("trucks", 0),
        "avg_miles": yd.get("avg", 0),
        "by_yard": {yard: yd},
    }


def _filter_speeding_to_yard(raw_events, yard):
    """Filter speeding events to single yard and re-process."""
    yard_events = [e for e in raw_events if e.get("yard") == yard]
    return process_speeding(yard_events)


def _filter_camera_to_yard(raw_events, yard):
    """Filter camera events to single yard and re-process."""
    yard_events = [e for e in raw_events if e.get("yard") == yard]
    return process_camera_events(yard_events)


def _filter_form_activity_to_yard(form_activity, yard):
    """Filter KPA form activity to single yard."""
    filtered = {}
    for form_id, fa in form_activity.items():
        yard_count = fa["by_yard"].get(yard, 0)
        filtered[form_id] = {
            "name": fa["name"],
            "total": yard_count,
            "by_yard": {yard: yard_count},
        }
    return filtered


def _filter_training_to_yard(training_compliance, yard):
    """Filter training compliance to single yard."""
    if not training_compliance:
        return {"employees": [], "overall_pct": 0, "total_employees": 0,
                "compliant_count": 0, "overdue_count": 0}
    yard_emps = [e for e in training_compliance.get("employees", [])
                 if e.get("yard") == yard]
    total = len(yard_emps)
    compliant = sum(1 for e in yard_emps if e["percent_complete"] >= 100)
    overdue = sum(1 for e in yard_emps if e["status"] == "Overdue")
    overall_pct = round(compliant / total * 100, 1) if total > 0 else 0
    return {
        "employees": yard_emps,
        "overall_pct": overall_pct,
        "total_employees": total,
        "compliant_count": compliant,
        "overdue_count": overdue,
    }


def _filter_incidents_to_yard(incident_classification, yard):
    """Filter classified incident data to a single yard."""
    if not incident_classification or not incident_classification.get("detail"):
        return {"total": 0, "recordable": 0, "first_aid": 0, "near_miss": 0,
                "vehicle_at_fault": 0, "vehicle_not_at_fault": 0,
                "equipment_damage": 0, "report_only": 0, "personal_illness": 0,
                "by_yard": {}, "detail": []}
    yard_detail = [d for d in incident_classification["detail"] if d.get("yard") == yard]
    ic = {
        "total": len(yard_detail),
        "recordable": sum(1 for d in yard_detail if d.get("is_recordable")),
        "first_aid": sum(1 for d in yard_detail if "First Aid" in d.get("type", "")),
        "near_miss": sum(1 for d in yard_detail if "Near Miss" in d.get("type", "")),
        "vehicle_at_fault": sum(1 for d in yard_detail if "At-Fault" in d.get("type", "") and "Not" not in d.get("type", "")),
        "vehicle_not_at_fault": sum(1 for d in yard_detail if "Not At-Fault" in d.get("type", "") or "Not-At-Fault" in d.get("type", "")),
        "equipment_damage": sum(1 for d in yard_detail if "Equipment" in d.get("type", "") or "Property" in d.get("type", "")),
        "report_only": sum(1 for d in yard_detail if "Report Only" in d.get("type", "")),
        "personal_illness": sum(1 for d in yard_detail if "Personal Illness" in d.get("type", "")),
        "by_yard": {yard: len(yard_detail)} if yard_detail else {},
        "detail": yard_detail,
    }
    return ic


def _filter_rcas_to_yard(rcas, yard):
    """Filter RCA list to a single yard."""
    if not rcas:
        return []
    return [r for r in rcas if r.get("_yard", "") == yard]


# ==============================================================================
# MOM DATA PERSISTENCE
# ==============================================================================

def save_month_data(month_str, mileage, speeding, camera, form_activity, assessments,
                    observations, incidents, inspections, training_compliance,
                    incident_classification=None, field_labels=None,
                    raw_assessments=None):
    """Save current month data as JSON for future MoM comparison."""
    os.makedirs(DATA_DIR, exist_ok=True)
    tc = training_compliance or {}
    aa = assessments  # assessment_analysis dict
    ic = incident_classification or {}
    fl = field_labels or {}

    recordable = ic.get("recordable", 0)
    first_aid = ic.get("first_aid", 0)
    near_miss = ic.get("near_miss", 0)
    total_inc = ic.get("total", len(incidents) if isinstance(incidents, list) else 0)
    trir = round(recordable * 200000 / MONTHLY_MAN_HOURS, 2) if MONTHLY_MAN_HOURS > 0 else 0
    dart = trir
    fafr = round(first_aid * 200000 / MONTHLY_MAN_HOURS, 2) if MONTHLY_MAN_HOURS > 0 and first_aid > 0 else 0

    # RCA stats
    rca_with = sum(1 for d in ic.get("detail", [])
                   if "No RCA" not in d.get("rca_status", "No RCA"))
    turnarounds = [d.get("rca_turnaround") for d in ic.get("detail", [])
                   if d.get("rca_turnaround") is not None]
    avg_turn = round(sum(turnarounds) / len(turnarounds)) if turnarounds else 0

    # Build per-yard observation type breakdowns
    obs_list = observations if isinstance(observations, list) else []
    obs_by_yard_type = defaultdict(Counter)
    obs_by_yard_total = Counter()
    for r in obs_list:
        y = r.get("_yard", "Unassigned")
        t = r.get(OBS_TYPE_HASH, "").strip()
        if t:
            obs_by_yard_type[y][t] += 1
        obs_by_yard_total[y] += 1

    # Build per-yard form activity totals
    fa_by_yard = Counter()
    for fa_entry in form_activity.values():
        for y, cnt in fa_entry.get("by_yard", {}).items():
            fa_by_yard[y] += cnt

    # Build per-yard assessment totals
    assess_by_yard = aa.get("by_yard", {}) if isinstance(aa, dict) else {}

    # Build per-yard inspection totals
    insp_list = inspections if isinstance(inspections, list) else []
    insp_by_yard = Counter(i.get("_yard", "Unassigned") for i in insp_list)

    # Build per-yard incident classification
    inc_by_yard_detail = defaultdict(list)
    for d in ic.get("detail", []):
        inc_by_yard_detail[d.get("yard", "Unassigned")].append(d)

    by_yard_data = {}
    for y in YARD_ORDER:
        y_man_hours = MONTHLY_MAN_HOURS_BY_YARD.get(y, 0)
        y_detail = inc_by_yard_detail.get(y, [])
        y_rec = sum(1 for d in y_detail if d.get("is_recordable"))
        y_fa = sum(1 for d in y_detail if "First Aid" in d.get("type", ""))
        y_nm = sum(1 for d in y_detail if "Near Miss" in d.get("type", ""))
        y_vaf = sum(1 for d in y_detail if "At-Fault" in d.get("type", "") and "Not" not in d.get("type", ""))
        y_trir = round(y_rec * 200000 / y_man_hours, 2) if y_man_hours > 0 else 0
        y_fafr = round(y_fa * 200000 / y_man_hours, 2) if y_man_hours > 0 and y_fa > 0 else 0

        yd_mileage = mileage["by_yard"].get(y, {})
        yd_speeding = speeding.get("by_yard", {}).get(y, 0)
        yd_camera = camera.get("by_yard", {}).get(y, {})

        by_yard_data[y] = {
            "mileage": {
                "total_miles": yd_mileage.get("miles", 0),
                "active_trucks": yd_mileage.get("trucks", 0),
                "avg_miles": yd_mileage.get("avg", 0),
            },
            "speeding": {"total": yd_speeding, "critical": 0},
            "camera": {
                "total": yd_camera.get("total", 0) if isinstance(yd_camera, dict) else 0,
                "red": yd_camera.get("red", 0) if isinstance(yd_camera, dict) else 0,
                "orange": yd_camera.get("orange", 0) if isinstance(yd_camera, dict) else 0,
                "yellow": yd_camera.get("yellow", 0) if isinstance(yd_camera, dict) else 0,
            },
            "observations_total": obs_by_yard_total.get(y, 0),
            "observations_by_type": dict(obs_by_yard_type.get(y, {})),
            "assessments_total": assess_by_yard.get(y, 0),
            "inspections_total": insp_by_yard.get(y, 0),
            "form_activity_total": fa_by_yard.get(y, 0),
            "incidents_total": len(y_detail),
            "incidents_recordable": y_rec,
            "incidents_first_aid": y_fa,
            "incidents_near_miss": y_nm,
            "incidents_vehicle_at_fault": y_vaf,
            "trir": y_trir,
            "dart": y_trir,
            "fafr": y_fafr,
        }

    # Distinct employee counts (by Name field, not KPA submitter)
    obs_employees = Counter(_obs_employee(r) for r in obs_list)
    distinct_observers = len(obs_employees)
    obs_by_yard_employees = defaultdict(set)
    for r in obs_list:
        y = r.get("_yard", "Unassigned")
        obs_by_yard_employees[y].add(_obs_employee(r))
    distinct_observers_by_yard = {y: len(s) for y, s in obs_by_yard_employees.items()}

    # Assessment finding categories (for repeat findings tracker)
    META_FIELDS_CACHE = {
        'report number', 'date', 'observer', 'status', 'link', 'kpa_link',
        'name', 'Name', 'form', 'form_id', 'updated_at', 'created_at',
        'report', 'id', 'response_id', '_yard', '_observer',
        '7vj2l992y7fwqhwz', 'yard', 'location',
        'updated', 'updated_time', 'version', 'observer-emp-num',
        'duration', 'latitude', 'longitude', 'temperature', 'wind-speed',
        'weather', 'parentrepnum', 'parentlink', 'surrogate',
        'select-yes', 'select-no', 'select-n/a', 'select-na',
    }
    assessment_finding_cats = Counter()
    raw_assess_list = raw_assessments if isinstance(raw_assessments, list) else []
    for row in raw_assess_list:
        for key, val in row.items():
            if not val or not isinstance(val, str):
                continue
            if key.lower() in META_FIELDS_CACHE:
                continue
            vl = val.strip().lower()
            if vl == "no" or any(kw in vl for kw in ("unsatisfactory", "fail", "deficien",
                                                       "corrective", "needs", "damage", "broken")):
                label = fl.get(key.lower(), key)
                if label in ("select-no", "select-yes", "select-n/a"):
                    continue
                assessment_finding_cats[label] += 1

    # Speeding unassigned count
    speeding_unassigned = speeding.get("unassigned", {}).get("events", 0) if speeding.get("unassigned") else 0

    # Drowsiness events
    drowsiness_events_cache = camera.get("drowsiness_events", [])

    # Build action items list for prior month status tracking
    _tc_cache = tc
    _aa_cache = aa if isinstance(aa, dict) else {"total": 0, "with_findings": 0, "clean": 0, "by_yard": {}, "by_rep": {}}
    _ic_cache = ic
    _obs_cache = obs_list
    _insp_cache = insp_list
    _cam_cache = camera
    _speed_cache = speeding
    _mil_cache = mileage
    _pm_cache = {}  # no nested MoM for action items
    cache_takeaways = _generate_takeaways(
        _mil_cache, _speed_cache, _cam_cache, _aa_cache,
        _obs_cache, _ic_cache or incidents, _insp_cache, _tc_cache, _pm_cache,
    )
    report_gen = datetime.now()
    cache_due = (report_gen + timedelta(days=30)).strftime("%B %d, %Y")
    action_items_cache = []
    for tw in cache_takeaways:
        owner = "Division Management"
        # Extract owner from text if present
        for y_name, y_info in YARD_INFO.items():
            if y_info.get("manager", "") in tw:
                owner = y_info["manager"]
                break
            if y_info.get("safety_reps", "") in tw:
                owner = y_info["safety_reps"]
                break
        action_items_cache.append({
            "item": tw[:120],
            "owner": owner,
            "due_date": cache_due,
            "status": "Open",
        })

    data = {
        "month": month_str,
        "mileage": {
            "total_miles": mileage["total_miles"],
            "active_trucks": mileage["active_trucks"],
            "avg_miles": mileage["avg_miles"],
            "by_yard": mileage["by_yard"],
        },
        "speeding": {
            "total": speeding["total"],
            "critical": speeding.get("critical", 0),
            "by_yard": speeding.get("by_yard", {}),
        },
        "camera": {
            "total": camera["total"],
            "red": camera.get("red", 0),
            "orange": camera.get("orange", 0),
            "yellow": camera.get("yellow", 0),
            "by_yard": camera.get("by_yard", {}),
        },
        "form_activity_total": sum(f["total"] for f in form_activity.values()),
        "assessments_total": aa.get("total", 0) if isinstance(aa, dict) else 0,
        "observations_total": len(obs_list),
        "observations_by_type": dict(Counter(
            r.get(OBS_TYPE_HASH, "").strip() for r in obs_list
            if r.get(OBS_TYPE_HASH, "").strip()
        )),
        "incidents_total": total_inc,
        "incidents_recordable": recordable,
        "incidents_first_aid": first_aid,
        "incidents_near_miss": near_miss,
        "incidents_vehicle_at_fault": ic.get("vehicle_at_fault", 0),
        "inspections_total": len(insp_list),
        "training_compliance_pct": tc.get("overall_pct", 0),
        "trir": trir,
        "dart": dart,
        "fafr": fafr,
        "rca_completion": f"{rca_with}/{total_inc}" if total_inc > 0 else "0/0",
        "rca_avg_turnaround": avg_turn,
        "man_hours": MONTHLY_MAN_HOURS,
        "headcount": EMPLOYEE_COUNT,
        "headcount_by_yard": dict(HEADCOUNT_BY_YARD),
        "distinct_observers": distinct_observers,
        "distinct_observers_by_yard": distinct_observers_by_yard,
        "speeding_unassigned": speeding_unassigned,
        "drowsiness_events": drowsiness_events_cache,
        "assessment_finding_categories": dict(assessment_finding_cats),
        "action_items": action_items_cache,
        "by_yard": by_yard_data,
    }
    path = os.path.join(DATA_DIR, f"{month_str}.json")
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    print(f"  Saved month data: {path}")


def load_month_data(month_str):
    """Load cached month data for MoM comparison."""
    path = os.path.join(DATA_DIR, f"{month_str}.json")
    if os.path.exists(path):
        with open(path) as f:
            data = json.load(f)
        print(f"  Loaded previous month data: {path}")
        return data
    return None


def _extract_yard_data(cached_data, yard):
    """Extract yard-specific comparison data from cached month JSON.

    If the cached data has a 'by_yard' key (new format), use yard-specific
    metrics. Otherwise fall back to company-wide data (old format).
    """
    if not cached_data:
        return None
    by_yard = cached_data.get("by_yard", {})
    yd = by_yard.get(yard)
    if yd:
        # New format -- yard-specific data available
        return yd
    # Old format fallback -- return company-wide (better than nothing)
    return cached_data


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(description="Casing Division Monthly HSE Recap")
    parser.add_argument("--month", required=True, help="Report month (YYYY-MM)")
    parser.add_argument("--output", help="Output DOCX path (default: auto-generated)")
    parser.add_argument("--skip-motive", action="store_true",
                        help="Skip Motive API calls (KPA only)")
    parser.add_argument("--skip-yard-reports", action="store_true",
                        help="Skip per-yard report generation")
    args = parser.parse_args()

    # Validate environment
    if not args.skip_motive and not MOTIVE_API_KEY:
        print("ERROR: MOTIVE_API_KEY not set. Use --skip-motive for KPA-only mode.")
        sys.exit(1)
    if not KPA_API_TOKEN:
        print("ERROR: KPA_API_TOKEN not set.")
        sys.exit(1)

    start_ct, end_ct = parse_month(args.month)
    print(f"\nCasing Division Monthly HSE Recap: {month_label(start_ct, end_ct)}")
    print(f"  Central Time window: {start_ct} to {end_ct}")

    output_path = args.output or os.path.join(
        SCRIPT_DIR, f"Casing_Monthly_HSE_Recap_{args.month}.docx"
    )

    # ---- MOTIVE DATA ----
    if not args.skip_motive:
        print("\n[1/11] Fetching Casing vehicles...")
        vehicles = get_casing_vehicles()

        print("\n[2/11] Pulling fleet mileage...")
        mileage_raw = get_casing_mileage(vehicles, start_ct, end_ct)

        print("\n[3/11] Pulling speeding events...")
        speeding_events = get_casing_speeding(vehicles, start_ct, end_ct)

        print("\n[4/11] Pulling camera events...")
        camera_events = get_casing_camera_events(vehicles, start_ct, end_ct)
    else:
        print("\n[1-4/11] Skipping Motive API (--skip-motive)")
        vehicles = {}
        mileage_raw = {"total_miles": 0, "active_trucks": 0, "avg_miles": 0, "by_yard": {}}
        speeding_events = []
        camera_events = []

    speeding_summary = process_speeding(speeding_events)
    camera_summary = process_camera_events(camera_events)

    # ---- KPA DATA ----
    print("\n[5/11] Pulling KPA form activity...")
    form_activity = get_kpa_form_activity(start_ct, end_ct)

    print("\n[6a/11] Pulling observations, incidents, field assessments...")
    observations = get_kpa_observations(start_ct, end_ct)
    incidents = get_kpa_incidents(start_ct, end_ct)
    assessments = get_kpa_assessments(start_ct, end_ct)
    assessment_analysis = analyze_assessments(assessments)

    print("\n[6b/11] Pulling RCA reports...")
    rcas = get_casing_rcas(start_ct, end_ct)

    print("\n[6c/11] Classifying incidents...")
    incident_classification = classify_incidents(incidents)
    cross_reference_rcas(incident_classification, rcas)
    rec = incident_classification.get("recordable", 0)
    fa = incident_classification.get("first_aid", 0)
    nm = incident_classification.get("near_miss", 0)
    tot = incident_classification.get("total", 0)
    print(f"    Total: {tot}  |  Recordable: {rec}  |  First Aid: {fa}  |  Near Miss: {nm}")
    rca_matched = sum(1 for d in incident_classification.get("detail", [])
                      if "No RCA" not in d.get("rca_status", "No RCA"))
    print(f"    RCAs matched: {rca_matched} of {tot} incidents")

    print("\n[7/11] Pulling pre/post trip inspections...")
    inspections = get_casing_vehicle_inspections(start_ct, end_ct)

    print("\n[8/11] Fetching assessment field labels...")
    field_labels = get_form_field_labels(FIELD_ASSESSMENT_FORM)
    print(f"    Field labels: {len(field_labels)} fields mapped")

    print("\n[9/11] Pulling training compliance data...")
    end_date_str = end_ct.strftime("%Y-%m-%d")
    training_compliance = get_casing_training_compliance(end_date_str)

    # Use live headcount from KPA active employees (not hardcoded)
    global HEADCOUNT_BY_YARD, EMPLOYEE_COUNT
    live_hc = training_compliance.get("headcount_by_yard", {})
    live_total = training_compliance.get("headcount_total", 0)
    if live_hc and live_total > 0:
        HEADCOUNT_BY_YARD = dict(live_hc)
        EMPLOYEE_COUNT = live_total
        print(f"    Live headcount: {EMPLOYEE_COUNT} total")

    # ---- PREVIOUS MONTH for MoM + YoY ----
    print("\n[10/11] Loading comparison data (MoM + YoY)...")
    prev_m = prev_month_str(args.month)
    prev_data = load_month_data(prev_m)
    if not prev_data:
        print(f"  No cached data for {prev_m}. MoM columns will show N/A.")
        print(f"  Run this script for {prev_m} first to generate baseline data.")

    yoy_m = yoy_month_str(args.month)
    yoy_data = load_month_data(yoy_m)
    if not yoy_data:
        print(f"  No cached data for {yoy_m}. YoY columns will show N/A.")
        print(f"  Run this script for {yoy_m} first to generate YoY baseline.")

    # ---- GENERATE COMPANY-WIDE REPORT ----
    print("\n[11/11] Generating company-wide DOCX report...")
    generate_report(
        args.month, mileage_raw, speeding_summary, camera_summary,
        form_activity, assessments, assessment_analysis,
        observations, incidents, inspections, training_compliance,
        prev_data, output_path, field_labels=field_labels,
        incident_classification=incident_classification, rcas=rcas,
        yoy_data=yoy_data,
    )

    # ---- SAVE CURRENT MONTH DATA ----
    print("\nSaving month data for future MoM comparison...")
    save_month_data(args.month, mileage_raw, speeding_summary, camera_summary,
                    form_activity, assessment_analysis, observations, incidents,
                    inspections, training_compliance,
                    incident_classification=incident_classification,
                    field_labels=field_labels,
                    raw_assessments=assessments)

    # ---- GENERATE PER-YARD REPORTS ----
    if not args.skip_yard_reports:
        print("\nGenerating per-yard reports...")
        for yard in YARD_ORDER:
            print(f"\n  --- {yard} ---")
            yard_output = os.path.join(
                SCRIPT_DIR, f"Casing_{yard}_Monthly_HSE_Recap_{args.month}.docx"
            )

            # Filter all data to this yard
            yard_mileage = _filter_mileage_to_yard(mileage_raw, yard)
            yard_speeding = _filter_speeding_to_yard(speeding_events, yard)
            yard_camera = _filter_camera_to_yard(camera_events, yard)
            yard_fa = _filter_form_activity_to_yard(form_activity, yard)
            yard_assessments = [a for a in assessments if a.get("_yard") == yard]
            yard_aa = analyze_assessments(yard_assessments)
            yard_obs = [o for o in observations if o.get("_yard") == yard]
            yard_inc = [i for i in incidents if i.get("_yard") == yard]
            yard_insp = [i for i in inspections if i.get("_yard") == yard]
            yard_tc = _filter_training_to_yard(training_compliance, yard)

            # Per-yard incident classification and RCA filtering
            yard_ic = _filter_incidents_to_yard(incident_classification, yard)
            yard_rcas = _filter_rcas_to_yard(rcas, yard)
            cross_reference_rcas(yard_ic, yard_rcas)

            # Extract yard-specific MoM/YoY comparison data
            yard_prev = _extract_yard_data(prev_data, yard)
            yard_yoy = _extract_yard_data(yoy_data, yard)

            generate_report(
                args.month, yard_mileage, yard_speeding, yard_camera,
                yard_fa, yard_assessments, yard_aa,
                yard_obs, yard_inc, yard_insp, yard_tc,
                yard_prev, yard_output, field_labels=field_labels, yard=yard,
                incident_classification=yard_ic, rcas=yard_rcas,
                yoy_data=yard_yoy,
            )

    print("\nDone!")


if __name__ == "__main__":
    main()
