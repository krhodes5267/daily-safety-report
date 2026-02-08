#!/usr/bin/env python3
"""
DAILY CASING CAMERA EVENTS REPORT - AUTOMATED (GitHub Actions)
===============================================================
Runs daily at 5:15 AM Central via GitHub Actions.

Pulls camera-detected safety events from Motive AI dashcams for the
CASING DIVISION ONLY (yesterday's full day, Central Time) and generates:
- Word document (.docx) in LANDSCAPE, grouped by Casing yard
- HTML email with the same structure, sent via Gmail SMTP

Uses /v2/driver_performance_events endpoint.
Events are wrapped as {"driver_performance_events": [{"driver_performance_event": {data}}]}.
Cross-references /v1/vehicles for driver names and Casing group membership.
All API speeds that arrive in km/h are converted to mph (* 0.621371).

Tier Classification:
- RED:    Immediate Action (distraction, cell phone, drowsiness, close following,
          forward collision warning, collision, near collision, stop sign,
          unsafe lane change, lane swerving)
- ORANGE: Coaching Required (hard brake, seatbelt, camera obstruction, smoking)
          Also the default for any unknown event types.
- YELLOW: Monitoring (hard accel, hard corner, speed violation)
"""

import requests
import smtplib
import os
import sys
import json
from datetime import datetime, timedelta, timezone
from html import escape as html_escape
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from collections import Counter, OrderedDict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

try:
    from zoneinfo import ZoneInfo
    CENTRAL_TZ = ZoneInfo("America/Chicago")
except Exception:
    CENTRAL_TZ = timezone(timedelta(hours=-6))  # Fallback to CST

# ==============================================================================
# CONFIGURATION
# ==============================================================================

MOTIVE_API_KEY = os.environ.get("MOTIVE_API_KEY")
if not MOTIVE_API_KEY:
    print("ERROR: MOTIVE_API_KEY environment variable is not set.")
    sys.exit(1)

MOTIVE_BASE_URL_V1 = "https://api.gomotive.com/v1"
MOTIVE_BASE_URL_V2 = "https://api.gomotive.com/v2"
KMH_TO_MPH = 0.621371
LOGOS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logos")
CALIBRI = "Calibri"

# ==============================================================================
# CASING GROUP IDS (from Motive Groups API)
# ==============================================================================

CASING_GROUP_IDS = {
    167175: "Midland",
    169090: "Bryan",
    169092: "Kilgore",
    186740: "Hobbs",
    169091: "Jourdanton",
    186739: "Laredo",
    186741: "San Angelo",
    186746: "",  # Parent "Casing" group (no specific yard)
}

ALL_CASING_GROUP_IDS = set(CASING_GROUP_IDS.keys())

# ==============================================================================
# YARD CONFIGURATION
# ==============================================================================

YARD_ORDER = ["Midland", "Bryan", "Kilgore", "Hobbs", "Jourdanton", "Laredo", "San Angelo"]

YARD_INFO = {
    "Midland": {
        "safety_reps": "Michael Hancock & Michael Salazar",
        "manager": "Richie Bentley",
    },
    "Bryan": {
        "safety_reps": "Justin Conrad",
        "manager": "Danny Lohse",
    },
    "Kilgore": {
        "safety_reps": "James Barnett (J.P.)",
        "manager": "Frankie Balderas",
    },
    "Hobbs": {
        "safety_reps": "Allen Batts",
        "manager": "Clifton Eaves",
    },
    "Jourdanton": {
        "safety_reps": "Joey Speyrer",
        "manager": "Enrique Flores",
    },
    "Laredo": {
        "safety_reps": "Joey Speyrer",
        "manager": "Chris Jacobo",
    },
    "San Angelo": {
        "safety_reps": "Michael Hancock & Michael Salazar",
        "manager": "Jeremy Jones",
    },
}

# ==============================================================================
# EVENT TYPE CLASSIFICATION
# ==============================================================================

# Normalize event type strings from API to canonical names
EVENT_TYPE_NORMALIZE = {
    # RED types
    "distraction": "distraction",
    "distracted_driving": "distraction",
    "driver_distraction": "distraction",
    "cell_phone": "cell_phone",
    "cell_phone_usage": "cell_phone",
    "phone_use": "cell_phone",
    "cellphone": "cell_phone",
    "phone_usage": "cell_phone",
    "drowsiness": "drowsiness",
    "drowsy": "drowsiness",
    "drowsy_driving": "drowsiness",
    "fatigue": "drowsiness",
    "driver_drowsiness": "drowsiness",
    "close_following": "close_following",
    "following_distance": "close_following",
    "tailgating": "close_following",
    "forward_collision_warning": "forward_collision_warning",
    "forward_collision": "forward_collision_warning",
    "fcw": "forward_collision_warning",
    "collision": "collision",
    "crash": "collision",
    "near_collision": "near_collision",
    "near_crash": "near_collision",
    "stop_sign_violation": "stop_sign_violation",
    "stop_sign": "stop_sign_violation",
    "ran_stop_sign": "stop_sign_violation",
    "unsafe_lane_change": "unsafe_lane_change",
    "lane_change": "unsafe_lane_change",
    "aggregated_lane_swerving": "lane_swerving",
    "lane_swerving": "lane_swerving",
    "lane_swerve": "lane_swerving",
    # ORANGE types
    "hard_brake": "hard_brake",
    "hard_braking": "hard_brake",
    "harsh_brake": "hard_brake",
    "harsh_braking": "hard_brake",
    "seat_belt_violation": "seat_belt_violation",
    "seatbelt": "seat_belt_violation",
    "seatbelt_violation": "seat_belt_violation",
    "no_seatbelt": "seat_belt_violation",
    "seat_belt": "seat_belt_violation",
    "camera_obstruction": "camera_obstruction",
    "obstruction": "camera_obstruction",
    "camera_blocked": "camera_obstruction",
    "smoking": "smoking",
    "vaping": "smoking",
    # YELLOW types
    "hard_accel": "hard_accel",
    "hard_acceleration": "hard_accel",
    "harsh_acceleration": "hard_accel",
    "rapid_acceleration": "hard_accel",
    "hard_corner": "hard_corner",
    "hard_cornering": "hard_corner",
    "hard_turn": "hard_corner",
    "harsh_cornering": "hard_corner",
    "harsh_turn": "hard_corner",
    "speed_violation": "speed_violation",
    "speeding": "speed_violation",
}

RED_TYPES = {
    "distraction", "cell_phone", "drowsiness", "close_following",
    "forward_collision_warning", "collision", "near_collision",
    "stop_sign_violation", "unsafe_lane_change", "lane_swerving",
}

ORANGE_TYPES = {
    "hard_brake", "seat_belt_violation", "camera_obstruction", "smoking",
}

YELLOW_TYPES = {
    "hard_accel", "hard_corner", "speed_violation",
}

# Severity ordering within tiers (lower number = more severe)
EVENT_SEVERITY_ORDER = {
    "collision": 1,
    "near_collision": 2,
    "forward_collision_warning": 3,
    "distraction": 4,
    "cell_phone": 5,
    "drowsiness": 6,
    "stop_sign_violation": 7,
    "unsafe_lane_change": 8,
    "lane_swerving": 8,   # Same severity as unsafe lane change (pre-crash indicator)
    "close_following": 9,
    "hard_brake": 10,
    "seat_belt_violation": 11,
    "camera_obstruction": 12,
    "smoking": 13,
    "hard_accel": 14,
    "hard_corner": 15,
    "speed_violation": 16,
}

# Human-readable display names
EVENT_DISPLAY_NAMES = {
    "distraction": "Distraction",
    "cell_phone": "Cell Phone",
    "drowsiness": "Drowsiness",
    "close_following": "Close Following",
    "forward_collision_warning": "Forward Collision Warning",
    "collision": "Collision",
    "near_collision": "Near Collision",
    "stop_sign_violation": "Stop Sign Violation",
    "unsafe_lane_change": "Unsafe Lane Change",
    "lane_swerving": "Lane Swerving",
    "hard_brake": "Hard Brake",
    "seat_belt_violation": "Seatbelt Violation",
    "camera_obstruction": "Camera Obstruction",
    "smoking": "Smoking",
    "hard_accel": "Hard Acceleration",
    "hard_corner": "Hard Corner",
    "speed_violation": "Speed Violation",
}


def _normalize_event_type(raw_type):
    """Normalize an event type string from the API to a canonical name."""
    if not raw_type:
        return "unknown"
    key = raw_type.lower().strip().replace(" ", "_").replace("-", "_")
    return EVENT_TYPE_NORMALIZE.get(key, key)


def _classify_tier(event_type):
    """Classify a normalized event type into RED/ORANGE/YELLOW."""
    if event_type in RED_TYPES:
        return "RED"
    elif event_type in ORANGE_TYPES:
        return "ORANGE"
    elif event_type in YELLOW_TYPES:
        return "YELLOW"
    else:
        return "ORANGE"  # Unknown types default to ORANGE


def _event_display_name(event_type, raw_type=""):
    """Get human-readable display name for an event type."""
    name = EVENT_DISPLAY_NAMES.get(event_type)
    if name:
        return name
    # For unknown types, make the raw type readable
    display = raw_type or event_type
    return display.replace("_", " ").title()


def _event_sort_key(event):
    """Sort key: tier order (RED=0, ORANGE=1, YELLOW=2) then severity within tier."""
    tier_order = {"RED": 0, "ORANGE": 1, "YELLOW": 2}
    severity = EVENT_SEVERITY_ORDER.get(event["event_type"], 50)
    return (tier_order.get(event["tier"], 1), severity)


# ==============================================================================
# MOTIVE API - VEHICLE LOOKUP (Casing only)
# ==============================================================================

def build_casing_vehicle_lookup():
    """Fetch all vehicles from Motive and build lookup for Casing vehicles.

    Returns:
        vehicle_drivers:      {vehicle_number: driver_name}  (Casing only)
        vehicle_yards:        {vehicle_number: yard_name}     (Casing only)
        all_casing_vehicles:  set of vehicle numbers in Casing
    """
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    vehicle_drivers = {}
    vehicle_yards = {}
    all_casing_vehicles = set()
    page = 1

    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_URL_V1}/vehicles",
                headers=headers,
                params={"per_page": 100, "page_no": page},
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            vehicles = data.get("vehicles", [])
            if not vehicles:
                break

            for wrapper in vehicles:
                v = wrapper.get("vehicle", wrapper)
                num = v.get("number", "")
                if not num:
                    continue

                # Check if this vehicle belongs to a Casing group
                group_ids = v.get("group_ids", [])
                yard = None
                for gid in group_ids:
                    if gid in CASING_GROUP_IDS:
                        yard = CASING_GROUP_IDS[gid]
                        break

                if yard is None:
                    continue  # Not a Casing vehicle

                all_casing_vehicles.add(num)
                if yard:
                    vehicle_yards[num] = yard

                # Driver: prefer current_driver, then permanent_driver
                driver_name = None
                for field in ("current_driver", "permanent_driver"):
                    d = v.get(field)
                    if d and isinstance(d, dict):
                        name = f"{d.get('first_name', '')} {d.get('last_name', '')}".strip()
                        if name:
                            driver_name = name
                            break
                if driver_name:
                    vehicle_drivers[num] = driver_name

            pag = data.get("pagination", {})
            if page * 100 >= pag.get("total", 0):
                break
            page += 1

        except Exception as e:
            print(f"    Warning: vehicle lookup page {page} failed: {e}")
            break

    return vehicle_drivers, vehicle_yards, all_casing_vehicles


# ==============================================================================
# MOTIVE API - CAMERA EVENTS
# ==============================================================================

def _format_duration(seconds):
    """Format duration in seconds to a readable string."""
    if not seconds or not isinstance(seconds, (int, float)):
        return "N/A"
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes = seconds // 60
    secs = seconds % 60
    if secs:
        return f"{minutes}m {secs}s"
    return f"{minutes}m"


def _utc_to_central(timestamp_str):
    """Convert UTC timestamp string to Central Time formatted string."""
    try:
        utc_dt = datetime.fromisoformat(timestamp_str.replace("Z", "+00:00"))
        central_dt = utc_dt.astimezone(CENTRAL_TZ)
        return central_dt.strftime("%m/%d/%Y %I:%M %p CT")
    except Exception:
        return str(timestamp_str)


def get_camera_events_for_date(report_date, vehicle_drivers, vehicle_yards, casing_vehicles):
    """Pull all camera events for Casing vehicles for a specific date.

    report_date: datetime in Central Time - pulls from 00:00:00 to 23:59:59.
    Uses /v2/driver_performance_events endpoint.
    Events wrapped as {"driver_performance_events": [{"driver_performance_event": {data}}]}.

    Tests date filtering (start_date/end_date vs start_time/end_time).
    Falls back to client-side filtering if API ignores date params.
    """
    if hasattr(report_date, 'date'):
        rd = report_date.date()
    else:
        rd = report_date

    start_central = datetime(rd.year, rd.month, rd.day, 0, 0, 0, tzinfo=CENTRAL_TZ)
    end_central = datetime(rd.year, rd.month, rd.day, 23, 59, 59, tzinfo=CENTRAL_TZ)

    start_utc = start_central.astimezone(timezone.utc)
    end_utc = end_central.astimezone(timezone.utc)

    api_start_date = start_utc.strftime("%Y-%m-%d")
    api_end_date = end_utc.strftime("%Y-%m-%d")

    print(f"    Pulling camera events for {rd.strftime('%A, %B %d, %Y')} (Central Time)")
    print(f"    Central window: {start_central.strftime('%m/%d/%Y %I:%M:%S %p')} to {end_central.strftime('%m/%d/%Y %I:%M:%S %p')}")
    print(f"    UTC equivalent: {start_utc.strftime('%Y-%m-%dT%H:%M:%SZ')} to {end_utc.strftime('%Y-%m-%dT%H:%M:%SZ')}")
    print(f"    API date filter: start_date={api_start_date}, end_date={api_end_date}")

    headers = {"X-Api-Key": MOTIVE_API_KEY}
    api_url = f"{MOTIVE_BASE_URL_V2}/driver_performance_events"

    # --- Test what date params work ---
    print(f"\n    Testing date parameter support...")
    print(f"    API URL: {api_url}")

    # Test 1: with start_date/end_date
    try:
        test_resp = requests.get(
            api_url, headers=headers,
            params={"per_page": 1, "page_no": 1, "start_date": api_start_date, "end_date": api_end_date},
            timeout=30,
        )
        test_resp.raise_for_status()
        test_data = test_resp.json()
        total_with_dates = test_data.get("pagination", {}).get("total", len(test_data.get("driver_performance_events", [])))
        print(f"    With start_date/end_date: {total_with_dates} total events")
    except Exception as e:
        total_with_dates = -1
        print(f"    With start_date/end_date: failed ({e})")

    # Test 2: with start_time/end_time (ISO format)
    try:
        test_resp2 = requests.get(
            api_url, headers=headers,
            params={
                "per_page": 1, "page_no": 1,
                "start_time": start_utc.strftime("%Y-%m-%dT%H:%M:%SZ"),
                "end_time": end_utc.strftime("%Y-%m-%dT%H:%M:%SZ"),
            },
            timeout=30,
        )
        test_resp2.raise_for_status()
        test_data2 = test_resp2.json()
        total_with_times = test_data2.get("pagination", {}).get("total", len(test_data2.get("driver_performance_events", [])))
        print(f"    With start_time/end_time: {total_with_times} total events")
    except Exception as e:
        total_with_times = -1
        print(f"    With start_time/end_time: failed ({e})")

    # Test 3: without any date params
    try:
        test_resp3 = requests.get(
            api_url, headers=headers,
            params={"per_page": 1, "page_no": 1},
            timeout=30,
        )
        test_resp3.raise_for_status()
        test_data3 = test_resp3.json()
        total_no_dates = test_data3.get("pagination", {}).get("total", len(test_data3.get("driver_performance_events", [])))
        print(f"    Without date params:       {total_no_dates} total events")
    except Exception as e:
        total_no_dates = -1
        print(f"    Without date params: failed ({e})")

    # Decide which params to use
    use_start_date = True  # default: use start_date/end_date
    if total_with_dates >= 0 and total_no_dates >= 0:
        if total_with_dates == total_no_dates and total_with_dates > 0:
            print(f"    WARNING: start_date/end_date appears IGNORED (same count {total_with_dates} vs {total_no_dates}).")
            if total_with_times >= 0 and total_with_times != total_no_dates:
                print(f"    start_time/end_time works! Using that instead.")
                use_start_date = False
            else:
                print(f"    Will fetch all events and client-side filter to Central Time window.")
        else:
            print(f"    start_date/end_date filtering appears to work ({total_with_dates} vs {total_no_dates}).")

    print()

    # --- Fetch all events ---
    raw_events = []
    page = 1
    page_cursor = None

    while True:
        params = {"per_page": 100}

        if use_start_date:
            params["start_date"] = api_start_date
            params["end_date"] = api_end_date
        else:
            params["start_time"] = start_utc.strftime("%Y-%m-%dT%H:%M:%SZ")
            params["end_time"] = end_utc.strftime("%Y-%m-%dT%H:%M:%SZ")

        # Support both page-number and cursor-based pagination
        if page_cursor:
            params["page_cursor"] = page_cursor
        else:
            params["page_no"] = page

        if page == 1 and not page_cursor:
            print(f"    Fetching with params: {params}")

        try:
            resp = requests.get(api_url, headers=headers, params=params, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            events = data.get("driver_performance_events", [])
            if not events:
                break

            if page == 1 and not page_cursor:
                pag_info = data.get("pagination", {})
                print(f"    First page: {len(events)} events, pagination: {pag_info}")

                # Dump first event's full JSON for field discovery/debugging
                if events:
                    first_evt = events[0]
                    unwrapped = first_evt.get("driver_performance_event", first_evt)
                    print(f"\n    === FIRST RAW EVENT (full JSON for field discovery) ===")
                    # Print all fields except large telemetry arrays
                    debug_copy = {}
                    for k, v in unwrapped.items():
                        if isinstance(v, list) and len(v) > 5:
                            debug_copy[k] = f"[array of {len(v)} items]"
                        else:
                            debug_copy[k] = v
                    print(json.dumps(debug_copy, indent=2, default=str))
                    # Specifically highlight video and speed fields
                    print(f"\n    Speed fields: start_speed={unwrapped.get('start_speed')}, "
                          f"max_speed={unwrapped.get('max_speed')}, "
                          f"end_speed={unwrapped.get('end_speed')}")
                    print(f"    camera_media={unwrapped.get('camera_media')}")
                    print(f"    === END FIRST EVENT ===\n")

            raw_events.extend(events)

            # Handle pagination
            pag = data.get("pagination", {})
            next_cursor = pag.get("next_cursor") or pag.get("next_page_cursor")
            if next_cursor:
                page_cursor = next_cursor
                continue

            total = pag.get("total", 0)
            if total and page * 100 >= total:
                break
            if not total and len(events) < 100:
                break
            page += 1

        except Exception as e:
            print(f"    Error fetching camera events page {page}: {e}")
            break

    print(f"    Raw events fetched: {len(raw_events)}")

    # --- Unwrap, filter, classify ---
    filtered = []
    non_casing_count = 0
    outside_window_count = 0
    raw_event_types = Counter()

    for wrapper in raw_events:
        evt = wrapper.get("driver_performance_event", wrapper)

        # Track all raw event types for console log
        raw_type = (
            evt.get("type", "")
            or evt.get("event_type", "")
            or evt.get("behavior_type", "")
            or ""
        )
        raw_event_types[raw_type] += 1

        # Time filter: only keep events within yesterday Central Time
        evt_time_str = (
            evt.get("start_time", "")
            or evt.get("event_time", "")
            or evt.get("created_at", "")
        )
        try:
            evt_utc = datetime.fromisoformat(evt_time_str.replace("Z", "+00:00"))
            evt_central = evt_utc.astimezone(CENTRAL_TZ)
            if not (start_central <= evt_central <= end_central):
                outside_window_count += 1
                continue
        except Exception:
            pass  # Can't parse time - include to avoid silently dropping

        # Vehicle filter: only keep Casing vehicles
        vehicle_obj = evt.get("vehicle", {})
        if isinstance(vehicle_obj, dict):
            vehicle_number = vehicle_obj.get("number", "")
        else:
            vehicle_number = str(vehicle_obj) if vehicle_obj else ""

        if vehicle_number and casing_vehicles and vehicle_number not in casing_vehicles:
            non_casing_count += 1
            continue

        enriched = enrich_camera_event(evt, vehicle_drivers, vehicle_yards, raw_type)
        filtered.append(enriched)

    # Print all unique event types found (required for first-run analysis)
    if raw_event_types:
        type_summary = ", ".join(f"{t} ({c})" for t, c in raw_event_types.most_common())
        print(f"    Event types found: {type_summary}")
    else:
        print(f"    Event types found: (none)")

    print(f"    After filtering: {len(filtered)} Casing event{'s' if len(filtered) != 1 else ''}")
    print(f"    Dropped: {non_casing_count} non-Casing, {outside_window_count} outside time window")

    # Sort by tier (RED first) then by severity within tier
    return sorted(filtered, key=_event_sort_key)


def enrich_camera_event(event, vehicle_drivers, vehicle_yards, raw_type):
    """Classify and enrich a single camera event."""
    # --- Event Type ---
    event_type = _normalize_event_type(raw_type)
    tier = _classify_tier(event_type)
    display_name = _event_display_name(event_type, raw_type)

    # --- Vehicle ---
    vehicle_obj = event.get("vehicle", {})
    if isinstance(vehicle_obj, dict):
        vehicle_number = vehicle_obj.get("number", "Unknown")
    else:
        vehicle_number = str(vehicle_obj) if vehicle_obj else "Unknown"

    # --- Driver ---
    driver_name = vehicle_drivers.get(vehicle_number)
    if not driver_name:
        drv = event.get("driver")
        if drv and isinstance(drv, dict):
            name = f"{drv.get('first_name', '')} {drv.get('last_name', '')}".strip()
            if name:
                driver_name = name
    if not driver_name:
        # Try parsing from vehicle number (e.g. "5010C John Smith")
        if " " in vehicle_number:
            candidate = vehicle_number.split(" ", 1)[1].strip().lstrip("- ")
            while candidate and candidate.split(" ", 1)[0].replace("-", "").isdigit():
                if " " in candidate:
                    candidate = candidate.split(" ", 1)[1].strip().lstrip("- ")
                else:
                    candidate = ""
            if len(candidate) > 2 and any(c.isalpha() for c in candidate):
                driver_name = candidate
    if not driver_name:
        driver_name = "Unknown"

    # --- Yard ---
    yard = vehicle_yards.get(vehicle_number, "")

    # --- Speed (API returns start_speed/max_speed in km/h, convert to mph) ---
    speed_kmh = event.get("start_speed") or event.get("max_speed") or event.get("end_speed") or 0
    try:
        speed_mph = round(float(speed_kmh) * KMH_TO_MPH, 1) if speed_kmh else 0
    except (ValueError, TypeError):
        speed_mph = 0

    # --- Duration ---
    duration_raw = event.get("duration") or event.get("duration_seconds") or 0
    duration_str = _format_duration(duration_raw)

    # --- Time ---
    timestamp = (
        event.get("start_time")
        or event.get("event_time")
        or event.get("created_at", "")
    )
    formatted_time = _utc_to_central(timestamp)

    # --- Video URL (API uses 'camera_media' field) ---
    video_url = ""
    camera_media = event.get("camera_media")
    if camera_media:
        if isinstance(camera_media, str):
            video_url = camera_media
        elif isinstance(camera_media, dict):
            video_url = (
                camera_media.get("url", "")
                or camera_media.get("video_url", "")
                or camera_media.get("s3_url", "")
                or camera_media.get("media_url", "")
                or camera_media.get("recording_url", "")
            )
            # Check nested 'video' or 'media' sub-objects
            if not video_url:
                for sub_key in ("video", "media", "recording"):
                    sub = camera_media.get(sub_key)
                    if sub and isinstance(sub, dict):
                        video_url = sub.get("url", "") or sub.get("video_url", "")
                        if video_url:
                            break
        elif isinstance(camera_media, list) and camera_media:
            first = camera_media[0]
            if isinstance(first, dict):
                video_url = (
                    first.get("url", "")
                    or first.get("video_url", "")
                    or first.get("s3_url", "")
                )
            elif isinstance(first, str):
                video_url = first

    # --- Location (API uses 'lat'/'lon') ---
    lat = event.get("lat") or event.get("start_lat") or event.get("latitude")
    lon = event.get("lon") or event.get("start_lon") or event.get("longitude")
    location = f"{lat:.4f}, {lon:.4f}" if lat and lon else ""

    return {
        "driver": driver_name,
        "vehicle": vehicle_number,
        "event_type": event_type,
        "raw_type": raw_type,
        "display_name": display_name,
        "tier": tier,
        "speed": speed_mph,
        "duration": duration_str,
        "time": formatted_time,
        "video_url": video_url,
        "location": location,
        "yard": yard,
    }


# ==============================================================================
# DATA ORGANIZATION
# ==============================================================================

def get_repeat_offenders(events):
    """Find drivers with 2+ camera events (exclude Unknown).

    Returns OrderedDict of {name: {count, types, events}} sorted by count desc.
    """
    driver_events = {}
    for e in events:
        if e["driver"] == "Unknown":
            continue
        driver_events.setdefault(e["driver"], []).append(e)

    repeats = {}
    for name, evts in driver_events.items():
        if len(evts) >= 2:
            type_counts = Counter(e["display_name"] for e in evts)
            type_summary = ", ".join(f"{t} x{c}" for t, c in type_counts.most_common())
            repeats[name] = {"count": len(evts), "types": type_summary, "events": evts}

    def sort_key(item):
        name, info = item
        worst_tier = min(
            ({"RED": 0, "ORANGE": 1, "YELLOW": 2}.get(e["tier"], 1) for e in info["events"]),
            default=1,
        )
        return (-info["count"], worst_tier)

    return OrderedDict(sorted(repeats.items(), key=sort_key))


def group_events_by_yard(events):
    """Group events by Casing yard. Returns OrderedDict following YARD_ORDER.
    Only includes yards that have events.
    """
    raw = {}
    for e in events:
        yard = e["yard"] or "Unassigned"
        raw.setdefault(yard, []).append(e)

    grouped = OrderedDict()
    for yard in YARD_ORDER:
        if yard in raw:
            grouped[yard] = sorted(raw[yard], key=_event_sort_key)

    # Any yards not in YARD_ORDER (shouldn't happen, but just in case)
    for yard in sorted(raw.keys()):
        if yard not in grouped and yard != "Unassigned":
            grouped[yard] = sorted(raw[yard], key=_event_sort_key)

    # Unassigned last
    if "Unassigned" in raw:
        grouped["Unassigned"] = sorted(raw["Unassigned"], key=_event_sort_key)

    return grouped


# ==============================================================================
# BUILD WORD DOCUMENT
# ==============================================================================

def _set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, size_pt=8, bold=False, color=None, italic=False):
    """Apply Calibri font and formatting to a run."""
    run.font.name = CALIBRI
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_logo(doc):
    """Add Butchs.jpg logo centered at top. Skip gracefully if missing."""
    logo_path = os.path.join(LOGOS_DIR, "Butchs.jpg")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        try:
            para.add_run().add_picture(logo_path, width=Inches(2.0))
            return True
        except Exception:
            pass
    # Fallback text
    run = para.add_run("BRHAS Casing Division")
    _set_run_font(run, 16, bold=True, color=RGBColor(192, 0, 0))
    return False


def _tier_bg_hex(tier):
    """Return background color hex for tier row shading."""
    if tier == "RED":
        return "FFE0E0"
    elif tier == "ORANGE":
        return "FFF0E0"
    else:
        return "FFFFF0"


def _add_event_table(doc, events):
    """Add a camera events table to the document."""
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    headers = ["Tier", "Driver", "Vehicle", "Event Type", "Speed", "Duration", "Time", "Video"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        _set_run_font(run, 8, bold=True)

    for evt in events:
        cells = table.add_row().cells

        bg = _tier_bg_hex(evt["tier"])
        for cell in cells:
            _set_cell_shading(cell, bg)

        cells[0].text = evt["tier"]
        tier_run = cells[0].paragraphs[0].runs[0]
        tier_color = (
            RGBColor(255, 0, 0) if evt["tier"] == "RED"
            else RGBColor(255, 140, 0) if evt["tier"] == "ORANGE"
            else RGBColor(204, 153, 0)
        )
        _set_run_font(tier_run, 8, bold=True, color=tier_color)

        cells[1].text = evt["driver"]
        cells[2].text = evt["vehicle"]
        cells[3].text = evt["display_name"]
        cells[4].text = f"{evt['speed']} mph" if evt["speed"] else "N/A"
        cells[5].text = evt["duration"]
        cells[6].text = evt["time"]
        cells[7].text = "Video" if evt["video_url"] else ""

        for c in cells[1:]:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_run_font(r, 8)


def _add_horizontal_rule(doc):
    """Add a visible horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="C00000"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _tier_breakdown_text(events_list, type_names):
    """Build a breakdown string like 'Distraction: 3 | Cell Phone: 1'."""
    type_counts = Counter(e["display_name"] for e in events_list)
    parts = []
    for dtype in type_names:
        count = type_counts.get(dtype, 0)
        if count > 0:
            parts.append(f"{dtype}: {count}")
    # Include any unlisted types
    for dtype, count in type_counts.items():
        if dtype not in type_names:
            parts.append(f"{dtype}: {count}")
    return " | ".join(parts)


def create_word_document(events, grouped, yesterday_date):
    """Build the full camera events report Word document in landscape."""
    doc = Document()

    # --- Landscape orientation ---
    for section in doc.sections:
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- Logo ---
    _add_logo(doc)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BRHAS CASING DIVISION")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAILY CAMERA EVENTS REPORT")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(yesterday_date.strftime("%A, %B %d, %Y"))
    _set_run_font(run, 11, italic=True)

    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}")
    _set_run_font(run, 11, color=RGBColor(128, 0, 0))

    doc.add_paragraph()

    # --- Executive Summary ---
    red_events = [e for e in events if e["tier"] == "RED"]
    orange_events = [e for e in events if e["tier"] == "ORANGE"]
    yellow_events = [e for e in events if e["tier"] == "YELLOW"]
    repeats = get_repeat_offenders(events)

    p = doc.add_paragraph()
    run = p.add_run("EXECUTIVE SUMMARY")
    _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    run = p.add_run(f"Total Camera Events: {len(events)}")
    _set_run_font(run, 11, bold=True)

    if red_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  RED \u2014 Immediate Action: {len(red_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(255, 0, 0))

        breakdown = _tier_breakdown_text(red_events, [
            "Distraction", "Cell Phone", "Drowsiness", "Close Following",
            "Forward Collision Warning", "Collision", "Near Collision",
            "Stop Sign Violation", "Unsafe Lane Change", "Lane Swerving",
        ])
        if breakdown:
            p = doc.add_paragraph()
            run = p.add_run(f"    {breakdown}")
            _set_run_font(run, 9)

    if orange_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  ORANGE \u2014 Coaching Required: {len(orange_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(255, 140, 0))

        breakdown = _tier_breakdown_text(orange_events, [
            "Hard Brake", "Seatbelt Violation", "Camera Obstruction", "Smoking",
        ])
        if breakdown:
            p = doc.add_paragraph()
            run = p.add_run(f"    {breakdown}")
            _set_run_font(run, 9)

    if yellow_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  YELLOW \u2014 Monitoring: {len(yellow_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(204, 153, 0))

        breakdown = _tier_breakdown_text(yellow_events, [
            "Hard Acceleration", "Hard Corner", "Speed Violation",
        ])
        if breakdown:
            p = doc.add_paragraph()
            run = p.add_run(f"    {breakdown}")
            _set_run_font(run, 9)

    if not events:
        p = doc.add_paragraph()
        run = p.add_run(f"No camera events for {yesterday_date.strftime('%A, %B %d, %Y')}")
        _set_run_font(run, 11, bold=True, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    # --- Top 5 Most Critical Events ---
    if events:
        p = doc.add_paragraph()
        run = p.add_run("TOP 5 MOST CRITICAL EVENTS")
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        top5 = sorted(events, key=_event_sort_key)[:5]
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        table.autofit = True

        for i, h in enumerate(["Driver", "Vehicle", "Event Type", "Speed", "Time", "Yard", "Video"]):
            table.rows[0].cells[i].text = h
            _set_run_font(table.rows[0].cells[i].paragraphs[0].runs[0], 9, bold=True)

        for evt in top5:
            cells = table.add_row().cells
            bg = _tier_bg_hex(evt["tier"])
            for c in cells:
                _set_cell_shading(c, bg)

            cells[0].text = evt["driver"]
            cells[1].text = evt["vehicle"]
            cells[2].text = evt["display_name"]
            cells[3].text = f"{evt['speed']} mph" if evt["speed"] else "N/A"
            cells[4].text = evt["time"]
            cells[5].text = evt["yard"] if evt["yard"] else "\u2014"
            cells[6].text = "Video" if evt["video_url"] else ""

            for c in cells:
                for para in c.paragraphs:
                    for r in para.runs:
                        _set_run_font(r, 9)

        doc.add_paragraph()

    # --- Repeat Offenders ---
    if repeats:
        p = doc.add_paragraph()
        run = p.add_run("REPEAT OFFENDERS (2+ events in day)")
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        for name, info in repeats.items():
            p = doc.add_paragraph()
            count = info["count"]
            run = p.add_run(f"  {name}: {count} event{'s' if count != 1 else ''}")
            _set_run_font(run, 10, bold=True, color=RGBColor(192, 0, 0))
            run2 = p.add_run(f" (types: {info['types']})")
            _set_run_font(run2, 10)

        doc.add_paragraph()

    # --- Yard Sections ---
    for yard, yard_events in grouped.items():
        _add_horizontal_rule(doc)

        info = YARD_INFO.get(yard, {})
        safety_reps = info.get("safety_reps", "")
        manager = info.get("manager", "")

        p = doc.add_paragraph()
        run = p.add_run(f"{yard.upper()} YARD")
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        if safety_reps:
            p = doc.add_paragraph()
            run = p.add_run(f"Safety Rep(s): {safety_reps}")
            _set_run_font(run, 10, italic=True)

        if manager:
            p = doc.add_paragraph()
            run = p.add_run(f"Manager: {manager}")
            _set_run_font(run, 10, italic=True)

        yard_red = len([e for e in yard_events if e["tier"] == "RED"])
        yard_orange = len([e for e in yard_events if e["tier"] == "ORANGE"])
        yard_yellow = len([e for e in yard_events if e["tier"] == "YELLOW"])
        total = len(yard_events)

        p = doc.add_paragraph()
        run = p.add_run(f"{total} event{'s' if total != 1 else ''}")
        _set_run_font(run, 10, bold=True)
        run2 = p.add_run(f" (RED: {yard_red} | ORANGE: {yard_orange} | YELLOW: {yard_yellow})")
        _set_run_font(run2, 10)

        doc.add_paragraph()
        _add_event_table(doc, yard_events)
        doc.add_paragraph()

    # --- Footer ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END OF REPORT")
    _set_run_font(run, 10, italic=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Butch's Rat Hole & Anchor Service Inc. | Casing Division | HSE Department")
    _set_run_font(run, 9, color=RGBColor(128, 0, 0))

    return doc


# ==============================================================================
# BUILD HTML EMAIL
# ==============================================================================

C_RED = "#C00000"
C_DARK = "#800000"
C_AMBER = "#FF8C00"
C_YELLOW_DARK = "#CC9900"
C_GREEN = "#008000"


def _h(text):
    """HTML-escape text safely."""
    return html_escape(str(text)) if text else ""


def _tier_colors(tier):
    """Return (text_color, bg_color) for a tier."""
    if tier == "RED":
        return "#FF0000", "#FFE0E0"
    elif tier == "ORANGE":
        return C_AMBER, "#FFF0E0"
    else:
        return C_YELLOW_DARK, "#FFFFF0"


def build_html_report(events, grouped, yesterday_date):
    """Build HTML email body organized by Casing yard."""
    red_events = [e for e in events if e["tier"] == "RED"]
    orange_events = [e for e in events if e["tier"] == "ORANGE"]
    yellow_events = [e for e in events if e["tier"] == "YELLOW"]
    repeats = get_repeat_offenders(events)

    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    parts = []

    # --- Wrapper + Header ---
    parts.append(f"""<html><head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#f4f4f4;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f4f4;">
<tr><td align="center">
<table width="700" cellpadding="0" cellspacing="0" style="background:#ffffff;border:1px solid #ddd;margin:20px auto;font-family:Calibri,Arial,Helvetica,sans-serif;font-size:14px;color:#333;">

<tr><td style="background:{C_RED};padding:30px 40px;text-align:center;">
  <div style="font-size:16px;font-weight:bold;color:#ffffff;letter-spacing:1px;">BRHAS CASING DIVISION</div>
  <div style="font-size:28px;font-weight:bold;color:#ffffff;margin:10px 0;">DAILY CAMERA EVENTS REPORT</div>
  <div style="font-size:13px;font-style:italic;color:#ffcccc;">Motive AI Dashcam Safety Events</div>
  <div style="font-size:12px;color:#ffffff;margin-top:8px;">Report Date: {yesterday_date.strftime('%A, %B %d, %Y')}</div>
  <div style="font-size:10px;color:#ffcccc;margin-top:4px;">Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}</div>
</td></tr>""")

    # --- Executive Summary ---
    summary = f"<b>Total Camera Events: {len(events)}</b><br><br>"

    if red_events:
        red_type_counts = Counter(e["display_name"] for e in red_events)
        red_breakdown = " | ".join(f"{t}: {c}" for t, c in red_type_counts.most_common())
        summary += f'<div style="color:#FF0000;font-weight:bold;margin:4px 0 4px 20px;">RED \u2014 Immediate Action: {len(red_events)}</div>'
        summary += f'<div style="font-size:12px;margin:2px 0 8px 40px;">{_h(red_breakdown)}</div>'

    if orange_events:
        orange_type_counts = Counter(e["display_name"] for e in orange_events)
        orange_breakdown = " | ".join(f"{t}: {c}" for t, c in orange_type_counts.most_common())
        summary += f'<div style="color:{C_AMBER};font-weight:bold;margin:4px 0 4px 20px;">ORANGE \u2014 Coaching Required: {len(orange_events)}</div>'
        summary += f'<div style="font-size:12px;margin:2px 0 8px 40px;">{_h(orange_breakdown)}</div>'

    if yellow_events:
        yellow_type_counts = Counter(e["display_name"] for e in yellow_events)
        yellow_breakdown = " | ".join(f"{t}: {c}" for t, c in yellow_type_counts.most_common())
        summary += f'<div style="color:{C_YELLOW_DARK};font-weight:bold;margin:4px 0 4px 20px;">YELLOW \u2014 Monitoring: {len(yellow_events)}</div>'
        summary += f'<div style="font-size:12px;margin:2px 0 8px 40px;">{_h(yellow_breakdown)}</div>'

    if not events:
        summary += f'<b style="color:{C_GREEN};">No camera events for {yesterday_date.strftime("%A, %B %d, %Y")}!</b>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">EXECUTIVE SUMMARY</h2>
  {summary}
</td></tr>""")

    # --- Top 5 Most Critical Events ---
    if events:
        top5 = sorted(events, key=_event_sort_key)[:5]
        top5_html = ""
        for e in top5:
            tc, bg = _tier_colors(e["tier"])
            top5_html += f'<div style="background:{bg};border-left:4px solid {tc};padding:10px 15px;margin:8px 0;">'
            top5_html += f'<b style="color:{tc};">{_h(e["display_name"])}</b>'
            if e["speed"]:
                top5_html += f' at {e["speed"]} mph'
            top5_html += '<br>'
            top5_html += f'<b>Driver:</b> {_h(e["driver"])} | <b>Vehicle:</b> {_h(e["vehicle"])}'
            if e["yard"]:
                top5_html += f' | <b>Yard:</b> {_h(e["yard"])}'
            top5_html += f' | <b>Time:</b> {_h(e["time"])}'
            if e["video_url"]:
                top5_html += f' | <a href="{_h(e["video_url"])}">Video</a>'
            top5_html += "</div>"

        parts.append(f"""
<tr><td style="padding:20px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">TOP 5 MOST CRITICAL EVENTS</h2>
  {top5_html}
</td></tr>""")

    # --- Repeat Offenders ---
    if repeats:
        repeat_html = ""
        for name, info in repeats.items():
            count = info["count"]
            repeat_html += f'<div style="background:#fff5f5;border-left:4px solid {C_RED};padding:10px 15px;margin:8px 0;">'
            repeat_html += f'<b style="color:{C_RED};">{_h(name)}: {count} event{"s" if count != 1 else ""}</b>'
            repeat_html += f' (types: {_h(info["types"])})'
            repeat_html += "</div>"

        parts.append(f"""
<tr><td style="padding:20px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">REPEAT OFFENDERS (2+ events in day)</h2>
  {repeat_html}
</td></tr>""")

    # --- Yard Sections ---
    for yard, yard_events in grouped.items():
        info = YARD_INFO.get(yard, {})
        safety_reps = info.get("safety_reps", "")
        manager = info.get("manager", "")

        yard_red = len([e for e in yard_events if e["tier"] == "RED"])
        yard_orange = len([e for e in yard_events if e["tier"] == "ORANGE"])
        yard_yellow = len([e for e in yard_events if e["tier"] == "YELLOW"])
        total = len(yard_events)

        parts.append(f"""
<tr><td style="padding:0 40px;"><hr style="border:none;border-top:3px solid {C_RED};margin:20px 0 0 0;"></td></tr>
<tr><td style="padding:15px 40px;">
  <h2 style="color:{C_RED};margin:0;font-size:20px;">{_h(yard.upper())} YARD</h2>
  {"<div style='font-size:12px;font-style:italic;color:#666;margin:4px 0;'>Safety Rep(s): " + _h(safety_reps) + "</div>" if safety_reps else ""}
  {"<div style='font-size:12px;font-style:italic;color:#666;margin:2px 0;'>Manager: " + _h(manager) + "</div>" if manager else ""}
  <div style="background:#f8f0f0;border-left:4px solid {C_RED};padding:10px 15px;margin:10px 0;font-size:13px;">
    <b>{total}</b> event{"s" if total != 1 else ""}
    (RED: {yard_red} | ORANGE: {yard_orange} | YELLOW: {yard_yellow})
  </div>
</td></tr>""")

        table_rows = ""
        for e in yard_events:
            tc, bg = _tier_colors(e["tier"])
            video_cell = f'<a href="{_h(e["video_url"])}" style="font-size:11px;">Video</a>' if e["video_url"] else ""
            table_rows += f"""<tr style="background:{bg};">
  <td style="padding:5px 6px;border:1px solid #ddd;"><b style="color:{tc};">{e["tier"]}</b></td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{_h(e["driver"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{_h(e["vehicle"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{_h(e["display_name"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;">{e["speed"] if e["speed"] else "N/A"}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;">{_h(e["duration"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;font-size:11px;">{_h(e["time"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{video_cell}</td>
</tr>"""

        parts.append(f"""
<tr><td style="padding:5px 40px 15px 40px;">
  <table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:12px;">
    <tr style="background:{C_RED};">
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Tier</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Driver</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Vehicle</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Event Type</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Speed</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Duration</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Time</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Video</th>
    </tr>
    {table_rows}
  </table>
</td></tr>""")

    # --- Footer ---
    parts.append(f"""
<tr><td style="background:{C_DARK};padding:20px 40px;text-align:center;">
  <div style="color:#ffffff;font-size:11px;font-style:italic;">END OF REPORT</div>
  <div style="color:#ffcccc;font-size:10px;margin-top:4px;">Butch's Rat Hole &amp; Anchor Service Inc. | Casing Division | HSE Department</div>
</td></tr>

</table>
</td></tr></table>
</body></html>""")

    return "\n".join(parts)


# ==============================================================================
# SEND EMAIL
# ==============================================================================

def send_email_report(html_body, docx_path, yesterday_date):
    """Send report via Gmail SMTP. Fails gracefully."""
    gmail_address = os.environ.get("GMAIL_ADDRESS", "")
    gmail_app_password = os.environ.get("GMAIL_APP_PASSWORD", "")
    recipient = os.environ.get("REPORT_RECIPIENT", "")

    if not gmail_address or not gmail_app_password or not recipient:
        print("  Email skipped \u2014 GMAIL_ADDRESS, GMAIL_APP_PASSWORD, or REPORT_RECIPIENT not set.")
        return

    subject = f"Casing Camera Events Report - {yesterday_date.strftime('%B %d, %Y')}"

    try:
        msg = MIMEMultipart("mixed")
        msg["From"] = gmail_address
        msg["To"] = recipient
        msg["Subject"] = subject

        msg.attach(MIMEText(html_body, "html"))

        if os.path.exists(docx_path):
            with open(docx_path, "rb") as f:
                part = MIMEBase(
                    "application",
                    "vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f'attachment; filename="{os.path.basename(docx_path)}"',
            )
            msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_address, gmail_app_password)
            server.sendmail(gmail_address, recipient, msg.as_string())

        print(f"  Email sent to {recipient}")
    except Exception as e:
        print(f"  Email failed: {e}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    today = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    yesterday = today - timedelta(days=1)

    print("\n" + "=" * 80)
    print("DAILY CASING CAMERA EVENTS REPORT - AUTOMATED")
    print(f"Report for: {yesterday.strftime('%A, %B %d, %Y')}")
    print("=" * 80)
    print("\n  Tier Classification:")
    print("    RED:    Immediate Action (distraction, cell phone, drowsiness, close following,")
    print("            collision, near collision, forward collision warning, stop sign,")
    print("            unsafe lane change, lane swerving)")
    print("    ORANGE: Coaching Required (hard brake, seatbelt, camera obstruction, smoking)")
    print("    YELLOW: Monitoring (hard accel, hard corner, speed violation)")
    print("    Unknown event types default to ORANGE")
    print("    Repeat: 2+ events flagged\n")

    print("[1] Building Casing vehicle/driver lookup from Motive...")
    vehicle_drivers, vehicle_yards, casing_vehicles = build_casing_vehicle_lookup()
    print(f"    {len(casing_vehicles)} Casing vehicles found")
    print(f"    {len(vehicle_drivers)} with driver names")
    print(f"    {len(vehicle_yards)} with yard assignments")

    print("\n[2] Fetching camera events from Motive...")
    events = get_camera_events_for_date(yesterday, vehicle_drivers, vehicle_yards, casing_vehicles)
    print(f"    Found {len(events)} Casing camera event{'s' if len(events) != 1 else ''}")

    if events:
        red = len([e for e in events if e["tier"] == "RED"])
        orange = len([e for e in events if e["tier"] == "ORANGE"])
        yellow = len([e for e in events if e["tier"] == "YELLOW"])
        known = len([e for e in events if e["driver"] != "Unknown"])
        repeats = get_repeat_offenders(events)
        print(f"    RED: {red} | ORANGE: {orange} | YELLOW: {yellow}")
        print(f"    Drivers identified: {known}/{len(events)}")
        if repeats:
            print(f"    Repeat offenders ({len(repeats)}): {', '.join(f'{n} ({info['count']}x)' for n, info in repeats.items())}")

    print("\n[3] Grouping events by Casing yard...")
    grouped = group_events_by_yard(events)
    for yard, yard_events in grouped.items():
        print(f"    {yard}: {len(yard_events)} event{'s' if len(yard_events) != 1 else ''}")

    print("\n[4] Creating Word document (landscape)...")
    doc = create_word_document(events, grouped, yesterday)

    date_str = yesterday.strftime("%Y-%m-%d")
    output_file = f"DailyCasingCameraReport_{date_str}.docx"
    doc.save(output_file)
    print(f"    Saved: {output_file}")

    print("\n[5] Building HTML email...")
    html_body = build_html_report(events, grouped, yesterday)

    print("[6] Sending email...")
    send_email_report(html_body, output_file, yesterday)

    print("\n" + "=" * 80)
    print("COMPLETE")
    print("=" * 80 + "\n")


if __name__ == "__main__":
    main()
