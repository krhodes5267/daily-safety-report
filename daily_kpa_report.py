"""
KPA DAILY SAFETY REPORT - AUTOMATED (GitHub Actions)
=====================================================
Runs daily at 5:00 AM Central via GitHub Actions.

CRITICAL: Observer Name Handling
- ALWAYS uses 'Name' field (the actual person observed)
- For James Barnett paper forms: Name = Ruben Lopez, Alfonso Orozco, etc.
- Never shows James Barnett as the person (he's only the data entry person)

Structure: Critical items first, only shows sections with data
- Safety Streak Metrics
- Executive Summary
- Action Items
- Near Misses (detailed)
- Open Items Tracking (Conditions & Procedures only - NOT Near Misses)
- Data Quality Alerts
- Hotspot Analysis
- Timing Analysis
- Assessment & Audit Analysis (NEW - assessor details, compliance by yard,
  critical findings, corrective actions, trends, leadership recommendations)
- Conditions (Top 10)
- Recognition Stars
- Other Forms
"""

import requests
import csv
from datetime import datetime, timedelta
import os
import sys
import smtplib
from io import StringIO
from html import escape as html_escape
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter

# ==============================================================================
# SETUP - API keys from environment variables
# ==============================================================================

API_TOKEN = os.environ.get("KPA_API_TOKEN")
if not API_TOKEN:
    print("ERROR: KPA_API_TOKEN environment variable is not set.")
    sys.exit(1)

MOTIVE_API_KEY = os.environ.get("MOTIVE_API_KEY", "")

API_BASE = "https://api.kpaehs.com/v1"

FORMS = {
    151085: "Observation Cards",
    151622: "Incident Report",
    180243: "Root Cause Analysis",
    381707: "CSG - Safety Casing Field Assessment",
    152018: "Vehicle Inspection Checklist",
    385365: "TD - Rig Inspection",
    484193: "TD - Observation Card",
    226217: "WS - Poly Pipe Field Assessment",
    386087: "WS - Pit Lining Field Assessment",
    172295: "Construction - Site Safety Review",
    153181: "RH - Rathole Field Assessment",
    152034: "HSE - Workplace Inspection Checklist",
    229645: "CSG - Pre/Post Trip Inspection"
}

OTHER_FORMS = [
    (381707, "CSG - Safety Casing Field Assessment"),
    (152018, "Vehicle Inspection Checklist"),
    (385365, "TD - Rig Inspection"),
    (226217, "WS - Poly Pipe Field Assessment"),
    (386087, "WS - Pit Lining Field Assessment"),
    (172295, "Construction - Site Safety Review"),
    (153181, "RH - Rathole Field Assessment"),
    (152034, "HSE - Workplace Inspection Checklist"),
    (229645, "CSG - Pre/Post Trip Inspection")
]

COLORS = {
    'primary': RGBColor(192, 0, 0),
    'secondary': RGBColor(128, 0, 0),
    'accent': RGBColor(0, 0, 0),
    'critical': RGBColor(192, 0, 0),
    'warning': RGBColor(192, 128, 0),
    'safe': RGBColor(0, 128, 0),
}

# Logos are optional - they exist on local machines but not on CI runners
LOGOS_PATH = os.path.expanduser("~/Downloads")
LOGOS = ['Butchs.jpg', 'ButchTrucking.jpg', 'Permian.jpg', 'Hutchs.png', 'Transcend.jpg', 'Valor.jpg']

# Assessment/Audit forms with metadata for deep analysis
ASSESSMENT_FORMS = {
    381707: {"name": "CSG - Safety Casing Field Assessment", "type": "Field Assessment", "division": "Casing", "company": "BRHAS"},
    152018: {"name": "Vehicle Inspection Checklist", "type": "Inspection", "division": "All", "company": "All"},
    385365: {"name": "TD - Rig Inspection", "type": "Rig Inspection", "division": "Transcend", "company": "Transcend"},
    206674: {"name": "TD - Transcend Field Assessment", "type": "Field Assessment", "division": "Transcend", "company": "Transcend"},
    226217: {"name": "WS - Poly Pipe Field Assessment", "type": "Field Assessment", "division": "Poly Pipe", "company": "BRHAS"},
    386087: {"name": "WS - Pit Lining Field Assessment", "type": "Field Assessment", "division": "Pit Lining", "company": "BRHAS"},
    172295: {"name": "Construction - Site Safety Review", "type": "Site Review", "division": "Construction", "company": "BRHAS"},
    153181: {"name": "RH - Rathole Field Assessment", "type": "Field Assessment", "division": "Rathole", "company": "BRHAS"},
    229645: {"name": "CSG - Pre/Post Trip Inspection", "type": "Inspection", "division": "Casing", "company": "BRHAS"},
    152034: {"name": "HSE - Workplace Inspection Checklist", "type": "Inspection", "division": "HSE", "company": "BRHAS"},
}

# Forms that should get deep analytics (compliance, findings, quality scores)
# For shared forms (152018), analytics are built per-company from the entry's company field.
ANALYTICS_FORMS = {
    381707: {"company": "BRHAS", "division": "Casing"},
    152018: {"company": "All", "division": "All"},     # split by entry company field
    153181: {"company": "BRHAS", "division": "Rathole"},
    385365: {"company": "Transcend", "division": "Transcend"},
    206674: {"company": "Transcend", "division": "Transcend"},
}

KPA_RESPONSE_URL = "https://brhas-ees.kpaehs.com/forms/responses/view"

# Service line hash fields across different KPA forms
SERVICE_LINE_HASHES = [
    "64c7upqkyt79zhh1", "sha7vur5q2l6d6gq", "68jbiriixiurowou",
    "hxy6pwclvjke1sln", "b5dqbf3qgxga92fi", "7cl5rryw636f6wbx",
    "service_line", "Service Line", "division", "Division",
]

def _get_service_line(row):
    """Extract service line value from a KPA row, trying multiple hash fields."""
    for key in SERVICE_LINE_HASHES:
        val = row.get(key, "").strip()
        if val and val.lower() not in ("", "service line", "division", "n/a", "na"):
            return val
    return ""

# Keywords for smart field detection in assessment CSV headers
COMPLIANCE_KEYWORDS = ['compliance', 'rating', 'satisfactory', 'pass', 'fail', 'acceptable',
                       'result', 'score', 'compliant', 'conformance']
FINDING_KEYWORDS = ['issue', 'finding', 'non-conformance', 'corrective', 'deficiency',
                    'violation', 'hazard', 'concern', 'recommendation', 'comment',
                    'note', 'detail']
YARD_KEYWORDS = ['yard', 'location', 'site', 'field office', 'facility', 'area']

# ==============================================================================
# API CALL
# ==============================================================================

def call_kpa(endpoint, params):
    """Make request to KPA API"""
    url = f"{API_BASE}/{endpoint}"
    payload = {"token": API_TOKEN}
    payload.update(params)

    try:
        response = requests.post(url, json=payload, timeout=30)
        return response.text
    except Exception as e:
        print(f"ERROR: {e}")
        return None


# ==============================================================================
# PULL FORM DATA - YESTERDAY ONLY
# ==============================================================================

def pull_form_data(form_id, form_name):
    """Pull incidents from YESTERDAY ONLY"""
    today = datetime.now()
    yesterday_start = today.replace(hour=0, minute=0, second=0, microsecond=0) - timedelta(days=1)
    today_start = today.replace(hour=0, minute=0, second=0, microsecond=0)

    yesterday_start_ms = int(yesterday_start.timestamp() * 1000)
    today_start_ms = int(today_start.timestamp() * 1000)

    params = {
        "form_id": form_id,
        "format": "csv",
        "updated_after": yesterday_start_ms
    }

    csv_text = call_kpa("responses.flat", params)

    if not csv_text or csv_text.strip() == "":
        return None

    try:
        csv_file = StringIO(csv_text)
        reader = csv.DictReader(csv_file)
        rows = list(reader)
        if len(rows) == 0:
            return None

        filtered_rows = []
        for row in rows:
            if row.get('report number') == 'Report Number':
                continue

            date_str = row.get('date', '')
            try:
                row_date = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                row_date_ms = int(row_date.timestamp() * 1000)

                if yesterday_start_ms <= row_date_ms < today_start_ms:
                    filtered_rows.append(row)
            except (ValueError, TypeError, KeyError):
                continue

        if len(filtered_rows) == 0:
            return None

        return {
            'headers': reader.fieldnames if reader.fieldnames else [],
            'rows': filtered_rows,
            'count': len(filtered_rows)
        }
    except Exception as e:
        print(f"Error parsing {form_name}: {e}")
        return None


# ==============================================================================
# HELPERS - GET ACTUAL OBSERVER NAME (NOT DATA ENTRY PERSON)
# ==============================================================================

def get_actual_observer_name(obs):
    """
    Get the ACTUAL person's name from the observation form

    CRITICAL: This field represents who actually DID the observation/work
    NOT who entered it into the system

    For paper forms submitted by James Barnett:
    - 'observer' field = James Barnett (system entry person - IGNORE)
    - 'Name' or 'name' field = Ruben Lopez, Alfonso Orozco, etc. (ACTUAL person - USE THIS)
    """

    # PRIMARY: Check 'Name' field (capital N)
    name = obs.get('Name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name

    # Try lowercase 'name' field as well
    name = obs.get('name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name

    # FALLBACK: observer field (only if Name is truly missing)
    observer = obs.get('observer', '').strip()
    if observer and observer.lower() not in ['unknown', 'none', '']:
        return observer

    return 'Unknown'


def get_observation_type(obs):
    """Get observation type"""
    obs_type = obs.get('bff8m4x6xbc033kg', 'Other')
    return obs_type.strip() if obs_type else 'Other'


def get_shift(date_str):
    """Determine shift from time"""
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
        hour = dt.hour
        if 0 <= hour < 8:
            return "Overnight (0-8 AM)"
        elif 8 <= hour < 16:
            return "Day Shift (8 AM-4 PM)"
        elif 16 <= hour < 24:
            return "Night Shift (4 PM-Midnight)"
    except (ValueError, TypeError):
        return "Unknown"


def analyze_observations(obs_data):
    """Analyze observations and group by type"""
    if not obs_data:
        return None

    observations_by_type = {}
    miscategorized = []

    for obs in obs_data['rows']:
        obs_type = get_observation_type(obs)
        if obs_type not in observations_by_type:
            observations_by_type[obs_type] = []
        observations_by_type[obs_type].append(obs)

        # Check for miscategorization
        text = obs.get('uncbcge9x8vow9pn', '').lower()
        if obs_type == 'At-Risk Condition':
            if ('good' in text or 'no issue' in text or 'no problem' in text or 'excellent' in text or 'perfect' in text) and len(text) < 100:
                miscategorized.append({
                    'report_num': obs.get('report number'),
                    'type': obs_type,
                    'actual_type': 'Recognition',
                    'description': text[:80],
                    'observer': get_actual_observer_name(obs)
                })

    total = sum(len(v) for v in observations_by_type.values())

    return {
        'total': total,
        'by_type': observations_by_type,
        'type_counts': {k: len(v) for k, v in observations_by_type.items()},
        'miscategorized': miscategorized
    }


# ==============================================================================
# ASSESSMENT & AUDIT ANALYSIS FUNCTIONS
# ==============================================================================

def detect_field_columns(headers):
    """Detect key columns from assessment CSV headers using keyword matching"""
    fields = {
        'compliance': [],
        'findings': [],
        'yard': [],
        'severity': [],
        'assessor': [],
        'corrective_action': [],
    }

    if not headers:
        return fields

    for header in headers:
        h_lower = header.lower()

        # Skip standard metadata fields
        if h_lower in ['report number', 'date', 'link', 'observer', 'name']:
            continue

        if any(kw in h_lower for kw in COMPLIANCE_KEYWORDS):
            fields['compliance'].append(header)
        if any(kw in h_lower for kw in FINDING_KEYWORDS):
            fields['findings'].append(header)
        if any(kw in h_lower for kw in YARD_KEYWORDS):
            fields['yard'].append(header)
        if any(kw in h_lower for kw in ['severity', 'priority', 'risk level', 'critical']):
            fields['severity'].append(header)
        if any(kw in h_lower for kw in ['assessor', 'inspector', 'auditor', 'reviewer', 'conducted by']):
            fields['assessor'].append(header)
        if any(kw in h_lower for kw in ['corrective', 'action required', 'action taken', 'follow up', 'follow-up']):
            fields['corrective_action'].append(header)

    return fields


def classify_compliance_value(value):
    """Classify a field value as compliant, non-compliant, or unknown"""
    if not value:
        return 'unknown'
    v = value.strip().lower()

    non_compliant_terms = ['fail', 'unsatisfactory', 'non-compliant', 'unacceptable',
                           'poor', 'deficient', 'inadequate', 'needs improvement', 'not met']
    compliant_terms = ['pass', 'yes', 'satisfactory', 'compliant', 'acceptable', 'good',
                       'meets', 'adequate', 'ok', 'n/a', 'not applicable']

    for term in non_compliant_terms:
        if term in v:
            return 'non_compliant'
    for term in compliant_terms:
        if term in v:
            return 'compliant'
    return 'unknown'


def classify_severity(text):
    """Classify the severity level of a finding based on its text"""
    if not text:
        return 'low'
    t = text.lower()

    critical_terms = ['critical', 'immediate', 'danger', 'life-threatening', 'fatal',
                      'imminent', 'emergency', 'severe', 'death']
    high_terms = ['high', 'serious', 'major', 'significant', 'injury potential',
                  'non-compliant', 'violation', 'failed']
    medium_terms = ['medium', 'moderate', 'minor damage', 'needs repair', 'worn', 'missing']

    for term in critical_terms:
        if term in t:
            return 'critical'
    for term in high_terms:
        if term in t:
            return 'high'
    for term in medium_terms:
        if term in t:
            return 'medium'

    return 'low'


def get_assessor_name(row):
    """Get assessor/observer name from assessment form row"""
    for field_name, value in row.items():
        if any(kw in field_name.lower() for kw in ['assessor', 'inspector', 'auditor', 'conducted by', 'reviewer']):
            if value and value.strip() and value.strip().lower() not in ['none', 'unknown', '']:
                return value.strip()

    return get_actual_observer_name(row)


def get_yard_from_row(row, detected_fields):
    """Extract yard/location from a row using detected fields"""
    for field in detected_fields.get('yard', []):
        val = row.get(field, '').strip()
        if val and val.lower() not in ['n/a', 'none', 'unknown', '']:
            return val

    for key in ['7vj2l992y7fwqhwz', 'lg5pnj4chjadnv46']:
        val = row.get(key, '').strip()
        if val and val.lower() not in ['n/a', 'none', 'unknown', '']:
            return val

    for field_name, value in row.items():
        if ('yard' in field_name.lower() or 'location' in field_name.lower()):
            if value and value.strip() and value.strip().lower() not in ['n/a', 'none', 'unknown', '']:
                return value.strip()

    return 'Unknown'


def get_kpa_link(report_num):
    """Build clickable KPA link from report number"""
    if report_num and report_num not in ['Report Number', '']:
        return f"{KPA_RESPONSE_URL}/{report_num}"
    return ''


def analyze_assessments(all_data):
    """Analyze all assessment/audit form data for the daily report"""
    analysis = {
        'activity_summary': [],
        'findings_by_severity': {'critical': [], 'high': [], 'medium': [], 'low': []},
        'compliance_by_yard': {},
        'assessor_stats': {},
        'corrective_actions': [],
        'all_findings': [],
        'trends': [],
        'recommendations': {'immediate': [], 'this_week': [], 'monthly': []},
        'total_assessments': 0,
        'total_findings': 0,
        'has_data': False,
    }

    for form_id, form_info in ASSESSMENT_FORMS.items():
        data = all_data.get(f"form_{form_id}")
        if not data or data['count'] == 0:
            continue

        analysis['has_data'] = True
        analysis['total_assessments'] += data['count']

        detected = detect_field_columns(data['headers'])

        form_assessors = set()
        form_findings = []
        form_compliant = 0
        form_non_compliant = 0

        for row in data['rows']:
            assessor = get_assessor_name(row)
            form_assessors.add(assessor)

            if assessor not in analysis['assessor_stats']:
                analysis['assessor_stats'][assessor] = {
                    'total': 0, 'forms': set(), 'divisions': set(), 'findings_found': 0
                }
            analysis['assessor_stats'][assessor]['total'] += 1
            analysis['assessor_stats'][assessor]['forms'].add(form_info['name'])
            analysis['assessor_stats'][assessor]['divisions'].add(form_info['division'])

            yard = get_yard_from_row(row, detected)

            if yard not in analysis['compliance_by_yard']:
                analysis['compliance_by_yard'][yard] = {
                    'total': 0, 'compliant': 0, 'non_compliant': 0,
                    'findings': [], 'forms_used': set()
                }
            analysis['compliance_by_yard'][yard]['total'] += 1
            analysis['compliance_by_yard'][yard]['forms_used'].add(form_info['name'])

            # Check compliance fields
            row_compliant = True
            for comp_field in detected['compliance']:
                val = row.get(comp_field, '')
                result = classify_compliance_value(val)
                if result == 'non_compliant':
                    row_compliant = False
                    break

            if row_compliant:
                form_compliant += 1
                analysis['compliance_by_yard'][yard]['compliant'] += 1
            else:
                form_non_compliant += 1
                analysis['compliance_by_yard'][yard]['non_compliant'] += 1

            # Extract findings
            for finding_field in detected['findings']:
                finding_text = row.get(finding_field, '').strip()
                if finding_text and len(finding_text) > 3 and finding_text.lower() not in ['n/a', 'none', 'no', 'na']:
                    severity = classify_severity(finding_text)

                    for sev_field in detected['severity']:
                        sev_val = row.get(sev_field, '').strip()
                        if sev_val:
                            severity = classify_severity(sev_val)
                            break

                    finding = {
                        'form_name': form_info['name'],
                        'division': form_info['division'],
                        'assessor': assessor,
                        'yard': yard,
                        'description': finding_text[:200],
                        'severity': severity,
                        'report_num': row.get('report number', ''),
                        'date': row.get('date', ''),
                        'link': get_kpa_link(row.get('report number', '')),
                    }

                    form_findings.append(finding)
                    analysis['findings_by_severity'][severity].append(finding)
                    analysis['compliance_by_yard'][yard]['findings'].append(finding)
                    analysis['all_findings'].append(finding)
                    analysis['total_findings'] += 1
                    analysis['assessor_stats'][assessor]['findings_found'] += 1

            # Extract corrective actions
            for ca_field in detected['corrective_action']:
                ca_text = row.get(ca_field, '').strip()
                if ca_text and len(ca_text) > 3 and ca_text.lower() not in ['n/a', 'none', 'no', 'na']:
                    analysis['corrective_actions'].append({
                        'form_name': form_info['name'],
                        'description': ca_text[:200],
                        'assessor': assessor,
                        'yard': yard,
                        'date': row.get('date', ''),
                        'report_num': row.get('report number', ''),
                        'link': get_kpa_link(row.get('report number', '')),
                        'status': 'Open',
                    })

        # Activity summary for this form
        compliance_rate = (form_compliant / data['count'] * 100) if data['count'] > 0 else 0

        assessment_analysis_item = {
            'form_name': form_info['name'],
            'form_type': form_info['type'],
            'division': form_info['division'],
            'count': data['count'],
            'assessors': sorted(form_assessors - {'Unknown'}),
            'findings_count': len(form_findings),
            'compliance_rate': compliance_rate,
            'compliant': form_compliant,
            'non_compliant': form_non_compliant,
        }
        # Only add assessors list if "Unknown" was the only one
        if not assessment_analysis_item['assessors'] and 'Unknown' in form_assessors:
            assessment_analysis_item['assessors'] = ['Unknown']

        analysis['activity_summary'].append(assessment_analysis_item)

    if analysis['has_data']:
        _generate_assessment_trends(analysis)
        _generate_assessment_recommendations(analysis)

    return analysis


def _generate_assessment_trends(analysis):
    """Generate trend observations from assessment data"""
    trends = []

    # Yards with multiple findings
    problem_yards = {yard: info for yard, info in analysis['compliance_by_yard'].items()
                     if len(info['findings']) >= 2}
    if problem_yards:
        for yard, info in sorted(problem_yards.items(), key=lambda x: len(x[1]['findings']), reverse=True):
            trends.append(f"{yard}: {len(info['findings'])} findings across {info['total']} assessments")

    # Common safety terms across findings
    finding_words = Counter()
    safety_terms = ['ppe', 'housekeeping', 'guarding', 'electrical', 'fall', 'fire',
                    'chemical', 'ergonomic', 'noise', 'ventilation', 'lighting',
                    'signage', 'barricade', 'grounding', 'lockout', 'tagout',
                    'harness', 'helmet', 'goggles', 'gloves', 'boots']
    for finding in analysis['all_findings']:
        words = finding['description'].lower().split()
        for word in words:
            if word in safety_terms:
                finding_words[word] += 1

    for term, count in finding_words.most_common(3):
        if count >= 2:
            trends.append(f"{term.upper()} issues noted in {count} assessments")

    # Division activity
    division_counts = Counter()
    for summary in analysis['activity_summary']:
        division_counts[summary['division']] += summary['count']

    if division_counts:
        most_active = division_counts.most_common(1)[0]
        trends.append(f"Most active division: {most_active[0]} ({most_active[1]} assessments)")

    # Clean assessments (positive trend)
    clean_count = sum(1 for s in analysis['activity_summary'] if s['findings_count'] == 0 and s['count'] > 0)
    if clean_count > 0:
        trends.append(f"{clean_count} form type(s) had zero findings - strong compliance")

    analysis['trends'] = trends


# Corrective Action field hashes for form 381707 (from casing_field_assessment_audit.py)
_381707_CA_FIELDS = {
    "1hmw2ia3zskyvejc", "5kdjo1sgtqk062kg", "8cjgdcimxgwgey3p",
    "9plzfzwedqeowmwg", "9y64nqctlfvt8cr9", "b1pn8n8a7q6aeqzw",
    "duqwjgpgqteg7lyk", "dvildqli27bjrqfy", "dxfoh88ikco1hh7l",
    "ig57lp6ouclrouhv", "is0fukv57b4jkgeb", "ix4ronsemj7tjfzt",
    "jujj36vyu4olv7pz", "lxt0mr2nwaad4huo", "n1nkzcisgwshsuuw",
    "nl1rlpr2zo521daa", "ok6tqh93ihdg8dvk", "p27ejw88a0yl0g3u",
    "qruw4o2dcqoow49t", "ra1wwm8bg0u78ab1", "rbgep7tnf0crrwy3",
    "rvx8cq5j48i31q34", "szz5pmdbnuh8a4x2", "t59rttx125h8cxz7",
    "u11h7t0zy2ta7n9z", "ue7h1hb4l16rlmdf", "vcy9kurjo6s8bybe",
    "yu9niadtq5rwadsu", "zozn7dzjkajlolso",
}

# Meta fields to skip when counting narrative depth
_381707_META_FIELDS = {
    "report number", "date", "observer", "observer-emp-num", "status",
    "link", "kpa_link", "parentlink", "parentrepnum",
    "name", "form", "form_id", "updated_at", "created_at",
    "report", "id", "response_id", "_yard", "_observer", "_date", "_ca_fields",
    "7vj2l992y7fwqhwz", "yard", "location",
    "k6qke9eoh052z0eh", "tm4zqob5uficucju",
    "tm4zqob5uficucju-lat", "tm4zqob5uficucju-lon",
    "latitude", "longitude", "temperature", "weather", "wind-speed",
    "updated", "updated_time", "duration", "version", "surrogate",
    "select-dot", "select-flush mount spider", "select-hse", "select-n/a",
    "select-no", "select-quality", "select-safety", "select-slips",
    "select-spider", "select-yes",
}


_YARD_NAMES = ["Midland", "Bryan", "Kilgore", "Hobbs", "Jourdanton", "Laredo"]

# Standard metadata fields to skip in narrative analysis (generic)
_GENERIC_META_FIELDS = {
    "report number", "date", "link", "observer", "name",
    "select-yes", "select-no", "select-n/a", "select-spider",
}


def _normalize_yard(raw_yard):
    """Normalize yard names like 'Midland Yukon' to 'Midland'."""
    for yard in _YARD_NAMES:
        if yard.lower() in raw_yard.lower():
            return yard
    return raw_yard


def _extract_generic_analytics(row):
    """Extract analytics from any form with select-yes/no and -notes fields.

    Works for Vehicle Inspection Checklist, Rathole FA, TD Rig Inspection,
    TD Transcend Field Assessment, etc.
    """
    # Compliance tallies
    try:
        sel_yes = int(row.get("select-yes", 0) or 0)
    except (ValueError, TypeError):
        sel_yes = 0
    try:
        sel_no = int(row.get("select-no", 0) or 0)
    except (ValueError, TypeError):
        sel_no = 0
    try:
        sel_na = int(row.get("select-n/a", 0) or 0)
    except (ValueError, TypeError):
        sel_na = 0

    # CAPAs created (from -followups fields, format "X of Y")
    capas_created = 0
    for key, val in row.items():
        if not key.endswith("-followups"):
            continue
        if not val or not isinstance(val, str):
            continue
        v = val.strip()
        if v and v != "0":
            try:
                parts = v.split(" of ")
                if len(parts) == 2 and int(parts[1]) > 0:
                    capas_created += int(parts[1])
            except (ValueError, IndexError):
                if v not in ("0", ""):
                    capas_created += 1

    # Notes with content (from -notes fields)
    notes_count = 0
    for key, val in row.items():
        if not key.endswith("-notes"):
            continue
        if not val or not isinstance(val, str):
            continue
        if val.strip():
            notes_count += 1

    # Finding detection: any select-no > 0 or capas > 0
    has_finding = sel_no > 0 or capas_created > 0

    return {
        "select_yes": sel_yes,
        "select_no": sel_no,
        "select_na": sel_na,
        "capas_created": capas_created,
        "notes_count": notes_count,
        "has_finding": has_finding,
    }


def _extract_381707_analytics(row):
    """Extract deep analytics fields from a form 381707 row."""
    # Compliance tallies
    try:
        sel_yes = int(row.get("select-yes", 0) or 0)
    except (ValueError, TypeError):
        sel_yes = 0
    try:
        sel_no = int(row.get("select-no", 0) or 0)
    except (ValueError, TypeError):
        sel_no = 0
    try:
        sel_na = int(row.get("select-n/a", 0) or 0)
    except (ValueError, TypeError):
        sel_na = 0

    # CAs marked (count of CA_FIELDS with value >= 1)
    cas_marked = 0
    for field in _381707_CA_FIELDS:
        val = row.get(field, "").strip()
        if val and val != "0":
            try:
                if int(val) > 0:
                    cas_marked += 1
            except ValueError:
                pass

    # CAPAs created (from -followups fields, format "X of Y")
    capas_created = 0
    for key, val in row.items():
        if not key.endswith("-followups"):
            continue
        if not val or not isinstance(val, str):
            continue
        v = val.strip()
        if v and v != "0":
            try:
                parts = v.split(" of ")
                if len(parts) == 2 and int(parts[1]) > 0:
                    capas_created += int(parts[1])
            except (ValueError, IndexError):
                if v not in ("0", ""):
                    capas_created += 1

    # Notes with content (from -notes fields)
    notes_count = 0
    for key, val in row.items():
        if not key.endswith("-notes"):
            continue
        if not val or not isinstance(val, str):
            continue
        v = val.strip()
        if v and v != "0":
            notes_count += 1

    # Narrative depth (fields with 20+ char real text)
    narrative_depth = 0
    for key, val in row.items():
        if key.lower() in _381707_META_FIELDS or key in _381707_CA_FIELDS:
            continue
        if key.startswith("select-") or key.endswith(("-followup-ids", "-followups", "-notes")):
            continue
        if not val or not isinstance(val, str):
            continue
        v = val.strip()
        if len(v) >= 20 and "kpaehs.com" not in v:
            narrative_depth += 1

    has_finding = cas_marked > 0 or capas_created > 0

    return {
        "select_yes": sel_yes,
        "select_no": sel_no,
        "select_na": sel_na,
        "cas_marked": cas_marked,
        "capas_created": capas_created,
        "notes_count": notes_count,
        "narrative_depth": narrative_depth,
        "has_finding": has_finding,
        "customer": row.get("0kcg8hpjaysw1jx8", "").strip(),
        "rig": row.get("wgfyefhpyd2x5pyi", "").strip(),
        "areas_of_concern": row.get("mhmo9hj3tgaad0hw", "").strip()[:300],
        "positive_obs": row.get("s5alwdtplf5gs9ik", "").strip()[:300],
    }


def pull_assessment_history():
    """Pull ALL assessment data from 2025-01-01 onward for all assessment forms.

    Returns a dict with month-by-month and assessor-level aggregation for the
    dashboard's Safety Rep Accountability section. This replaces the yesterday-only
    view with full historical + YTD data.
    """
    from collections import defaultdict
    import json as _json

    start_date = datetime(2025, 1, 1)
    start_ms = int(start_date.timestamp() * 1000)

    print("\n  Pulling assessment history (2025-01-01 to present)...")

    all_assessments = []
    import time as _time

    for form_id, form_info in ASSESSMENT_FORMS.items():
        base_params = {
            "form_id": form_id,
            "format": "csv",
            "updated_after": start_ms,
        }

        # Paginated fetch -- KPA returns max 100 responses per page
        # (each response can have multiple sub-rows for crew/repeating sections)
        # Pre/Post Trip (229645) is very high volume; cap at 10 pages
        all_rows = []
        fieldnames = None
        page = 1
        max_pages = 10 if form_id == 229645 else 50
        while page <= max_pages:
            params = dict(base_params)
            params["page"] = page
            csv_text = call_kpa("responses.flat", params)
            _time.sleep(0.3)
            if not csv_text or csv_text.strip() == "":
                break
            try:
                reader = csv.DictReader(StringIO(csv_text))
                rows = list(reader)
                data = [r for r in rows if r.get("date", "") != "Date"]
                if not data:
                    break
                if fieldnames is None:
                    fieldnames = reader.fieldnames
                # Count unique report numbers to detect pagination boundary
                page_reports = set()
                for r in data:
                    rpt = r.get("report number", "").strip()
                    if rpt:
                        page_reports.add(rpt)
                all_rows.extend(data)
                # KPA returns 100 responses per page; if fewer, we're done
                if len(page_reports) < 100:
                    break
                page += 1
            except Exception:
                break

        if not all_rows:
            print(f"    {form_info['name']}: 0 rows")
            continue

        detected = detect_field_columns(fieldnames or [])

        # Deduplicate by report number (sub-rows share same report number)
        seen_reports = set()
        form_count = 0
        for row in all_rows:
            rpt = row.get("report number", "").strip()
            if not rpt or rpt == "Report Number":
                continue

            date_str = row.get("date", "").strip()
            if not date_str:
                continue  # sub-row

            if rpt in seen_reports:
                continue
            seen_reports.add(rpt)

            # Parse date
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d %H:%M:%S")
            except (ValueError, TypeError):
                try:
                    dt = datetime.strptime(date_str[:10], "%Y-%m-%d")
                except (ValueError, TypeError):
                    continue

            if dt < start_date:
                continue

            assessor = get_assessor_name(row)
            yard = get_yard_from_row(row, detected)
            service_line = _get_service_line(row)

            # Form 381707 has yard in hash field k6qke9eoh052z0eh
            # Values like "Midland Yukon" -- normalize to base yard name
            if form_id == 381707 and yard == "Unknown":
                raw_yard = row.get("k6qke9eoh052z0eh", "").strip()
                if raw_yard:
                    yard = _normalize_yard(raw_yard)

            # Override division/company from form fields for shared forms
            division = form_info["division"]
            company = form_info.get("company", "")
            if form_id == 152018:
                # Vehicle Inspection Checklist has company/service_line fields
                raw_co = row.get("ge09m6h1ne6po6x9", "").strip()
                raw_svc = row.get("hxy6pwclvjke1sln", "").strip()
                if "trucking" in raw_co.lower() or "trucking" in raw_svc.lower():
                    company = "BTI"
                    division = "BTI"
                elif "rat hole" in raw_co.lower() or "rat hole" in raw_svc.lower() or "rathole" in raw_svc.lower():
                    company = "BRHAS"
                    division = "Rathole"
                elif "casing" in raw_svc.lower():
                    company = "BRHAS"
                    division = "Casing"
                else:
                    company = "BRHAS"
                    division = raw_svc or "Unknown"

            entry = {
                "date": dt.strftime("%Y-%m-%d"),
                "month": dt.strftime("%Y-%m"),
                "assessor": assessor,
                "form_id": form_id,
                "form_name": form_info["name"],
                "division": division,
                "company": company,
                "yard": yard,
                "service_line": service_line,
                "report_number": rpt,
            }

            # Deep analytics extraction per form
            if form_id == 381707:
                entry.update(_extract_381707_analytics(row))
            elif form_id in ANALYTICS_FORMS:
                entry.update(_extract_generic_analytics(row))

            all_assessments.append(entry)
            form_count += 1

        print(f"    {form_info['name']}: {form_count} assessments ({len(all_rows)} raw rows, {page} page(s))")

    print(f"  Total assessments (all forms, 2025+): {len(all_assessments)}")

    # Aggregate
    now = datetime.now()
    current_month = now.strftime("%Y-%m")
    current_year = now.year

    # By assessor -- current month
    mtd_by_assessor = defaultdict(lambda: {"count": 0, "yards": set(), "forms": set()})
    # By assessor -- current year
    ytd_by_assessor = defaultdict(lambda: {"count": 0, "yards": set(), "forms": set()})
    # By month (for trend)
    by_month = defaultdict(int)
    # By form -- current month
    mtd_by_form = defaultdict(int)
    # By yard -- current month
    mtd_by_yard = defaultdict(int)
    # By form -- all time
    by_form_all = defaultdict(int)

    for a in all_assessments:
        by_month[a["month"]] += 1
        by_form_all[a["form_name"]] += 1

        if a["date"][:4] == str(current_year):
            ytd_by_assessor[a["assessor"]]["count"] += 1
            ytd_by_assessor[a["assessor"]]["yards"].add(a["yard"])
            ytd_by_assessor[a["assessor"]]["forms"].add(a["form_name"])

        if a["month"] == current_month:
            mtd_by_assessor[a["assessor"]]["count"] += 1
            mtd_by_assessor[a["assessor"]]["yards"].add(a["yard"])
            mtd_by_assessor[a["assessor"]]["forms"].add(a["form_name"])
            mtd_by_form[a["form_name"]] += 1
            mtd_by_yard[a["yard"]] += 1

    # Convert sets to lists for JSON serialization
    mtd_assessors = {}
    for name, data in mtd_by_assessor.items():
        mtd_assessors[name] = {
            "count": data["count"],
            "yards": sorted(data["yards"]),
            "forms": sorted(data["forms"]),
        }

    ytd_assessors = {}
    for name, data in ytd_by_assessor.items():
        ytd_assessors[name] = {
            "count": data["count"],
            "yards": sorted(data["yards"]),
            "forms": sorted(data["forms"]),
        }

    # === DEEP ANALYTICS for form 381707 (Casing Field Assessment) ===
    # Exclude San Angelo (no longer active) and Justin Conrad
    _EXCLUDED_YARDS = {"San Angelo"}
    _EXCLUDED_ASSESSORS = {"Justin Conrad", "Ricky Rhine", "David Dudley"}
    casing_assessments = [
        a for a in all_assessments
        if a.get("form_id") == 381707
        and a.get("yard") not in _EXCLUDED_YARDS
        and a.get("assessor") not in _EXCLUDED_ASSESSORS
    ]
    print(f"  Building analytics for {len(casing_assessments)} casing field assessments...")

    # Junk phrases for recent findings filter
    _JUNK_CONCERNS = {
        "n/a", "na", "none", "none observed", "no unsafe practices observed",
        "no unsafe acts observed", "no unsafe acts were observed",
    }

    def _is_junk_concern(text):
        """Return True if concern text is a non-finding placeholder."""
        if not text:
            return True
        low = text.strip().lower().rstrip(".")
        if low in _JUNK_CONCERNS:
            return True
        # Catch all "no unsafe" / "no issues" / "did not" patterns
        _JUNK_PHRASES = [
            "no unsafe", "did not note", "did not witness", "did not observe",
            "no one safe practices", "no issues noted", "no issues observed",
            "no concerns noted", "no concerns observed", "no issues were",
            "i did not note", "i did not witness", "i did not observe",
            "since the crew was only", "no other issues noted",
            "none were observed", "no issues with", "not observed",
            "were not observed", "no findings", "nothing noted",
            "no areas of concern", "no safety concerns",
        ]
        for phrase in _JUNK_PHRASES:
            if phrase in low:
                return True
        return False

    # Active rep roster with yard assignments and monthly targets
    _ACTIVE_REPS = {
        "Michael Salazar": {"yards": "Midland / Jourdanton", "target": 8},
        "Michael Hancock": {"yards": "Midland", "target": 8},
        "Joseph Speyrer": {"yards": "Jourdanton / Laredo", "target": 6},
        "James Barnett": {"yards": "Kilgore", "target": 6},
        "Allen Batts": {"yards": "Hobbs", "target": 7},
    }

    analytics = None
    if casing_assessments:
        # YTD assessments only for KPIs and scorecard (consistent time window)
        ytd_casing = [a for a in casing_assessments if a["date"][:4] == str(current_year)]
        n_ytd = len(ytd_casing)

        # Overall KPIs (YTD)
        total_yes = sum(a.get("select_yes", 0) for a in ytd_casing)
        total_no = sum(a.get("select_no", 0) for a in ytd_casing)
        total_responses = total_yes + total_no
        compliance_rate = round(total_yes / total_responses * 100, 1) if total_responses > 0 else 0
        finding_count = sum(1 for a in ytd_casing if a.get("has_finding"))
        finding_rate = round(finding_count / n_ytd * 100, 1) if n_ytd > 0 else 0
        avg_capas = round(sum(a.get("capas_created", 0) for a in ytd_casing) / n_ytd, 2) if n_ytd > 0 else 0
        avg_notes = round(sum(a.get("notes_count", 0) for a in ytd_casing) / n_ytd, 2) if n_ytd > 0 else 0
        avg_narrative = round(sum(a.get("narrative_depth", 0) for a in ytd_casing) / n_ytd, 1) if n_ytd > 0 else 0

        # By month (all-time for trend chart)
        analytics_by_month = defaultdict(lambda: {
            "count": 0, "yes": 0, "no": 0, "findings": 0,
            "capas": 0, "notes": 0, "narrative": 0
        })
        for a in casing_assessments:
            m = a["month"]
            analytics_by_month[m]["count"] += 1
            analytics_by_month[m]["yes"] += a.get("select_yes", 0)
            analytics_by_month[m]["no"] += a.get("select_no", 0)
            analytics_by_month[m]["findings"] += 1 if a.get("has_finding") else 0
            analytics_by_month[m]["capas"] += a.get("capas_created", 0)
            analytics_by_month[m]["notes"] += a.get("notes_count", 0)
            analytics_by_month[m]["narrative"] += a.get("narrative_depth", 0)

        by_month_analytics = {}
        for m, d in sorted(analytics_by_month.items()):
            total_resp = d["yes"] + d["no"]
            by_month_analytics[m] = {
                "count": d["count"],
                "compliance_rate": round(d["yes"] / total_resp * 100, 1) if total_resp > 0 else 0,
                "finding_rate": round(d["findings"] / d["count"] * 100, 1) if d["count"] > 0 else 0,
                "avg_capas": round(d["capas"] / d["count"], 2) if d["count"] > 0 else 0,
                "avg_notes": round(d["notes"] / d["count"], 2) if d["count"] > 0 else 0,
            }

        # By assessor (YTD for quality scoring)
        assessor_rows = defaultdict(list)
        for a in ytd_casing:
            assessor_rows[a["assessor"]].append(a)

        # Compute dataset-wide YTD averages for benchmark (not Kelly Rhodes outlier)
        bench_capas = avg_capas if avg_capas > 0 else 0.5
        bench_notes = avg_notes if avg_notes > 0 else 0.5
        bench_narr = avg_narrative if avg_narrative > 0 else 4.0
        # Scale benchmarks to 2x dataset avg (top performer target)
        bench_capas_target = max(bench_capas * 2, 1.0)
        bench_notes_target = max(bench_notes * 2, 1.0)
        bench_narr_target = max(bench_narr * 1.5, 4.0)

        analytics_by_assessor = {}
        for name, rows in assessor_rows.items():
            n = len(rows)
            a_yes = sum(r.get("select_yes", 0) for r in rows)
            a_no = sum(r.get("select_no", 0) for r in rows)
            a_resp = a_yes + a_no
            a_findings = sum(1 for r in rows if r.get("has_finding"))
            a_capas = sum(r.get("capas_created", 0) for r in rows)
            a_notes = sum(r.get("notes_count", 0) for r in rows)
            a_narr = sum(r.get("narrative_depth", 0) for r in rows)

            # Quality score (0-85) with dataset-relative benchmarks
            # CAPAs: 29pts, Notes: 25pts, Narrative: 17pts, CA Rate: 14pts
            capa_score = min((a_capas / n) / bench_capas_target, 1.0) * 29 if n > 0 else 0
            notes_score = min((a_notes / n) / bench_notes_target, 1.0) * 25 if n > 0 else 0
            narr_score = min((a_narr / n) / bench_narr_target, 1.0) * 17 if n > 0 else 0
            find_score = min((a_findings / n) / 0.5, 1.0) * 14 if n > 0 else 0
            quality = round(min(capa_score + notes_score + narr_score + find_score, 85))

            # Red flag detection
            flags = []
            # Consecutive clean assessments (no findings)
            sorted_rows = sorted(rows, key=lambda x: x["date"])
            consecutive_clean = 0
            max_clean = 0
            for r in sorted_rows:
                if not r.get("has_finding"):
                    consecutive_clean += 1
                    max_clean = max(max_clean, consecutive_clean)
                else:
                    consecutive_clean = 0
            if max_clean >= 10:
                flags.append(f"{max_clean} consecutive assessments with no findings")

            if n >= 3 and a_capas == 0:
                flags.append(f"0 CAPAs created across {n} assessments")

            if n >= 3 and a_notes == 0:
                flags.append(f"0 observation notes across {n} assessments")

            analytics_by_assessor[name] = {
                "count": n,
                "compliance_rate": round(a_yes / a_resp * 100, 1) if a_resp > 0 else 0,
                "finding_rate": round(a_findings / n * 100, 1) if n > 0 else 0,
                "avg_capas": round(a_capas / n, 2) if n > 0 else 0,
                "avg_notes": round(a_notes / n, 2) if n > 0 else 0,
                "avg_narrative": round(a_narr / n, 1) if n > 0 else 0,
                "quality_score": quality,
                "red_flags": flags,
            }

        # Coverage gap detection -- reps with 0 assessments this month
        for rep_name, rep_info in _ACTIVE_REPS.items():
            mtd_count = sum(1 for a in ytd_casing if a["assessor"] == rep_name and a["month"] == current_month)
            day_of_month = now.day
            if mtd_count == 0 and day_of_month >= 10:
                if rep_name not in analytics_by_assessor:
                    analytics_by_assessor[rep_name] = {
                        "count": 0, "compliance_rate": 0, "finding_rate": 0,
                        "avg_capas": 0, "avg_notes": 0, "avg_narrative": 0,
                        "quality_score": 0, "red_flags": [],
                    }
                analytics_by_assessor[rep_name]["red_flags"].append(
                    f"0 assessments filed in {current_month} ({day_of_month} days in)"
                )
            elif mtd_count > 0 and day_of_month >= 15:
                # Check if on pace for monthly target
                target = rep_info["target"]
                pace = round(mtd_count / day_of_month * 30)
                if pace < target * 0.5:
                    if rep_name in analytics_by_assessor:
                        analytics_by_assessor[rep_name]["red_flags"].append(
                            f"On pace for ~{pace} assessments this month (target: {target})"
                        )

        # By yard (YTD, with rep attribution)
        _YARD_TO_REP = {
            "Midland": "Salazar / Hancock",
            "Bryan": "",
            "Kilgore": "Barnett",
            "Hobbs": "Batts",
            "Jourdanton": "Speyrer / Salazar",
            "Laredo": "Speyrer",
        }
        yard_data = defaultdict(lambda: {
            "count": 0, "yes": 0, "no": 0, "findings": 0, "capas": 0,
            "assessors": set()
        })
        for a in ytd_casing:
            y = a.get("yard", "Unknown")
            yard_data[y]["count"] += 1
            yard_data[y]["yes"] += a.get("select_yes", 0)
            yard_data[y]["no"] += a.get("select_no", 0)
            yard_data[y]["findings"] += 1 if a.get("has_finding") else 0
            yard_data[y]["capas"] += a.get("capas_created", 0)
            yard_data[y]["assessors"].add(a.get("assessor", "Unknown"))

        analytics_by_yard = {}
        for y, d in yard_data.items():
            total_resp = d["yes"] + d["no"]
            analytics_by_yard[y] = {
                "count": d["count"],
                "compliance_rate": round(d["yes"] / total_resp * 100, 1) if total_resp > 0 else 0,
                "finding_rate": round(d["findings"] / d["count"] * 100, 1) if d["count"] > 0 else 0,
                "avg_capas": round(d["capas"] / d["count"], 2) if d["count"] > 0 else 0,
                "rep": _YARD_TO_REP.get(y, ""),
                "assessors": sorted(d["assessors"]),
            }
        # Ensure all active yards appear even if 0 assessments
        for y, rep in _YARD_TO_REP.items():
            if y not in analytics_by_yard:
                analytics_by_yard[y] = {
                    "count": 0, "compliance_rate": 0, "finding_rate": 0,
                    "avg_capas": 0, "rep": rep or "NO REP ASSIGNED",
                    "assessors": [],
                }

        # Recent findings (last 20, filtered to real concerns only)
        with_findings = [
            a for a in casing_assessments
            if a.get("areas_of_concern") and not _is_junk_concern(a.get("areas_of_concern", ""))
        ]
        # Also include assessments with CAs or CAPAs even if concern text is junk
        ca_findings = [
            a for a in casing_assessments
            if (a.get("cas_marked", 0) > 0 or a.get("capas_created", 0) > 0)
            and a not in with_findings
        ]
        all_findings_pool = with_findings + ca_findings
        all_findings_pool.sort(key=lambda x: x["date"], reverse=True)
        # Deduplicate by report_number
        seen_rpts = set()
        recent_findings = []
        for a in all_findings_pool:
            rpt = a.get("report_number", "")
            if rpt in seen_rpts:
                continue
            seen_rpts.add(rpt)
            concern = a.get("areas_of_concern", "")
            if _is_junk_concern(concern):
                concern = "(Finding noted via CA/CAPA -- no narrative provided)"
            recent_findings.append({
                "date": a["date"],
                "assessor": a["assessor"],
                "yard": a.get("yard", "Unknown"),
                "customer": a.get("customer", ""),
                "rig": a.get("rig", ""),
                "concern": concern,
                "positive": a.get("positive_obs", ""),
                "cas_marked": a.get("cas_marked", 0),
                "capas_created": a.get("capas_created", 0),
            })
            if len(recent_findings) >= 20:
                break

        # Collect all red flags
        all_red_flags = []
        for name, data in analytics_by_assessor.items():
            for flag in data.get("red_flags", []):
                all_red_flags.append({"assessor": name, "detail": flag})

        analytics = {
            "overall": {
                "total_ytd": n_ytd,
                "total_all_time": len(casing_assessments),
                "compliance_rate": compliance_rate,
                "finding_rate": finding_rate,
                "avg_capas": avg_capas,
                "avg_notes": avg_notes,
                "avg_narrative": avg_narrative,
            },
            "by_month": by_month_analytics,
            "by_assessor": analytics_by_assessor,
            "by_yard": analytics_by_yard,
            "recent_findings": recent_findings,
            "red_flags": all_red_flags,
        }

    # === GENERIC ANALYTICS for other forms (BTI, Rathole, Transcend) ===
    form_analytics = {}

    def _build_form_analytics(key, label, company, division, all_rows, ytd_rows):
        """Build analytics dict for a set of form rows."""
        n_ytd = len(ytd_rows)
        if n_ytd == 0:
            return
        f_yes = sum(a.get("select_yes", 0) for a in ytd_rows)
        f_no = sum(a.get("select_no", 0) for a in ytd_rows)
        f_resp = f_yes + f_no
        f_compliance = round(f_yes / f_resp * 100, 1) if f_resp > 0 else 0
        f_findings = sum(1 for a in ytd_rows if a.get("has_finding"))
        f_finding_rate = round(f_findings / n_ytd * 100, 1) if n_ytd > 0 else 0
        f_capas = sum(a.get("capas_created", 0) for a in ytd_rows)
        f_notes = sum(a.get("notes_count", 0) for a in ytd_rows)
        f_avg_capas = round(f_capas / n_ytd, 2) if n_ytd > 0 else 0
        f_avg_notes = round(f_notes / n_ytd, 2) if n_ytd > 0 else 0

        # By month (all-time for trend)
        f_by_month = defaultdict(lambda: {"count": 0, "yes": 0, "no": 0, "findings": 0})
        for a in all_rows:
            m = a["month"]
            f_by_month[m]["count"] += 1
            f_by_month[m]["yes"] += a.get("select_yes", 0)
            f_by_month[m]["no"] += a.get("select_no", 0)
            f_by_month[m]["findings"] += 1 if a.get("has_finding") else 0
        f_month_analytics = {}
        for m, d in sorted(f_by_month.items()):
            tr = d["yes"] + d["no"]
            f_month_analytics[m] = {
                "count": d["count"],
                "compliance_rate": round(d["yes"] / tr * 100, 1) if tr > 0 else 0,
                "finding_rate": round(d["findings"] / d["count"] * 100, 1) if d["count"] > 0 else 0,
            }

        # By assessor (YTD)
        f_assessor_rows = defaultdict(list)
        for a in ytd_rows:
            f_assessor_rows[a["assessor"]].append(a)
        f_by_assessor = {}
        for aname, arows in f_assessor_rows.items():
            an = len(arows)
            ay = sum(r.get("select_yes", 0) for r in arows)
            ano = sum(r.get("select_no", 0) for r in arows)
            aresp = ay + ano
            af = sum(1 for r in arows if r.get("has_finding"))
            ac = sum(r.get("capas_created", 0) for r in arows)
            anotes = sum(r.get("notes_count", 0) for r in arows)
            flags = []
            sorted_a = sorted(arows, key=lambda x: x["date"])
            consec = 0
            max_consec = 0
            for r in sorted_a:
                if not r.get("has_finding"):
                    consec += 1
                    max_consec = max(max_consec, consec)
                else:
                    consec = 0
            if max_consec >= 10:
                flags.append(f"{max_consec} consecutive with no findings")
            if an >= 3 and ac == 0:
                flags.append(f"0 CAPAs across {an} assessments")
            if an >= 3 and anotes == 0:
                flags.append(f"0 notes across {an} assessments")
            f_by_assessor[aname] = {
                "count": an,
                "compliance_rate": round(ay / aresp * 100, 1) if aresp > 0 else 0,
                "finding_rate": round(af / an * 100, 1) if an > 0 else 0,
                "avg_capas": round(ac / an, 2) if an > 0 else 0,
                "avg_notes": round(anotes / an, 2) if an > 0 else 0,
                "red_flags": flags,
            }

        f_red_flags = []
        for aname, adata in f_by_assessor.items():
            for flag in adata.get("red_flags", []):
                f_red_flags.append({"assessor": aname, "detail": flag})

        form_analytics[key] = {
            "form_name": label,
            "company": company,
            "division": division,
            "overall": {
                "total_ytd": n_ytd,
                "total_all_time": len(all_rows),
                "compliance_rate": f_compliance,
                "finding_rate": f_finding_rate,
                "avg_capas": f_avg_capas,
                "avg_notes": f_avg_notes,
            },
            "by_month": f_month_analytics,
            "by_assessor": f_by_assessor,
            "red_flags": f_red_flags,
        }
        print(f"  {label} ({company}/{division}): {n_ytd} YTD, compliance={f_compliance}%, findings={f_finding_rate}%")

    for af_id, af_info in ANALYTICS_FORMS.items():
        if af_id == 381707:
            continue  # Casing handled above with deep analysis
        form_rows = [a for a in all_assessments if a.get("form_id") == af_id]
        if not form_rows:
            continue
        form_name = ASSESSMENT_FORMS[af_id]["name"]

        if af_info["company"] == "All":
            # Shared form -- split by company from entry data
            companies_seen = set(a.get("company", "") for a in form_rows)
            for co in sorted(companies_seen):
                if not co or co == "Unknown":
                    continue
                co_rows = [a for a in form_rows if a.get("company") == co]
                co_div = co_rows[0].get("division", co) if co_rows else co
                co_ytd = [a for a in co_rows if a["date"][:4] == str(current_year)]
                key = f"{af_id}_{co}"
                label = f"{form_name} ({co})"
                _build_form_analytics(key, label, co, co_div, co_rows, co_ytd)
        else:
            ytd_rows = [a for a in form_rows if a["date"][:4] == str(current_year)]
            _build_form_analytics(
                str(af_id), form_name, af_info["company"], af_info["division"],
                form_rows, ytd_rows
            )

    result = {
        "current_month": current_month,
        "current_year": current_year,
        "total_all_time": len(all_assessments),
        "total_mtd": sum(1 for a in all_assessments if a["month"] == current_month),
        "total_ytd": sum(1 for a in all_assessments if a["date"][:4] == str(current_year)),
        "mtd_by_assessor": mtd_assessors,
        "ytd_by_assessor": ytd_assessors,
        "mtd_by_form": dict(mtd_by_form),
        "mtd_by_yard": dict(mtd_by_yard),
        "by_month": dict(sorted(by_month.items())),
        "by_form_all": dict(by_form_all),
        "assessment_analytics": analytics,
        "form_analytics": form_analytics,
    }

    print(f"  MTD assessments: {result['total_mtd']}")
    print(f"  YTD assessments: {result['total_ytd']}")
    if analytics:
        print(f"  Casing compliance rate: {analytics['overall']['compliance_rate']}%")
        print(f"  Casing finding rate: {analytics['overall']['finding_rate']}%")
        print(f"  Red flags: {len(analytics['red_flags'])}")

    return result


def _generate_assessment_recommendations(analysis):
    """Generate leadership recommendations based on assessment analysis"""
    recs = analysis['recommendations']

    # IMMEDIATE: Critical and high findings
    critical = analysis['findings_by_severity']['critical']
    if critical:
        yards = set(f['yard'] for f in critical)
        recs['immediate'].append(f"Address {len(critical)} critical finding(s) in: {', '.join(yards)}")

    high = analysis['findings_by_severity']['high']
    if high:
        recs['immediate'].append(f"Review {len(high)} high-severity finding(s) requiring prompt attention")

    # THIS WEEK: Non-compliant yards and corrective actions
    non_compliant_yards = [yard for yard, info in analysis['compliance_by_yard'].items()
                           if info['non_compliant'] > 0]
    if non_compliant_yards:
        recs['this_week'].append(f"Follow up on non-compliant assessments at: {', '.join(non_compliant_yards[:5])}")

    open_cas = [ca for ca in analysis['corrective_actions'] if ca['status'] == 'Open']
    if open_cas:
        recs['this_week'].append(f"Track {len(open_cas)} open corrective action(s) to closure")

    # MONTHLY: Recognition and coverage
    if analysis['assessor_stats']:
        top_assessors = sorted(analysis['assessor_stats'].items(),
                               key=lambda x: x[1]['total'], reverse=True)[:3]
        names = [a[0] for a in top_assessors if a[0] != 'Unknown']
        if names:
            recs['monthly'].append(f"Recognize top assessors: {', '.join(names)}")

    active_divisions = set(s['division'] for s in analysis['activity_summary'])
    all_divisions = set(info['division'] for info in ASSESSMENT_FORMS.values())
    missing = all_divisions - active_divisions
    if missing:
        recs['monthly'].append(f"No assessments from: {', '.join(missing)} - consider scheduling")

    recs['monthly'].append("Review assessment frequency targets vs. actual completion rates")


# ==============================================================================
# ASSESSMENT AUDIT SUMMARY (replaces old "OTHER SAFETY FORMS SUMMARY")
# ==============================================================================

def _get_customer_from_row(row):
    """Extract customer/client name from a form row"""
    for field_name, value in row.items():
        if any(kw in field_name.lower() for kw in ['customer', 'client', 'company', 'operator', 'contractor']):
            if value and value.strip() and value.strip().lower() not in ['n/a', 'none', 'unknown', '']:
                return value.strip()
    return ''


def _get_issue_from_row(row, detected_fields):
    """Extract the primary issue/finding text from a form row"""
    # Try detected finding fields first
    for field in detected_fields.get('findings', []):
        val = row.get(field, '').strip()
        if val and len(val) > 3 and val.lower() not in ['n/a', 'none', 'no', 'na', 'no issues']:
            return val[:120]

    # Try corrective action fields (often contain the issue description)
    for field in detected_fields.get('corrective_action', []):
        val = row.get(field, '').strip()
        if val and len(val) > 3 and val.lower() not in ['n/a', 'none', 'no', 'na']:
            return val[:120]

    # Try observation description field used by observation cards
    for key in ['uncbcge9x8vow9pn']:
        val = row.get(key, '').strip()
        if val and len(val) > 3 and val.lower() not in ['n/a', 'none', 'no', 'na']:
            return val[:120]

    return 'None noted'


def extract_assessment_details(all_data):
    """Extract assessor, location, customer, form_id, and issues from each assessment row.

    Returns a list of dicts, one per form type in OTHER_FORMS.
    Each dict has: form_name, form_id, count, rows (list of detail dicts).
    Forms with 0 assessments still appear with count=0 and empty rows.
    """
    results = []

    for form_id, form_name in OTHER_FORMS:
        data = all_data.get(f"form_{form_id}")
        entry = {
            'form_name': form_name,
            'form_id': form_id,
            'count': data['count'] if data else 0,
            'rows': [],
        }

        if data and data['count'] > 0:
            detected = detect_field_columns(data['headers'])

            for row in data['rows']:
                report_num = row.get('report number', '')
                entry['rows'].append({
                    'assessor': get_assessor_name(row),
                    'location': get_yard_from_row(row, detected),
                    'customer': _get_customer_from_row(row),
                    'form_id': report_num,
                    'link': get_kpa_link(report_num),
                    'issue': _get_issue_from_row(row, detected),
                    'service_line': _get_service_line(row),
                })

        results.append(entry)

    return results


def add_assessment_audit_summary(doc, assessment_details):
    """Create a Word table summarizing all assessment/audit forms.

    Replaces the old 'OTHER SAFETY FORMS SUMMARY' with a 6-column table:
    Form Type | Assessor | Location | Customer | Form ID | Issue Found
    """
    doc.add_page_break()
    add_heading(doc, "ASSESSMENT & AUDIT SUMMARY", 1)
    doc.add_paragraph()

    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    # Check if there are any rows at all
    total_rows = sum(entry['count'] for entry in assessment_details)

    if total_rows == 0:
        p = doc.add_paragraph()
        p.add_run("No assessment or audit forms were completed yesterday.").font.italic = True

        # Still show the form list with counts
        doc.add_paragraph()
        for entry in assessment_details:
            p = doc.add_paragraph()
            run = p.add_run(f"{entry['form_name']}: ")
            run.font.bold = True
            p.add_run("0")
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Dark header row
    headers = ['Form Type', 'Assessor', 'Location', 'Customer', 'Form ID', 'Issue Found']
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr_cells[i].text = txt
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading = _OE('w:shd')
        shading.set(_qn('w:fill'), '800000')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    for entry in assessment_details:
        if entry['count'] == 0:
            # Show a single row with 0 count
            row_cells = table.add_row().cells
            row_cells[0].text = entry['form_name']
            row_cells[1].text = '-'
            row_cells[2].text = '-'
            row_cells[3].text = '-'
            row_cells[4].text = '-'
            row_cells[5].text = '0 assessments'
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(128, 128, 128)
        else:
            for detail in entry['rows']:
                row_cells = table.add_row().cells
                row_cells[0].text = entry['form_name']
                row_cells[1].text = detail['assessor']
                row_cells[2].text = detail['location']
                row_cells[3].text = detail['customer'] or '-'
                row_cells[4].text = str(detail['form_id'])
                row_cells[5].text = detail['issue']

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

                # Color-code the issue cell
                issue_text = detail['issue'].lower()
                if issue_text != 'none noted':
                    for paragraph in row_cells[5].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = COLORS['warning']

                # Make Form ID a clickable link if available
                if detail['link']:
                    for paragraph in row_cells[4].paragraphs:
                        paragraph.clear()
                    p = row_cells[4].paragraphs[0]
                    add_hyperlink(p, detail['link'], str(detail['form_id']))

    # Summary line
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Total: {total_rows} assessments/audits completed").font.bold = True


def build_assessment_html(assessment_details):
    """Build an HTML table for the assessment/audit summary in email.

    Returns an HTML string with a styled table matching BRHAS color scheme.
    """
    total_rows = sum(entry['count'] for entry in assessment_details)

    if total_rows == 0:
        html = '<p style="font-style:italic;">No assessment or audit forms were completed yesterday.</p>'
        html += '<ul style="margin:5px 0;color:#888;">'
        for entry in assessment_details:
            html += f'<li><b>{_h(entry["form_name"])}:</b> 0</li>'
        html += '</ul>'
        return html

    html = '<table width="100%" cellpadding="5" cellspacing="0" '
    html += 'style="border-collapse:collapse;font-size:12px;margin-bottom:10px;">'

    # Header
    html += f'<tr style="background:{HTML_COLORS["secondary"]};color:#ffffff;">'
    for hdr in ['Form Type', 'Assessor', 'Location', 'Customer', 'Form ID', 'Issue Found']:
        html += f'<th style="text-align:left;padding:8px;border:1px solid #600000;">{hdr}</th>'
    html += '</tr>'

    row_idx = 0
    for entry in assessment_details:
        if entry['count'] == 0:
            bg = '#f9f9f9' if row_idx % 2 == 0 else '#ffffff'
            html += f'<tr style="background:{bg};color:#999;">'
            html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(entry["form_name"])}</td>'
            for _ in range(4):
                html += '<td style="border:1px solid #ddd;padding:6px;text-align:center;">-</td>'
            html += '<td style="border:1px solid #ddd;padding:6px;">0 assessments</td>'
            html += '</tr>'
            row_idx += 1
        else:
            for detail in entry['rows']:
                bg = '#f9f9f9' if row_idx % 2 == 0 else '#ffffff'
                html += f'<tr style="background:{bg};">'
                html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(entry["form_name"])}</td>'
                html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(detail["assessor"])}</td>'
                html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(detail["location"])}</td>'
                html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(detail["customer"]) or "-"}</td>'

                # Form ID with link
                if detail['link']:
                    html += f'<td style="border:1px solid #ddd;padding:6px;">'
                    html += f'<a href="{_h(detail["link"])}" style="color:#0563C1;">{_h(detail["form_id"])}</a></td>'
                else:
                    html += f'<td style="border:1px solid #ddd;padding:6px;">{_h(detail["form_id"])}</td>'

                # Issue with color
                issue = detail['issue']
                if issue.lower() != 'none noted':
                    html += f'<td style="border:1px solid #ddd;padding:6px;color:{HTML_COLORS["warning"]};">{_h(issue)}</td>'
                else:
                    html += f'<td style="border:1px solid #ddd;padding:6px;color:{HTML_COLORS["safe"]};">{_h(issue)}</td>'

                html += '</tr>'
                row_idx += 1

    html += '</table>'
    html += f'<p><b>Total: {total_rows} assessments/audits completed</b></p>'

    return html


# ==============================================================================
# DOCUMENT HELPERS
# ==============================================================================

def add_heading(doc, text, level=1, color=None):
    """Add formatted heading"""
    p = doc.add_paragraph()
    run = p.add_run(text)

    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = color or COLORS['primary']
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = color or COLORS['secondary']
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = color or COLORS['accent']

    return p


# ==============================================================================
# ASSESSMENT & AUDIT ANALYSIS - WORD DOCUMENT SECTION
# ==============================================================================

def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a Word document paragraph"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_assessment_analysis_section(doc, assessment_data):
    """Add Assessment & Audit Analysis section to the Word document.

    Inserts after Incident Timing Analysis, before At-Risk Conditions.
    Shows assessor activity, compliance by yard, critical findings,
    corrective actions, trends, and leadership recommendations.
    """
    if not assessment_data or not assessment_data.get('has_data'):
        return

    doc.add_page_break()
    add_heading(doc, "ASSESSMENT & AUDIT ANALYSIS", 1, COLORS['primary'])

    p = doc.add_paragraph()
    p.add_run(f"Total Assessments Completed: ").font.bold = True
    p.add_run(f"{assessment_data['total_assessments']}")
    p.add_run(f"  |  ")
    p.add_run(f"Total Findings: ").font.bold = True
    p.add_run(f"{assessment_data['total_findings']}")
    doc.add_paragraph()

    # --- 1. Activity Summary Table ---
    add_heading(doc, "Assessment Activity Summary", 2)

    if assessment_data['activity_summary']:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, txt in enumerate(['Form', 'Count', 'Assessor(s)', 'Findings', 'Compliance']):
            hdr_cells[i].text = txt
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Dark header background
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OE
            shading = _OE('w:shd')
            shading.set(_qn('w:fill'), '800000')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

        for summary in assessment_data['activity_summary']:
            row_cells = table.add_row().cells
            row_cells[0].text = summary['form_name']
            row_cells[1].text = str(summary['count'])

            assessor_text = ', '.join(summary['assessors'][:3])
            if len(summary['assessors']) > 3:
                assessor_text += f" +{len(summary['assessors']) - 3} more"
            row_cells[2].text = assessor_text

            row_cells[3].text = str(summary['findings_count'])

            if summary['count'] > 0:
                rate = summary['compliance_rate']
                if rate >= 90:
                    row_cells[4].text = f"\u2705 {rate:.0f}%"
                elif rate >= 70:
                    row_cells[4].text = f"\U0001f7e1 {rate:.0f}%"
                else:
                    row_cells[4].text = f"\U0001f534 {rate:.0f}%"
            else:
                row_cells[4].text = "N/A"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()

    # --- 2. Compliance Dashboard by Yard ---
    if assessment_data['compliance_by_yard']:
        add_heading(doc, "Compliance by Yard", 2)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, txt in enumerate(['Yard/Location', 'Assessments', 'Compliant', 'Non-Compliant', 'Status']):
            hdr_cells[i].text = txt
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OE
            shading = _OE('w:shd')
            shading.set(_qn('w:fill'), '800000')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

        for yard, info in sorted(assessment_data['compliance_by_yard'].items(),
                                  key=lambda x: x[1]['non_compliant'], reverse=True):
            row_cells = table.add_row().cells
            row_cells[0].text = yard
            row_cells[1].text = str(info['total'])
            row_cells[2].text = str(info['compliant'])
            row_cells[3].text = str(info['non_compliant'])

            if info['total'] > 0:
                rate = info['compliant'] / info['total'] * 100
                if rate >= 90:
                    row_cells[4].text = f"\u2705 {rate:.0f}%"
                elif rate >= 70:
                    row_cells[4].text = f"\U0001f7e1 {rate:.0f}%"
                else:
                    row_cells[4].text = f"\U0001f534 {rate:.0f}%"
            else:
                row_cells[4].text = "N/A"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    # --- 3. Critical Findings ---
    critical = assessment_data['findings_by_severity']['critical']
    high = assessment_data['findings_by_severity']['high']

    if critical or high:
        add_heading(doc, "Critical Findings - Immediate Attention Required", 2, COLORS['critical'])

        for finding in critical:
            p = doc.add_paragraph()
            run = p.add_run("\U0001f534 CRITICAL: ")
            run.font.bold = True
            run.font.color.rgb = COLORS['critical']
            p.add_run(finding['description'])

            doc.add_paragraph(
                f"Form: {finding['form_name']} | Assessor: {finding['assessor']}",
                style='List Bullet'
            )
            doc.add_paragraph(
                f"Yard: {finding['yard']} | Date: {finding['date']}",
                style='List Bullet'
            )

            if finding['link']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("View in KPA: ")
                add_hyperlink(p, finding['link'], finding['link'])

            doc.add_paragraph()

        for finding in high[:5]:
            p = doc.add_paragraph()
            run = p.add_run("\U0001f7e1 HIGH: ")
            run.font.bold = True
            run.font.color.rgb = COLORS['warning']
            p.add_run(finding['description'])

            doc.add_paragraph(
                f"Form: {finding['form_name']} | Yard: {finding['yard']}",
                style='List Bullet'
            )

            if finding['link']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("View in KPA: ")
                add_hyperlink(p, finding['link'], finding['link'])

            doc.add_paragraph()

        if len(high) > 5:
            p = doc.add_paragraph()
            run = p.add_run(f"... and {len(high) - 5} more high-severity findings")
            run.font.italic = True
    else:
        add_heading(doc, "Findings Summary", 2)
        medium = assessment_data['findings_by_severity']['medium']
        low = assessment_data['findings_by_severity']['low']

        if medium or low:
            p = doc.add_paragraph()
            p.add_run("No critical or high-severity findings. ").font.bold = True
            p.add_run(f"{len(medium)} medium, {len(low)} low-severity items noted.")
        else:
            p = doc.add_paragraph("\u2705 No findings - All assessments passed!")
            p.runs[0].font.color.rgb = COLORS['safe']
            p.runs[0].font.bold = True

    doc.add_paragraph()

    # --- 4. Top Performing Assessors ---
    if assessment_data['assessor_stats']:
        add_heading(doc, "Top Performing Assessors", 2, COLORS['safe'])

        sorted_assessors = sorted(
            assessment_data['assessor_stats'].items(),
            key=lambda x: x[1]['total'], reverse=True
        )

        rank = 0
        for name, stats in sorted_assessors[:10]:
            if name == 'Unknown':
                continue
            rank += 1

            p = doc.add_paragraph()
            prefix = "\u2B50 " if rank <= 3 else "   "
            run = p.add_run(f"{prefix}{rank}. {name}")
            run.font.bold = True

            divisions = ', '.join(stats['divisions']) if stats['divisions'] else 'N/A'
            detail = f" - {stats['total']} assessment(s) | Divisions: {divisions}"
            if stats['findings_found'] > 0:
                detail += f" | {stats['findings_found']} finding(s) identified"
            p.add_run(detail)

        doc.add_paragraph()

    # --- 5. Corrective Actions Tracker ---
    if assessment_data['corrective_actions']:
        add_heading(doc, "Corrective Actions Tracker", 2, COLORS['warning'])

        p = doc.add_paragraph()
        p.add_run(f"Open Corrective Actions: {len(assessment_data['corrective_actions'])}").font.bold = True
        doc.add_paragraph()

        for i, ca in enumerate(assessment_data['corrective_actions'][:10], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {ca['description']}")
            run.font.bold = True

            doc.add_paragraph(
                f"Form: {ca['form_name']} | Yard: {ca['yard']}",
                style='List Bullet'
            )
            doc.add_paragraph(
                f"Identified by: {ca['assessor']} on {ca['date']}",
                style='List Bullet'
            )

            if ca['link']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("View: ")
                add_hyperlink(p, ca['link'], ca['link'])

        if len(assessment_data['corrective_actions']) > 10:
            p = doc.add_paragraph()
            run = p.add_run(
                f"... and {len(assessment_data['corrective_actions']) - 10} more corrective actions"
            )
            run.font.italic = True

        doc.add_paragraph()

    # --- 6. Trends & Patterns ---
    if assessment_data['trends']:
        add_heading(doc, "Trends & Patterns", 2)

        for trend in assessment_data['trends']:
            doc.add_paragraph(f"\U0001F4CA {trend}", style='List Bullet')

        doc.add_paragraph()

    # --- 7. Recommended Actions for Leadership ---
    recs = assessment_data['recommendations']
    if any([recs['immediate'], recs['this_week'], recs['monthly']]):
        add_heading(doc, "Recommended Actions for Leadership", 2, COLORS['primary'])

        if recs['immediate']:
            p = doc.add_paragraph()
            run = p.add_run("\U0001f534 IMMEDIATE:")
            run.font.bold = True
            run.font.color.rgb = COLORS['critical']
            for rec in recs['immediate']:
                doc.add_paragraph(rec, style='List Bullet')

        if recs['this_week']:
            p = doc.add_paragraph()
            run = p.add_run("\U0001f7e1 THIS WEEK:")
            run.font.bold = True
            run.font.color.rgb = COLORS['warning']
            for rec in recs['this_week']:
                doc.add_paragraph(rec, style='List Bullet')

        if recs['monthly']:
            p = doc.add_paragraph()
            run = p.add_run("\U0001F4CA MONTH-OVER-MONTH:")
            run.font.bold = True
            for rec in recs['monthly']:
                doc.add_paragraph(rec, style='List Bullet')


# ==============================================================================
# BUILD WORD DOCUMENT
# ==============================================================================

def build_word_document(all_data, yesterday_date):
    """Build HSE director daily report"""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ========================================================================
    # HEADER
    # ========================================================================

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logos_added = 0
    for logo_filename in LOGOS:
        logo_path = os.path.join(LOGOS_PATH, logo_filename)
        if os.path.exists(logo_path):
            try:
                run = p.add_run()
                run.add_picture(logo_path, width=Inches(1.0))
                logos_added += 1
            except Exception:
                pass

    if logos_added == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BRHAS Safety Companies")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLORS['primary']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAILY SAFETY REPORT")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLORS['primary']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HSE Management Summary")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = COLORS['secondary']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Report Date: {yesterday_date.strftime('%A, %B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLORS['accent']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS['secondary']

    doc.add_paragraph()

    # ========================================================================
    # SAFETY STREAK METRICS
    # ========================================================================

    add_heading(doc, "SAFETY STREAK METRICS", 1, COLORS['primary'])

    p = doc.add_paragraph()
    p.add_run("Days Since Lost-Time Injury: ").font.bold = True
    p.add_run("127 days ✅")

    p = doc.add_paragraph()
    p.add_run("Days Since Recordable Incident: ").font.bold = True
    p.add_run("89 days ✅")

    if 'incident_reports' in all_data and all_data['incident_reports']:
        inc_data = all_data['incident_reports']
        real_incidents = [inc for inc in inc_data['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            p = doc.add_paragraph()
            p.add_run("Days Since Any Incident: ").font.bold = True
            run = p.add_run("0 days (New incident reported)")
            run.font.color.rgb = COLORS['critical']

    p = doc.add_paragraph()
    p.add_run("Days Since Near-Miss Report: ").font.bold = True

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']
        near_miss = obs_analysis['type_counts'].get('Near Miss', 0)
        if near_miss > 0:
            run = p.add_run("0 days (Early warning system active) ✅")
            run.font.color.rgb = COLORS['safe']
        else:
            p.add_run("N/A")

    doc.add_paragraph()

    # ========================================================================
    # EXECUTIVE SUMMARY
    # ========================================================================

    add_heading(doc, "EXECUTIVE SUMMARY", 1)

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']

        p = doc.add_paragraph()
        p.add_run(f"Total Observations: ").font.bold = True
        p.add_run(f"{obs_analysis['total']}")

        near_miss_count = obs_analysis['type_counts'].get('Near Miss', 0)
        at_risk_behavior_count = obs_analysis['type_counts'].get('At-Risk Behavior', 0)
        at_risk_condition_count = obs_analysis['type_counts'].get('At-Risk Condition', 0)
        at_risk_procedure_count = obs_analysis['type_counts'].get('At-Risk Procedure', 0)
        recognition_count = obs_analysis['type_counts'].get('Recognition', 0)

        p = doc.add_paragraph()
        p.add_run("Summary: ").font.bold = True

        if near_miss_count > 0:
            run = doc.add_paragraph(f"🔴 NEAR MISSES: {near_miss_count}", style='List Bullet').runs[0]
            run.font.color.rgb = COLORS['critical']

        if at_risk_behavior_count > 0:
            run = doc.add_paragraph(f"🔴 AT-RISK BEHAVIOR: {at_risk_behavior_count}", style='List Bullet').runs[0]
            run.font.color.rgb = COLORS['critical']

        if at_risk_condition_count > 0:
            doc.add_paragraph(f"🟡 AT-RISK CONDITIONS: {at_risk_condition_count}", style='List Bullet')

        if at_risk_procedure_count > 0:
            doc.add_paragraph(f"🟡 AT-RISK PROCEDURES: {at_risk_procedure_count}", style='List Bullet')

        if recognition_count > 0:
            run = doc.add_paragraph(f"✅ SAFETY RECOGNITION: {recognition_count}", style='List Bullet').runs[0]
            run.font.color.rgb = COLORS['safe']
    else:
        p = doc.add_paragraph()
        p.add_run(f"Total Observations: ").font.bold = True
        p.add_run("0 - Safe day!")

    if 'incident_reports' in all_data and all_data['incident_reports']:
        inc_data = all_data['incident_reports']
        real_incidents = [inc for inc in inc_data['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            run = doc.add_paragraph(f"⚠️ INCIDENT REPORTS: {len(real_incidents)}", style='List Bullet').runs[0]
            run.font.color.rgb = COLORS['critical']

    doc.add_paragraph()

    # ========================================================================
    # ACTION ITEMS FOR TODAY
    # ========================================================================

    add_heading(doc, "ACTION ITEMS FOR TODAY", 1, COLORS['critical'])

    action_count = 0

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']

        near_misses = obs_analysis['by_type'].get('Near Miss', [])
        at_risk_behavior = obs_analysis['by_type'].get('At-Risk Behavior', [])

        if near_misses:
            action_count += len(near_misses)
            p = doc.add_paragraph()
            p.add_run(f"1. NEAR MISSES - Contact {len(near_misses)} for incident investigation").font.bold = True
            for nm in near_misses:
                actual_name = get_actual_observer_name(nm)
                doc.add_paragraph(
                    f"• Report #{nm.get('report number')} - {actual_name} - {nm.get('date')}",
                    style='List Bullet 2'
                )

        if at_risk_behavior:
            action_count += len(at_risk_behavior)
            p = doc.add_paragraph()
            p.add_run(f"2. AT-RISK BEHAVIORS - Schedule coaching for {len(at_risk_behavior)}").font.bold = True
            for arb in at_risk_behavior:
                actual_name = get_actual_observer_name(arb)
                doc.add_paragraph(
                    f"• Report #{arb.get('report number')} - {actual_name} - {arb.get('date')}",
                    style='List Bullet 2'
                )

    if 'incident_reports' in all_data and all_data['incident_reports']:
        inc_data = all_data['incident_reports']
        real_incidents = [inc for inc in inc_data['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            action_count += 1
            p = doc.add_paragraph()
            p.add_run(f"3. INCIDENT - Review and assess").font.bold = True
            for inc in real_incidents:
                doc.add_paragraph(
                    f"• {inc.get('nojcquy0tfl9hqih', 'Incident')} - {inc.get('date')}",
                    style='List Bullet 2'
                )

    if action_count == 0:
        p = doc.add_paragraph("✅ No immediate action items - Safe day!")
        p.runs[0].font.color.rgb = COLORS['safe']
        p.runs[0].font.bold = True

    doc.add_paragraph()

    # ========================================================================
    # CRITICAL ITEMS (Incidents, RCA, Near Misses) - ONLY IF THEY EXIST
    # ========================================================================

    # INCIDENT REPORTS
    if 'incident_reports' in all_data and all_data['incident_reports']:
        inc_data = all_data['incident_reports']
        real_incidents = [inc for inc in inc_data['rows'] if inc.get('report number') != 'Report Number']

        if real_incidents:
            doc.add_page_break()
            add_heading(doc, f"INCIDENT REPORTS ({len(real_incidents)}) - CRITICAL", 1, COLORS['critical'])
            doc.add_paragraph()

            for i, inc in enumerate(real_incidents, 1):
                add_heading(doc, f"Incident #{i}: Report #{inc.get('report number')}", 2, COLORS['critical'])

                p = doc.add_paragraph()
                p.add_run("Date: ").font.bold = True
                p.add_run(inc.get('date', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Type: ").font.bold = True
                p.add_run(inc.get('nojcquy0tfl9hqih', inc.get('report', 'N/A')))

                p = doc.add_paragraph()
                p.add_run("Location: ").font.bold = True
                p.add_run(inc.get('pk6qj0kiu9vek20v', 'N/A'))

                desc = inc.get('313e9txgrof0uute', '')
                if desc:
                    p = doc.add_paragraph()
                    p.add_run("Description:\n").font.bold = True
                    p.add_run(desc)

                link = inc.get('link', '')
                if link and link != 'Link':
                    p = doc.add_paragraph()
                    p.add_run("Link: ").font.bold = True
                    p.add_run(link)

                doc.add_paragraph()

    # ROOT CAUSE ANALYSIS
    if 'rca' in all_data and all_data['rca']:
        rca_data = all_data['rca']
        real_rca = [r for r in rca_data['rows'] if r.get('report number') != 'Report Number']

        if real_rca:
            doc.add_page_break()
            add_heading(doc, f"ROOT CAUSE ANALYSIS ({len(real_rca)})", 1, COLORS['critical'])
            doc.add_paragraph()

            for i, rca in enumerate(real_rca, 1):
                add_heading(doc, f"RCA #{i}: Report #{rca.get('report number')}", 2, COLORS['critical'])

                p = doc.add_paragraph()
                p.add_run("Date: ").font.bold = True
                p.add_run(rca.get('date', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Description: ").font.bold = True
                p.add_run(rca.get('description', 'N/A'))

                link = rca.get('link', '')
                if link and link != 'Link':
                    p = doc.add_paragraph()
                    p.add_run("Link: ").font.bold = True
                    p.add_run(link)

                doc.add_paragraph()

    # NEAR MISSES
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']
        near_misses = obs_analysis['by_type'].get('Near Miss', [])

        if near_misses:
            doc.add_page_break()
            add_heading(doc, f"NEAR MISSES ({len(near_misses)}) - IMMEDIATE ACTION REQUIRED", 1, COLORS['critical'])
            doc.add_paragraph()

            for i, nm in enumerate(near_misses, 1):
                actual_name = get_actual_observer_name(nm)
                add_heading(doc, f"{i}. Report #{nm.get('report number')} - {actual_name}", 3, COLORS['critical'])

                p = doc.add_paragraph()
                p.add_run("Date: ").font.bold = True
                p.add_run(nm.get('date', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Yard: ").font.bold = True
                p.add_run(nm.get('7vj2l992y7fwqhwz', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Location: ").font.bold = True
                p.add_run(nm.get('lg5pnj4chjadnv46', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Description: ").font.bold = True
                p.add_run(nm.get('uncbcge9x8vow9pn', 'No description'))

                corrective = nm.get('dpy2klalngsr7ek9', '')
                if corrective and corrective.strip():
                    p = doc.add_paragraph()
                    p.add_run("Status: ").font.bold = True
                    p.add_run("CLOSED")
                else:
                    p = doc.add_paragraph()
                    p.add_run("Status: ").font.bold = True
                    run = p.add_run("OPEN - ACTION REQUIRED")
                    run.font.color.rgb = COLORS['critical']

                link = nm.get('link', '')
                if link and link != 'Link':
                    p = doc.add_paragraph()
                    p.add_run("Link: ").font.bold = True
                    p.add_run(link)

                doc.add_paragraph()

    # ========================================================================
    # OPEN ITEMS TRACKING (At-Risk Conditions & Procedures ONLY)
    # ========================================================================

    add_heading(doc, "OPEN ITEMS TRACKING - CORRECTIVE ACTIONS NEEDED", 1, COLORS['warning'])

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']

        # Only At-Risk Conditions and Procedures (NOT Near Misses - they have their own section)
        pending_items = []
        for obs_type, obs_list in obs_analysis['by_type'].items():
            if obs_type in ['At-Risk Condition', 'At-Risk Procedure']:
                for obs in obs_list:
                    corrective = obs.get('dpy2klalngsr7ek9', '')
                    if not corrective or not corrective.strip():
                        pending_items.append({
                            'type': obs_type,
                            'report_num': obs.get('report number'),
                            'person': get_actual_observer_name(obs),
                            'date': obs.get('date'),
                            'yard': obs.get('7vj2l992y7fwqhwz', 'Unknown'),
                            'location': obs.get('lg5pnj4chjadnv46', 'Unknown'),
                            'description': obs.get('uncbcge9x8vow9pn', 'No description')[:80],
                            'link': obs.get('link', '')
                        })

        if pending_items:
            p = doc.add_paragraph()
            p.add_run(f"Pending Corrective Actions: {len(pending_items)} items").font.bold = True
            doc.add_paragraph()

            for item in pending_items:
                p = doc.add_paragraph()
                run = p.add_run(f"Report #{item['report_num']} - {item['type']}")
                run.font.bold = True
                run.font.color.rgb = COLORS['critical']

                doc.add_paragraph(f"Person: {item['person']}", style='List Bullet')
                doc.add_paragraph(f"Date: {item['date']}", style='List Bullet')
                doc.add_paragraph(f"Yard: {item['yard']}", style='List Bullet')
                doc.add_paragraph(f"Location: {item['location']}", style='List Bullet')
                doc.add_paragraph(f"Issue: {item['description']}", style='List Bullet')
                doc.add_paragraph(f"Assigned To: TBD | Deadline: TBD", style='List Bullet')

                if item['link']:
                    doc.add_paragraph(f"Link: {item['link']}", style='List Bullet')

                doc.add_paragraph()
        else:
            p = doc.add_paragraph("✅ All corrective actions completed!")
            p.runs[0].font.color.rgb = COLORS['safe']

    doc.add_paragraph()

    # ========================================================================
    # DATA QUALITY ALERT
    # ========================================================================

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']
        miscategorized = obs_analysis.get('miscategorized', [])

        if miscategorized:
            add_heading(doc, f"⚠️ DATA QUALITY ALERT - {len(miscategorized)} MISCATEGORIZED", 1, COLORS['warning'])
            doc.add_paragraph("These observations were filed as the wrong type:")
            doc.add_paragraph()

            for item in miscategorized:
                p = doc.add_paragraph()
                run = p.add_run(f"Report #{item['report_num']}")
                run.font.bold = True

                doc.add_paragraph(f"Current Type: {item['type']}", style='List Bullet')
                doc.add_paragraph(f"Should Be: {item['actual_type']}", style='List Bullet')
                doc.add_paragraph(f"Text: '{item['description']}'", style='List Bullet')
                doc.add_paragraph(f"Person: {item['observer']}", style='List Bullet')
                doc.add_paragraph(f"Action: Reclassify in KPA", style='List Bullet')

                doc.add_paragraph()

            doc.add_paragraph()

    # ========================================================================
    # HOTSPOT ANALYSIS - Uses ACTUAL observer name (Name field), not system observer
    # ========================================================================

    add_heading(doc, "HOTSPOT ANALYSIS", 1)

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']

        # CRITICAL: Use get_actual_observer_name() for ACTUAL person observed
        # NOT the system observer field (which includes James Barnett, Shelly Batts, etc. who are just data entry)
        names = []
        for obs_list in obs_analysis['by_type'].values():
            for obs in obs_list:
                actual_name = get_actual_observer_name(obs)
                if actual_name and actual_name != 'Unknown':
                    names.append(actual_name)

        name_counts = Counter(names)

        if name_counts:
            p = doc.add_paragraph()
            p.add_run("Most Active Observers (based on actual Name field):").font.bold = True
            for name, count in name_counts.most_common(5):
                if name and name != 'Unknown':
                    doc.add_paragraph(f"{name}: {count} observations ⭐", style='List Bullet')

    doc.add_paragraph()

    # ========================================================================
    # INCIDENT TIMING
    # ========================================================================

    add_heading(doc, "INCIDENT TIMING ANALYSIS", 1)

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']

        shift_counts = {'Day Shift (8 AM-4 PM)': 0, 'Night Shift (4 PM-Midnight)': 0, 'Overnight (0-8 AM)': 0}

        for obs_list in obs_analysis['by_type'].values():
            for obs in obs_list:
                shift = get_shift(obs.get('date', ''))
                if shift in shift_counts:
                    shift_counts[shift] += 1

        for shift, count in shift_counts.items():
            if count > 0:
                doc.add_paragraph(f"{shift}: {count} observations", style='List Bullet')

    doc.add_paragraph()

    # ========================================================================
    # ASSESSMENT & AUDIT ANALYSIS (after Timing, before At-Risk Conditions)
    # ========================================================================

    if 'assessment_analysis' in all_data and all_data['assessment_analysis']:
        try:
            add_assessment_analysis_section(doc, all_data['assessment_analysis'])
        except Exception as e:
            print(f"Warning: Assessment analysis section error: {e}")
            # Continue building report even if this section fails

    # ========================================================================
    # AT-RISK CONDITIONS
    # ========================================================================

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']
        conditions = obs_analysis['by_type'].get('At-Risk Condition', [])

        if conditions:
            doc.add_page_break()
            display_count = min(10, len(conditions))
            add_heading(doc, f"AT-RISK CONDITIONS (Top {display_count} of {len(conditions)})", 1, COLORS['warning'])
            doc.add_paragraph()

            for i, cond in enumerate(conditions[:10], 1):
                actual_name = get_actual_observer_name(cond)
                add_heading(doc, f"{i}. Report #{cond.get('report number')} - {actual_name}", 3)

                p = doc.add_paragraph()
                p.add_run("Date: ").font.bold = True
                p.add_run(cond.get('date', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Location: ").font.bold = True
                p.add_run(cond.get('lg5pnj4chjadnv46', 'N/A'))

                p = doc.add_paragraph()
                p.add_run("Condition: ").font.bold = True
                p.add_run(cond.get('uncbcge9x8vow9pn', 'No description'))

                corrective = cond.get('dpy2klalngsr7ek9', '')
                if corrective and corrective.strip():
                    p = doc.add_paragraph()
                    p.add_run("Status: ").font.bold = True
                    run = p.add_run("CORRECTED")
                    run.font.color.rgb = COLORS['safe']
                else:
                    p = doc.add_paragraph()
                    p.add_run("Status: ").font.bold = True
                    run = p.add_run("PENDING ACTION")
                    run.font.color.rgb = COLORS['warning']

                link = cond.get('link', '')
                if link and link != 'Link':
                    p = doc.add_paragraph()
                    p.add_run("Link: ").font.bold = True
                    p.add_run(link)

                doc.add_paragraph()

            if len(conditions) > 10:
                p = doc.add_paragraph()
                run = p.add_run(f"... and {len(conditions) - 10} more conditions in KPA")
                run.font.italic = True

    # ========================================================================
    # RECOGNITION
    # ========================================================================

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs_analysis = all_data['observation_analysis']
        recognition = obs_analysis['by_type'].get('Recognition', [])

        if recognition:
            doc.add_page_break()
            add_heading(doc, f"SAFETY RECOGNITION - STARS ({len(recognition)})", 1, COLORS['safe'])
            doc.add_paragraph()

            recognition_names = []
            for rec in recognition:
                recognition_names.append({
                    'name': get_actual_observer_name(rec),
                    'description': rec.get('uncbcge9x8vow9pn'),
                })

            name_counter = Counter([r['name'] for r in recognition_names])

            for name, count in name_counter.most_common(10):
                if name and name != 'Unknown':
                    p = doc.add_paragraph()
                    run = p.add_run(f"✅ {name}")
                    run.font.bold = True
                    p.add_run(f" - {count} recognition(s)")

                    for rec in recognition_names:
                        if rec['name'] == name:
                            doc.add_paragraph(f"'{rec['description']}'", style='List Bullet')
                            break

    # ========================================================================
    # ASSESSMENT & AUDIT SUMMARY (detailed table replacing old "Other Forms")
    # ========================================================================

    if 'assessment_details' in all_data:
        try:
            add_assessment_audit_summary(doc, all_data['assessment_details'])
        except Exception as e:
            print(f"Warning: Assessment audit summary table error: {e}")
            # Fallback to simple count list
            doc.add_page_break()
            add_heading(doc, "OTHER SAFETY FORMS SUMMARY", 1)
            doc.add_paragraph()
            for form_id, form_name in OTHER_FORMS:
                data = all_data.get(f"form_{form_id}")
                count = data['count'] if data else 0
                p = doc.add_paragraph()
                run = p.add_run(f"{form_name}: ")
                run.font.bold = True
                p.add_run(f"{count}")
    else:
        # Fallback if assessment_details not generated
        doc.add_page_break()
        add_heading(doc, "OTHER SAFETY FORMS SUMMARY", 1)
        doc.add_paragraph()
        for form_id, form_name in OTHER_FORMS:
            data = all_data.get(f"form_{form_id}")
            count = data['count'] if data else 0
            p = doc.add_paragraph()
            run = p.add_run(f"{form_name}: ")
            run.font.bold = True
            p.add_run(f"{count}")

    doc.add_paragraph()

    # ========================================================================
    # FOOTER
    # ========================================================================

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END OF REPORT")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = COLORS['primary']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Butch's Rat Hole & Anchor Service Inc. | HSE Department")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS['secondary']

    return doc


# ==============================================================================
# BUILD HTML EMAIL BODY
# ==============================================================================

HTML_COLORS = {
    'primary': '#C00000',
    'secondary': '#800000',
    'accent': '#000000',
    'critical': '#C00000',
    'warning': '#C08000',
    'safe': '#008000',
}


def _h(text):
    """HTML-escape text safely"""
    return html_escape(str(text)) if text else ''


def build_html_report(all_data, yesterday_date):
    """Build HTML version of the report for email body"""
    sections = []

    # --- Wrapper start ---
    sections.append(f"""<html><head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#f4f4f4;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f4f4;">
<tr><td align="center">
<table width="700" cellpadding="0" cellspacing="0" style="background:#ffffff;border:1px solid #ddd;margin:20px auto;font-family:Arial,Helvetica,sans-serif;font-size:14px;color:#333;">""")

    # --- HEADER ---
    sections.append(f"""
<tr><td style="background:{HTML_COLORS['primary']};padding:30px 40px;text-align:center;">
  <div style="font-size:16px;font-weight:bold;color:#ffffff;letter-spacing:1px;">BRHAS Safety Companies</div>
  <div style="font-size:28px;font-weight:bold;color:#ffffff;margin:10px 0;">DAILY SAFETY REPORT</div>
  <div style="font-size:13px;font-style:italic;color:#ffcccc;">HSE Management Summary</div>
  <div style="font-size:12px;color:#ffffff;margin-top:8px;">Report Date: {yesterday_date.strftime('%A, %B %d, %Y')}</div>
  <div style="font-size:10px;color:#ffcccc;margin-top:4px;">Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}</div>
</td></tr>""")

    # --- SAFETY STREAK METRICS ---
    streak_rows = []
    streak_rows.append('<b>Days Since Lost-Time Injury:</b> 127 days &#9989;')
    streak_rows.append('<b>Days Since Recordable Incident:</b> 89 days &#9989;')

    if 'incident_reports' in all_data and all_data['incident_reports']:
        real_incidents = [inc for inc in all_data['incident_reports']['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            streak_rows.append(f'<b>Days Since Any Incident:</b> <span style="color:{HTML_COLORS["critical"]};">0 days (New incident reported)</span>')

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        near_miss = all_data['observation_analysis']['type_counts'].get('Near Miss', 0)
        if near_miss > 0:
            streak_rows.append(f'<b>Days Since Near-Miss Report:</b> <span style="color:{HTML_COLORS["safe"]};">0 days (Early warning system active) &#9989;</span>')
        else:
            streak_rows.append('<b>Days Since Near-Miss Report:</b> N/A')

    sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">SAFETY STREAK METRICS</h2>
  {'<br>'.join(streak_rows)}
</td></tr>""")

    # --- EXECUTIVE SUMMARY ---
    summary_html = ''
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs = all_data['observation_analysis']
        summary_html += f'<b>Total Observations:</b> {obs["total"]}<br><br>'

        near_miss_count = obs['type_counts'].get('Near Miss', 0)
        at_risk_behavior_count = obs['type_counts'].get('At-Risk Behavior', 0)
        at_risk_condition_count = obs['type_counts'].get('At-Risk Condition', 0)
        at_risk_procedure_count = obs['type_counts'].get('At-Risk Procedure', 0)
        recognition_count = obs['type_counts'].get('Recognition', 0)

        if near_miss_count > 0:
            summary_html += f'<div style="color:{HTML_COLORS["critical"]};margin:4px 0 4px 20px;">&#128308; NEAR MISSES: {near_miss_count}</div>'
        if at_risk_behavior_count > 0:
            summary_html += f'<div style="color:{HTML_COLORS["critical"]};margin:4px 0 4px 20px;">&#128308; AT-RISK BEHAVIOR: {at_risk_behavior_count}</div>'
        if at_risk_condition_count > 0:
            summary_html += f'<div style="color:{HTML_COLORS["warning"]};margin:4px 0 4px 20px;">&#128992; AT-RISK CONDITIONS: {at_risk_condition_count}</div>'
        if at_risk_procedure_count > 0:
            summary_html += f'<div style="color:{HTML_COLORS["warning"]};margin:4px 0 4px 20px;">&#128992; AT-RISK PROCEDURES: {at_risk_procedure_count}</div>'
        if recognition_count > 0:
            summary_html += f'<div style="color:{HTML_COLORS["safe"]};margin:4px 0 4px 20px;">&#9989; SAFETY RECOGNITION: {recognition_count}</div>'
    else:
        summary_html += '<b>Total Observations:</b> 0 - Safe day!'

    if 'incident_reports' in all_data and all_data['incident_reports']:
        real_incidents = [inc for inc in all_data['incident_reports']['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            summary_html += f'<div style="color:{HTML_COLORS["critical"]};margin:4px 0 4px 20px;">&#9888;&#65039; INCIDENT REPORTS: {len(real_incidents)}</div>'

    sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">EXECUTIVE SUMMARY</h2>
  {summary_html}
</td></tr>""")

    # --- ACTION ITEMS ---
    action_html = ''
    action_count = 0

    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs = all_data['observation_analysis']
        near_misses = obs['by_type'].get('Near Miss', [])
        at_risk_behavior = obs['by_type'].get('At-Risk Behavior', [])

        if near_misses:
            action_count += len(near_misses)
            action_html += f'<b>1. NEAR MISSES - Contact {len(near_misses)} for incident investigation</b><ul style="margin:5px 0 15px 0;">'
            for nm in near_misses:
                action_html += f'<li>Report #{_h(nm.get("report number"))} - {_h(get_actual_observer_name(nm))} - {_h(nm.get("date"))}</li>'
            action_html += '</ul>'

        if at_risk_behavior:
            action_count += len(at_risk_behavior)
            action_html += f'<b>2. AT-RISK BEHAVIORS - Schedule coaching for {len(at_risk_behavior)}</b><ul style="margin:5px 0 15px 0;">'
            for arb in at_risk_behavior:
                action_html += f'<li>Report #{_h(arb.get("report number"))} - {_h(get_actual_observer_name(arb))} - {_h(arb.get("date"))}</li>'
            action_html += '</ul>'

    if 'incident_reports' in all_data and all_data['incident_reports']:
        real_incidents = [inc for inc in all_data['incident_reports']['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            action_count += 1
            action_html += '<b>3. INCIDENT - Review and assess</b><ul style="margin:5px 0 15px 0;">'
            for inc in real_incidents:
                action_html += f'<li>{_h(inc.get("nojcquy0tfl9hqih", "Incident"))} - {_h(inc.get("date"))}</li>'
            action_html += '</ul>'

    if action_count == 0:
        action_html = f'<b style="color:{HTML_COLORS["safe"]};">&#9989; No immediate action items - Safe day!</b>'

    sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['critical']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['critical']};padding-bottom:5px;">ACTION ITEMS FOR TODAY</h2>
  {action_html}
</td></tr>""")

    # --- INCIDENT REPORTS (only if they exist) ---
    if 'incident_reports' in all_data and all_data['incident_reports']:
        real_incidents = [inc for inc in all_data['incident_reports']['rows'] if inc.get('report number') != 'Report Number']
        if real_incidents:
            inc_html = ''
            for i, inc in enumerate(real_incidents, 1):
                inc_html += f'<div style="background:#fff5f5;border-left:4px solid {HTML_COLORS["critical"]};padding:12px 15px;margin:10px 0;">'
                inc_html += f'<b style="color:{HTML_COLORS["critical"]};font-size:15px;">Incident #{i}: Report #{_h(inc.get("report number"))}</b><br>'
                inc_html += f'<b>Date:</b> {_h(inc.get("date", "N/A"))}<br>'
                inc_html += f'<b>Type:</b> {_h(inc.get("nojcquy0tfl9hqih", inc.get("report", "N/A")))}<br>'
                inc_html += f'<b>Location:</b> {_h(inc.get("pk6qj0kiu9vek20v", "N/A"))}<br>'
                desc = inc.get('313e9txgrof0uute', '')
                if desc:
                    inc_html += f'<b>Description:</b> {_h(desc)}<br>'
                link = inc.get('link', '')
                if link and link != 'Link':
                    inc_html += f'<b>Link:</b> <a href="{_h(link)}">{_h(link)}</a><br>'
                inc_html += '</div>'

            sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['critical']};">
  <h2 style="color:{HTML_COLORS['critical']};margin:0 0 15px 0;font-size:18px;">INCIDENT REPORTS ({len(real_incidents)}) - CRITICAL</h2>
  {inc_html}
</td></tr>""")

    # --- ROOT CAUSE ANALYSIS (only if exists) ---
    if 'rca' in all_data and all_data['rca']:
        real_rca = [r for r in all_data['rca']['rows'] if r.get('report number') != 'Report Number']
        if real_rca:
            rca_html = ''
            for i, rca in enumerate(real_rca, 1):
                rca_html += f'<div style="background:#fff5f5;border-left:4px solid {HTML_COLORS["critical"]};padding:12px 15px;margin:10px 0;">'
                rca_html += f'<b style="color:{HTML_COLORS["critical"]};">RCA #{i}: Report #{_h(rca.get("report number"))}</b><br>'
                rca_html += f'<b>Date:</b> {_h(rca.get("date", "N/A"))}<br>'
                rca_html += f'<b>Description:</b> {_h(rca.get("description", "N/A"))}<br>'
                link = rca.get('link', '')
                if link and link != 'Link':
                    rca_html += f'<b>Link:</b> <a href="{_h(link)}">{_h(link)}</a><br>'
                rca_html += '</div>'

            sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['critical']};">
  <h2 style="color:{HTML_COLORS['critical']};margin:0 0 15px 0;font-size:18px;">ROOT CAUSE ANALYSIS ({len(real_rca)})</h2>
  {rca_html}
</td></tr>""")

    # --- NEAR MISSES (only if exist) ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        near_misses = all_data['observation_analysis']['by_type'].get('Near Miss', [])
        if near_misses:
            nm_html = ''
            for i, nm in enumerate(near_misses, 1):
                actual_name = get_actual_observer_name(nm)
                corrective = nm.get('dpy2klalngsr7ek9', '')
                if corrective and corrective.strip():
                    status = '<span style="color:#008000;"><b>CLOSED</b></span>'
                else:
                    status = f'<span style="color:{HTML_COLORS["critical"]};"><b>OPEN - ACTION REQUIRED</b></span>'

                nm_html += f'<div style="background:#fff5f5;border-left:4px solid {HTML_COLORS["critical"]};padding:12px 15px;margin:10px 0;">'
                nm_html += f'<b style="color:{HTML_COLORS["critical"]};">{i}. Report #{_h(nm.get("report number"))} - {_h(actual_name)}</b><br>'
                nm_html += f'<b>Date:</b> {_h(nm.get("date", "N/A"))}<br>'
                nm_html += f'<b>Yard:</b> {_h(nm.get("7vj2l992y7fwqhwz", "N/A"))}<br>'
                nm_html += f'<b>Location:</b> {_h(nm.get("lg5pnj4chjadnv46", "N/A"))}<br>'
                nm_html += f'<b>Description:</b> {_h(nm.get("uncbcge9x8vow9pn", "No description"))}<br>'
                nm_html += f'<b>Status:</b> {status}<br>'
                link = nm.get('link', '')
                if link and link != 'Link':
                    nm_html += f'<b>Link:</b> <a href="{_h(link)}">{_h(link)}</a><br>'
                nm_html += '</div>'

            sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['critical']};">
  <h2 style="color:{HTML_COLORS['critical']};margin:0 0 15px 0;font-size:18px;">NEAR MISSES ({len(near_misses)}) - IMMEDIATE ACTION REQUIRED</h2>
  {nm_html}
</td></tr>""")

    # --- OPEN ITEMS TRACKING ---
    open_html = ''
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs = all_data['observation_analysis']
        pending_items = []
        for obs_type, obs_list in obs['by_type'].items():
            if obs_type in ['At-Risk Condition', 'At-Risk Procedure']:
                for o in obs_list:
                    corrective = o.get('dpy2klalngsr7ek9', '')
                    if not corrective or not corrective.strip():
                        pending_items.append({
                            'type': obs_type,
                            'report_num': o.get('report number'),
                            'person': get_actual_observer_name(o),
                            'date': o.get('date'),
                            'yard': o.get('7vj2l992y7fwqhwz', 'Unknown'),
                            'location': o.get('lg5pnj4chjadnv46', 'Unknown'),
                            'description': o.get('uncbcge9x8vow9pn', 'No description')[:80],
                            'link': o.get('link', '')
                        })

        if pending_items:
            open_html += f'<b>Pending Corrective Actions: {len(pending_items)} items</b><br><br>'
            for item in pending_items:
                open_html += f'<div style="background:#fffbf0;border-left:4px solid {HTML_COLORS["warning"]};padding:12px 15px;margin:10px 0;">'
                open_html += f'<b style="color:{HTML_COLORS["critical"]};">Report #{_h(item["report_num"])} - {_h(item["type"])}</b><br>'
                open_html += f'Person: {_h(item["person"])} | Date: {_h(item["date"])}<br>'
                open_html += f'Yard: {_h(item["yard"])} | Location: {_h(item["location"])}<br>'
                open_html += f'Issue: {_h(item["description"])}<br>'
                open_html += f'Assigned To: TBD | Deadline: TBD<br>'
                if item['link']:
                    open_html += f'<a href="{_h(item["link"])}">View in KPA</a><br>'
                open_html += '</div>'
        else:
            open_html = f'<b style="color:{HTML_COLORS["safe"]};">&#9989; All corrective actions completed!</b>'

    sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['warning']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['warning']};padding-bottom:5px;">OPEN ITEMS TRACKING - CORRECTIVE ACTIONS NEEDED</h2>
  {open_html}
</td></tr>""")

    # --- DATA QUALITY ALERT (only if exists) ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        miscategorized = all_data['observation_analysis'].get('miscategorized', [])
        if miscategorized:
            dq_html = '<p>These observations were filed as the wrong type:</p>'
            for item in miscategorized:
                dq_html += f'<div style="background:#fffbf0;border-left:4px solid {HTML_COLORS["warning"]};padding:12px 15px;margin:10px 0;">'
                dq_html += f'<b>Report #{_h(item["report_num"])}</b><br>'
                dq_html += f'Current Type: {_h(item["type"])} | Should Be: {_h(item["actual_type"])}<br>'
                dq_html += f'Text: \'{_h(item["description"])}\'<br>'
                dq_html += f'Person: {_h(item["observer"])} | Action: Reclassify in KPA<br>'
                dq_html += '</div>'

            sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['warning']};margin:0 0 15px 0;font-size:18px;">&#9888;&#65039; DATA QUALITY ALERT - {len(miscategorized)} MISCATEGORIZED</h2>
  {dq_html}
</td></tr>""")

    # --- HOTSPOT ANALYSIS ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs = all_data['observation_analysis']
        names = []
        for obs_list in obs['by_type'].values():
            for o in obs_list:
                actual_name = get_actual_observer_name(o)
                if actual_name and actual_name != 'Unknown':
                    names.append(actual_name)
        name_counts = Counter(names)

        if name_counts:
            hotspot_html = '<b>Most Active Observers:</b><ul style="margin:5px 0;">'
            for name, count in name_counts.most_common(5):
                if name and name != 'Unknown':
                    hotspot_html += f'<li>{_h(name)}: {count} observations &#11088;</li>'
            hotspot_html += '</ul>'

            sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">HOTSPOT ANALYSIS</h2>
  {hotspot_html}
</td></tr>""")

    # --- INCIDENT TIMING ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        obs = all_data['observation_analysis']
        shift_counts = {'Day Shift (8 AM-4 PM)': 0, 'Night Shift (4 PM-Midnight)': 0, 'Overnight (0-8 AM)': 0}
        for obs_list in obs['by_type'].values():
            for o in obs_list:
                shift = get_shift(o.get('date', ''))
                if shift in shift_counts:
                    shift_counts[shift] += 1

        active_shifts = {k: v for k, v in shift_counts.items() if v > 0}
        if active_shifts:
            timing_html = '<ul style="margin:5px 0;">'
            for shift, count in active_shifts.items():
                timing_html += f'<li>{_h(shift)}: {count} observations</li>'
            timing_html += '</ul>'

            sections.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">INCIDENT TIMING ANALYSIS</h2>
  {timing_html}
</td></tr>""")

    # --- ASSESSMENT & AUDIT ANALYSIS ---
    if 'assessment_analysis' in all_data and all_data['assessment_analysis']:
        try:
            aa = all_data['assessment_analysis']
            if aa.get('has_data'):
                aa_html = ''

                # Header stats
                aa_html += f'<b>Total Assessments:</b> {aa["total_assessments"]} | '
                aa_html += f'<b>Total Findings:</b> {aa["total_findings"]}<br><br>'

                # Activity Summary Table
                if aa['activity_summary']:
                    aa_html += f'<h3 style="color:{HTML_COLORS["secondary"]};margin:10px 0 8px 0;font-size:15px;">Assessment Activity Summary</h3>'
                    aa_html += '<table width="100%" cellpadding="6" cellspacing="0" style="border-collapse:collapse;font-size:13px;margin-bottom:15px;">'
                    aa_html += f'<tr style="background:{HTML_COLORS["secondary"]};color:#ffffff;">'
                    aa_html += '<th style="text-align:left;padding:8px;">Form</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Count</th>'
                    aa_html += '<th style="text-align:left;padding:8px;">Assessor(s)</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Findings</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Compliance</th></tr>'

                    for i, s in enumerate(aa['activity_summary']):
                        bg = '#f9f9f9' if i % 2 == 0 else '#ffffff'
                        assessor_text = _h(', '.join(s['assessors'][:3]))
                        if len(s['assessors']) > 3:
                            assessor_text += f' +{len(s["assessors"]) - 3}'

                        rate = s['compliance_rate']
                        if rate >= 90:
                            comp_text = f'<span style="color:{HTML_COLORS["safe"]};">&#9989; {rate:.0f}%</span>'
                        elif rate >= 70:
                            comp_text = f'<span style="color:{HTML_COLORS["warning"]};">&#128993; {rate:.0f}%</span>'
                        else:
                            comp_text = f'<span style="color:{HTML_COLORS["critical"]};">&#128308; {rate:.0f}%</span>'

                        aa_html += f'<tr style="background:{bg};">'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;">{_h(s["form_name"])}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{s["count"]}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;">{assessor_text}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{s["findings_count"]}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{comp_text}</td></tr>'

                    aa_html += '</table>'

                # Compliance by Yard Table
                if aa['compliance_by_yard']:
                    aa_html += f'<h3 style="color:{HTML_COLORS["secondary"]};margin:15px 0 8px 0;font-size:15px;">Compliance by Yard</h3>'
                    aa_html += '<table width="100%" cellpadding="6" cellspacing="0" style="border-collapse:collapse;font-size:13px;margin-bottom:15px;">'
                    aa_html += f'<tr style="background:{HTML_COLORS["secondary"]};color:#ffffff;">'
                    aa_html += '<th style="text-align:left;padding:8px;">Yard</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Total</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Compliant</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Non-Compliant</th>'
                    aa_html += '<th style="text-align:center;padding:8px;">Status</th></tr>'

                    sorted_yards = sorted(aa['compliance_by_yard'].items(),
                                          key=lambda x: x[1]['non_compliant'], reverse=True)
                    for i, (yard, info) in enumerate(sorted_yards):
                        bg = '#f9f9f9' if i % 2 == 0 else '#ffffff'
                        if info['total'] > 0:
                            rate = info['compliant'] / info['total'] * 100
                            if rate >= 90:
                                status = f'<span style="color:{HTML_COLORS["safe"]};">&#9989; {rate:.0f}%</span>'
                            elif rate >= 70:
                                status = f'<span style="color:{HTML_COLORS["warning"]};">&#128993; {rate:.0f}%</span>'
                            else:
                                status = f'<span style="color:{HTML_COLORS["critical"]};">&#128308; {rate:.0f}%</span>'
                        else:
                            status = 'N/A'

                        aa_html += f'<tr style="background:{bg};">'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;">{_h(yard)}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{info["total"]}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{info["compliant"]}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{info["non_compliant"]}</td>'
                        aa_html += f'<td style="border-bottom:1px solid #eee;padding:6px;text-align:center;">{status}</td></tr>'

                    aa_html += '</table>'

                # Critical Findings
                critical = aa['findings_by_severity']['critical']
                high = aa['findings_by_severity']['high']
                if critical or high:
                    aa_html += f'<h3 style="color:{HTML_COLORS["critical"]};margin:15px 0 8px 0;font-size:15px;">Critical Findings - Immediate Attention</h3>'

                    for f in critical:
                        aa_html += f'<div style="background:#fff5f5;border-left:4px solid {HTML_COLORS["critical"]};padding:12px 15px;margin:8px 0;">'
                        aa_html += f'<b style="color:{HTML_COLORS["critical"]};">&#128308; CRITICAL:</b> {_h(f["description"])}<br>'
                        aa_html += f'Form: {_h(f["form_name"])} | Assessor: {_h(f["assessor"])} | Yard: {_h(f["yard"])}<br>'
                        if f['link']:
                            aa_html += f'<a href="{_h(f["link"])}">View in KPA</a>'
                        aa_html += '</div>'

                    for f in high[:5]:
                        aa_html += f'<div style="background:#fffbf0;border-left:4px solid {HTML_COLORS["warning"]};padding:12px 15px;margin:8px 0;">'
                        aa_html += f'<b style="color:{HTML_COLORS["warning"]};">&#128993; HIGH:</b> {_h(f["description"])}<br>'
                        aa_html += f'Form: {_h(f["form_name"])} | Yard: {_h(f["yard"])}<br>'
                        if f['link']:
                            aa_html += f'<a href="{_h(f["link"])}">View in KPA</a>'
                        aa_html += '</div>'

                    if len(high) > 5:
                        aa_html += f'<p style="font-style:italic;">... and {len(high) - 5} more high-severity findings</p>'
                else:
                    medium = aa['findings_by_severity']['medium']
                    low = aa['findings_by_severity']['low']
                    if medium or low:
                        aa_html += f'<p><b>No critical or high-severity findings.</b> {len(medium)} medium, {len(low)} low-severity items noted.</p>'
                    else:
                        aa_html += f'<p style="color:{HTML_COLORS["safe"]};"><b>&#9989; No findings - All assessments passed!</b></p>'

                # Top Assessors
                if aa['assessor_stats']:
                    aa_html += f'<h3 style="color:{HTML_COLORS["safe"]};margin:15px 0 8px 0;font-size:15px;">Top Performing Assessors</h3>'
                    sorted_a = sorted(aa['assessor_stats'].items(), key=lambda x: x[1]['total'], reverse=True)
                    rank = 0
                    for name, stats in sorted_a[:10]:
                        if name == 'Unknown':
                            continue
                        rank += 1
                        star = '&#11088; ' if rank <= 3 else ''
                        divs = ', '.join(stats['divisions']) if stats['divisions'] else 'N/A'
                        finding_note = f' | {stats["findings_found"]} finding(s)' if stats['findings_found'] > 0 else ''
                        aa_html += f'<div style="margin:4px 0 4px 15px;">{star}<b>{_h(name)}</b> - {stats["total"]} assessment(s) | {_h(divs)}{finding_note}</div>'

                # Corrective Actions
                if aa['corrective_actions']:
                    aa_html += f'<h3 style="color:{HTML_COLORS["warning"]};margin:15px 0 8px 0;font-size:15px;">Corrective Actions ({len(aa["corrective_actions"])} open)</h3>'
                    for i, ca in enumerate(aa['corrective_actions'][:5], 1):
                        aa_html += f'<div style="background:#fffbf0;border-left:4px solid {HTML_COLORS["warning"]};padding:10px 15px;margin:6px 0;">'
                        aa_html += f'<b>{i}. {_h(ca["description"])}</b><br>'
                        aa_html += f'{_h(ca["form_name"])} | {_h(ca["yard"])} | By: {_h(ca["assessor"])}<br>'
                        if ca['link']:
                            aa_html += f'<a href="{_h(ca["link"])}">View in KPA</a>'
                        aa_html += '</div>'
                    if len(aa['corrective_actions']) > 5:
                        aa_html += f'<p style="font-style:italic;">... and {len(aa["corrective_actions"]) - 5} more</p>'

                # Trends
                if aa['trends']:
                    aa_html += f'<h3 style="color:{HTML_COLORS["primary"]};margin:15px 0 8px 0;font-size:15px;">Trends &amp; Patterns</h3>'
                    aa_html += '<ul style="margin:5px 0;">'
                    for trend in aa['trends']:
                        aa_html += f'<li>&#128202; {_h(trend)}</li>'
                    aa_html += '</ul>'

                # Recommendations
                recs = aa['recommendations']
                if any([recs['immediate'], recs['this_week'], recs['monthly']]):
                    aa_html += f'<h3 style="color:{HTML_COLORS["primary"]};margin:15px 0 8px 0;font-size:15px;">Recommended Actions for Leadership</h3>'

                    if recs['immediate']:
                        aa_html += f'<div style="margin:5px 0;"><b style="color:{HTML_COLORS["critical"]};">&#128308; IMMEDIATE:</b></div>'
                        aa_html += '<ul style="margin:3px 0;">'
                        for r in recs['immediate']:
                            aa_html += f'<li>{_h(r)}</li>'
                        aa_html += '</ul>'

                    if recs['this_week']:
                        aa_html += f'<div style="margin:5px 0;"><b style="color:{HTML_COLORS["warning"]};">&#128993; THIS WEEK:</b></div>'
                        aa_html += '<ul style="margin:3px 0;">'
                        for r in recs['this_week']:
                            aa_html += f'<li>{_h(r)}</li>'
                        aa_html += '</ul>'

                    if recs['monthly']:
                        aa_html += '<div style="margin:5px 0;"><b>&#128202; MONTH-OVER-MONTH:</b></div>'
                        aa_html += '<ul style="margin:3px 0;">'
                        for r in recs['monthly']:
                            aa_html += f'<li>{_h(r)}</li>'
                        aa_html += '</ul>'

                sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['primary']};">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">ASSESSMENT &amp; AUDIT ANALYSIS</h2>
  {aa_html}
</td></tr>""")
        except Exception as e:
            print(f"Warning: HTML assessment analysis error: {e}")

    # --- AT-RISK CONDITIONS (top 10) ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        conditions = all_data['observation_analysis']['by_type'].get('At-Risk Condition', [])
        if conditions:
            display_count = min(10, len(conditions))
            cond_html = ''
            for i, cond in enumerate(conditions[:10], 1):
                actual_name = get_actual_observer_name(cond)
                corrective = cond.get('dpy2klalngsr7ek9', '')
                if corrective and corrective.strip():
                    status = f'<span style="color:{HTML_COLORS["safe"]};"><b>CORRECTED</b></span>'
                else:
                    status = f'<span style="color:{HTML_COLORS["warning"]};"><b>PENDING ACTION</b></span>'

                cond_html += f'<div style="background:#fffbf0;border-left:4px solid {HTML_COLORS["warning"]};padding:12px 15px;margin:10px 0;">'
                cond_html += f'<b>{i}. Report #{_h(cond.get("report number"))} - {_h(actual_name)}</b><br>'
                cond_html += f'Date: {_h(cond.get("date", "N/A"))} | Location: {_h(cond.get("lg5pnj4chjadnv46", "N/A"))}<br>'
                cond_html += f'Condition: {_h(cond.get("uncbcge9x8vow9pn", "No description"))}<br>'
                cond_html += f'Status: {status}<br>'
                link = cond.get('link', '')
                if link and link != 'Link':
                    cond_html += f'<a href="{_h(link)}">View in KPA</a><br>'
                cond_html += '</div>'

            if len(conditions) > 10:
                cond_html += f'<p style="font-style:italic;">... and {len(conditions) - 10} more conditions in KPA</p>'

            sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['warning']};">
  <h2 style="color:{HTML_COLORS['warning']};margin:0 0 15px 0;font-size:18px;">AT-RISK CONDITIONS (Top {display_count} of {len(conditions)})</h2>
  {cond_html}
</td></tr>""")

    # --- RECOGNITION ---
    if 'observation_analysis' in all_data and all_data['observation_analysis']:
        recognition = all_data['observation_analysis']['by_type'].get('Recognition', [])
        if recognition:
            recognition_names = [{'name': get_actual_observer_name(rec), 'description': rec.get('uncbcge9x8vow9pn', '')} for rec in recognition]
            name_counter = Counter([r['name'] for r in recognition_names])

            rec_html = ''
            for name, count in name_counter.most_common(10):
                if name and name != 'Unknown':
                    rec_html += f'<div style="background:#f0fff0;border-left:4px solid {HTML_COLORS["safe"]};padding:12px 15px;margin:10px 0;">'
                    rec_html += f'<b style="color:{HTML_COLORS["safe"]};">&#9989; {_h(name)}</b> - {count} recognition(s)<br>'
                    for rec in recognition_names:
                        if rec['name'] == name:
                            rec_html += f'<i>\'{_h(rec["description"])}\'</i><br>'
                            break
                    rec_html += '</div>'

            sections.append(f"""
<tr><td style="padding:25px 40px;border-top:3px solid {HTML_COLORS['safe']};">
  <h2 style="color:{HTML_COLORS['safe']};margin:0 0 15px 0;font-size:18px;">SAFETY RECOGNITION - STARS ({len(recognition)})</h2>
  {rec_html}
</td></tr>""")

    # --- ASSESSMENT & AUDIT SUMMARY (replaces old "Other Forms Summary") ---
    if 'assessment_details' in all_data:
        try:
            audit_table_html = build_assessment_html(all_data['assessment_details'])
        except Exception as e:
            print(f"Warning: HTML assessment summary table error: {e}")
            audit_table_html = ''
            for form_id, form_name in OTHER_FORMS:
                data = all_data.get(f"form_{form_id}")
                count = data['count'] if data else 0
                audit_table_html += f'<b>{_h(form_name)}:</b> {count}<br>'
    else:
        audit_table_html = ''
        for form_id, form_name in OTHER_FORMS:
            data = all_data.get(f"form_{form_id}")
            count = data['count'] if data else 0
            audit_table_html += f'<b>{_h(form_name)}:</b> {count}<br>'

    sections.append(f"""
<tr><td style="padding:25px 40px;border-top:2px solid #ddd;">
  <h2 style="color:{HTML_COLORS['primary']};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {HTML_COLORS['primary']};padding-bottom:5px;">ASSESSMENT &amp; AUDIT SUMMARY</h2>
  {audit_table_html}
</td></tr>""")

    # --- FOOTER ---
    sections.append(f"""
<tr><td style="background:{HTML_COLORS['secondary']};padding:20px 40px;text-align:center;">
  <div style="color:#ffffff;font-size:11px;font-style:italic;">END OF REPORT</div>
  <div style="color:#ffcccc;font-size:10px;margin-top:4px;">Butch's Rat Hole &amp; Anchor Service Inc. | HSE Department</div>
</td></tr>""")

    # --- Wrapper end ---
    sections.append("""
</table>
</td></tr></table>
</body></html>""")

    return '\n'.join(sections)


# ==============================================================================
# SEND EMAIL
# ==============================================================================

def send_email_report(html_body, docx_path, yesterday_date):
    """Send report via Gmail SMTP. Fails gracefully - prints error, does not crash."""
    gmail_address = os.environ.get("GMAIL_ADDRESS", "")
    gmail_app_password = os.environ.get("GMAIL_APP_PASSWORD", "")
    recipient = os.environ.get("REPORT_RECIPIENT", "")

    if not gmail_address or not gmail_app_password or not recipient:
        print("  Email skipped -- GMAIL_ADDRESS, GMAIL_APP_PASSWORD, or REPORT_RECIPIENT not set.")
        return

    subject = f"Daily Safety Report - {yesterday_date.strftime('%B %d, %Y')}"

    try:
        msg = MIMEMultipart('mixed')
        msg['From'] = gmail_address
        msg['To'] = recipient
        msg['Subject'] = subject

        # HTML body
        msg.attach(MIMEText(html_body, 'html'))

        # .docx attachment
        if os.path.exists(docx_path):
            with open(docx_path, 'rb') as f:
                part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(docx_path)}"')
            msg.attach(part)

        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(gmail_address, gmail_app_password)
            server.sendmail(gmail_address, recipient, msg.as_string())

        print(f"[OK] Email sent to {recipient}")
    except Exception as e:
        print(f"  Email failed: {e}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    today = datetime.now()
    yesterday = today - timedelta(days=1)

    print("\n" + "="*80)
    print("KPA DAILY SAFETY REPORT - AUTOMATED")
    print(f"Report for: {yesterday.strftime('%A, %B %d, %Y')}")
    print("="*80)
    print("\n[+] Name field ONLY (actual observer, NOT James Barnett)")
    print("[+] Critical items first (Incidents, RCA, Near Misses)")
    print("[+] No blank sections - only shows data that exists")
    print("[+] Open Items excludes Near Misses (they have own section)")
    print("[+] Data quality alerts for miscategorization")
    print("[+] Assessment & Audit Analysis with compliance, findings, trends")
    print("[+] Dated filename\n")

    all_data = {}

    print("Pulling data from KPA...\n")

    for form_id, form_name in FORMS.items():
        data = pull_form_data(form_id, form_name)

        if form_id == 151085:
            obs_analysis = analyze_observations(data)
            all_data['observation_analysis'] = obs_analysis
            if obs_analysis:
                print(f"[+]Observation Cards: {obs_analysis['total']} total")
            else:
                print(f"[+]Observation Cards: 0")
        elif form_id == 151622:
            all_data['incident_reports'] = data
            if data:
                print(f"[+]Incident Reports: {data['count']}")
            else:
                print(f"[+]Incident Reports: 0")
        elif form_id == 180243:
            all_data['rca'] = data
            if data:
                print(f"[+]Root Cause Analysis: {data['count']}")
            else:
                print(f"[+]Root Cause Analysis: 0")
        else:
            all_data[f"form_{form_id}"] = data
            if data:
                print(f"[+]{form_name}: {data['count']}")
            else:
                print(f"[+]{form_name}: 0")

    # Analyze assessment/audit forms for the deep-analysis section
    print("\nAnalyzing assessment & audit data...")
    try:
        assessment_analysis = analyze_assessments(all_data)
        all_data['assessment_analysis'] = assessment_analysis
        if assessment_analysis['has_data']:
            print(f"[+]Assessment Analysis: {assessment_analysis['total_assessments']} assessments, "
                  f"{assessment_analysis['total_findings']} findings")
        else:
            print("[+]Assessment Analysis: No assessment data for yesterday")
    except Exception as e:
        print(f"  WARNING: Assessment analysis failed (non-fatal): {e}")
        all_data['assessment_analysis'] = None

    # Extract per-row assessment details for the summary table
    try:
        assessment_details = extract_assessment_details(all_data)
        all_data['assessment_details'] = assessment_details
        detail_count = sum(entry['count'] for entry in assessment_details)
        print(f"[+]Assessment Details: {detail_count} form rows extracted for summary table")
    except Exception as e:
        print(f"  WARNING: Assessment details extraction failed (non-fatal): {e}")
        all_data['assessment_details'] = None

    # Pull full assessment history (2025+) for dashboard accountability section
    print("\nPulling assessment history for dashboard...")
    try:
        assessment_history = pull_assessment_history()
        all_data['assessment_history'] = assessment_history
    except Exception as e:
        print(f"  WARNING: Assessment history pull failed (non-fatal): {e}")
        all_data['assessment_history'] = None

    print("\nGenerating report...")
    doc = build_word_document(all_data, yesterday)

    # Output to current working directory (works on both local and CI)
    date_str = yesterday.strftime('%Y-%m-%d')
    output_file = f"DailyKPAReport_{date_str}.docx"

    doc.save(output_file)

    print(f"\n[OK] Report saved: {output_file}")
    print(f"   Full path: {os.path.abspath(output_file)}")

    # --- Write JSON for dashboard ---
    import json as _json
    os.makedirs("output", exist_ok=True)
    obs = all_data.get('observation_analysis')
    inc = all_data.get('incident_reports')
    assess = all_data.get('assessment_analysis')
    assess_detail = all_data.get('assessment_details')

    # Build near misses list from observations
    near_misses = []
    if obs and obs.get('by_type', {}).get('Near Miss'):
        for nm in obs['by_type']['Near Miss']:
            near_misses.append({
                "report_number": nm.get('report number', ''),
                "date": nm.get('date', ''),
                "observer": get_actual_observer_name(nm),
                "description": nm.get('uncbcge9x8vow9pn', ''),
                "location": nm.get('lg5pnj4chjadnv46', ''),
                "type": nm.get('bff8m4x6xbc033kg', ''),
                "service_line": nm.get('64c7upqkyt79zhh1', ''),
            })

    # Build observations summary
    obs_summary = []
    if obs:
        for obs_type, obs_list in obs.get('by_type', {}).items():
            for o in obs_list:
                obs_summary.append({
                    "report_number": o.get('report number', ''),
                    "date": o.get('date', ''),
                    "observer": get_actual_observer_name(o),
                    "type": obs_type,
                    "description": o.get('uncbcge9x8vow9pn', ''),
                    "location": o.get('lg5pnj4chjadnv46', ''),
                    "service_line": o.get('64c7upqkyt79zhh1', ''),
                })

    # Merge form 484193 (TD - Observation Card) into observations + near misses
    td_obs_data = all_data.get('form_484193')
    if td_obs_data and td_obs_data.get('count', 0) > 0:
        for row in td_obs_data['rows']:
            obs_type = row.get('bff8m4x6xbc033kg', '').strip()
            sl = _get_service_line(row) or 'Drilling'
            entry = {
                "report_number": row.get('report number', ''),
                "date": row.get('date', ''),
                "observer": get_actual_observer_name(row),
                "type": obs_type,
                "description": row.get('uncbcge9x8vow9pn', ''),
                "location": row.get('lg5pnj4chjadnv46', ''),
                "service_line": sl,
            }
            obs_summary.append(entry)
            if obs_type == 'Near Miss':
                near_misses.append(entry)
        print(f"  Merged {td_obs_data['count']} TD Observation Card(s) into observations")

    # Update observation totals to include merged 484193 entries
    if obs:
        obs['total'] = len(obs_summary)
        type_counts = {}
        for entry in obs_summary:
            t = entry.get('type', 'Other') or 'Other'
            type_counts[t] = type_counts.get(t, 0) + 1
        obs['type_counts'] = type_counts

    # Build incidents list
    incidents_list = []
    if inc:
        for row in inc.get('rows', []):
            incidents_list.append({
                "report_number": row.get('report number', ''),
                "date": row.get('date', ''),
                "type": row.get('nojcquy0tfl9hqih', ''),
                "employee": row.get('55gg4nkoemnnfo2a', ''),
                "description": row.get('313e9txgrof0uute', ''),
                "location": row.get('9ohdd2lwvl7p0oc6', ''),
                "service_line": row.get('sha7vur5q2l6d6gq', ''),
            })

    json_data = {
        "report_date": yesterday.strftime("%Y-%m-%d"),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "observations": {
            "total": obs['total'] if obs else 0,
            "by_type": obs['type_counts'] if obs else {},
            "details": obs_summary,
        },
        "near_misses": near_misses,
        "incidents": incidents_list,
        "assessments": {
            "total": assess['total_assessments'] if assess and assess.get('has_data') else 0,
            "findings": assess['total_findings'] if assess and assess.get('has_data') else 0,
            "by_assessor": assess.get('by_assessor', {}) if assess else {},
            "details": assess_detail if assess_detail else [],
        },
        "assessment_activity": all_data.get('assessment_history'),
    }
    json_path = os.path.join("output", "kpa_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        _json.dump(json_data, f, indent=2, default=str)
    print(f"    JSON: {json_path}")

    # Build HTML and send email
    print("\nBuilding HTML email...")
    html_body = build_html_report(all_data, yesterday)

    if "--no-email" not in sys.argv:
        print("Sending email...")
        send_email_report(html_body, output_file, yesterday)
    else:
        print("Email skipped (--no-email flag)")
    print()

if __name__ == "__main__":
    main()
