#!/usr/bin/env python3
"""
DAILY SPEEDING REPORT - AUTOMATED (GitHub Actions)
===================================================
Runs daily at 5:05 AM Central via GitHub Actions.

Pulls speeding events from Motive API (last 24 hours) and generates:
- Word document (.docx) in LANDSCAPE, grouped by division/yard
- HTML email with the same structure, sent via Gmail SMTP

Thresholds (whichever is worse wins):
- RED:    20+ mph over posted limit OR 90+ mph absolute (immediate action)
- ORANGE: 15-19 mph over posted limit (coaching required)
- YELLOW: 10-14 mph over posted limit (monitoring)

Uses /v1/speeding_events endpoint (events wrapped as {speeding_event: {data}}).
Cross-references /v1/vehicles for driver names and /v1/groups for division/yard.
All API speeds are km/h - converted to mph (* 0.621371).
"""

import requests
import smtplib
import os
import re
import sys
from datetime import datetime, timedelta, timezone
from html import escape as html_escape
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from collections import Counter, OrderedDict
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from zoneinfo import ZoneInfo
    CENTRAL_TZ = ZoneInfo("America/Chicago")
except Exception:
    CENTRAL_TZ = timezone(timedelta(hours=-6))  # Fallback to CST

# ==============================================================================
# CONFIGURATION
# ==============================================================================

MOTIVE_API_KEY = os.environ.get("MOTIVE_API_KEY")
if not MOTIVE_API_KEY:
    print("ERROR: MOTIVE_API_KEY environment variable is not set.")
    sys.exit(1)

MOTIVE_BASE_URL = "https://api.gomotive.com/v1"
KMH_TO_MPH = 0.621371
LOGOS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logos")
CALIBRI = "Calibri"

# ==============================================================================
# GROUP ID -> (DIVISION, YARD) MAPPING
# Built from actual Motive groups API response.
# ==============================================================================

GROUP_ID_MAP = {
    # Rathole yards
    266026: ("Rathole", "Midland"),
    266025: ("Rathole", "Levelland"),
    266024: ("Rathole", "Barstow"),
    265996: ("Rathole", "Jourdanton"),
    290472: ("Rathole", "Jourdanton"),      # "Jourdanton Rathole All Vehicles"
    265998: ("Rathole", "Oklahoma"),         # "Shawnee Rat Hole"
    266028: ("Rathole", "Ohio"),             # "Wintersville Rat Hole"
    266027: ("Rathole", "Pennsylvania"),     # "Towanda Rat Hole"
    265997: ("Rathole", "North Dakota"),     # "Dickinson Rat Hole"
    265988: ("Rathole", ""),                 # Parent "Rat Hole" group

    # Casing yards
    167175: ("Casing", "Midland"),
    169090: ("Casing", "Bryan"),
    169092: ("Casing", "Kilgore"),
    186740: ("Casing", "Hobbs"),
    169091: ("Casing", "Jourdanton"),
    186739: ("Casing", "Laredo"),
    186741: ("Casing", "San Angelo"),
    186746: ("Casing", ""),                  # Parent "Casing" group

    # Other divisions (no yard breakdown)
    265993: ("Poly Pipe", ""),
    296040: ("Poly Pipe", ""),               # "Poly Crew"
    296036: ("Poly Pipe", ""),               # "Poly OM"
    296017: ("Poly Pipe", ""),               # "pumps & Gens"
    296020: ("Poly Pipe", ""),               # "Supervisors" (under Poly)
    265992: ("Pit Lining", ""),
    265983: ("Construction", ""),
    265987: ("Environmental", ""),
    265991: ("Fencing", ""),
    265982: ("Anchors", ""),                 # "Anchor"
    265989: ("Butch's Trucking", ""),        # "BTI"
    265986: ("Transcend Drilling", ""),       # "Transcend"
    265985: ("Valor Energy Services", ""),    # "Valor"
    265984: ("Sales/Admin", ""),              # "Sales & Support"
    290471: ("Rathole", "Jourdanton"),        # "Manager/Supervisor" under Jourdanton
}

# ==============================================================================
# VEHICLE PREFIX -> (DIVISION, YARD) FALLBACK
# Used when a vehicle has no group_ids.
# ==============================================================================

VEHICLE_PREFIX_MAP = [
    ("LL-RAT",  "Rathole",              "Levelland"),
    ("MID-RAT", "Rathole",              "Midland"),
    ("WIN-RAT", "Rathole",              "Ohio"),
    ("BAR-RAT", "Rathole",              "Barstow"),
    ("JOU-RAT", "Rathole",              "Jourdanton"),
    ("TOW-RAT", "Rathole",              "Pennsylvania"),
    ("OK-RAT",  "Rathole",              "Oklahoma"),
    ("DS-RAT",  "Rathole",              ""),
    ("BTI-",    "Butch's Trucking",     ""),
    ("VAL-",    "Valor Energy Services", ""),
    ("TD-",     "Transcend Drilling",   ""),
    ("POL-",    "Poly Pipe",            ""),
    ("ENV-",    "Environmental",        ""),
    ("FEN-",    "Fencing",              ""),
    ("ANC-",    "Anchors",              ""),
    ("PIT-",    "Pit Lining",           ""),
    ("CON-",    "Construction",         ""),
    ("SALES",   "Sales/Admin",          ""),
]

_CASING_RE = re.compile(r"^\d+C\b")  # e.g. "5036C", "19107C"

# Vehicles that should ALWAYS stay in Sales/Admin regardless of prefix
SALES_ADMIN_VEHICLES = {
    "Safety WIN-RAT-2529 Sean Fry",
}


def _division_from_prefix(vehicle_number):
    """Determine (division, yard) from vehicle number prefix."""
    # Check explicit Sales/Admin overrides first
    if vehicle_number in SALES_ADMIN_VEHICLES:
        return ("Sales/Admin", "")

    vn = vehicle_number.upper()

    # Check "SHAWNEE" anywhere in name -> Oklahoma
    if "SHAWNEE" in vn:
        return ("Rathole", "Oklahoma")

    for prefix, div, yard in VEHICLE_PREFIX_MAP:
        if vn.startswith(prefix.upper()):
            return (div, yard)
    if _CASING_RE.match(vn):
        return ("Casing", "")
    return ("Unassigned", "")


# ==============================================================================
# SAFETY REP ASSIGNMENTS
# ==============================================================================

SAFETY_REPS = {
    ("Rathole", "Midland"):       "John Snodgrass",
    ("Rathole", "Levelland"):     "John Snodgrass",
    ("Rathole", "Barstow"):       "Wes Franklin",
    ("Rathole", "Wink"):          "John Snodgrass",
    ("Rathole", "Jourdanton"):    "Leean Benevides",
    ("Rathole", "Ohio"):          "Sean Fry",
    ("Rathole", "Pennsylvania"):  "Sean Fry",
    ("Rathole", "Oklahoma"):      "Wes Franklin",
    ("Rathole", "North Dakota"):  "Wes Franklin",
    ("Casing", "Midland"):        "Michael Hancock & Michael Salazar",
    ("Casing", "Bryan"):          "Justin Conrad",
    ("Casing", "Kilgore"):        "James Barnett",
    ("Casing", "Hobbs"):          "Allen Batts",
    ("Casing", "Jourdanton"):     "Joey Speyrer",
    ("Casing", "Laredo"):         "Joey Speyrer",
    ("Casing", "San Angelo"):     "Michael Hancock & Michael Salazar",
    ("Butch's Trucking", ""):     "Bernard Bradley",
    ("Transcend Drilling", ""):   "John Snodgrass",
    ("Valor Energy Services", ""): "John Snodgrass",
    ("Poly Pipe", ""):            "Jose Romero",
    ("Pit Lining", ""):           "Jose Romero",
    ("Construction", ""):         "Jose Romero",
    ("Downhole Tools", ""):       "Jose Romero",
    ("Environmental", ""):        "John Snodgrass",
    ("Fencing", ""):              "John Snodgrass",
    ("Anchors", ""):              "John Snodgrass",
}

DIVISION_ORDER = [
    "Rathole", "Casing",
    "Butch's Trucking", "Transcend Drilling", "Valor Energy Services",
    "Poly Pipe", "Pit Lining", "Environmental", "Fencing", "Anchors",
    "Construction", "Downhole Tools", "Sales/Admin", "Unassigned",
]

YARD_ORDER = {
    "Rathole": [
        "Midland", "Levelland", "Barstow", "Wink", "Jourdanton",
        "Ohio", "Pennsylvania", "Oklahoma", "North Dakota",
    ],
    "Casing": [
        "Midland", "Bryan", "Kilgore", "Hobbs",
        "Jourdanton", "Laredo", "San Angelo",
    ],
}

DIVISION_REPS_SUMMARY = {
    "Rathole": "John Snodgrass, Wes Franklin, Leean Benevides, Sean Fry",
    "Casing": "Hancock/Salazar, Conrad, Barnett, Batts, Speyrer",
    "Butch's Trucking": "Bernard Bradley",
    "Transcend Drilling": "John Snodgrass",
    "Valor Energy Services": "John Snodgrass",
    "Poly Pipe": "Jose Romero",
    "Pit Lining": "Jose Romero",
    "Environmental": "John Snodgrass",
    "Fencing": "John Snodgrass",
    "Anchors": "John Snodgrass",
    "Construction": "Jose Romero",
    "Downhole Tools": "Jose Romero",
}

# Notes to display at the top of specific division sections
DIVISION_NOTES = {
    "Sales/Admin": "Note: This section includes non-field personnel (sales representatives, administrative staff, and safety vehicles). These are not operational drivers.",
}

# ==============================================================================
# MOTIVE API — VEHICLE + DRIVER LOOKUP
# ==============================================================================

def build_vehicle_lookup():
    """Fetch all vehicles from Motive and build lookup dicts.

    Returns:
        vehicle_drivers: {vehicle_number: driver_name}
        vehicle_groups:  {vehicle_number: (division, yard)}
    """
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    vehicle_drivers = {}
    vehicle_groups = {}
    page = 1

    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_URL}/vehicles",
                headers=headers,
                params={"per_page": 100, "page_no": page},
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            vehicles = data.get("vehicles", [])
            if not vehicles:
                break

            for wrapper in vehicles:
                v = wrapper.get("vehicle", wrapper)
                num = v.get("number", "")
                if not num:
                    continue

                # Driver: prefer current_driver, then permanent_driver
                driver_name = None
                for field in ("current_driver", "permanent_driver"):
                    d = v.get(field)
                    if d and isinstance(d, dict):
                        name = f"{d.get('first_name', '')} {d.get('last_name', '')}".strip()
                        if name:
                            driver_name = name
                            break
                if driver_name:
                    vehicle_drivers[num] = driver_name

                # Groups: use first matching group_id
                group_ids = v.get("group_ids", [])
                for gid in group_ids:
                    if gid in GROUP_ID_MAP:
                        vehicle_groups[num] = GROUP_ID_MAP[gid]
                        break

            pag = data.get("pagination", {})
            if page * 100 >= pag.get("total", 0):
                break
            page += 1

        except Exception as e:
            print(f"    Warning: vehicle lookup page {page} failed: {e}")
            break

    return vehicle_drivers, vehicle_groups


# ==============================================================================
# MOTIVE API — SPEEDING EVENTS
# ==============================================================================

def _format_duration(seconds):
    """Format duration in seconds to a readable string."""
    if not seconds or not isinstance(seconds, (int, float)):
        return "N/A"
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes = seconds // 60
    secs = seconds % 60
    if secs:
        return f"{minutes}m {secs}s"
    return f"{minutes}m"


def _utc_to_central(timestamp_str):
    """Convert UTC timestamp string to Central Time formatted string."""
    try:
        utc_dt = datetime.fromisoformat(timestamp_str.replace("Z", "+00:00"))
        central_dt = utc_dt.astimezone(CENTRAL_TZ)
        return central_dt.strftime("%m/%d/%Y %I:%M %p CT")
    except Exception:
        return str(timestamp_str)


def get_24h_speeding_events(vehicle_drivers, vehicle_groups):
    """Pull all speeding events from the last 24 hours.

    Uses /v1/speeding_events endpoint. Each event is wrapped as
    {"speeding_event": {actual data}}.
    """
    end_time = datetime.now(timezone.utc)
    start_time = end_time - timedelta(hours=24)

    start_iso = start_time.strftime("%Y-%m-%dT%H:%M:%SZ")
    end_iso = end_time.strftime("%Y-%m-%dT%H:%M:%SZ")

    print(f"    API date range: {start_iso} to {end_iso}")

    headers = {"X-Api-Key": MOTIVE_API_KEY}
    all_events = []
    page = 1

    while True:
        params = {
            "per_page": 100,
            "page_no": page,
            "start_time": start_iso,
            "end_time": end_iso,
        }

        try:
            resp = requests.get(
                f"{MOTIVE_BASE_URL}/speeding_events",
                headers=headers,
                params=params,
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            events = data.get("speeding_events", [])
            if not events:
                break

            if page == 1:
                print(f"    API reports total: {data.get('total', '?')} events")

            for wrapper in events:
                evt = wrapper.get("speeding_event", wrapper)
                enriched = enrich_event(evt, vehicle_drivers, vehicle_groups)
                all_events.append(enriched)

            total = data.get("total", 0)
            if page * 100 >= total:
                break
            page += 1

        except Exception as e:
            print(f"    Error fetching speeding page {page}: {e}")
            break

    # Sort by overspeed descending (worst violations first)
    return sorted(all_events, key=lambda x: x["overspeed"], reverse=True)


def enrich_event(event, vehicle_drivers, vehicle_groups):
    """Classify and enrich a single speeding event."""
    # --- Speeds (km/h -> mph) ---
    max_speed_kmh = event.get("max_vehicle_speed") or event.get("avg_vehicle_speed") or 0
    max_speed = round(max_speed_kmh * KMH_TO_MPH, 1)

    posted_kmh = event.get("min_posted_speed_limit_in_kph") or 0
    posted_speed = round(posted_kmh * KMH_TO_MPH, 1)

    over_kmh = event.get("max_over_speed_in_kph") or event.get("avg_over_speed_in_kph") or 0
    overspeed = round(over_kmh * KMH_TO_MPH, 1)

    # --- Tier classification ---
    if overspeed >= 20 or max_speed >= 90:
        tier = "RED"
    elif overspeed >= 15:
        tier = "ORANGE"
    elif overspeed >= 10:
        tier = "YELLOW"
    else:
        tier = "YELLOW"  # API events are already 6+ over

    # --- Vehicle ---
    vehicle_obj = event.get("vehicle", {})
    if isinstance(vehicle_obj, dict):
        vehicle_number = vehicle_obj.get("number", "Unknown")
    else:
        vehicle_number = str(vehicle_obj)

    # --- Driver (cross-reference vehicle lookup, then event, then parse) ---
    driver_name = vehicle_drivers.get(vehicle_number)
    if not driver_name:
        drv = event.get("driver")
        if drv and isinstance(drv, dict):
            name = f"{drv.get('first_name', '')} {drv.get('last_name', '')}".strip()
            if name:
                driver_name = name
    if not driver_name:
        # Parse from vehicle number: "TD-TD33171 Nick Sanchez" -> "Nick Sanchez"
        if " " in vehicle_number:
            candidate = vehicle_number.split(" ", 1)[1].strip().lstrip("- ")
            while candidate and candidate.split(" ", 1)[0].replace("-", "").isdigit():
                if " " in candidate:
                    candidate = candidate.split(" ", 1)[1].strip().lstrip("- ")
                else:
                    candidate = ""
            if len(candidate) > 2 and any(c.isalpha() for c in candidate):
                driver_name = candidate
    if not driver_name:
        driver_name = "Unknown"

    # --- Division / Yard ---
    # Check explicit Sales/Admin override first
    if vehicle_number in SALES_ADMIN_VEHICLES:
        division, yard = "Sales/Admin", ""
    else:
        div_yard = vehicle_groups.get(vehicle_number)
        if not div_yard:
            div_yard = _division_from_prefix(vehicle_number)
        division, yard = div_yard

    # --- Duration ---
    duration_str = _format_duration(event.get("duration", 0))

    # --- Severity ---
    metadata = event.get("metadata", {})
    severity = metadata.get("severity", "unknown") if isinstance(metadata, dict) else "unknown"

    # --- Time (convert to Central) ---
    timestamp = event.get("start_time") or event.get("end_time", "")
    formatted_time = _utc_to_central(timestamp)

    # --- Location / Map ---
    lat = event.get("start_lat")
    lon = event.get("start_lon")
    maps_link = f"https://www.google.com/maps?q={lat},{lon}" if lat and lon else ""
    location = f"{lat:.4f}, {lon:.4f}" if lat and lon else "Unknown"

    return {
        "driver": driver_name,
        "vehicle": vehicle_number,
        "speed": max_speed,
        "posted_speed": posted_speed,
        "overspeed": overspeed,
        "duration": duration_str,
        "severity": severity,
        "time": formatted_time,
        "location": location,
        "maps_link": maps_link,
        "tier": tier,
        "division": division,
        "yard": yard,
    }


# ==============================================================================
# DATA ORGANIZATION
# ==============================================================================

def get_repeat_offenders(events):
    """Find drivers with 3+ speeding events (exclude Unknown).

    Returns at most 10 offenders, sorted by worst single overspeed.
    """
    driver_counts = Counter(e["driver"] for e in events if e["driver"] != "Unknown")
    # Only 3+ events
    repeats_3plus = {n: c for n, c in driver_counts.items() if c >= 3}

    if not repeats_3plus:
        return {}

    # Sort by worst single overspeed event, take top 10
    def worst_over(name):
        return max(e["overspeed"] for e in events if e["driver"] == name)

    sorted_names = sorted(repeats_3plus.keys(), key=worst_over, reverse=True)[:10]
    return {n: repeats_3plus[n] for n in sorted_names}


def group_events(events):
    """Group events by division -> yard -> list of events.

    Returns OrderedDict following DIVISION_ORDER and YARD_ORDER.
    Only includes divisions/yards that have events.
    Events within each yard are sorted by overspeed descending.
    """
    raw = {}
    for e in events:
        div = e["division"]
        yard = e["yard"]
        raw.setdefault(div, {}).setdefault(yard, []).append(e)

    grouped = OrderedDict()
    for div in DIVISION_ORDER:
        if div not in raw:
            continue
        yards_data = raw[div]
        ordered_yards = OrderedDict()

        if div in YARD_ORDER:
            for y in YARD_ORDER[div]:
                if y in yards_data:
                    ordered_yards[y] = sorted(yards_data[y], key=lambda x: x["overspeed"], reverse=True)
            for y in sorted(yards_data.keys()):
                if y not in ordered_yards:
                    ordered_yards[y] = sorted(yards_data[y], key=lambda x: x["overspeed"], reverse=True)
        else:
            all_evts = []
            for y_evts in yards_data.values():
                all_evts.extend(y_evts)
            ordered_yards[""] = sorted(all_evts, key=lambda x: x["overspeed"], reverse=True)

        grouped[div] = ordered_yards

    for div in sorted(raw.keys()):
        if div not in grouped:
            yards_data = raw[div]
            all_evts = []
            for y_evts in yards_data.values():
                all_evts.extend(y_evts)
            grouped[div] = OrderedDict([("", sorted(all_evts, key=lambda x: x["overspeed"], reverse=True))])

    return grouped


def _yard_label(division, yard):
    """Build display label for a yard section."""
    rep = SAFETY_REPS.get((division, yard), "")
    if division in YARD_ORDER and yard:
        label = f"{yard} Yard"
    else:
        label = division
    return label, rep


# ==============================================================================
# BUILD WORD DOCUMENT
# ==============================================================================

def _set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, size_pt=8, bold=False, color=None, italic=False):
    """Apply Calibri font and formatting to a run."""
    run.font.name = CALIBRI
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_logo_row(doc):
    """Add company logos across the top. Skip missing logos gracefully."""
    logo_files = [
        "Butchs.jpg", "ButchTrucking.jpg", "Permian.jpg",
        "Hutchs.png", "Transcend.jpg", "Valor.jpg",
    ]
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    added = 0
    for lf in logo_files:
        path = os.path.join(LOGOS_DIR, lf)
        if os.path.exists(path):
            try:
                para.add_run().add_picture(path, width=Inches(1.3))
                para.add_run("  ")
                added += 1
            except Exception:
                pass
    if added == 0:
        run = para.add_run("BRHAS Safety Companies")
        _set_run_font(run, 16, bold=True, color=RGBColor(192, 0, 0))
    return added > 0


def _tier_bg_hex(tier):
    """Return background color hex for tier row shading."""
    if tier == "RED":
        return "FFE0E0"
    elif tier == "ORANGE":
        return "FFF0E0"
    else:
        return "FFFFF0"


def _add_event_table(doc, events):
    """Add a speeding events table to the document."""
    table = doc.add_table(rows=1, cols=9)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    headers = ["Tier", "Driver", "Vehicle", "Speed", "Limit", "Over", "Dur.", "Time", "Map"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        _set_run_font(run, 8, bold=True)

    for evt in events:
        cells = table.add_row().cells

        # Apply row background shading
        bg = _tier_bg_hex(evt["tier"])
        for cell in cells:
            _set_cell_shading(cell, bg)

        cells[0].text = evt["tier"]
        tier_run = cells[0].paragraphs[0].runs[0]
        tier_color = RGBColor(255, 0, 0) if evt["tier"] == "RED" else (
            RGBColor(255, 140, 0) if evt["tier"] == "ORANGE" else RGBColor(204, 153, 0))
        _set_run_font(tier_run, 8, bold=True, color=tier_color)

        cells[1].text = evt["driver"]
        cells[2].text = evt["vehicle"]
        cells[3].text = f"{evt['speed']}"
        cells[4].text = f"{evt['posted_speed']}"
        cells[5].text = f"+{evt['overspeed']}"
        cells[6].text = evt["duration"]
        cells[7].text = evt["time"]
        cells[8].text = "Map" if evt["maps_link"] else ""

        for c in cells[1:]:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_run_font(r, 8)


def _add_horizontal_rule(doc):
    """Add a visible horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="C00000"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def create_word_document(events, grouped, yesterday_date):
    """Build the full speeding report Word document in landscape."""
    doc = Document()

    # --- Landscape orientation ---
    for section in doc.sections:
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- Logos ---
    _add_logo_row(doc)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BRHAS SAFETY COMPANIES")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAILY SPEEDING REPORT")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(yesterday_date.strftime("%A, %B %d, %Y"))
    _set_run_font(run, 11, italic=True)

    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}")
    _set_run_font(run, 11, color=RGBColor(128, 0, 0))

    doc.add_paragraph()

    # --- Executive Summary ---
    red_events = [e for e in events if e["tier"] == "RED"]
    orange_events = [e for e in events if e["tier"] == "ORANGE"]
    yellow_events = [e for e in events if e["tier"] == "YELLOW"]
    repeats = get_repeat_offenders(events)

    p = doc.add_paragraph()
    run = p.add_run("EXECUTIVE SUMMARY")
    _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    run = p.add_run(f"Total Speeding Events: {len(events)}")
    _set_run_font(run, 11, bold=True)

    if red_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  RED — Immediate Action (20+ over or 90+ mph): {len(red_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(255, 0, 0))

    if orange_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  ORANGE — Coaching Required (15-19 over): {len(orange_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(255, 140, 0))

    if yellow_events:
        p = doc.add_paragraph()
        run = p.add_run(f"  YELLOW — Monitoring (10-14 over): {len(yellow_events)}")
        _set_run_font(run, 11, bold=True, color=RGBColor(204, 153, 0))

    if not events:
        p = doc.add_paragraph()
        run = p.add_run("No speeding events in the last 24 hours")
        _set_run_font(run, 11, bold=True, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    # --- Top 5 Worst Violations ---
    if events:
        p = doc.add_paragraph()
        run = p.add_run("TOP 5 WORST VIOLATIONS")
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        top5 = sorted(events, key=lambda x: x["overspeed"], reverse=True)[:5]
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        table.autofit = True

        for i, h in enumerate(["Driver", "Vehicle", "Speed", "Limit", "Over", "Division", "Yard"]):
            table.rows[0].cells[i].text = h
            _set_run_font(table.rows[0].cells[i].paragraphs[0].runs[0], 9, bold=True)

        for evt in top5:
            cells = table.add_row().cells
            _set_cell_shading(cells[0], _tier_bg_hex(evt["tier"]))
            cells[0].text = evt["driver"]
            cells[1].text = evt["vehicle"]
            cells[2].text = f"{evt['speed']} mph"
            cells[3].text = f"{evt['posted_speed']} mph"
            cells[4].text = f"+{evt['overspeed']} mph"
            cells[5].text = evt["division"]
            cells[6].text = evt["yard"] if evt["yard"] else "—"
            for c in cells:
                _set_cell_shading(c, _tier_bg_hex(evt["tier"]))
                for para in c.paragraphs:
                    for r in para.runs:
                        _set_run_font(r, 9)

        doc.add_paragraph()

    # --- Repeat Offenders ---
    if repeats:
        p = doc.add_paragraph()
        run = p.add_run(f"REPEAT OFFENDERS (3+ events — showing top {len(repeats)})")
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        for name, count in repeats.items():
            driver_evts = [e for e in events if e["driver"] == name]
            worst = max(driver_evts, key=lambda x: x["overspeed"])
            p = doc.add_paragraph()
            run = p.add_run(f"  {name}: {count} events")
            _set_run_font(run, 10, bold=True, color=RGBColor(192, 0, 0))
            run2 = p.add_run(f" (worst: +{worst['overspeed']} over at {worst['speed']} mph)")
            _set_run_font(run2, 10)

        doc.add_paragraph()

    # --- Division Sections ---
    for div, yards_data in grouped.items():
        _add_horizontal_rule(doc)

        p = doc.add_paragraph()
        run = p.add_run(div.upper())
        _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))

        rep_summary = DIVISION_REPS_SUMMARY.get(div, "")
        if rep_summary:
            p = doc.add_paragraph()
            run = p.add_run(f"Safety Rep(s): {rep_summary}")
            _set_run_font(run, 10, italic=True)

        # Division note (e.g. Sales/Admin)
        note = DIVISION_NOTES.get(div)
        if note:
            p = doc.add_paragraph()
            run = p.add_run(note)
            _set_run_font(run, 9, italic=True, color=RGBColor(100, 100, 100))

        div_total = sum(len(evts) for evts in yards_data.values())
        div_red = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "RED")
        div_orange = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "ORANGE")
        div_yellow = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "YELLOW")
        p = doc.add_paragraph()
        run = p.add_run(f"{div_total} events")
        _set_run_font(run, 10, bold=True)
        run2 = p.add_run(f" (RED: {div_red} | ORANGE: {div_orange} | YELLOW: {div_yellow})")
        _set_run_font(run2, 10)

        doc.add_paragraph()

        for yard, yard_events in yards_data.items():
            label, rep = _yard_label(div, yard)
            if div in YARD_ORDER and yard:
                p = doc.add_paragraph()
                header_text = f"{label.upper()}"
                if rep:
                    header_text += f" ({rep})"
                header_text += f" — {len(yard_events)} event{'s' if len(yard_events) != 1 else ''}"
                run = p.add_run(header_text)
                _set_run_font(run, 11, bold=True, color=RGBColor(64, 0, 0))

            _add_event_table(doc, yard_events)
            doc.add_paragraph()

    # --- Footer ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END OF REPORT")
    _set_run_font(run, 10, italic=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Butch's Rat Hole & Anchor Service Inc. | HSE Department")
    _set_run_font(run, 9, color=RGBColor(128, 0, 0))

    return doc


# ==============================================================================
# BUILD HTML EMAIL
# ==============================================================================

C_RED = "#C00000"
C_DARK = "#800000"
C_AMBER = "#FF8C00"
C_YELLOW_DARK = "#CC9900"
C_GREEN = "#008000"


def _h(text):
    """HTML-escape text safely."""
    return html_escape(str(text)) if text else ""


def _tier_colors(tier):
    """Return (text_color, bg_color) for a tier."""
    if tier == "RED":
        return "#FF0000", "#FFE0E0"
    elif tier == "ORANGE":
        return C_AMBER, "#FFF0E0"
    else:
        return C_YELLOW_DARK, "#FFFFF0"


def build_html_report(events, grouped, yesterday_date):
    """Build HTML email body organized by division/yard."""
    red_events = [e for e in events if e["tier"] == "RED"]
    orange_events = [e for e in events if e["tier"] == "ORANGE"]
    yellow_events = [e for e in events if e["tier"] == "YELLOW"]
    repeats = get_repeat_offenders(events)

    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    parts = []

    # --- Wrapper + Header ---
    parts.append(f"""<html><head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#f4f4f4;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f4f4;">
<tr><td align="center">
<table width="700" cellpadding="0" cellspacing="0" style="background:#ffffff;border:1px solid #ddd;margin:20px auto;font-family:Calibri,Arial,Helvetica,sans-serif;font-size:14px;color:#333;">

<tr><td style="background:{C_RED};padding:30px 40px;text-align:center;">
  <div style="font-size:16px;font-weight:bold;color:#ffffff;letter-spacing:1px;">BRHAS SAFETY COMPANIES</div>
  <div style="font-size:28px;font-weight:bold;color:#ffffff;margin:10px 0;">DAILY SPEEDING REPORT</div>
  <div style="font-size:13px;font-style:italic;color:#ffcccc;">HSE Management Summary</div>
  <div style="font-size:12px;color:#ffffff;margin-top:8px;">Report Date: {yesterday_date.strftime('%A, %B %d, %Y')}</div>
  <div style="font-size:10px;color:#ffcccc;margin-top:4px;">Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}</div>
</td></tr>""")

    # --- Executive Summary ---
    summary = f"<b>Total Speeding Events: {len(events)}</b><br><br>"
    if red_events:
        summary += f'<div style="color:#FF0000;font-weight:bold;margin:4px 0 4px 20px;">RED — Immediate Action (20+ over or 90+ mph): {len(red_events)}</div>'
    if orange_events:
        summary += f'<div style="color:{C_AMBER};font-weight:bold;margin:4px 0 4px 20px;">ORANGE — Coaching Required (15-19 over): {len(orange_events)}</div>'
    if yellow_events:
        summary += f'<div style="color:{C_YELLOW_DARK};font-weight:bold;margin:4px 0 4px 20px;">YELLOW — Monitoring (10-14 over): {len(yellow_events)}</div>'
    if not events:
        summary += f'<b style="color:{C_GREEN};">No speeding events in the last 24 hours!</b>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">EXECUTIVE SUMMARY</h2>
  {summary}
</td></tr>""")

    # --- Top 5 Worst Violations ---
    if events:
        top5 = sorted(events, key=lambda x: x["overspeed"], reverse=True)[:5]
        top5_html = ""
        for e in top5:
            tc, bg = _tier_colors(e["tier"])
            top5_html += f'<div style="background:{bg};border-left:4px solid {tc};padding:10px 15px;margin:8px 0;">'
            top5_html += f'<b style="color:{tc};">+{e["overspeed"]} mph over</b> '
            top5_html += f'({e["speed"]} in a {e["posted_speed"]} zone)<br>'
            top5_html += f'<b>Driver:</b> {_h(e["driver"])} | <b>Vehicle:</b> {_h(e["vehicle"])}<br>'
            top5_html += f'<b>Division:</b> {_h(e["division"])}'
            if e["yard"]:
                top5_html += f' / {_h(e["yard"])}'
            top5_html += f' | <b>Time:</b> {_h(e["time"])}'
            if e["maps_link"]:
                top5_html += f' | <a href="{_h(e["maps_link"])}">Map</a>'
            top5_html += "</div>"

        parts.append(f"""
<tr><td style="padding:20px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">TOP 5 WORST VIOLATIONS</h2>
  {top5_html}
</td></tr>""")

    # --- Repeat Offenders ---
    if repeats:
        repeat_html = ""
        for name, count in repeats.items():
            driver_evts = [e for e in events if e["driver"] == name]
            worst = max(driver_evts, key=lambda x: x["overspeed"])
            repeat_html += f'<div style="background:#fff5f5;border-left:4px solid {C_RED};padding:10px 15px;margin:8px 0;">'
            repeat_html += f'<b style="color:{C_RED};">{_h(name)}: {count} events</b>'
            repeat_html += f' (worst: +{worst["overspeed"]} over at {worst["speed"]} mph)'
            repeat_html += "</div>"

        parts.append(f"""
<tr><td style="padding:20px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">REPEAT OFFENDERS (3+ events — top {len(repeats)})</h2>
  {repeat_html}
</td></tr>""")

    # --- Division Sections ---
    for div, yards_data in grouped.items():
        div_total = sum(len(evts) for evts in yards_data.values())
        div_red = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "RED")
        div_orange = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "ORANGE")
        div_yellow = sum(1 for evts in yards_data.values() for e in evts if e["tier"] == "YELLOW")

        rep_summary = DIVISION_REPS_SUMMARY.get(div, "")
        note = DIVISION_NOTES.get(div, "")

        parts.append(f"""
<tr><td style="padding:0 40px;"><hr style="border:none;border-top:3px solid {C_RED};margin:20px 0 0 0;"></td></tr>
<tr><td style="padding:15px 40px;">
  <h2 style="color:{C_RED};margin:0;font-size:20px;">{_h(div.upper())}</h2>
  {"<div style='font-size:12px;font-style:italic;color:#666;margin:4px 0;'>Safety Rep(s): " + _h(rep_summary) + "</div>" if rep_summary else ""}
  {"<div style='font-size:11px;font-style:italic;color:#888;margin:4px 0;'>" + _h(note) + "</div>" if note else ""}
  <div style="background:#f8f0f0;border-left:4px solid {C_RED};padding:10px 15px;margin:10px 0;font-size:13px;">
    <b>{_h(div)}</b> had <b>{div_total}</b> speeding event{"s" if div_total != 1 else ""} today
    (RED: {div_red} | ORANGE: {div_orange} | YELLOW: {div_yellow})
  </div>
</td></tr>""")

        for yard, yard_events in yards_data.items():
            label, rep = _yard_label(div, yard)
            if div in YARD_ORDER and yard:
                yard_header = f"{_h(label.upper())}"
                if rep:
                    yard_header += f" <span style='font-weight:normal;font-size:12px;'>({_h(rep)})</span>"
                yard_header += f" — {len(yard_events)} event{'s' if len(yard_events) != 1 else ''}"
                parts.append(f"""
<tr><td style="padding:10px 40px 5px 40px;">
  <h3 style="color:{C_DARK};margin:0;font-size:15px;">{yard_header}</h3>
</td></tr>""")

            table_rows = ""
            for e in yard_events:
                tc, bg = _tier_colors(e["tier"])
                map_cell = f'<a href="{_h(e["maps_link"])}" style="font-size:11px;">Map</a>' if e["maps_link"] else ""
                table_rows += f"""<tr style="background:{bg};">
  <td style="padding:5px 6px;border:1px solid #ddd;"><b style="color:{tc};">{e["tier"]}</b></td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{_h(e["driver"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{_h(e["vehicle"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;font-weight:bold;">{e["speed"]}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;">{e["posted_speed"]}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;font-weight:bold;color:{tc};">+{e["overspeed"]}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;text-align:center;">{_h(e["duration"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;font-size:11px;">{_h(e["time"])}</td>
  <td style="padding:5px 6px;border:1px solid #ddd;">{map_cell}</td>
</tr>"""

            parts.append(f"""
<tr><td style="padding:5px 40px 15px 40px;">
  <table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:12px;">
    <tr style="background:{C_RED};">
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Tier</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Driver</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Vehicle</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Speed</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Limit</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Over</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Dur.</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Time</th>
      <th style="padding:6px;color:#fff;border:1px solid #ddd;">Map</th>
    </tr>
    {table_rows}
  </table>
</td></tr>""")

    # --- Footer ---
    parts.append(f"""
<tr><td style="background:{C_DARK};padding:20px 40px;text-align:center;">
  <div style="color:#ffffff;font-size:11px;font-style:italic;">END OF REPORT</div>
  <div style="color:#ffcccc;font-size:10px;margin-top:4px;">Butch's Rat Hole &amp; Anchor Service Inc. | HSE Department</div>
</td></tr>

</table>
</td></tr></table>
</body></html>""")

    return "\n".join(parts)


# ==============================================================================
# SEND EMAIL
# ==============================================================================

def send_email_report(html_body, docx_path, yesterday_date):
    """Send report via Gmail SMTP. Fails gracefully."""
    gmail_address = os.environ.get("GMAIL_ADDRESS", "")
    gmail_app_password = os.environ.get("GMAIL_APP_PASSWORD", "")
    recipient = os.environ.get("REPORT_RECIPIENT", "")

    if not gmail_address or not gmail_app_password or not recipient:
        print("  Email skipped — GMAIL_ADDRESS, GMAIL_APP_PASSWORD, or REPORT_RECIPIENT not set.")
        return

    subject = f"Daily Speeding Report - {yesterday_date.strftime('%B %d, %Y')}"

    try:
        msg = MIMEMultipart("mixed")
        msg["From"] = gmail_address
        msg["To"] = recipient
        msg["Subject"] = subject

        msg.attach(MIMEText(html_body, "html"))

        if os.path.exists(docx_path):
            with open(docx_path, "rb") as f:
                part = MIMEBase(
                    "application",
                    "vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f'attachment; filename="{os.path.basename(docx_path)}"',
            )
            msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_address, gmail_app_password)
            server.sendmail(gmail_address, recipient, msg.as_string())

        print(f"  Email sent to {recipient}")
    except Exception as e:
        print(f"  Email failed: {e}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    today = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    yesterday = today - timedelta(days=1)

    print("\n" + "=" * 80)
    print("DAILY SPEEDING REPORT - AUTOMATED")
    print(f"Report for: {yesterday.strftime('%A, %B %d, %Y')}")
    print("=" * 80)
    print("\n  Thresholds (whichever is worse wins):")
    print("    RED:    20+ over posted limit OR 90+ mph (immediate action)")
    print("    ORANGE: 15-19 over posted limit (coaching required)")
    print("    YELLOW: 10-14 over posted limit (monitoring)")
    print("    Repeat: 3+ events flagged\n")

    print("[1] Building vehicle/driver lookup from Motive...")
    vehicle_drivers, vehicle_groups = build_vehicle_lookup()
    print(f"    {len(vehicle_drivers)} vehicles with driver names")
    print(f"    {len(vehicle_groups)} vehicles with group assignments")

    print("\n[2] Fetching speeding events from Motive...")
    events = get_24h_speeding_events(vehicle_drivers, vehicle_groups)
    print(f"    Found {len(events)} events")

    if events:
        red = len([e for e in events if e["tier"] == "RED"])
        orange = len([e for e in events if e["tier"] == "ORANGE"])
        yellow = len([e for e in events if e["tier"] == "YELLOW"])
        known = len([e for e in events if e["driver"] != "Unknown"])
        repeats = get_repeat_offenders(events)
        print(f"    RED: {red} | ORANGE: {orange} | YELLOW: {yellow}")
        print(f"    Drivers identified: {known}/{len(events)} ({100*known//len(events)}%)")
        if repeats:
            print(f"    Repeat offenders ({len(repeats)}): {', '.join(f'{n} ({c}x)' for n, c in repeats.items())}")

    print("\n[3] Grouping events by division/yard...")
    grouped = group_events(events)
    for div, yards_data in grouped.items():
        total = sum(len(evts) for evts in yards_data.values())
        print(f"    {div}: {total} events")

    print("\n[4] Creating Word document (landscape)...")
    doc = create_word_document(events, grouped, yesterday)

    date_str = yesterday.strftime("%Y-%m-%d")
    output_file = f"DailySpeedingReport_{date_str}.docx"
    doc.save(output_file)
    print(f"    Saved: {output_file}")

    print("\n[5] Building HTML email...")
    html_body = build_html_report(events, grouped, yesterday)

    print("[6] Sending email...")
    send_email_report(html_body, output_file, yesterday)

    print("\n" + "=" * 80)
    print("COMPLETE")
    print("=" * 80 + "\n")


if __name__ == "__main__":
    main()
