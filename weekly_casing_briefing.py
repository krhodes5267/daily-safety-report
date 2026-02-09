#!/usr/bin/env python3
"""
WEEKLY CASING DIVISION SAFETY INTELLIGENCE BRIEFING
=====================================================
Runs Monday mornings at 6:00 AM Central via GitHub Actions.
Covers previous 7 days (Monday 00:00:00 CT through Sunday 23:59:59 CT).

Aggregates data from ALL three sources:
1. Motive Speeding API (/v1/speeding_events)
2. Motive Camera Events API (/v2/driver_performance_events)
3. KPA EHS API (/v1/reports)

Generates a meeting-ready Word document + HTML email for the Monday 2 PM safety meeting.

10 Sections:
1. Week at a Glance
2. Red Flag Drivers
3. Camera Event Summary
4. Speeding Summary
5. KPA Incidents & Observations
6. Open Action Items
7. Weekend Spotlight
8. Agenda Assignments
9. Vehicle Health Flags
10. Yard Comparison Scorecard
"""

import requests
import smtplib
import os
import sys
import csv
import json
from datetime import datetime, timedelta, timezone
from html import escape as html_escape
from io import StringIO
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from collections import Counter, OrderedDict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

try:
    from zoneinfo import ZoneInfo
    CENTRAL_TZ = ZoneInfo("America/Chicago")
except Exception:
    CENTRAL_TZ = timezone(timedelta(hours=-6))

# ==============================================================================
# CONFIGURATION
# ==============================================================================

MOTIVE_API_KEY = os.environ.get("MOTIVE_API_KEY")
if not MOTIVE_API_KEY:
    print("ERROR: MOTIVE_API_KEY environment variable is not set.")
    sys.exit(1)

KPA_API_TOKEN = os.environ.get("KPA_API_TOKEN", "")
KPA_AVAILABLE = bool(KPA_API_TOKEN)
if not KPA_AVAILABLE:
    print("WARNING: KPA_API_TOKEN not set. KPA sections will show 'data unavailable'.")

MOTIVE_BASE_URL_V1 = "https://api.gomotive.com/v1"
MOTIVE_BASE_URL_V2 = "https://api.gomotive.com/v2"
KPA_API_BASE = "https://api.kpaehs.com/v1"
KMH_TO_MPH = 0.621371
LOGOS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logos")
CALIBRI = "Calibri"

# ==============================================================================
# CASING GROUP IDS (from Motive Groups API)
# ==============================================================================

CASING_GROUP_IDS = {
    167175: "Midland",
    169090: "Bryan",
    169092: "Kilgore",
    186740: "Hobbs",
    169091: "Jourdanton",
    186739: "Laredo",
    186741: "San Angelo",
    186746: "",  # Parent "Casing" group
}

ALL_CASING_GROUP_IDS = set(CASING_GROUP_IDS.keys())

# ==============================================================================
# YARD CONFIGURATION
# ==============================================================================

YARD_ORDER = ["Midland", "Bryan", "Kilgore", "Hobbs", "Jourdanton", "Laredo", "San Angelo"]

YARD_INFO = {
    "Midland": {
        "safety_reps": "Michael Hancock & Michael Salazar",
        "safety_rep_emails": ["mhancock@brhas.com", "msalazar@brhas.com"],
        "manager": "Richie Bentley",
    },
    "Bryan": {
        "safety_reps": "Justin Conrad",
        "safety_rep_emails": ["jconrad@brhas.com"],
        "manager": "Danny Lohse",
    },
    "Kilgore": {
        "safety_reps": "James Barnett (J.P.)",
        "safety_rep_emails": ["jbarnett@brhas.com"],
        "manager": "Frankie Balderas",
    },
    "Hobbs": {
        "safety_reps": "Allen Batts",
        "safety_rep_emails": ["abatts@brhas.com"],
        "manager": "Clifton Eaves",
    },
    "Jourdanton": {
        "safety_reps": "Joey Speyrer",
        "safety_rep_emails": ["jspeyrer@brhas.com"],
        "manager": "Enrique Flores",
    },
    "Laredo": {
        "safety_reps": "Joey Speyrer",
        "safety_rep_emails": ["jspeyrer@brhas.com"],
        "manager": "Chris Jacobo",
    },
    "San Angelo": {
        "safety_reps": "Michael Hancock & Michael Salazar",
        "safety_rep_emails": ["mhancock@brhas.com", "msalazar@brhas.com"],
        "manager": "Jeremy Jones",
    },
}

# Map safety rep name to their yards for agenda assignments
SAFETY_REP_YARDS = OrderedDict([
    ("MICHAEL HANCOCK & MICHAEL SALAZAR", ["Midland", "San Angelo"]),
    ("JUSTIN CONRAD", ["Bryan"]),
    ("JAMES BARNETT (J.P.)", ["Kilgore"]),
    ("ALLEN BATTS", ["Hobbs"]),
    ("JOEY SPEYRER", ["Jourdanton", "Laredo"]),
])

OBS_TARGET_PER_YARD = 3

# ==============================================================================
# EVENT TYPE CLASSIFICATION (reused from camera report)
# ==============================================================================

EVENT_TYPE_NORMALIZE = {
    "distraction": "distraction", "distracted_driving": "distraction",
    "driver_distraction": "distraction",
    "cell_phone": "cell_phone", "cell_phone_usage": "cell_phone",
    "phone_use": "cell_phone", "cellphone": "cell_phone", "phone_usage": "cell_phone",
    "drowsiness": "drowsiness", "drowsy": "drowsiness",
    "drowsy_driving": "drowsiness", "fatigue": "drowsiness",
    "driver_drowsiness": "drowsiness",
    "close_following": "close_following", "following_distance": "close_following",
    "tailgating": "close_following",
    "forward_collision_warning": "forward_collision_warning",
    "forward_collision": "forward_collision_warning", "fcw": "forward_collision_warning",
    "collision": "collision", "crash": "collision",
    "near_collision": "near_collision", "near_crash": "near_collision",
    "stop_sign_violation": "stop_sign_violation", "stop_sign": "stop_sign_violation",
    "ran_stop_sign": "stop_sign_violation",
    "unsafe_lane_change": "unsafe_lane_change", "lane_change": "unsafe_lane_change",
    "aggregated_lane_swerving": "lane_swerving", "lane_swerving": "lane_swerving",
    "lane_swerve": "lane_swerving",
    "hard_brake": "hard_brake", "hard_braking": "hard_brake",
    "harsh_brake": "hard_brake", "harsh_braking": "hard_brake",
    "seat_belt_violation": "seat_belt_violation", "seatbelt": "seat_belt_violation",
    "seatbelt_violation": "seat_belt_violation", "no_seatbelt": "seat_belt_violation",
    "seat_belt": "seat_belt_violation",
    "camera_obstruction": "camera_obstruction", "obstruction": "camera_obstruction",
    "camera_blocked": "camera_obstruction",
    "driver_facing_cam_obstruction": "camera_obstruction",
    "road_facing_cam_obstruction": "camera_obstruction",
    "smoking": "smoking", "vaping": "smoking",
    "unsafe_parking": "unsafe_parking",
    "hard_accel": "hard_accel", "hard_acceleration": "hard_accel",
    "harsh_acceleration": "hard_accel", "rapid_acceleration": "hard_accel",
    "hard_corner": "hard_corner", "hard_cornering": "hard_corner",
    "hard_turn": "hard_corner", "harsh_cornering": "hard_corner", "harsh_turn": "hard_corner",
    "speed_violation": "speed_violation", "speeding": "speed_violation",
}

RED_TYPES = {
    "distraction", "cell_phone", "drowsiness", "close_following",
    "forward_collision_warning", "collision", "near_collision",
    "stop_sign_violation", "unsafe_lane_change", "lane_swerving",
}

ORANGE_TYPES = {
    "hard_brake", "seat_belt_violation", "camera_obstruction", "smoking", "unsafe_parking",
}

YELLOW_TYPES = {
    "hard_accel", "hard_corner", "speed_violation",
}

OBSTRUCTION_RAW_TYPES = {
    "camera_obstruction", "driver_facing_cam_obstruction",
    "road_facing_cam_obstruction", "obstruction", "camera_blocked",
}

EVENT_DISPLAY_NAMES = {
    "distraction": "Distraction", "cell_phone": "Cell Phone",
    "drowsiness": "Drowsiness", "close_following": "Close Following",
    "forward_collision_warning": "Forward Collision Warning",
    "collision": "Collision", "near_collision": "Near Collision",
    "stop_sign_violation": "Stop Sign Violation",
    "unsafe_lane_change": "Unsafe Lane Change", "lane_swerving": "Lane Swerving",
    "hard_brake": "Hard Brake", "seat_belt_violation": "Seatbelt Violation",
    "camera_obstruction": "Camera Obstruction", "smoking": "Smoking",
    "unsafe_parking": "Unsafe Parking",
    "hard_accel": "Hard Acceleration", "hard_corner": "Hard Corner",
    "speed_violation": "Speed Violation",
}

EVENT_SEVERITY_ORDER = {
    "collision": 1, "near_collision": 2, "forward_collision_warning": 3,
    "distraction": 4, "cell_phone": 5, "drowsiness": 6,
    "stop_sign_violation": 7, "unsafe_lane_change": 8, "lane_swerving": 8,
    "close_following": 9, "hard_brake": 10, "seat_belt_violation": 11,
    "camera_obstruction": 12, "smoking": 13, "unsafe_parking": 14,
    "hard_accel": 15, "hard_corner": 16, "speed_violation": 17,
}


def _normalize_event_type(raw_type):
    if not raw_type:
        return "unknown"
    key = raw_type.lower().strip().replace(" ", "_").replace("-", "_")
    return EVENT_TYPE_NORMALIZE.get(key, key)


def _classify_tier(event_type):
    if event_type in RED_TYPES:
        return "RED"
    elif event_type in ORANGE_TYPES:
        return "ORANGE"
    elif event_type in YELLOW_TYPES:
        return "YELLOW"
    return "ORANGE"


def _event_display_name(event_type, raw_type=""):
    name = EVENT_DISPLAY_NAMES.get(event_type)
    if name:
        return name
    display = raw_type or event_type
    return display.replace("_", " ").title()


def _event_sort_key(event):
    tier_order = {"RED": 0, "ORANGE": 1, "YELLOW": 2}
    severity = EVENT_SEVERITY_ORDER.get(event["event_type"], 50)
    return (tier_order.get(event["tier"], 1), severity)


# ==============================================================================
# UTILITY FUNCTIONS
# ==============================================================================

def _format_duration(seconds):
    if not seconds or not isinstance(seconds, (int, float)):
        return "N/A"
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes = seconds // 60
    secs = seconds % 60
    if secs:
        return f"{minutes}m {secs}s"
    return f"{minutes}m"


def _utc_to_central(timestamp_str):
    try:
        utc_dt = datetime.fromisoformat(timestamp_str.replace("Z", "+00:00"))
        central_dt = utc_dt.astimezone(CENTRAL_TZ)
        return central_dt.strftime("%m/%d/%Y %I:%M %p CT")
    except Exception:
        return str(timestamp_str)


def _utc_to_central_dt(timestamp_str):
    try:
        utc_dt = datetime.fromisoformat(timestamp_str.replace("Z", "+00:00"))
        return utc_dt.astimezone(CENTRAL_TZ)
    except Exception:
        return None


def _plural(count, singular, plural_form=None):
    if count == 1:
        return f"1 {singular}"
    return f"{count} {plural_form or singular + 's'}"


def _h(text):
    return html_escape(str(text)) if text else ""


# ==============================================================================
# DATE RANGE CALCULATION
# ==============================================================================

def get_week_range():
    """Calculate Mon-Sun range for the previous week.

    If run on Monday Feb 16, covers Mon Feb 9 00:00:00 CT through Sun Feb 15 23:59:59 CT.
    """
    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    today = now_central.date()

    # Find last Monday (if today is Monday, go back 7 days)
    days_since_monday = today.weekday()  # 0=Mon
    if days_since_monday == 0:
        last_monday = today - timedelta(days=7)
    else:
        last_monday = today - timedelta(days=days_since_monday + 7)

    last_sunday = last_monday + timedelta(days=6)

    start_ct = datetime(last_monday.year, last_monday.month, last_monday.day, 0, 0, 0, tzinfo=CENTRAL_TZ)
    end_ct = datetime(last_sunday.year, last_sunday.month, last_sunday.day, 23, 59, 59, tzinfo=CENTRAL_TZ)

    return start_ct, end_ct, last_monday, last_sunday


# ==============================================================================
# MOTIVE API - VEHICLE LOOKUP (Casing only)
# ==============================================================================

def build_casing_vehicle_lookup():
    """Fetch all vehicles from Motive and build lookup for Casing vehicles."""
    headers = {"X-Api-Key": MOTIVE_API_KEY}
    vehicle_drivers = {}
    vehicle_yards = {}
    all_casing_vehicles = set()
    yard_vehicle_counts = Counter()
    page = 1

    while True:
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_URL_V1}/vehicles",
                headers=headers,
                params={"per_page": 100, "page_no": page},
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            vehicles = data.get("vehicles", [])
            if not vehicles:
                break

            for wrapper in vehicles:
                v = wrapper.get("vehicle", wrapper)
                num = v.get("number", "")
                if not num:
                    continue

                group_ids = v.get("group_ids", [])
                yard = None
                for gid in group_ids:
                    if gid in CASING_GROUP_IDS:
                        yard = CASING_GROUP_IDS[gid]
                        break

                if yard is None:
                    continue

                all_casing_vehicles.add(num)
                if yard:
                    vehicle_yards[num] = yard
                    yard_vehicle_counts[yard] += 1

                driver_name = None
                for field in ("current_driver", "permanent_driver"):
                    d = v.get(field)
                    if d and isinstance(d, dict):
                        name = f"{d.get('first_name', '')} {d.get('last_name', '')}".strip()
                        if name:
                            driver_name = name
                            break
                if driver_name:
                    vehicle_drivers[num] = driver_name

            pag = data.get("pagination", {})
            if page * 100 >= pag.get("total", 0):
                break
            page += 1

        except Exception as e:
            print(f"    Warning: vehicle lookup page {page} failed: {e}")
            break

    return vehicle_drivers, vehicle_yards, all_casing_vehicles, dict(yard_vehicle_counts)


# ==============================================================================
# MOTIVE API - SPEEDING EVENTS (Weekly)
# ==============================================================================

def get_speeding_events_weekly(start_ct, end_ct, vehicle_drivers, vehicle_yards, casing_vehicles):
    """Pull speeding events for the full week, filtered to Casing vehicles."""
    start_utc = start_ct.astimezone(timezone.utc)
    end_utc = end_ct.astimezone(timezone.utc)

    api_start_date = start_utc.strftime("%Y-%m-%d")
    api_end_date = end_utc.strftime("%Y-%m-%d")

    print(f"    Central window: {start_ct.strftime('%m/%d/%Y %I:%M %p')} to {end_ct.strftime('%m/%d/%Y %I:%M %p')}")
    print(f"    API date filter: start_date={api_start_date}, end_date={api_end_date}")

    headers = {"X-Api-Key": MOTIVE_API_KEY}
    raw_events = []
    page = 1

    while True:
        params = {
            "per_page": 100, "page_no": page,
            "start_date": api_start_date, "end_date": api_end_date,
        }
        try:
            resp = requests.get(
                f"{MOTIVE_BASE_URL_V1}/speeding_events",
                headers=headers, params=params, timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            events = data.get("speeding_events", [])
            if not events:
                break
            raw_events.extend(events)
            total = data.get("total", 0)
            if page * 100 >= total:
                break
            page += 1
        except Exception as e:
            print(f"    Error fetching speeding page {page}: {e}")
            break

    print(f"    Raw speeding events fetched: {len(raw_events)}")

    filtered = []
    for wrapper in raw_events:
        evt = wrapper.get("speeding_event", wrapper)
        evt_time_str = evt.get("start_time", "")
        try:
            evt_utc = datetime.fromisoformat(evt_time_str.replace("Z", "+00:00"))
            evt_central = evt_utc.astimezone(CENTRAL_TZ)
            if not (start_ct <= evt_central <= end_ct):
                continue
        except Exception:
            pass

        vehicle_obj = evt.get("vehicle", {})
        vehicle_number = vehicle_obj.get("number", "") if isinstance(vehicle_obj, dict) else str(vehicle_obj)

        if vehicle_number and casing_vehicles and vehicle_number not in casing_vehicles:
            continue

        enriched = _enrich_speeding_event(evt, vehicle_drivers, vehicle_yards)
        filtered.append(enriched)

    print(f"    After Casing + Central Time filter: {_plural(len(filtered), 'event')}")
    return sorted(filtered, key=lambda x: x["overspeed"], reverse=True)


def _enrich_speeding_event(event, vehicle_drivers, vehicle_yards):
    max_speed_kmh = event.get("max_vehicle_speed") or event.get("avg_vehicle_speed") or 0
    max_speed = round(max_speed_kmh * KMH_TO_MPH, 1)

    posted_kmh = event.get("min_posted_speed_limit_in_kph") or 0
    posted_speed = round(posted_kmh * KMH_TO_MPH, 1)

    over_kmh = event.get("max_over_speed_in_kph") or event.get("avg_over_speed_in_kph") or 0
    overspeed = round(over_kmh * KMH_TO_MPH, 1)

    if overspeed >= 20 or max_speed >= 90:
        tier = "RED"
    elif overspeed >= 15:
        tier = "ORANGE"
    else:
        tier = "YELLOW"

    vehicle_obj = event.get("vehicle", {})
    vehicle_number = vehicle_obj.get("number", "Unknown") if isinstance(vehicle_obj, dict) else str(vehicle_obj)

    driver_name = vehicle_drivers.get(vehicle_number)
    if not driver_name:
        drv = event.get("driver")
        if drv and isinstance(drv, dict):
            name = f"{drv.get('first_name', '')} {drv.get('last_name', '')}".strip()
            if name:
                driver_name = name
    if not driver_name:
        if " " in vehicle_number:
            candidate = vehicle_number.split(" ", 1)[1].strip().lstrip("- ")
            while candidate and candidate.split(" ", 1)[0].replace("-", "").isdigit():
                if " " in candidate:
                    candidate = candidate.split(" ", 1)[1].strip().lstrip("- ")
                else:
                    candidate = ""
            if len(candidate) > 2 and any(c.isalpha() for c in candidate):
                driver_name = candidate
    if not driver_name:
        driver_name = "Unknown"

    yard = vehicle_yards.get(vehicle_number, "")
    duration_str = _format_duration(event.get("duration", 0))

    timestamp = event.get("start_time") or event.get("end_time", "")
    formatted_time = _utc_to_central(timestamp)
    central_dt = _utc_to_central_dt(timestamp)
    is_weekend = central_dt.weekday() >= 5 if central_dt else False

    lat = event.get("start_lat")
    lon = event.get("start_lon")
    maps_link = f"https://www.google.com/maps?q={lat},{lon}" if lat and lon else ""

    return {
        "driver": driver_name, "vehicle": vehicle_number,
        "speed": max_speed, "posted_speed": posted_speed, "overspeed": overspeed,
        "duration": duration_str, "time": formatted_time, "maps_link": maps_link,
        "tier": tier, "yard": yard, "is_weekend": is_weekend,
        "central_dt": central_dt,
    }


# ==============================================================================
# MOTIVE API - CAMERA EVENTS (Weekly)
# ==============================================================================

def get_camera_events_weekly(start_ct, end_ct, vehicle_drivers, vehicle_yards, casing_vehicles):
    """Pull camera events for the full week, filtered to Casing vehicles."""
    start_utc = start_ct.astimezone(timezone.utc)
    end_utc = end_ct.astimezone(timezone.utc)

    api_start_date = start_utc.strftime("%Y-%m-%d")
    api_end_date = end_utc.strftime("%Y-%m-%d")

    print(f"    Central window: {start_ct.strftime('%m/%d/%Y %I:%M %p')} to {end_ct.strftime('%m/%d/%Y %I:%M %p')}")
    print(f"    API date filter: start_date={api_start_date}, end_date={api_end_date}")

    headers = {"X-Api-Key": MOTIVE_API_KEY}
    api_url = f"{MOTIVE_BASE_URL_V2}/driver_performance_events"
    raw_events = []
    page = 1
    page_cursor = None

    while True:
        params = {"per_page": 100, "start_date": api_start_date, "end_date": api_end_date}
        if page_cursor:
            params["page_cursor"] = page_cursor
        else:
            params["page_no"] = page

        try:
            resp = requests.get(api_url, headers=headers, params=params, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            events = data.get("driver_performance_events", [])
            if not events:
                break

            raw_events.extend(events)

            pag = data.get("pagination", {})
            next_cursor = pag.get("next_cursor") or pag.get("next_page_cursor")
            if next_cursor:
                page_cursor = next_cursor
                continue

            total = pag.get("total", 0)
            if total and page * 100 >= total:
                break
            if not total and len(events) < 100:
                break
            page += 1
        except Exception as e:
            print(f"    Error fetching camera events page {page}: {e}")
            break

    print(f"    Raw camera events fetched: {len(raw_events)}")

    filtered = []
    raw_event_types = Counter()

    for wrapper in raw_events:
        evt = wrapper.get("driver_performance_event", wrapper)
        raw_type = evt.get("type", "") or evt.get("event_type", "") or evt.get("behavior_type", "") or ""
        raw_event_types[raw_type] += 1

        evt_time_str = evt.get("start_time", "") or evt.get("event_time", "") or evt.get("created_at", "")
        try:
            evt_utc = datetime.fromisoformat(evt_time_str.replace("Z", "+00:00"))
            evt_central = evt_utc.astimezone(CENTRAL_TZ)
            if not (start_ct <= evt_central <= end_ct):
                continue
        except Exception:
            pass

        vehicle_obj = evt.get("vehicle", {})
        vehicle_number = vehicle_obj.get("number", "") if isinstance(vehicle_obj, dict) else str(vehicle_obj) if vehicle_obj else ""

        if vehicle_number and casing_vehicles and vehicle_number not in casing_vehicles:
            continue

        enriched = _enrich_camera_event(evt, vehicle_drivers, vehicle_yards, raw_type)
        filtered.append(enriched)

    if raw_event_types:
        type_summary = ", ".join(f"{t} ({c})" for t, c in raw_event_types.most_common())
        print(f"    Event types found: {type_summary}")

    print(f"    After Casing + Central Time filter: {_plural(len(filtered), 'event')}")
    return sorted(filtered, key=_event_sort_key)


def _enrich_camera_event(event, vehicle_drivers, vehicle_yards, raw_type):
    event_type = _normalize_event_type(raw_type)
    tier = _classify_tier(event_type)
    display_name = _event_display_name(event_type, raw_type)

    vehicle_obj = event.get("vehicle", {})
    vehicle_number = vehicle_obj.get("number", "Unknown") if isinstance(vehicle_obj, dict) else str(vehicle_obj) if vehicle_obj else "Unknown"

    driver_name = vehicle_drivers.get(vehicle_number)
    if not driver_name:
        drv = event.get("driver")
        if drv and isinstance(drv, dict):
            name = f"{drv.get('first_name', '')} {drv.get('last_name', '')}".strip()
            if name:
                driver_name = name
    if not driver_name:
        if " " in vehicle_number:
            candidate = vehicle_number.split(" ", 1)[1].strip().lstrip("- ")
            while candidate and candidate.split(" ", 1)[0].replace("-", "").isdigit():
                if " " in candidate:
                    candidate = candidate.split(" ", 1)[1].strip().lstrip("- ")
                else:
                    candidate = ""
            if len(candidate) > 2 and any(c.isalpha() for c in candidate):
                driver_name = candidate
    if not driver_name:
        driver_name = "Unknown"

    yard = vehicle_yards.get(vehicle_number, "")

    speed_kmh = event.get("start_speed") or event.get("max_speed") or event.get("end_speed") or 0
    try:
        speed_mph = round(float(speed_kmh) * KMH_TO_MPH, 1) if speed_kmh else 0
    except (ValueError, TypeError):
        speed_mph = 0

    duration_raw = event.get("duration") or event.get("duration_seconds") or 0
    duration_str = _format_duration(duration_raw)

    timestamp = event.get("start_time") or event.get("event_time") or event.get("created_at", "")
    formatted_time = _utc_to_central(timestamp)
    central_dt = _utc_to_central_dt(timestamp)
    is_weekend = central_dt.weekday() >= 5 if central_dt else False

    is_obstruction = raw_type.lower().strip().replace(" ", "_").replace("-", "_") in OBSTRUCTION_RAW_TYPES or event_type == "camera_obstruction"

    return {
        "driver": driver_name, "vehicle": vehicle_number,
        "event_type": event_type, "raw_type": raw_type,
        "display_name": display_name, "tier": tier,
        "speed": speed_mph, "duration": duration_str,
        "duration_raw": duration_raw,
        "time": formatted_time, "yard": yard,
        "is_weekend": is_weekend, "is_obstruction": is_obstruction,
        "central_dt": central_dt,
    }


# ==============================================================================
# KPA EHS API (Weekly)
# ==============================================================================

KPA_FORMS = {
    151085: "Observation Cards",
    151622: "Incident Report",
    381707: "CSG - Safety Casing Field Assessment",
}

ASSESSMENT_TARGET_PER_YARD = 3

# Map safety rep last names to their primary yard (for field assessment yard detection)
_OBSERVER_TO_YARD = {
    "salazar": "Midland",
    "hancock": "Midland",
    "conrad": "Bryan",
    "barnett": "Kilgore",
    "batts": "Hobbs",
    "speyrer": "Jourdanton",
}

# Map safety rep last names to their SAFETY_REP_YARDS key (for accountability)
_OBSERVER_TO_REP = {
    "salazar": "MICHAEL HANCOCK & MICHAEL SALAZAR",
    "hancock": "MICHAEL HANCOCK & MICHAEL SALAZAR",
    "conrad": "JUSTIN CONRAD",
    "barnett": "JAMES BARNETT (J.P.)",
    "batts": "ALLEN BATTS",
    "speyrer": "JOEY SPEYRER",
}

# Fields to skip when scanning for findings (metadata, not content)
_KPA_META_FIELDS = {
    'report number', 'date', 'observer', 'status', 'link', 'kpa_link',
    'name', 'Name', 'form', 'form_id', 'updated_at', 'created_at',
    'report', 'id', 'response_id',
}

# Phrases that indicate a POSITIVE observation (skip these)
_POSITIVE_PHRASES = [
    'no unsafe practices', 'good hand placement', 'good communication',
    'all good', 'no issues', 'doing a good job', 'no findings',
    'good job', 'no concerns', 'satisfactory', 'in compliance',
    'properly worn', 'properly secured', 'no deficiencies',
    'good housekeeping', 'well maintained', 'good condition',
    'no hazards', 'no violations', 'no corrective',
    'safe practices', 'following procedure', 'good practice',
    'proper hand placement', 'no members placing',
    'keeping hands out of hazardous', 'maintaining proper body positioning',
]

# Prefixes that indicate a purely positive observation (skip if no corrective keyword)
_POSITIVE_PREFIXES = ['proper ', 'no members ', 'no unsafe ']

# Words that confirm something was actually WRONG or CORRECTED (not just positive)
_CORRECTIVE_KEYWORDS = [
    'corrected', 'found', 'issue', 'should not', 'replaced', 'reminded',
    'had to', 'violated', 'failed', 'missing', 'damaged', 'broken',
    'not worn', 'not completed', 'not following', 'not in compliance',
    'improper', 'unsafe', 'need to', 'needs', 'disappointed',
    'hazard', 'deficien', 'violation', 'incomplete', 'expired',
]

# Keywords that indicate a finding vs. a clean pass
_FINDING_KEYWORDS = [
    'corrective', 'corrected', 'finding', 'found', 'hazard', 'deficien',
    'violation', 'issue', 'damaged', 'broken', 'missing', 'expired',
    'not completed', 'not worn', 'failed', 'needs repair', 'need to',
    'needs replacement', 'out of date', 'should not', 'replaced',
    'disappointed', 'no jsa', 'no ppe', 'not in compliance',
    'not following', 'improper', 'unsafe', 'leak', 'spill', 'trip',
    'not secured', 'unsecured', 'obstruct', 'crack', 'worn', 'frayed',
    'no fire', 'no extinguisher', 'no permit', 'incomplete',
]

# Keywords for categorizing findings
_EQUIP_KEYWORDS = ['equipment', 'vehicle', 'truck', 'trailer', 'tire', 'brake',
                    'light', 'engine', 'hydraulic', 'pump', 'hose', 'chain',
                    'tool', 'wrench', 'damaged', 'broken', 'repair', 'maintenance',
                    'defect', 'mechanical', 'gauge', 'pressure']
_BEHAVIOR_KEYWORDS = ['ppe', 'hard hat', 'glasses', 'gloves', 'vest', 'jsa',
                       'procedure', 'shortcut', 'compliance', 'not worn',
                       'not completed', 'not following', 'behavior', 'seatbelt',
                       'cell phone', 'speed', 'horseplay', 'training', 'cert']
_HOUSEKEEPING_KEYWORDS = ['housekeeping', 'clean', 'messy', 'trip', 'slip',
                          'rigging', 'spill', 'debris', 'clutter', 'organized',
                          'stacked', 'stored', 'site condition', 'ground',
                          'walk', 'path', 'access']
_DOC_KEYWORDS = ['permit', 'certification', 'paperwork', 'document', 'expired',
                 'inspection', 'checklist', 'log', 'record', 'sign', 'posted',
                 'label', 'sds', 'msds']


def _call_kpa(endpoint, params):
    url = f"{KPA_API_BASE}/{endpoint}"
    payload = {"token": KPA_API_TOKEN}
    payload.update(params)
    try:
        response = requests.post(url, json=payload, timeout=30)
        return response.text
    except Exception as e:
        print(f"    KPA API error: {e}")
        return None


def get_kpa_data_weekly(start_ct, end_ct):
    """Pull KPA incidents, observation cards, and field assessments for the week."""
    if not KPA_AVAILABLE:
        return {"incidents": [], "observations": [], "assessments": []}

    start_ms = int(start_ct.timestamp() * 1000)
    end_ms = int(end_ct.timestamp() * 1000)

    results = {"incidents": [], "observations": [], "assessments": []}

    for form_id, form_name in KPA_FORMS.items():
        print(f"    Pulling KPA {form_name} (form {form_id})...")

        params = {
            "form_id": form_id,
            "format": "csv",
            "updated_after": start_ms,
        }

        csv_text = _call_kpa("responses.flat", params)
        if not csv_text or csv_text.strip() == "":
            print(f"      No data returned for {form_name}")
            continue

        try:
            csv_file = StringIO(csv_text)
            reader = csv.DictReader(csv_file)
            rows = list(reader)

            filtered = []
            for row in rows:
                if row.get('report number') == 'Report Number':
                    continue
                date_str = row.get('date', '')
                try:
                    row_date = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                    row_date_ms = int(row_date.timestamp() * 1000)
                    if start_ms <= row_date_ms <= end_ms:
                        report_num = row.get('report number', '')
                        if report_num:
                            row['kpa_link'] = f"https://brhas-ees.kpaehs.com/forms/responses/view/{report_num}"
                        filtered.append(row)
                except Exception:
                    continue

            print(f"      {form_name}: {len(filtered)} in date range")

            if form_id == 151622:
                results["incidents"] = filtered
            elif form_id == 151085:
                results["observations"] = filtered
            elif form_id == 381707:
                results["assessments"] = filtered
                # Debug: print raw CSV headers and first 2 rows
                if filtered:
                    all_keys = list(filtered[0].keys())
                    print(f"      [DEBUG] Assessment CSV headers ({len(all_keys)} columns):")
                    for ki, k in enumerate(all_keys):
                        print(f"        [{ki}] {k}")
                    for ri, dbg_row in enumerate(filtered[:2]):
                        print(f"      [DEBUG] Row {ri} (report #{dbg_row.get('report number', '?')}):")
                        for k, v in dbg_row.items():
                            if v and str(v).strip():
                                print(f"        {k}: {str(v)[:120]}")
                    print(f"      [DEBUG] End assessment debug")

        except Exception as e:
            print(f"      Error parsing {form_name}: {e}")

    return results


def _is_casing_kpa(row):
    """Check if a KPA row belongs to Casing division."""
    for field_val in row.values():
        if isinstance(field_val, str):
            val_lower = field_val.lower()
            if "casing" in val_lower or "csg" in val_lower:
                return True
    return False


def _get_kpa_yard(row):
    """Try to extract yard from KPA row."""
    yard_field = row.get('7vj2l992y7fwqhwz', '') or row.get('yard', '') or row.get('location', '')
    for yard in YARD_ORDER:
        if yard.lower() in yard_field.lower():
            return yard
    return ""


def _get_kpa_observer(row):
    """Get the actual observer name from KPA row."""
    name = row.get('Name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name
    name = row.get('name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name
    observer = row.get('observer', '').strip()
    if observer and observer.lower() not in ['unknown', 'none', '']:
        return observer
    return 'Unknown'


def _get_assessment_observer(row):
    """Get who FILED the assessment (the safety rep).

    For field assessments, 'observer' = the safety rep who conducted it.
    'Name' may be the person/area being assessed, so check observer first.
    """
    observer = row.get('observer', '').strip()
    if observer and observer.lower() not in ['unknown', 'none', '']:
        return observer
    name = row.get('Name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name
    name = row.get('name', '').strip()
    if name and name.lower() not in ['none', 'unknown', '']:
        return name
    return 'Unknown'


def _map_observer_to_yard(observer_name):
    """Map an observer name to their primary yard using last name matching."""
    if not observer_name:
        return ""
    obs_lower = observer_name.lower().strip()
    for last_name, yard in _OBSERVER_TO_YARD.items():
        if last_name in obs_lower:
            return yard
    return ""


def _map_observer_to_rep(observer_name):
    """Map an observer name to their SAFETY_REP_YARDS key."""
    if not observer_name:
        return observer_name
    obs_lower = observer_name.lower().strip()
    for last_name, rep_key in _OBSERVER_TO_REP.items():
        if last_name in obs_lower:
            return rep_key
    return observer_name


def _extract_findings(row):
    """Scan all fields in a KPA assessment row for finding-like content.

    Returns list of finding strings. Empty list = clean assessment.

    Skips: URLs, KPA field codes, company/operator names, positive observations.
    Only includes fields that indicate something WRONG or CORRECTED.
    """
    meta_lower = {k.lower() for k in _KPA_META_FIELDS}
    findings = []
    for key, val in row.items():
        if not val or not isinstance(val, str):
            continue

        key_lower = key.lower().strip()

        # Skip metadata fields
        if key_lower in meta_lower:
            continue

        val_stripped = val.strip()
        if not val_stripped or len(val_stripped) < 5:
            continue

        val_lower = val_stripped.lower()

        # Skip URLs (image uploads, KPA links, etc.)
        if 'http://' in val_lower or 'https://' in val_lower:
            continue

        # Skip purely numeric or date values
        if val_stripped.replace('.', '').replace('-', '').replace('/', '').replace(':', '').replace(' ', '').isdigit():
            continue

        # Skip KPA internal field codes (random alphanumeric strings like "vonz52oh7281f36pc831")
        if len(val_stripped) < 30 and val_stripped.replace('_', '').replace('-', '').isalnum() and not any(c == ' ' for c in val_stripped):
            # Short single-token alphanumeric = likely a field code or ID, not a finding
            if not any(kw in val_lower for kw in _FINDING_KEYWORDS):
                continue

        # Skip common non-finding single-word values
        if val_lower in ('yes', 'no', 'n/a', 'na', 'pass', 'ok', 'good', 'safe',
                         'true', 'false', 'compliant', 'satisfactory', 'acceptable',
                         'not applicable', 'none', 'no issues', 'no findings',
                         'casing', 'csg', 'brhas', 'butch'):
            continue

        # Skip positive observations (exact phrase match)
        if any(phrase in val_lower for phrase in _POSITIVE_PHRASES):
            continue

        # Skip text that starts with positive prefixes unless it also
        # contains a corrective keyword (e.g. "proper" alone = positive,
        # but "proper PPE was not worn" = finding)
        if any(val_lower.startswith(pfx) for pfx in _POSITIVE_PREFIXES):
            if not any(ck in val_lower for ck in _CORRECTIVE_KEYWORDS):
                continue

        # Skip short company/operator name fields (operator, company, division fields)
        if key_lower in ('operator', 'company', 'division', 'department', 'location',
                         'site', 'yard', 'area', 'unit', 'rig', 'crew', 'shift'):
            continue

        # ONLY include if it has a corrective keyword — something was WRONG or CORRECTED
        has_corrective = any(ck in val_lower for ck in _CORRECTIVE_KEYWORDS)
        if has_corrective:
            findings.append(val_stripped)

    return findings


def _categorize_finding(text):
    """Categorize a finding into one of 4 categories."""
    text_lower = text.lower()
    scores = {
        "EQUIPMENT/VEHICLE ISSUES": sum(1 for kw in _EQUIP_KEYWORDS if kw in text_lower),
        "BEHAVIORAL/COMPLIANCE": sum(1 for kw in _BEHAVIOR_KEYWORDS if kw in text_lower),
        "HOUSEKEEPING/SITE CONDITIONS": sum(1 for kw in _HOUSEKEEPING_KEYWORDS if kw in text_lower),
        "DOCUMENTATION": sum(1 for kw in _DOC_KEYWORDS if kw in text_lower),
    }
    best = max(scores, key=scores.get)
    if scores[best] > 0:
        return best
    return "BEHAVIORAL/COMPLIANCE"  # default


def _get_assessment_status(row):
    """Determine assessment finding status."""
    for key, val in row.items():
        if not val or not isinstance(val, str):
            continue
        val_lower = val.lower().strip()
        if 'corrected on site' in val_lower or 'corrected on-site' in val_lower:
            return 'Corrected on site'
        if 'follow up' in val_lower or 'follow-up' in val_lower or 'requires follow' in val_lower:
            return 'Requires follow-up'
    status = (row.get('status', '') or '').strip()
    if status:
        return status
    return 'Open'


def analyze_field_assessments(assessments):
    """Analyze field assessments into with-findings and clean.

    Uses observer name to determine yard and safety rep for accountability.

    Returns dict with:
        with_findings: list of {yard, rep, date, report_num, link, findings, categories, status}
        clean: list of {yard, date, report_num}
        by_yard: {yard: count}
        by_rep: {rep_name: count}
    """
    with_findings = []
    clean = []
    by_yard = Counter()
    by_rep = Counter()

    for row in assessments:
        # Use assessment-specific observer (who filed it = the safety rep)
        rep = _get_assessment_observer(row)
        # Map observer to yard — use observer name, fallback to KPA field
        yard = _get_kpa_yard(row) or _map_observer_to_yard(rep)
        date = row.get('date', 'N/A')
        report_num = row.get('report number', 'N/A')
        link = row.get('kpa_link', '')

        by_yard[yard or 'Unknown'] += 1
        # Map observer to SAFETY_REP_YARDS key for accountability
        rep_key = _map_observer_to_rep(rep)
        by_rep[rep_key] += 1

        print(f"    [DEBUG] Assessment #{report_num}: observer='{rep}' -> yard='{yard}', rep_key='{rep_key}'")

        findings = _extract_findings(row)
        if findings:
            categorized = {}
            for f in findings:
                cat = _categorize_finding(f)
                categorized.setdefault(cat, []).append(f)

            status = _get_assessment_status(row)
            with_findings.append({
                "yard": yard or "Unknown", "rep": rep,
                "date": date, "report_num": report_num, "link": link,
                "findings": findings, "categories": categorized,
                "status": status, "observer": rep,
            })
        else:
            clean.append({
                "yard": yard or "Unknown", "date": date,
                "report_num": report_num,
            })

    return {
        "with_findings": with_findings,
        "clean": clean,
        "by_yard": dict(by_yard),
        "by_rep": dict(by_rep),
    }


# ==============================================================================
# ANALYSIS FUNCTIONS
# ==============================================================================

def analyze_red_flag_drivers(camera_events, speeding_events, kpa_incidents):
    """Cross-reference drivers across data sources."""
    driver_data = {}

    for evt in camera_events:
        name = evt["driver"]
        if name == "Unknown":
            continue
        driver_data.setdefault(name, {"camera": [], "speeding": [], "kpa": [], "vehicle": "", "yard": ""})
        driver_data[name]["camera"].append(evt)
        if evt["vehicle"]:
            driver_data[name]["vehicle"] = evt["vehicle"]
        if evt["yard"]:
            driver_data[name]["yard"] = evt["yard"]

    for evt in speeding_events:
        name = evt["driver"]
        if name == "Unknown":
            continue
        driver_data.setdefault(name, {"camera": [], "speeding": [], "kpa": [], "vehicle": "", "yard": ""})
        driver_data[name]["speeding"].append(evt)
        if evt["vehicle"]:
            driver_data[name]["vehicle"] = evt["vehicle"]
        if evt["yard"]:
            driver_data[name]["yard"] = evt["yard"]

    flagged = []
    for name, data in driver_data.items():
        cam_count = len(data["camera"])
        spd_count = len(data["speeding"])
        kpa_count = len(data["kpa"])

        is_flagged = False
        reasons = []

        if cam_count > 0 and spd_count > 0:
            is_flagged = True
            reasons.append("appears in both camera and speeding events")
        if cam_count >= 3:
            is_flagged = True
            reasons.append(f"{cam_count} camera events")
        if spd_count >= 5:
            is_flagged = True
            reasons.append(f"{spd_count} speeding events")
        if cam_count > 0 and kpa_count > 0:
            is_flagged = True
            reasons.append("camera event + KPA incident")

        if is_flagged:
            cam_types = Counter(e["display_name"] for e in data["camera"])
            cam_summary = ", ".join(f"{t} x{c}" for t, c in cam_types.most_common())

            spd_worst = max(data["speeding"], key=lambda x: x["overspeed"]) if data["speeding"] else None
            spd_summary = ""
            if spd_worst:
                spd_summary = f"{_plural(spd_count, 'event')}, worst: +{spd_worst['overspeed']} over at {spd_worst['speed']} mph"

            # Auto-generate recommended action
            action = _generate_action(data)

            total_events = cam_count + spd_count + kpa_count
            flagged.append({
                "name": name, "vehicle": data["vehicle"], "yard": data["yard"],
                "camera_count": cam_count, "camera_summary": cam_summary,
                "speeding_count": spd_count, "speeding_summary": spd_summary,
                "kpa_count": kpa_count, "total": total_events,
                "action": action,
            })

    return sorted(flagged, key=lambda x: x["total"], reverse=True)


def _generate_action(driver_data):
    cam_types = Counter(e["event_type"] for e in driver_data["camera"])
    fatigue_types = {"drowsiness", "lane_swerving"}
    distraction_types = {"distraction", "cell_phone"}

    if any(t in fatigue_types for t in cam_types):
        return "Pattern: fatigue — address scheduling and rest compliance"
    if any(t in distraction_types for t in cam_types):
        return "Pattern: distraction — formal coaching required"
    if len(driver_data["speeding"]) >= 5:
        return "Pattern: speed non-compliance — formal coaching required"
    if len(driver_data["camera"]) >= 3 and len(driver_data["speeding"]) > 0:
        return "Multiple safety categories — supervisor meeting required"
    return "Cross-source flags — safety rep to review and coach"


def analyze_time_buckets(camera_events):
    """Count camera events by time bucket."""
    buckets = {"6AM-12PM": 0, "12PM-6PM": 0, "6PM-12AM": 0, "12AM-6AM": 0}
    type_buckets = {}

    for evt in camera_events:
        dt = evt.get("central_dt")
        if not dt:
            continue
        hour = dt.hour
        if 6 <= hour < 12:
            bucket = "6AM-12PM"
        elif 12 <= hour < 18:
            bucket = "12PM-6PM"
        elif 18 <= hour < 24:
            bucket = "6PM-12AM"
        else:
            bucket = "12AM-6AM"
        buckets[bucket] += 1
        type_buckets.setdefault(evt["event_type"], Counter())[bucket] += 1

    notes = []
    # Check for drowsiness pattern in PM
    drowsy_pm = type_buckets.get("drowsiness", Counter())
    pm_count = drowsy_pm.get("12PM-6PM", 0) + drowsy_pm.get("6PM-12AM", 0)
    am_count = drowsy_pm.get("6AM-12PM", 0) + drowsy_pm.get("12AM-6AM", 0)
    if pm_count > am_count and pm_count >= 2:
        notes.append("Drowsiness events concentrated in afternoon/evening — consider scheduling adjustments")

    return buckets, notes


# ==============================================================================
# WORD DOCUMENT HELPERS
# ==============================================================================

def _set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, size_pt=8, bold=False, color=None, italic=False):
    run.font.name = CALIBRI
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="C00000"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_section_header(doc, title):
    _add_horizontal_rule(doc)
    p = doc.add_paragraph()
    run = p.add_run(title)
    _set_run_font(run, 14, bold=True, color=RGBColor(192, 0, 0))
    return p


def _tier_bg_hex(tier):
    if tier == "RED":
        return "FFE0E0"
    elif tier == "ORANGE":
        return "FFF0E0"
    return "FFFFF0"


def _tier_color(tier):
    if tier == "RED":
        return RGBColor(255, 0, 0)
    elif tier == "ORANGE":
        return RGBColor(255, 140, 0)
    return RGBColor(204, 153, 0)


# ==============================================================================
# BUILD WORD DOCUMENT
# ==============================================================================

def create_word_document(camera_events, speeding_events, kpa_data, yard_vehicle_counts,
                         start_date, end_date, assessment_analysis=None):
    doc = Document()

    # Landscape orientation
    for section in doc.sections:
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Logo - Butchs.jpg only, centered
    logo_path = os.path.join(LOGOS_DIR, "Butchs.jpg")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        try:
            para.add_run().add_picture(logo_path, width=Inches(2.0))
        except Exception:
            run = para.add_run("BRHAS Casing Division")
            _set_run_font(run, 16, bold=True, color=RGBColor(192, 0, 0))
    else:
        run = para.add_run("BRHAS Casing Division")
        _set_run_font(run, 16, bold=True, color=RGBColor(192, 0, 0))

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BRHAS CASING DIVISION")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WEEKLY SAFETY INTELLIGENCE BRIEFING")
    _set_run_font(run, 18, bold=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Week of {start_date.strftime('%B %d, %Y')} \u2013 {end_date.strftime('%B %d, %Y')}")
    _set_run_font(run, 12, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Monday Safety Meeting \u2014 2:00 PM CT")
    _set_run_font(run, 11, bold=True, color=RGBColor(128, 0, 0))

    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}")
    _set_run_font(run, 9, color=RGBColor(128, 0, 0))

    doc.add_paragraph()

    kpa_incidents = kpa_data.get("incidents", [])
    kpa_observations = kpa_data.get("observations", [])
    casing_incidents = [r for r in kpa_incidents if _is_casing_kpa(r)]
    casing_observations = [r for r in kpa_observations if _is_casing_kpa(r)]
    aa = assessment_analysis or {"with_findings": [], "clean": [], "by_yard": {}, "by_rep": {}}

    red_flags = analyze_red_flag_drivers(camera_events, speeding_events, casing_incidents)

    # ===== SECTION 1: WEEK AT A GLANCE =====
    _add_section_header(doc, "SECTION 1 \u2014 WEEK AT A GLANCE")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    headers = ["Yard", "Camera Events", "Speeding Events", "KPA Incidents", "Field Assessments", "Obs Cards"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        if table.rows[0].cells[i].paragraphs[0].runs:
            _set_run_font(table.rows[0].cells[i].paragraphs[0].runs[0], 8, bold=True)

    for yard in YARD_ORDER:
        yard_cam = [e for e in camera_events if e["yard"] == yard]
        yard_spd = [e for e in speeding_events if e["yard"] == yard]
        yard_inc = [r for r in casing_incidents if _get_kpa_yard(r) == yard]
        yard_obs = [r for r in casing_observations if _get_kpa_yard(r) == yard]
        yard_assess_ct = aa["by_yard"].get(yard, 0)

        cam_red = len([e for e in yard_cam if e["tier"] == "RED"])
        spd_red = len([e for e in yard_spd if e["tier"] == "RED"])

        cells = table.add_row().cells
        cells[0].text = yard

        if yard_cam:
            cells[1].text = f"{cam_red}/{len(yard_cam)}"
            _set_cell_shading(cells[1], _tier_bg_hex("RED") if cam_red else "FFFFFF")
        else:
            cells[1].text = "\u2014"

        if yard_spd:
            cells[2].text = f"{spd_red}/{len(yard_spd)}"
            _set_cell_shading(cells[2], _tier_bg_hex("RED") if spd_red else "FFFFFF")
        else:
            cells[2].text = "\u2014"

        cells[3].text = str(len(yard_inc)) if yard_inc else "0"
        cells[4].text = str(yard_assess_ct) if yard_assess_ct else "0"
        if yard_assess_ct < ASSESSMENT_TARGET_PER_YARD:
            _set_cell_shading(cells[4], "FFF0E0")
        obs_count = len(yard_obs)
        cells[5].text = str(obs_count)

        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_run_font(r, 8)

    doc.add_paragraph()

    # ===== SECTION 2: RED FLAG DRIVERS =====
    _add_section_header(doc, "SECTION 2 \u2014 RED FLAG DRIVERS")

    if red_flags:
        for flag in red_flags:
            p = doc.add_paragraph()
            run = p.add_run(f"{flag['name']}")
            _set_run_font(run, 11, bold=True, color=RGBColor(192, 0, 0))
            run2 = p.add_run(f" \u2014 {flag['vehicle']} \u2014 {flag['yard'] or 'Unknown Yard'}")
            _set_run_font(run2, 10)

            if flag["camera_count"]:
                p = doc.add_paragraph()
                run = p.add_run(f"  Camera: {_plural(flag['camera_count'], 'event')} \u2014 {flag['camera_summary']}")
                _set_run_font(run, 9)
            if flag["speeding_count"]:
                p = doc.add_paragraph()
                run = p.add_run(f"  Speeding: {flag['speeding_summary']}")
                _set_run_font(run, 9)
            if flag["kpa_count"]:
                p = doc.add_paragraph()
                run = p.add_run(f"  KPA Items: {flag['kpa_count']}")
                _set_run_font(run, 9)

            p = doc.add_paragraph()
            run = p.add_run(f"  Recommended: {flag['action']}")
            _set_run_font(run, 9, italic=True, color=RGBColor(128, 0, 0))
            doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run("No cross-source red flag drivers this week.")
        _set_run_font(run, 10, color=RGBColor(0, 128, 0))

    # ===== SECTION 3: CAMERA EVENT SUMMARY =====
    _add_section_header(doc, "SECTION 3 \u2014 CAMERA EVENT SUMMARY")

    cam_red = [e for e in camera_events if e["tier"] == "RED"]
    cam_orange = [e for e in camera_events if e["tier"] == "ORANGE"]
    cam_yellow = [e for e in camera_events if e["tier"] == "YELLOW"]

    p = doc.add_paragraph()
    run = p.add_run(f"Total Camera Events: {len(camera_events)}")
    _set_run_font(run, 11, bold=True)
    run2 = p.add_run(f" (RED: {len(cam_red)} | ORANGE: {len(cam_orange)} | YELLOW: {len(cam_yellow)})")
    _set_run_font(run2, 10)

    # Events by type table
    if camera_events:
        type_counts = Counter(e["display_name"] for e in camera_events)
        type_tiers = {}
        for e in camera_events:
            type_tiers[e["display_name"]] = e["tier"]

        p = doc.add_paragraph()
        run = p.add_run("Events by Type:")
        _set_run_font(run, 10, bold=True)

        etable = doc.add_table(rows=1, cols=3)
        etable.style = "Light Grid Accent 1"
        etable.autofit = True
        for i, h in enumerate(["Event Type", "Count", "Tier"]):
            etable.rows[0].cells[i].text = h
            _set_run_font(etable.rows[0].cells[i].paragraphs[0].runs[0], 8, bold=True)

        tier_order = {"RED": 0, "ORANGE": 1, "YELLOW": 2}
        sorted_types = sorted(type_counts.items(), key=lambda x: (tier_order.get(type_tiers.get(x[0], "ORANGE"), 1), -x[1]))

        for dtype, count in sorted_types:
            cells = etable.add_row().cells
            cells[0].text = dtype
            cells[1].text = str(count)
            tier = type_tiers.get(dtype, "ORANGE")
            cells[2].text = tier
            _set_cell_shading(cells[2], _tier_bg_hex(tier))
            for c in cells:
                for p2 in c.paragraphs:
                    for r in p2.runs:
                        _set_run_font(r, 8)

        doc.add_paragraph()

        # Events by yard
        p = doc.add_paragraph()
        run = p.add_run("Events by Yard:")
        _set_run_font(run, 10, bold=True)

        for yard in YARD_ORDER:
            yard_evts = [e for e in camera_events if e["yard"] == yard]
            if yard_evts:
                info = YARD_INFO.get(yard, {})
                yard_red = len([e for e in yard_evts if e["tier"] == "RED"])
                p = doc.add_paragraph()
                run = p.add_run(f"  {yard} ({info.get('safety_reps', '')}): {_plural(len(yard_evts), 'event')} (RED: {yard_red})")
                _set_run_font(run, 9)

        doc.add_paragraph()

        # Time-of-day analysis
        buckets, notes = analyze_time_buckets(camera_events)
        p = doc.add_paragraph()
        run = p.add_run("Time-of-Day Analysis:")
        _set_run_font(run, 10, bold=True)

        for bucket, count in buckets.items():
            if count:
                p = doc.add_paragraph()
                run = p.add_run(f"  {bucket}: {_plural(count, 'event')}")
                _set_run_font(run, 9)

        for note in notes:
            p = doc.add_paragraph()
            run = p.add_run(f"  Note: {note}")
            _set_run_font(run, 9, italic=True, color=RGBColor(192, 0, 0))

    doc.add_paragraph()

    # ===== SECTION 4: SPEEDING SUMMARY =====
    _add_section_header(doc, "SECTION 4 \u2014 SPEEDING SUMMARY (Casing Only)")

    spd_red = [e for e in speeding_events if e["tier"] == "RED"]
    spd_orange = [e for e in speeding_events if e["tier"] == "ORANGE"]
    spd_yellow = [e for e in speeding_events if e["tier"] == "YELLOW"]

    p = doc.add_paragraph()
    run = p.add_run(f"Total Speeding Events: {len(speeding_events)}")
    _set_run_font(run, 11, bold=True)
    run2 = p.add_run(f" (RED: {len(spd_red)} | ORANGE: {len(spd_orange)} | YELLOW: {len(spd_yellow)})")
    _set_run_font(run2, 10)

    if speeding_events:
        # Events by yard
        p = doc.add_paragraph()
        run = p.add_run("Events by Yard:")
        _set_run_font(run, 10, bold=True)
        for yard in YARD_ORDER:
            yard_evts = [e for e in speeding_events if e["yard"] == yard]
            if yard_evts:
                yard_red = len([e for e in yard_evts if e["tier"] == "RED"])
                p = doc.add_paragraph()
                run = p.add_run(f"  {yard}: {_plural(len(yard_evts), 'event')} (RED: {yard_red})")
                _set_run_font(run, 9)

        doc.add_paragraph()

        # Repeat speeders (3+)
        driver_counts = Counter(e["driver"] for e in speeding_events if e["driver"] != "Unknown")
        repeats = {n: c for n, c in driver_counts.items() if c >= 3}
        if repeats:
            p = doc.add_paragraph()
            run = p.add_run("Repeat Speeders (3+ events):")
            _set_run_font(run, 10, bold=True, color=RGBColor(192, 0, 0))
            for name in sorted(repeats, key=lambda n: max(e["overspeed"] for e in speeding_events if e["driver"] == n), reverse=True):
                worst = max((e for e in speeding_events if e["driver"] == name), key=lambda x: x["overspeed"])
                p = doc.add_paragraph()
                run = p.add_run(f"  {name}: {_plural(repeats[name], 'event')}, worst: +{worst['overspeed']} over at {worst['speed']} mph")
                _set_run_font(run, 9)
            doc.add_paragraph()

        # Worst violation
        worst = speeding_events[0]  # already sorted by overspeed desc
        p = doc.add_paragraph()
        run = p.add_run("Worst Violation of the Week:")
        _set_run_font(run, 10, bold=True, color=RGBColor(192, 0, 0))
        p = doc.add_paragraph()
        run = p.add_run(f"  {worst['driver']} \u2014 +{worst['overspeed']} mph over ({worst['speed']} mph in a {worst['posted_speed']} zone) \u2014 {worst['yard']} \u2014 {worst['time']}")
        _set_run_font(run, 9)

    doc.add_paragraph()

    # ===== SECTION 5: FIELD ASSESSMENTS & KPA =====
    _add_section_header(doc, "SECTION 5 \u2014 FIELD ASSESSMENTS, INCIDENTS & OBSERVATIONS")

    if not KPA_AVAILABLE:
        p = doc.add_paragraph()
        run = p.add_run("KPA data unavailable \u2014 API token not configured.")
        _set_run_font(run, 10, italic=True, color=RGBColor(192, 0, 0))
    else:
        # --- Incidents (brief) ---
        if casing_incidents:
            p = doc.add_paragraph()
            run = p.add_run(f"Incident Reports: {len(casing_incidents)}")
            _set_run_font(run, 11, bold=True)
            for inc in casing_incidents:
                report_num = inc.get('report number', 'N/A')
                form_name = inc.get('nojcquy0tfl9hqih', inc.get('report', 'Incident'))
                date = inc.get('date', 'N/A')
                yard = _get_kpa_yard(inc) or 'Unknown'
                link = inc.get('kpa_link', '')
                p = doc.add_paragraph()
                run = p.add_run(f"  #{report_num} \u2014 {form_name} \u2014 {date} \u2014 {yard}")
                _set_run_font(run, 9)
                if link:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(f"    {link}")
                    _set_run_font(run2, 8, color=RGBColor(0, 0, 180))
            doc.add_paragraph()

        # --- PART A: FINDINGS THAT NEED DISCUSSION ---
        p = doc.add_paragraph()
        run = p.add_run("PART A \u2014 FIELD ASSESSMENT FINDINGS THAT NEED DISCUSSION")
        _set_run_font(run, 11, bold=True, color=RGBColor(192, 0, 0))

        if aa["with_findings"]:
            for af in aa["with_findings"]:
                # One entry per assessment
                p = doc.add_paragraph()
                run = p.add_run(f"  Assessment #{af['report_num']} \u2014 {af['yard']} \u2014 {af['date']} \u2014 {af['rep']}")
                _set_run_font(run, 9, bold=True)

                status_color = RGBColor(0, 128, 0) if 'corrected' in af['status'].lower() else RGBColor(192, 0, 0)
                p = doc.add_paragraph()
                run = p.add_run(f"    Status: {af['status']}")
                _set_run_font(run, 8, bold=True, color=status_color)

                # List all findings under this assessment with category tags
                for cat, findings_list in af["categories"].items():
                    for finding in findings_list:
                        p = doc.add_paragraph()
                        run = p.add_run(f"    [{cat}] {finding[:200]}")
                        _set_run_font(run, 8)

                if af["link"]:
                    p = doc.add_paragraph()
                    run = p.add_run(f"    {af['link']}")
                    _set_run_font(run, 7, color=RGBColor(0, 0, 180))

                doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            run = p.add_run("  No findings requiring discussion this week \u2014 all assessments clean.")
            _set_run_font(run, 9, color=RGBColor(0, 128, 0))

        doc.add_paragraph()

        # --- PART B: ASSESSMENT ACCOUNTABILITY ---
        p = doc.add_paragraph()
        run = p.add_run("PART B \u2014 ASSESSMENT ACCOUNTABILITY")
        _set_run_font(run, 11, bold=True, color=RGBColor(192, 0, 0))

        # Count per safety rep
        p = doc.add_paragraph()
        run = p.add_run("  Assessments filed per safety rep:")
        _set_run_font(run, 10, bold=True)

        for rep_name, rep_yards in SAFETY_REP_YARDS.items():
            rep_count = sum(aa["by_yard"].get(y, 0) for y in rep_yards)
            yard_label = "/".join(rep_yards)
            warn = " \u26a0\ufe0f" if rep_count < ASSESSMENT_TARGET_PER_YARD * len(rep_yards) else ""
            p = doc.add_paragraph()
            run = p.add_run(f"    {rep_name} ({yard_label}): {rep_count} filed{warn}")
            _set_run_font(run, 9)

        doc.add_paragraph()

        # Target
        p = doc.add_paragraph()
        run = p.add_run(f"  Target: {ASSESSMENT_TARGET_PER_YARD} field assessments per yard per week")
        _set_run_font(run, 9, bold=True)

        # Missing assessments warning
        for yard in YARD_ORDER:
            yard_ct = aa["by_yard"].get(yard, 0)
            if yard_ct < ASSESSMENT_TARGET_PER_YARD:
                info = YARD_INFO.get(yard, {})
                rep = info.get("safety_reps", "safety rep")
                p = doc.add_paragraph()
                run = p.add_run(f"  {yard}: {yard_ct} filed (below target) \u2014 {rep} to address")
                _set_run_font(run, 9, bold=True, color=RGBColor(192, 0, 0))

        doc.add_paragraph()

        # Clean assessments summary
        if aa["clean"]:
            clean_yards = Counter(c["yard"] for c in aa["clean"])
            clean_parts = ", ".join(f"{y} x{c}" for y, c in clean_yards.most_common())
            p = doc.add_paragraph()
            run = p.add_run(f"  {_plural(len(aa['clean']), 'assessment')} filed with no findings ({clean_parts}) \u2014 Good work.")
            _set_run_font(run, 9, color=RGBColor(0, 128, 0))

        doc.add_paragraph()

        # --- Observation Cards ---
        p = doc.add_paragraph()
        run = p.add_run(f"Observation Cards: {len(casing_observations)}")
        _set_run_font(run, 11, bold=True)

        p = doc.add_paragraph()
        run = p.add_run(f"  Target: {OBS_TARGET_PER_YARD} per yard per week")
        _set_run_font(run, 9, bold=True)

        for yard in YARD_ORDER:
            yard_obs = [r for r in casing_observations if _get_kpa_yard(r) == yard]
            if len(yard_obs) == 0:
                p = doc.add_paragraph()
                run = p.add_run(f"  {yard}: No observation cards filed this week")
                _set_run_font(run, 9, bold=True, color=RGBColor(192, 0, 0))

    doc.add_paragraph()

    # ===== SECTION 6: OPEN ACTION ITEMS =====
    _add_section_header(doc, "SECTION 6 \u2014 OPEN ACTION ITEMS")

    if not KPA_AVAILABLE:
        p = doc.add_paragraph()
        run = p.add_run("KPA data unavailable \u2014 API token not configured.")
        _set_run_font(run, 10, italic=True, color=RGBColor(192, 0, 0))
    else:
        open_items = []
        for inc in casing_incidents:
            status = (inc.get('status', '') or '').lower()
            if status in ('open', 'in progress', 'in_progress', ''):
                open_items.append(inc)

        if open_items:
            for item in open_items:
                report_num = item.get('report number', 'N/A')
                form_name = item.get('nojcquy0tfl9hqih', item.get('report', 'Item'))
                status = item.get('status', 'Open')
                link = item.get('kpa_link', '')
                p = doc.add_paragraph()
                run = p.add_run(f"  #{report_num} \u2014 {form_name} \u2014 Status: {status}")
                _set_run_font(run, 9, color=RGBColor(192, 0, 0))
                if link:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(f"    {link}")
                    _set_run_font(run2, 8, color=RGBColor(0, 0, 180))
        else:
            p = doc.add_paragraph()
            run = p.add_run("No open action items \u2014 all current.")
            _set_run_font(run, 10, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    # ===== SECTION 7: WEEKEND SPOTLIGHT =====
    _add_section_header(doc, "SECTION 7 \u2014 WEEKEND SPOTLIGHT")

    p = doc.add_paragraph()
    run = p.add_run("Weekend events may not have been addressed yet \u2014 discuss on this call.")
    _set_run_font(run, 9, italic=True, color=RGBColor(128, 0, 0))

    weekend_cam = [e for e in camera_events if e["is_weekend"]]
    weekend_spd = [e for e in speeding_events if e["is_weekend"] and e["tier"] in ("RED", "ORANGE")]

    if weekend_cam:
        p = doc.add_paragraph()
        run = p.add_run(f"Weekend Camera Events: {len(weekend_cam)}")
        _set_run_font(run, 10, bold=True)
        for evt in weekend_cam:
            p = doc.add_paragraph()
            run = p.add_run(f"  [{evt['tier']}] {evt['display_name']} \u2014 {evt['driver']} \u2014 {evt['vehicle']} \u2014 {evt['yard']} \u2014 {evt['time']}")
            _set_run_font(run, 8)
            _set_run_font(run, 8, color=_tier_color(evt['tier']) if evt['tier'] == 'RED' else None)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No weekend camera events.")
        _set_run_font(run, 9, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    if weekend_spd:
        p = doc.add_paragraph()
        run = p.add_run(f"Weekend Speeding Events (RED/ORANGE only): {len(weekend_spd)}")
        _set_run_font(run, 10, bold=True)
        for evt in weekend_spd:
            p = doc.add_paragraph()
            run = p.add_run(f"  [{evt['tier']}] {evt['driver']} \u2014 +{evt['overspeed']} over ({evt['speed']} mph) \u2014 {evt['yard']} \u2014 {evt['time']}")
            _set_run_font(run, 8)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No RED/ORANGE weekend speeding events.")
        _set_run_font(run, 9, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    # ===== SECTION 8: AGENDA ASSIGNMENTS =====
    _add_section_header(doc, "SECTION 8 \u2014 AGENDA ASSIGNMENTS")

    for rep_name, rep_yards in SAFETY_REP_YARDS.items():
        rep_cam = [e for e in camera_events if e["yard"] in rep_yards]
        rep_spd = [e for e in speeding_events if e["yard"] in rep_yards]
        rep_obs = [r for r in casing_observations if _get_kpa_yard(r) in rep_yards]
        rep_inc = [r for r in casing_incidents if _get_kpa_yard(r) in rep_yards]
        rep_flags = [f for f in red_flags if f["yard"] in rep_yards]
        rep_assess_count = sum(aa["by_yard"].get(y, 0) for y in rep_yards)
        rep_assess_target = ASSESSMENT_TARGET_PER_YARD * len(rep_yards)

        # Collect findings for this rep's yards
        rep_findings = []
        for item in aa.get("with_findings", []):
            if item.get("yard") in rep_yards:
                rep_findings.append(item)

        has_data = rep_cam or rep_spd or rep_obs or rep_inc or rep_flags or rep_assess_count or rep_findings
        if not has_data:
            continue

        yard_label = " / ".join(rep_yards)
        p = doc.add_paragraph()
        run = p.add_run(f"{rep_name} \u2014 {yard_label}")
        _set_run_font(run, 11, bold=True, color=RGBColor(192, 0, 0))

        for flag in rep_flags:
            p = doc.add_paragraph()
            run = p.add_run(f"  RED FLAG: {flag['name']} \u2014 {flag['action']}")
            _set_run_font(run, 9, bold=True, color=RGBColor(192, 0, 0))

        spd_red = len([e for e in rep_spd if e["tier"] == "RED"])
        if rep_spd:
            p = doc.add_paragraph()
            run = p.add_run(f"  {_plural(len(rep_spd), 'speeding event')} this week, {spd_red} RED")
            _set_run_font(run, 9)

        if rep_cam:
            cam_types = Counter(e["display_name"] for e in rep_cam)
            type_str = ", ".join(f"{t}: {c}" for t, c in cam_types.most_common(5))
            p = doc.add_paragraph()
            run = p.add_run(f"  {_plural(len(rep_cam), 'camera event')} this week \u2014 {type_str}")
            _set_run_font(run, 9)

        # Field assessments per rep
        assess_warn = " \u26a0\ufe0f" if rep_assess_count < rep_assess_target else ""
        p = doc.add_paragraph()
        run = p.add_run(f"  Field assessments filed: {rep_assess_count} (target: {rep_assess_target}){assess_warn}")
        _set_run_font(run, 9, bold=True if rep_assess_count < rep_assess_target else False)

        if rep_findings:
            finding_briefs = []
            for af in rep_findings[:5]:
                for ft in af.get("findings", []):
                    finding_briefs.append(str(ft))
            findings_str = "; ".join(finding_briefs) if finding_briefs else "None"
            p = doc.add_paragraph()
            run = p.add_run(f"  Findings to address: {findings_str}")
            _set_run_font(run, 9, color=RGBColor(192, 0, 0))
        else:
            p = doc.add_paragraph()
            run = p.add_run("  Findings to address: None \u2014 all clean")
            _set_run_font(run, 9, color=RGBColor(0, 128, 0))

        for yard in rep_yards:
            yard_obs_count = len([r for r in rep_obs if _get_kpa_yard(r) == yard])
            p = doc.add_paragraph()
            run = p.add_run(f"  Observation cards filed ({yard}): {yard_obs_count} (target: {OBS_TARGET_PER_YARD})")
            _set_run_font(run, 9)

        if rep_inc:
            p = doc.add_paragraph()
            run = p.add_run(f"  KPA Incidents: {len(rep_inc)}")
            _set_run_font(run, 9)

        doc.add_paragraph()

    # ===== SECTION 9: VEHICLE HEALTH FLAGS =====
    _add_section_header(doc, "SECTION 9 \u2014 VEHICLE HEALTH FLAGS")

    obstruction_events = [e for e in camera_events if e["is_obstruction"]]
    if obstruction_events:
        vehicle_obstructions = {}
        for evt in obstruction_events:
            key = evt["vehicle"]
            vehicle_obstructions.setdefault(key, {"driver": evt["driver"], "yard": evt["yard"], "count": 0, "duration_total": 0})
            vehicle_obstructions[key]["count"] += 1
            try:
                vehicle_obstructions[key]["duration_total"] += int(evt.get("duration_raw", 0) or 0)
            except (ValueError, TypeError):
                pass

        for veh, info in sorted(vehicle_obstructions.items(), key=lambda x: x[1]["count"], reverse=True):
            dur = _format_duration(info["duration_total"])
            p = doc.add_paragraph()
            run = p.add_run(f"  {veh} \u2014 {info['driver']} \u2014 {info['yard']} \u2014 {_plural(info['count'], 'obstruction event')} \u2014 Total duration: {dur}")
            _set_run_font(run, 9)

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Note: Camera obstruction may indicate intentional blocking (disciplinary) or equipment damage (maintenance). Safety rep to investigate.")
        _set_run_font(run, 9, italic=True, color=RGBColor(128, 0, 0))
    else:
        p = doc.add_paragraph()
        run = p.add_run("No camera obstruction events this week.")
        _set_run_font(run, 10, color=RGBColor(0, 128, 0))

    doc.add_paragraph()

    # ===== SECTION 10: YARD COMPARISON SCORECARD =====
    _add_section_header(doc, "SECTION 10 \u2014 YARD COMPARISON SCORECARD")

    yard_scores = []
    for yard in YARD_ORDER:
        cam_count = len([e for e in camera_events if e["yard"] == yard])
        spd_count = len([e for e in speeding_events if e["yard"] == yard])
        total = cam_count + spd_count
        vehicles = yard_vehicle_counts.get(yard, 0)
        rate = round(total / vehicles, 2) if vehicles > 0 else 0
        yard_scores.append({
            "yard": yard, "vehicles": vehicles,
            "camera": cam_count, "speeding": spd_count,
            "total": total, "rate": rate,
        })

    yard_scores.sort(key=lambda x: x["rate"], reverse=True)

    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, h in enumerate(["Rank", "Yard", "Vehicles", "Camera", "Speeding", "Total", "Events/Vehicle"]):
        table.rows[0].cells[i].text = h
        _set_run_font(table.rows[0].cells[i].paragraphs[0].runs[0], 8, bold=True)

    for rank, ys in enumerate(yard_scores, 1):
        cells = table.add_row().cells
        cells[0].text = str(rank)
        cells[1].text = ys["yard"]
        cells[2].text = str(ys["vehicles"])
        cells[3].text = str(ys["camera"])
        cells[4].text = str(ys["speeding"])
        cells[5].text = str(ys["total"])
        cells[6].text = f"{ys['rate']:.2f}"

        if rank <= 2 and ys["total"] > 0:
            for c in cells:
                _set_cell_shading(c, "FFE0E0")

        for c in cells:
            for p2 in c.paragraphs:
                for r in p2.runs:
                    _set_run_font(r, 8)

    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END OF BRIEFING")
    _set_run_font(run, 10, italic=True, color=RGBColor(192, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Butch's Rat Hole & Anchor Service Inc. | Casing Division | HSE Department")
    _set_run_font(run, 9, color=RGBColor(128, 0, 0))

    return doc


# ==============================================================================
# BUILD HTML EMAIL
# ==============================================================================

C_RED = "#C00000"
C_DARK = "#800000"
C_AMBER = "#FF8C00"
C_YELLOW_DARK = "#CC9900"
C_GREEN = "#008000"


def _tier_colors_html(tier):
    if tier == "RED":
        return "#FF0000", "#FFE0E0"
    elif tier == "ORANGE":
        return C_AMBER, "#FFF0E0"
    return C_YELLOW_DARK, "#FFFFF0"


def build_html_report(camera_events, speeding_events, kpa_data, yard_vehicle_counts,
                      start_date, end_date, assessment_analysis=None):
    now_central = datetime.now(timezone.utc).astimezone(CENTRAL_TZ)

    kpa_incidents = kpa_data.get("incidents", [])
    kpa_observations = kpa_data.get("observations", [])
    casing_incidents = [r for r in kpa_incidents if _is_casing_kpa(r)]
    casing_observations = [r for r in kpa_observations if _is_casing_kpa(r)]
    red_flags = analyze_red_flag_drivers(camera_events, speeding_events, casing_incidents)
    aa = assessment_analysis or {"with_findings": [], "clean": [], "by_yard": {}, "by_rep": {}}

    parts = []

    # Header
    parts.append(f"""<html><head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#f4f4f4;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f4f4;">
<tr><td align="center">
<table width="750" cellpadding="0" cellspacing="0" style="background:#ffffff;border:1px solid #ddd;margin:20px auto;font-family:Calibri,Arial,Helvetica,sans-serif;font-size:14px;color:#333;">

<tr><td style="background:{C_RED};padding:30px 40px;text-align:center;">
  <div style="font-size:16px;font-weight:bold;color:#ffffff;letter-spacing:1px;">BRHAS CASING DIVISION</div>
  <div style="font-size:26px;font-weight:bold;color:#ffffff;margin:10px 0;">WEEKLY SAFETY INTELLIGENCE BRIEFING</div>
  <div style="font-size:13px;color:#ffcccc;">Monday Safety Meeting &mdash; 2:00 PM CT</div>
  <div style="font-size:12px;color:#ffffff;margin-top:8px;">Week of {start_date.strftime('%B %d, %Y')} &ndash; {end_date.strftime('%B %d, %Y')}</div>
  <div style="font-size:10px;color:#ffcccc;margin-top:4px;">Generated: {now_central.strftime('%B %d, %Y at %I:%M %p CT')}</div>
</td></tr>""")

    # S1: Week at a Glance
    glance_rows = ""
    for yard in YARD_ORDER:
        yard_cam = [e for e in camera_events if e["yard"] == yard]
        yard_spd = [e for e in speeding_events if e["yard"] == yard]
        yard_inc = [r for r in casing_incidents if _get_kpa_yard(r) == yard]
        yard_obs = [r for r in casing_observations if _get_kpa_yard(r) == yard]
        yard_assess_ct = aa["by_yard"].get(yard, 0)
        cam_red = len([e for e in yard_cam if e["tier"] == "RED"])
        spd_red = len([e for e in yard_spd if e["tier"] == "RED"])
        cam_cell = f"{cam_red}/{len(yard_cam)}" if yard_cam else "&mdash;"
        spd_cell = f"{spd_red}/{len(yard_spd)}" if yard_spd else "&mdash;"
        obs_warn = " &#9888;&#65039;" if len(yard_obs) == 0 else ""
        assess_warn = " &#9888;&#65039;" if yard_assess_ct < ASSESSMENT_TARGET_PER_YARD else ""
        assess_bg = "#FFF0E0" if yard_assess_ct < ASSESSMENT_TARGET_PER_YARD else "#fff"
        glance_rows += f"""<tr>
  <td style="padding:5px 8px;border:1px solid #ddd;">{_h(yard)}</td>
  <td style="padding:5px 8px;border:1px solid #ddd;text-align:center;background:{'#FFE0E0' if cam_red else '#fff'};">{cam_cell}</td>
  <td style="padding:5px 8px;border:1px solid #ddd;text-align:center;background:{'#FFE0E0' if spd_red else '#fff'};">{spd_cell}</td>
  <td style="padding:5px 8px;border:1px solid #ddd;text-align:center;">{len(yard_inc)}</td>
  <td style="padding:5px 8px;border:1px solid #ddd;text-align:center;background:{assess_bg};">{yard_assess_ct}{assess_warn}</td>
  <td style="padding:5px 8px;border:1px solid #ddd;text-align:center;">{len(yard_obs)}{obs_warn}</td>
</tr>"""

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 1 &mdash; WEEK AT A GLANCE</h2>
  <table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:12px;">
    <tr style="background:{C_RED};"><th style="padding:6px;color:#fff;border:1px solid #ddd;">Yard</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Camera (RED/total)</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Speeding (RED/total)</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">KPA Incidents</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Field Assessments</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Obs Cards</th></tr>
    {glance_rows}
  </table>
</td></tr>""")

    # S2: Red Flag Drivers
    flags_html = ""
    if red_flags:
        for flag in red_flags:
            flags_html += f'<div style="background:#FFE0E0;border-left:4px solid #FF0000;padding:12px 15px;margin:8px 0;">'
            flags_html += f'<b style="color:#FF0000;">{_h(flag["name"])}</b> &mdash; {_h(flag["vehicle"])} &mdash; {_h(flag["yard"] or "Unknown")}<br>'
            if flag["camera_count"]:
                flags_html += f'Camera: {_plural(flag["camera_count"], "event")} &mdash; {_h(flag["camera_summary"])}<br>'
            if flag["speeding_count"]:
                flags_html += f'Speeding: {_h(flag["speeding_summary"])}<br>'
            if flag["kpa_count"]:
                flags_html += f'KPA Items: {flag["kpa_count"]}<br>'
            flags_html += f'<i style="color:{C_DARK};">Recommended: {_h(flag["action"])}</i>'
            flags_html += '</div>'
    else:
        flags_html = f'<span style="color:{C_GREEN};">No cross-source red flag drivers this week.</span>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 2 &mdash; RED FLAG DRIVERS</h2>
  {flags_html}
</td></tr>""")

    # S3: Camera Summary
    cam_red_ct = len([e for e in camera_events if e["tier"] == "RED"])
    cam_orange_ct = len([e for e in camera_events if e["tier"] == "ORANGE"])
    cam_yellow_ct = len([e for e in camera_events if e["tier"] == "YELLOW"])

    cam_html = f"<b>Total: {len(camera_events)}</b> (RED: {cam_red_ct} | ORANGE: {cam_orange_ct} | YELLOW: {cam_yellow_ct})<br><br>"

    if camera_events:
        type_counts = Counter(e["display_name"] for e in camera_events)
        type_tiers = {e["display_name"]: e["tier"] for e in camera_events}
        cam_html += '<table width="100%" style="border-collapse:collapse;font-size:12px;margin:10px 0;"><tr style="background:#eee;"><th style="padding:4px 8px;border:1px solid #ddd;">Type</th><th style="padding:4px 8px;border:1px solid #ddd;">Count</th><th style="padding:4px 8px;border:1px solid #ddd;">Tier</th></tr>'
        tier_order = {"RED": 0, "ORANGE": 1, "YELLOW": 2}
        for dtype, count in sorted(type_counts.items(), key=lambda x: (tier_order.get(type_tiers.get(x[0], "ORANGE"), 1), -x[1])):
            tier = type_tiers.get(dtype, "ORANGE")
            tc, bg = _tier_colors_html(tier)
            cam_html += f'<tr><td style="padding:4px 8px;border:1px solid #ddd;">{_h(dtype)}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{count}</td><td style="padding:4px 8px;border:1px solid #ddd;background:{bg};color:{tc};text-align:center;font-weight:bold;">{tier}</td></tr>'
        cam_html += '</table>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 3 &mdash; CAMERA EVENT SUMMARY</h2>
  {cam_html}
</td></tr>""")

    # S4: Speeding Summary
    spd_red_ct = len([e for e in speeding_events if e["tier"] == "RED"])
    spd_orange_ct = len([e for e in speeding_events if e["tier"] == "ORANGE"])
    spd_yellow_ct = len([e for e in speeding_events if e["tier"] == "YELLOW"])

    spd_html = f"<b>Total: {len(speeding_events)}</b> (RED: {spd_red_ct} | ORANGE: {spd_orange_ct} | YELLOW: {spd_yellow_ct})<br><br>"

    if speeding_events:
        worst = speeding_events[0]
        spd_html += f'<div style="background:#FFE0E0;border-left:4px solid #FF0000;padding:10px 15px;margin:8px 0;"><b style="color:#FF0000;">Worst: +{worst["overspeed"]} mph over</b> ({worst["speed"]} in a {worst["posted_speed"]} zone)<br>{_h(worst["driver"])} &mdash; {_h(worst["yard"])} &mdash; {_h(worst["time"])}</div>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 4 &mdash; SPEEDING SUMMARY</h2>
  {spd_html}
</td></tr>""")

    # S5: KPA Incidents, Field Assessments & Observations
    kpa_html = ""
    if not KPA_AVAILABLE:
        kpa_html = f'<i style="color:{C_RED};">KPA data unavailable &mdash; API token not configured.</i>'
    else:
        # Incidents
        kpa_html += f"<b>Incidents: {len(casing_incidents)}</b><br>"
        for inc in casing_incidents:
            link = inc.get('kpa_link', '')
            kpa_html += f'<div style="background:#fff5f5;border-left:4px solid {C_RED};padding:8px 12px;margin:4px 0;">#{_h(inc.get("report number", ""))} &mdash; {_h(inc.get("date", ""))} &mdash; {_h(_get_kpa_yard(inc))}'
            if link:
                kpa_html += f' &mdash; <a href="{_h(link)}">View</a>'
            kpa_html += '</div>'

        # PART A: Findings That Need Discussion
        kpa_html += f'<br><div style="background:#f8f0f0;border:2px solid {C_RED};padding:15px;margin:15px 0;">'
        kpa_html += f'<b style="color:{C_RED};font-size:14px;">PART A &mdash; FINDINGS THAT NEED DISCUSSION</b><br><br>'

        findings_with = aa.get("with_findings", [])
        if findings_with:
            print(f"  [DEBUG] HTML S5 Part A: {len(findings_with)} assessments with findings")
            for dbg_af in findings_with[:2]:
                print(f"  [DEBUG]   #{dbg_af.get('report_num')}: {len(dbg_af.get('findings', []))} findings, categories={list(dbg_af.get('categories', {}).keys())}")

            # One entry per assessment, with findings listed below
            for af in findings_with:
                yard = _h(af.get("yard", ""))
                date = _h(af.get("date", ""))
                rep = _h(af.get("rep", ""))
                status = _h(af.get("status", ""))
                link = af.get("link", "")
                report_num = _h(af.get("report_num", ""))
                status_color = C_GREEN if "corrected" in status.lower() else C_RED

                kpa_html += f'<div style="border-left:4px solid {status_color};padding:8px 12px;margin:8px 0;font-size:12px;background:#fff5f5;">'
                kpa_html += f'<b>Assessment #{report_num}</b> &mdash; {yard} &mdash; {date} &mdash; {rep}<br>'
                kpa_html += f'<b style="color:{status_color};">Status: {status}</b>'
                if link:
                    kpa_html += f' &mdash; <a href="{_h(link)}">View in KPA</a>'
                kpa_html += '<br>'

                for cat, findings_list in af.get("categories", {}).items():
                    for finding_text in findings_list:
                        kpa_html += f'<div style="margin:3px 0 3px 15px;font-size:11px;">[{_h(cat)}] <i>{_h(str(finding_text)[:200])}</i></div>'

                kpa_html += '</div>'
        else:
            kpa_html += f'<span style="color:{C_GREEN};">No findings requiring discussion this week &mdash; all assessments clean.</span><br>'

        kpa_html += '</div>'

        # PART B: Assessment Accountability
        kpa_html += f'<div style="background:#f0f4f8;border:2px solid {C_DARK};padding:15px;margin:15px 0;">'
        kpa_html += f'<b style="color:{C_DARK};font-size:14px;">PART B &mdash; ASSESSMENT ACCOUNTABILITY</b><br><br>'

        kpa_html += '<b>Assessments filed per safety rep:</b><br>'
        for rep_name, rep_yards in SAFETY_REP_YARDS.items():
            rep_count = sum(aa["by_yard"].get(y, 0) for y in rep_yards)
            target = ASSESSMENT_TARGET_PER_YARD * len(rep_yards)
            yard_label = "/".join(rep_yards)
            warn = f' <span style="color:red;font-weight:bold;">&#9888;&#65039; Below target</span>' if rep_count < target else ""
            kpa_html += f'{_h(rep_name)} ({_h(yard_label)}): {rep_count} filed (target: {target}){warn}<br>'

        kpa_html += f'<br><b>Target: {ASSESSMENT_TARGET_PER_YARD} field assessments per yard per week</b><br>'

        for yard in YARD_ORDER:
            yard_ct = aa["by_yard"].get(yard, 0)
            if yard_ct < ASSESSMENT_TARGET_PER_YARD:
                info = YARD_INFO.get(yard, {})
                rep = info.get("safety_reps", "safety rep")
                kpa_html += f'<span style="color:red;font-weight:bold;">{_h(yard)}: {yard_ct} filed (below target) &mdash; {_h(rep)} to address</span><br>'

        if aa["clean"]:
            clean_yards = Counter(c["yard"] for c in aa["clean"])
            clean_parts = ", ".join(f"{y} x{c}" for y, c in clean_yards.most_common())
            kpa_html += f'<br><span style="color:{C_GREEN};">{_plural(len(aa["clean"]), "assessment")} filed with no findings ({_h(clean_parts)}) &mdash; Good work.</span><br>'

        kpa_html += '</div>'

        # Observation Cards
        kpa_html += f"<br><b>Observation Cards: {len(casing_observations)}</b><br>"
        kpa_html += f"<b style='font-size:12px;'>Target: {OBS_TARGET_PER_YARD} per yard per week</b><br>"
        for yard in YARD_ORDER:
            yobs = [r for r in casing_observations if _get_kpa_yard(r) == yard]
            warn = f' <span style="color:red;font-weight:bold;">&#9888;&#65039; ZERO filed</span>' if not yobs else ""
            kpa_html += f"{_h(yard)}: {len(yobs)}{warn}<br>"

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 5 &mdash; KPA INCIDENTS, FIELD ASSESSMENTS &amp; OBSERVATIONS</h2>
  {kpa_html}
</td></tr>""")

    # S6: Open Action Items
    open_html = ""
    if not KPA_AVAILABLE:
        open_html = f'<i style="color:{C_RED};">KPA data unavailable.</i>'
    else:
        open_items = [inc for inc in casing_incidents if (inc.get('status', '') or '').lower() in ('open', 'in progress', 'in_progress', '')]
        if open_items:
            for item in open_items:
                link = item.get('kpa_link', '')
                open_html += f'<div style="background:#fff5f5;border-left:4px solid {C_RED};padding:8px 12px;margin:4px 0;">#{_h(item.get("report number", ""))} &mdash; {_h(item.get("status", "Open"))}'
                if link:
                    open_html += f' &mdash; <a href="{_h(link)}">View</a>'
                open_html += '</div>'
        else:
            open_html = f'<span style="color:{C_GREEN};">No open action items &mdash; all current.</span>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 6 &mdash; OPEN ACTION ITEMS</h2>
  {open_html}
</td></tr>""")

    # S7: Weekend Spotlight
    weekend_cam = [e for e in camera_events if e["is_weekend"]]
    weekend_spd = [e for e in speeding_events if e["is_weekend"] and e["tier"] in ("RED", "ORANGE")]

    wknd_html = f'<i style="color:{C_DARK};">Weekend events may not have been addressed yet &mdash; discuss on this call.</i><br><br>'
    if weekend_cam:
        wknd_html += f"<b>Camera Events ({len(weekend_cam)}):</b><br>"
        for evt in weekend_cam:
            tc, bg = _tier_colors_html(evt["tier"])
            wknd_html += f'<div style="background:{bg};border-left:3px solid {tc};padding:4px 10px;margin:3px 0;font-size:12px;">[{evt["tier"]}] {_h(evt["display_name"])} &mdash; {_h(evt["driver"])} &mdash; {_h(evt["yard"])} &mdash; {_h(evt["time"])}</div>'
    if weekend_spd:
        wknd_html += f"<br><b>Speeding (RED/ORANGE) ({len(weekend_spd)}):</b><br>"
        for evt in weekend_spd:
            tc, bg = _tier_colors_html(evt["tier"])
            wknd_html += f'<div style="background:{bg};border-left:3px solid {tc};padding:4px 10px;margin:3px 0;font-size:12px;">[{evt["tier"]}] {_h(evt["driver"])} +{evt["overspeed"]} over &mdash; {_h(evt["yard"])} &mdash; {_h(evt["time"])}</div>'
    if not weekend_cam and not weekend_spd:
        wknd_html += f'<span style="color:{C_GREEN};">No significant weekend events.</span>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 7 &mdash; WEEKEND SPOTLIGHT</h2>
  {wknd_html}
</td></tr>""")

    # S8: Agenda Assignments
    agenda_html = ""
    for rep_name, rep_yards in SAFETY_REP_YARDS.items():
        rep_cam = [e for e in camera_events if e["yard"] in rep_yards]
        rep_spd = [e for e in speeding_events if e["yard"] in rep_yards]
        rep_obs = [r for r in casing_observations if _get_kpa_yard(r) in rep_yards]
        rep_inc = [r for r in casing_incidents if _get_kpa_yard(r) in rep_yards]
        rep_flags_list = [f for f in red_flags if f["yard"] in rep_yards]
        rep_assess_count = sum(aa["by_yard"].get(y, 0) for y in rep_yards)
        rep_assess_target = ASSESSMENT_TARGET_PER_YARD * len(rep_yards)

        # Collect findings for this rep's yards
        rep_findings = []
        for item in aa.get("with_findings", []):
            if item.get("yard") in rep_yards:
                rep_findings.append(item)

        has_data = rep_cam or rep_spd or rep_obs or rep_inc or rep_flags_list or rep_assess_count or rep_findings
        if not has_data:
            continue

        yard_label = " / ".join(rep_yards)
        agenda_html += f'<div style="background:#f8f0f0;border:2px solid {C_RED};padding:12px 15px;margin:10px 0;">'
        agenda_html += f'<b style="color:{C_RED};font-size:14px;">{_h(rep_name)} &mdash; {_h(yard_label)}</b><ul style="margin:5px 0;">'

        for flag in rep_flags_list:
            agenda_html += f'<li style="color:#FF0000;"><b>{_h(flag["name"])}</b> &mdash; {_h(flag["action"])}</li>'
        spd_red_count = len([e for e in rep_spd if e["tier"] == "RED"])
        if rep_spd:
            agenda_html += f'<li>{_plural(len(rep_spd), "speeding event")}, {spd_red_count} RED</li>'
        if rep_cam:
            cam_types = Counter(e["display_name"] for e in rep_cam)
            type_str = ", ".join(f"{t}: {c}" for t, c in cam_types.most_common(5))
            agenda_html += f'<li>{_plural(len(rep_cam), "camera event")} &mdash; {_h(type_str)}</li>'

        # Field assessments per rep
        assess_warn = ' <span style="color:red;font-weight:bold;">&#9888;&#65039;</span>' if rep_assess_count < rep_assess_target else ""
        agenda_html += f'<li><b>Field assessments filed: {rep_assess_count} (target: {rep_assess_target})</b>{assess_warn}</li>'

        if rep_findings:
            finding_briefs = []
            for af in rep_findings[:5]:
                for ft in af.get("findings", []):
                    finding_briefs.append(_h(str(ft)))
            findings_str = "; ".join(finding_briefs) if finding_briefs else "None"
            agenda_html += f'<li style="color:{C_RED};">Findings to address: {findings_str}</li>'
        else:
            agenda_html += f'<li style="color:{C_GREEN};">Findings to address: None &mdash; all clean</li>'

        for yard in rep_yards:
            yard_obs_ct = len([r for r in rep_obs if _get_kpa_yard(r) == yard])
            agenda_html += f'<li>Observation cards ({_h(yard)}): {yard_obs_ct} (target: {OBS_TARGET_PER_YARD})</li>'
        if rep_inc:
            agenda_html += f'<li>KPA Incidents: {len(rep_inc)}</li>'

        agenda_html += '</ul></div>'

    if not agenda_html:
        agenda_html = f'<span style="color:{C_GREEN};">No agenda items this week.</span>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 8 &mdash; AGENDA ASSIGNMENTS</h2>
  {agenda_html}
</td></tr>""")

    # S9: Vehicle Health Flags
    obstruction_events = [e for e in camera_events if e["is_obstruction"]]
    obs_html = ""
    if obstruction_events:
        vehicle_obs = {}
        for evt in obstruction_events:
            key = evt["vehicle"]
            vehicle_obs.setdefault(key, {"driver": evt["driver"], "yard": evt["yard"], "count": 0, "dur": 0})
            vehicle_obs[key]["count"] += 1
            try:
                vehicle_obs[key]["dur"] += int(evt.get("duration_raw", 0) or 0)
            except (ValueError, TypeError):
                pass

        for veh, info in sorted(vehicle_obs.items(), key=lambda x: x[1]["count"], reverse=True):
            obs_html += f'<div style="background:#FFF0E0;border-left:3px solid {C_AMBER};padding:6px 12px;margin:4px 0;font-size:12px;">{_h(veh)} &mdash; {_h(info["driver"])} &mdash; {_h(info["yard"])} &mdash; {_plural(info["count"], "event")} &mdash; {_format_duration(info["dur"])}</div>'
        obs_html += f'<br><i style="color:{C_DARK};font-size:11px;">Camera obstruction may indicate intentional blocking (disciplinary) or equipment damage (maintenance). Safety rep to investigate.</i>'
    else:
        obs_html = f'<span style="color:{C_GREEN};">No camera obstruction events this week.</span>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 9 &mdash; VEHICLE HEALTH FLAGS</h2>
  {obs_html}
</td></tr>""")

    # S10: Yard Comparison
    yard_scores = []
    for yard in YARD_ORDER:
        cam_count = len([e for e in camera_events if e["yard"] == yard])
        spd_count = len([e for e in speeding_events if e["yard"] == yard])
        total = cam_count + spd_count
        vehicles = yard_vehicle_counts.get(yard, 0)
        rate = round(total / vehicles, 2) if vehicles > 0 else 0
        yard_scores.append({"yard": yard, "vehicles": vehicles, "camera": cam_count, "speeding": spd_count, "total": total, "rate": rate})
    yard_scores.sort(key=lambda x: x["rate"], reverse=True)

    score_rows = ""
    for rank, ys in enumerate(yard_scores, 1):
        bg = "#FFE0E0" if rank <= 2 and ys["total"] > 0 else "#fff"
        score_rows += f'<tr style="background:{bg};"><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{rank}</td><td style="padding:4px 8px;border:1px solid #ddd;">{_h(ys["yard"])}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{ys["vehicles"]}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{ys["camera"]}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{ys["speeding"]}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;">{ys["total"]}</td><td style="padding:4px 8px;border:1px solid #ddd;text-align:center;font-weight:bold;">{ys["rate"]:.2f}</td></tr>'

    parts.append(f"""
<tr><td style="padding:25px 40px;">
  <h2 style="color:{C_RED};margin:0 0 15px 0;font-size:18px;border-bottom:2px solid {C_RED};padding-bottom:5px;">SECTION 10 &mdash; YARD COMPARISON SCORECARD</h2>
  <table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:12px;">
    <tr style="background:{C_RED};"><th style="padding:6px;color:#fff;border:1px solid #ddd;">Rank</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Yard</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Vehicles</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Camera</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Speeding</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Total</th><th style="padding:6px;color:#fff;border:1px solid #ddd;">Events/Vehicle</th></tr>
    {score_rows}
  </table>
</td></tr>""")

    # Footer
    parts.append(f"""
<tr><td style="background:{C_DARK};padding:20px 40px;text-align:center;">
  <div style="color:#ffffff;font-size:11px;font-style:italic;">END OF BRIEFING</div>
  <div style="color:#ffcccc;font-size:10px;margin-top:4px;">Butch's Rat Hole &amp; Anchor Service Inc. | Casing Division | HSE Department</div>
</td></tr>

</table>
</td></tr></table>
</body></html>""")

    return "\n".join(parts)


# ==============================================================================
# SEND EMAIL
# ==============================================================================

def send_email_report(html_body, docx_path, start_date, end_date):
    gmail_address = os.environ.get("GMAIL_ADDRESS", "")
    gmail_app_password = os.environ.get("GMAIL_APP_PASSWORD", "")
    recipient = os.environ.get("REPORT_RECIPIENT", "")

    if not gmail_address or not gmail_app_password or not recipient:
        print("  Email skipped — GMAIL_ADDRESS, GMAIL_APP_PASSWORD, or REPORT_RECIPIENT not set.")
        return

    subject = f"Casing Weekly Safety Briefing - Week of {start_date.strftime('%B %d, %Y')}"

    try:
        msg = MIMEMultipart("mixed")
        msg["From"] = gmail_address
        msg["To"] = recipient
        msg["Subject"] = subject

        msg.attach(MIMEText(html_body, "html"))

        if os.path.exists(docx_path):
            with open(docx_path, "rb") as f:
                part = MIMEBase(
                    "application",
                    "vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f'attachment; filename="{os.path.basename(docx_path)}"',
            )
            msg.attach(part)

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_address, gmail_app_password)
            server.sendmail(gmail_address, recipient, msg.as_string())

        print(f"  Email sent to {recipient}")
    except Exception as e:
        print(f"  Email failed: {e}")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    start_ct, end_ct, start_date, end_date = get_week_range()

    print("\n" + "=" * 80)
    print("WEEKLY CASING DIVISION SAFETY INTELLIGENCE BRIEFING")
    print(f"Week of: {start_date.strftime('%A, %B %d, %Y')} through {end_date.strftime('%A, %B %d, %Y')}")
    print("=" * 80)
    print(f"\n  Monday 2 PM Safety Meeting Prep")
    print(f"  Data Sources: Motive Speeding + Motive Camera + KPA EHS")
    print(f"  Division: Casing Only (7 yards)\n")

    print("[1] Building Casing vehicle/driver lookup from Motive...")
    vehicle_drivers, vehicle_yards, casing_vehicles, yard_vehicle_counts = build_casing_vehicle_lookup()
    print(f"    {len(casing_vehicles)} Casing vehicles found")
    print(f"    {len(vehicle_drivers)} with driver names")
    print(f"    {len(vehicle_yards)} with yard assignments")
    for yard in YARD_ORDER:
        print(f"      {yard}: {yard_vehicle_counts.get(yard, 0)} vehicles")

    print("\n[2] Fetching speeding events from Motive (7-day window)...")
    speeding_events = get_speeding_events_weekly(start_ct, end_ct, vehicle_drivers, vehicle_yards, casing_vehicles)
    print(f"    {_plural(len(speeding_events), 'speeding event')} total")

    print("\n[3] Fetching camera events from Motive (7-day window)...")
    camera_events = get_camera_events_weekly(start_ct, end_ct, vehicle_drivers, vehicle_yards, casing_vehicles)
    print(f"    {_plural(len(camera_events), 'camera event')} total")

    print("\n[4] Fetching KPA EHS data (7-day window)...")
    kpa_data = get_kpa_data_weekly(start_ct, end_ct)
    kpa_incidents = kpa_data.get("incidents", [])
    kpa_observations = kpa_data.get("observations", [])
    casing_incidents = [r for r in kpa_incidents if _is_casing_kpa(r)]
    casing_observations = [r for r in kpa_observations if _is_casing_kpa(r)]
    print(f"    Casing incidents: {len(casing_incidents)}")
    print(f"    Casing observations: {len(casing_observations)}")

    print("\n[5] Analyzing red flag drivers...")
    red_flags = analyze_red_flag_drivers(camera_events, speeding_events, casing_incidents)
    print(f"    {_plural(len(red_flags), 'red flag driver')}")
    for flag in red_flags:
        print(f"      {flag['name']}: {flag['total']} total events — {flag['action']}")

    print("\n[5b] Analyzing field assessments...")
    kpa_assessments = kpa_data.get("assessments", [])
    casing_assessments = [r for r in kpa_assessments if _is_casing_kpa(r)]
    assessment_analysis = analyze_field_assessments(casing_assessments)
    print(f"    {_plural(len(casing_assessments), 'field assessment')} total")
    print(f"    {_plural(len(assessment_analysis['with_findings']), 'assessment')} with findings")
    print(f"    {_plural(len(assessment_analysis['clean']), 'assessment')} clean")

    print("\n[6] Creating Word document (landscape)...")
    doc = create_word_document(camera_events, speeding_events, kpa_data,
                               yard_vehicle_counts, start_date, end_date,
                               assessment_analysis=assessment_analysis)

    output_file = f"WeeklyCasingBriefing_{start_date.strftime('%Y-%m-%d')}.docx"
    doc.save(output_file)
    print(f"    Saved: {output_file}")

    print("\n[7] Building HTML email...")
    html_body = build_html_report(camera_events, speeding_events, kpa_data,
                                   yard_vehicle_counts, start_date, end_date,
                                   assessment_analysis=assessment_analysis)

    print("[8] Sending email...")
    send_email_report(html_body, output_file, start_date, end_date)

    print("\n" + "=" * 80)
    print("COMPLETE")
    print("=" * 80 + "\n")


if __name__ == "__main__":
    main()
